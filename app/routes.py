
from asyncio import open_connection
import calendar
import json
import time
import os
import socket
import ssl
from urllib import parse
from urllib.parse import urlsplit
from flask import jsonify, render_template, flash, redirect, render_template_string, session, url_for, request, send_file, Response, stream_with_context, abort, current_app
from flask_login import login_user, logout_user, current_user, login_required
import sqlalchemy as sa
from docx import Document
import re
from app import app, db, csrf
from app.forms import EventForm, LoginForm, PerfomanceTargetsForm, RegistrationForm, BookAllocationForm, ReportForm, WorkspaceForm, ProjectForm, TaskForm, CSVUploadForm, ChampionCSVUploadForm, ChampionSchoolForm, AkelloSimEventForm
from app.models import PerfomanceTargets, Scorecard, User, BookAllocations, BookAllocationRequest, Report, Workspace, Project, Task, ChampionSchool, ChampionSchoolRequest, SchoolVisitLog, Event, WeeklyReport, TaskA, TaskASubtask, ColumnA, ProjectA, TaskAttachment, TaskAComment, TaskAActivity, TaskALabel, TaskACustomFieldValue, ProjectACustomField, TaskADependency, AkelloSimEvent, UserActivity, ActiveSession, PageAnalytics, WorkspaceFile, Lesson, ActivityQuestion, CollateralItems, CollateralRequest, GameUser, Game, GameScore, Notification, Ticket, Ticket, AuditLog, AppSetting, taska_comment_mentions
from datetime import datetime, timezone, timedelta, date
from collections import Counter
from collections import defaultdict
from flask import Blueprint
from sqlalchemy import and_, create_engine, select, text
from sqlalchemy.orm import joinedload
import pymysql
from werkzeug.utils import secure_filename
import os
import io
from sshtunnel import SSHTunnelForwarder
from sqlalchemy.exc import IntegrityError
import paramiko
import logging
import csv
import pandas as pd
import secrets
import string
import uuid
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
import mysql.connector
from mysql.connector import Error
import seaborn as sns
import random
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
import pandas as pd
import plotly.graph_objs as go
import plotly.io as pio
import asyncio
import aiomysql
from dbutils.pooled_db import PooledDB
import atexit
from collections import defaultdict
# Add this import at the top of your file
from psycopg2 import sql
from flask_caching import Cache

# Email + token utils for password reset
import smtplib
from email.message import EmailMessage
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
import uuid
from zoneinfo import ZoneInfo
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from openai import OpenAI

# Load .env variables
# Try to load from project root directory
from pathlib import Path
env_path = Path(__file__).parent.parent / '.env'
load_dotenv(dotenv_path=env_path)
# Also try loading from current directory (fallback)
load_dotenv()




# # ---------- DIRECT CONNECTION POOL ----------
# # ---------- LAZY-LOADED CONNECTION POOLS ----------
# Initialize pools as None - will be created on first use
ruzivo_pool = None
direct_library_pool = None

def get_ruzivo_pool():
    """Get or create Ruzivo connection pool"""
    global ruzivo_pool
    if ruzivo_pool is None:
        try:
            ruzivo_pool = PooledDB(
                creator=pymysql,
                maxconnections=5,     # Reduced for stability
                mincached=0,          # No minimum cached connections
                maxcached=3,
                maxshared=2,
                maxusage=1000,
                blocking=False,
                setsession=[],
                ping=1,               # Only ping on checkout
                host='40.90.237.225',
                port=33000,
                user='qmashava',
                password='#Qhava@@!Fu11',
                database='ruzivo_2017',
                cursorclass=pymysql.cursors.DictCursor,
                autocommit=True,
                connect_timeout=30,
                read_timeout=60,
                write_timeout=60,
                charset='utf8mb4',
            )
            app.logger.info("Ruzivo connection pool initialized successfully")
        except Exception as e:
            app.logger.error(f"Failed to initialize Ruzivo pool: {e}")
            ruzivo_pool = None
    return ruzivo_pool

def get_library_pool():
    """Get or create Library connection pool"""
    global direct_library_pool
    if direct_library_pool is None:
        try:
            direct_library_pool = PooledDB(
                creator=pymysql,
                maxconnections=5,
                mincached=0,
                maxcached=3,
                maxshared=2,
                maxusage=1000,
                blocking=False,
                setsession=[],
                ping=1,
                host='40.88.149.15',
                port=33000,
                user='kmudzimuirema',
                password='Ak3110$2022',
                database='akello_library',
                cursorclass=pymysql.cursors.DictCursor,
                autocommit=True,
                connect_timeout=30,
                read_timeout=60,
                write_timeout=60,
                charset='utf8mb4',
            )
            app.logger.info("Library connection pool initialized successfully")
        except Exception as e:
            app.logger.error(f"Failed to initialize Library pool: {e}")
            direct_library_pool = None
    return direct_library_pool



# ---------- Helper functions ----------
# ---------- Improved Helper functions with error handling ----------
def get_direct_library_conn():
    """Get a library database connection - direct connection for help desk"""
    try:
        import pymysql
        
        # Use credentials that have access from this IP
        # Note: User qmashava doesn't have access from your current IP (77.246.50.217)
        host = '40.88.149.15'
        port = 33000
        user = 'kmudzimuirema'
        password = 'Ak3110$2022'
        database = 'akello_library'
        
        app.logger.info(f"Connecting to library database at {host}:{port}/{database} as {user}")
        
        return pymysql.connect(
            host=host,
            port=port,
            user=user,
            password=password,
            database=database,
            cursorclass=pymysql.cursors.DictCursor,
            autocommit=True,
            connect_timeout=30,
            read_timeout=60,
            write_timeout=60,
            charset='utf8mb4'
        )
    except Exception as e:
        app.logger.error(f"Failed to get library connection: {e}")
        raise

def get_ruzivo_conn():
    """Get a ruzivo database connection with error handling and timeout"""
    try:
        pool = get_ruzivo_pool()
        if pool is None:
            raise Exception("Ruzivo connection pool not available")
        
        conn = pool.connection()
        if conn is None:
            raise Exception("Failed to get ruzivo database connection from pool")
        return conn
    except Exception as e:
        app.logger.error(f"Error getting ruzivo database connection: {e}")
        # Try to create a direct connection as fallback
        try:
            import pymysql
            return pymysql.connect(
                host='40.90.237.225',
                port=33000,
                user='qmashava',
                password='#Qhava@@!Fu11',
                database='ruzivo_2017',
                cursorclass=pymysql.cursors.DictCursor,
                autocommit=True,
                connect_timeout=30,
                read_timeout=60,
                write_timeout=60,
                charset='utf8mb4'
            )
        except Exception as fallback_error:
            app.logger.error(f"Fallback connection also failed: {fallback_error}")
            raise Exception("Unable to establish ruzivo database connection")

def safe_db_execute(conn, query, params=None, fetch_one=False, fetch_all=True):
    """Safely execute database queries with proper error handling and connection cleanup"""
    cursor = None
    try:
        cursor = conn.cursor()
        cursor.execute(query, params)
        
        if fetch_one:
            return cursor.fetchone()
        elif fetch_all:
            return cursor.fetchall()
        else:
            return cursor.rowcount
            
    except Exception as e:
        app.logger.error(f"Database query error: {e}")
        if conn:
            conn.rollback()
        raise e
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

def write_debug_log(log_entry):
    """Write debug log entry to logs/debug.log"""
    try:
        # Get project root (akellodashboard directory)
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        logs_dir = os.path.join(project_root, 'logs')
        
        # Create logs directory if it doesn't exist
        os.makedirs(logs_dir, exist_ok=True)
        
        # Write to debug.log
        debug_log_path = os.path.join(logs_dir, 'debug.log')
        with open(debug_log_path, 'a', encoding='utf-8') as f:
            f.write(json.dumps(log_entry) + '\n')
    except Exception:
        # Silently fail - app.logger is already called separately
        pass



@app.before_request
def before_request():
    if current_user.is_authenticated:
        current_user.last_seen = datetime.now(timezone.utc)
        db.session.commit()



# global ruzivo_conn
# global library_conn
# Connections will be created on-demand to avoid startup failures
# ruzivo_conn = get_ruzivo_conn()
# library_conn = get_direct_library_conn()



# Initialize OpenAI client
# client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize cache for API endpoints
cache = Cache(app)

@app.route("/generate_lesson", methods=["POST"])
def generate_lesson():
    """Legacy endpoint for lesson plan generation - now uses Ollama instead of OpenAI"""
    data = request.get_json()
    topic = data.get("topic")
    objectives = data.get("objectives", [])
    age = data.get("age")
    subject = data.get("subject", "")
    duration = data.get("duration", 45)
    class_size = data.get("class_size", 25)
    notes = data.get("notes", "")

    if not topic or not age:
        return jsonify({"error": "Please provide topic and age"}), 400

    objectives_text = ", ".join(objectives) if isinstance(objectives, list) else objectives
    
    # Build subject context
    subject_context = f" for {subject}" if subject else ""
    
    # Build duration context
    duration_context = f"The lesson should be designed for {duration} minutes."
    
    # Build class size context
    class_context = f"Consider a class size of approximately {class_size} students."
    
    # Build notes context
    notes_context = f"\n\nAdditional requirements: {notes}" if notes else ""

    # Enhanced prompt construction
    prompt = f"""
    You are an experienced British teacher. Create a detailed, age-appropriate lesson plan{subject_context}.

    Topic: {topic}
    Learner age: {age} years old
    Subject: {subject if subject else "General"}
    Duration: {duration} minutes
    Class size: {class_size} students
    Learning objectives: {objectives_text if objectives_text else "To understand and explain the key concepts of the topic"}{notes_context}

    {duration_context} {class_context}

    The lesson plan should include:
    - A clear, engaging title
    - Learning objectives (specific, measurable, and age-appropriate)
    - An engaging introduction/warm-up activity (5-10% of lesson time)
    - Main content broken into logical sections with timing
    - At least 3 varied exercises or activities suitable for the class size
    - Assessment opportunities
    - A brief summary/plenary at the end
    - Required materials/resources
    - Differentiation strategies for different ability levels

    Use British spelling and educational terminology. Make it practical and actionable for teachers.

    Please return the result in structured JSON format like this:
    {{
      "title": "",
      "objectives": [],
      "introduction": "",
      "content": [],
      "exercises": [],
      "assessment": "",
      "summary": "",
      "materials": [],
      "differentiation": ""
    }}
    """

    # Ollama config
    base_url = os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434').rstrip('/')
    model = os.getenv('OLLAMA_MODEL') or 'llama3.1'

    # Try HTTP API first
    content = ''
    http_error = None
    headers = {}
    authz = os.getenv('OLLAMA_AUTHORIZATION')
    api_key = os.getenv('OLLAMA_API_KEY')
    if authz:
        headers['Authorization'] = authz
    elif api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    try:
        resp = requests.post(f"{base_url}/api/generate", json={
            'model': model,
            'prompt': prompt,
            'stream': False
        }, timeout=120, headers=headers)
        resp.raise_for_status()
        payload = resp.json()
        content = payload.get('response') or payload.get('message') or ''
        if not content:
            content = payload.get('data') or ''
    except Exception as e:
        http_error = e
    # If HTTP failed or returned empty, try HTTP chat endpoint
    if not content:
        try:
            resp2 = requests.post(f"{base_url}/api/chat", json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'You are a professional British lesson plan designer with expertise in creating engaging, age-appropriate educational content.'},
                    {'role': 'user', 'content': prompt}
            ],
                'stream': False
            }, timeout=120, headers=headers)
            resp2.raise_for_status()
            payload2 = resp2.json()
            content = payload2.get('message', {}).get('content') or ''
        except Exception:
            pass
    # If HTTP failed or returned empty, try Python library fallback
    if not content:
        try:
            from ollama import Client
            client = Client(host=base_url)
            res = client.generate(model=model, prompt=prompt, stream=False)
            content = res.get('response') or res.get('message', {}).get('content') or ''
            if not content:
                res2 = client.chat(model=model, messages=[
                    {'role': 'system', 'content': 'You are a professional British lesson plan designer with expertise in creating engaging, age-appropriate educational content.'},
                    {'role': 'user', 'content': prompt}
                ])
                content = res2.get('message', {}).get('content') or ''
        except Exception as e:
            if http_error:
                return jsonify({'error': f'Failed to generate from Ollama HTTP ({http_error}); and Python client also failed: {e}'}), 502
            return jsonify({'error': f'Failed to generate from Ollama Python client: {e}'}), 502

    if not content:
        return jsonify({'error': 'Failed to generate lesson plan from Ollama'}), 502

        return jsonify({"lesson": content})



@app.route('/analyze_users_csv', methods=['POST'])
@login_required
def analyze_users_csv():
    """Analyze user CSV for duplicates and validate columns before uploading"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if not file or not file.filename:
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.lower().endswith('.csv'):
        return jsonify({'error': 'Please upload a .csv file'}), 400

    try:
        # Parse CSV
        stream = io.TextIOWrapper(file.stream, encoding='utf-8')
        reader = csv.DictReader(stream)
        
        # Required columns for user upload
        required_columns = ['firstname', 'lastname', 'email', 'userRole', 'department', 'province']
        csv_columns = reader.fieldnames or []
        
        # Check for missing columns
        missing_columns = [col for col in required_columns if col not in csv_columns]
        if missing_columns:
            return jsonify({
                'error': 'Missing required columns',
                'missing_columns': missing_columns,
                'required_columns': required_columns,
                'found_columns': csv_columns
            }), 400
        
        # Track duplicates
        email_duplicates = []  # Emails that exist in DB
        csv_email_duplicates = []  # Duplicate emails within CSV
        seen_emails = {}
        
        users_to_create = []
        row_num = 1  # Start at 1 (header is row 0)
        
        for row in reader:
            row_num += 1
            email = (row.get('email') or '').strip().lower()
            firstname = (row.get('firstname') or '').strip()
            lastname = (row.get('lastname') or '').strip()
            
            if not email:
                continue
            
            # Check for duplicates within CSV
            if email in seen_emails:
                csv_email_duplicates.append({
                    'email': email,
                    'firstname': firstname,
                    'lastname': lastname,
                    'rows': [seen_emails[email], row_num]
                })
            else:
                seen_emails[email] = row_num
            
            # Check if email exists in database
            existing_user = User.query.filter_by(email=email).first()
            if existing_user:
                email_duplicates.append({
                    'email': email,
                    'firstname': firstname,
                    'lastname': lastname,
                    'existing_user': f"{existing_user.firstname} {existing_user.lastname}",
                    'existing_username': existing_user.username,
                    'row': row_num
                })
            else:
                users_to_create.append({
                    'email': email,
                    'firstname': firstname,
                    'lastname': lastname,
                    'userRole': row.get('userRole', '').strip(),
                    'department': row.get('department', '').strip(),
                    'province': row.get('province', '').strip()
                })
        
        has_duplicates = len(email_duplicates) > 0 or len(csv_email_duplicates) > 0
        
        return jsonify({
            'success': True,
            'has_duplicates': has_duplicates,
            'csv_email_duplicates': csv_email_duplicates,
            'existing_email_duplicates': email_duplicates,
            'total_new_users': len(users_to_create),
            'total_duplicates': len(email_duplicates)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/upload_users_confirmed', methods=['POST'])
@login_required
def upload_users_confirmed():
    """Upload users CSV with skip options"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    skip_emails = request.form.get('skip_emails', '[]')
    
    try:
        skip_emails = json.loads(skip_emails)
    except:
        skip_emails = []

    try:
        stream = io.TextIOWrapper(file.stream, encoding='utf-8')
        reader = csv.DictReader(stream)
        
        created = 0
        skipped = 0
        errors = []
        
        for row in reader:
            try:
                firstname = row.get('firstname', '').strip()
                lastname = row.get('lastname', '').strip()
                email = row.get('email', '').strip().lower()
                role = row.get('userRole', '').strip()
                department = row.get('department', '').strip()
                province = row.get('province', '').strip()

                if not (firstname and lastname and email and role and department and province):
                    continue
                
                # Check if email should be skipped
                if email in skip_emails:
                    skipped += 1
                    continue

                # Generate unique username
                base_username = f"{firstname.lower()}{lastname.lower()}"
                username = base_username
                count = 1
                while User.query.filter_by(username=username).first():
                    username = f"{base_username}{count}"
                    count += 1

                # Create user
                password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
                user = User(
                    username=username,
                    firstname=firstname,
                    lastname=lastname,
                    email=email,
                    userRole=role,
                    department=department,
                    province=province
                )
                user.set_password(password)
                db.session.add(user)
                created += 1
                
            except Exception as e:
                errors.append(str(e))
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Upload complete! Created: {created}, Skipped: {skipped}',
            'created': created,
            'skipped': skipped,
            'errors': errors
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/upload_users', methods=['GET', 'POST'])
@login_required
def upload_users():
    form = CSVUploadForm()
    if form.validate_on_submit():
        file = form.csv_file.data
        stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
        reader = csv.DictReader(stream)

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['First Name', 'Last Name', 'Email', 'Username', 'Password', 'Status'])

        for row in reader:
            try:
                firstname = row.get('firstname', '').strip()
                lastname = row.get('lastname', '').strip()
                email = row.get('email', '').strip()
                role = row.get('userRole', '').strip()
                department = row.get('department', '').strip()
                province = row.get('province', '').strip()

                if not (firstname and lastname and email and role and department and province):
                    writer.writerow([firstname, lastname, email, '', '', 'Missing Fields'])
                    continue

                base_username = f"{firstname.lower()}{lastname.lower()}"
                username = base_username
                count = 1
                while User.query.filter_by(username=username).first():
                    username = f"{base_username}{count}"
                    count += 1

                if User.query.filter_by(email=email).first():
                    writer.writerow([firstname, lastname, email, username, '', 'Email Exists'])
                    continue

                password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
                user = User(
                    username=username,
                    firstname=firstname,
                    lastname=lastname,
                    email=email,
                    userRole=role,
                    department=department,
                    province=province
                )
                user.set_password(password)
                print(user)
                db.session.add(user)
                db.session.commit()

                writer.writerow([firstname, lastname, email, username, password, 'Created'])
                print(f"[INFO] Created user: {username} ({email})")

            except Exception as e:
                db.session.rollback()
                print("Error while creating user:", e)
                traceback.print_exc()
                writer.writerow([firstname, lastname, email, '', '', f'Error: {str(e)}'])

        output.seek(0)
        return send_file(
            io.BytesIO(output.getvalue().encode()),
            mimetype='text/csv',
            as_attachment=True,
            download_name='registered_users_report.csv'
        )
    return redirect(url_for('settings'))




@app.route('/welcome', methods=['GET'])
@login_required
def welcome():
    return render_template('overview.html', title='Welcome')


@app.route('/index', methods=['GET', 'POST'])
@login_required
def index():
    today = datetime.now().date()
    all_events = Event.query.order_by(Event.start_date.desc()).all()
    updated = False

    for event in all_events:
        if event.status in ['Cancelled', 'Deleted', 'Event Ended']:
            continue  # skip protected statuses

        if event.start_date.date() <= today <= event.end_date.date():
            if event.status != "In Progress":
                event.status = "In Progress"
                updated = True
        elif today > event.end_date.date():
            if event.status != "Event Ended":
                event.status = "Event Ended"
                updated = True

    if updated:
        db.session.commit()

    form = EventForm()
    return render_template(
        'index.html',
        events=all_events,
        form=form,
        forms={event.id: EventForm(obj=event) for event in all_events},
        title='Home'
    )






@app.route('/api/events')
@login_required
def api_events():
    events = Event.query.all()
    event_list = []
    for event in events:
        event_list.append({
            'id': event.id,
            'title': event.title,
            'start': event.start_date.isoformat(),
            'end': event.end_date.isoformat(),
            'color': '#1E90FF' if event.status == 'Confirmed' else ('#FFA500' if event.status == 'In Progress' else '#DC143C')
        })
    return jsonify(event_list)



@app.route('/event/add', methods=['GET', 'POST'])
@login_required
def add_event():
    form = EventForm()
    if form.validate_on_submit():
        new_event = Event(
            title=form.title.data,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            status=form.status.data,
            request_collateral=form.request_collateral.data,
            added_by=current_user.username
        )
        db.session.add(new_event)
        db.session.commit()
        return redirect(url_for('index'))

@app.route('/event/edit/<int:event_id>', methods=['GET', 'POST'])
@login_required
def edit_event(event_id):
    event = Event.query.get_or_404(event_id)
    form = EventForm(obj=event)
    if form.validate_on_submit():
        event.title = form.title.data
        event.start_date = form.start_date.data
        event.end_date = form.end_date.data
        event.status = form.status.data
        event.request_collateral = form.request_collateral.data
        event.added_by = current_user.username
        db.session.commit()
        return redirect(url_for('index'))

    




@app.route("/admin/users/<int:user_id>/grant", methods=["POST"])
@login_required
def grant_privilege(user_id):
    if not current_user.userRole == "Admin":
        return jsonify({"error": "Unauthorized"}), 403

    privilege = request.json.get("privilege")
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    user.add_privilege(privilege)
    return jsonify({"message": f"Privilege '{privilege}' granted to {user.username}."})


@app.route("/admin/users/<int:user_id>/revoke", methods=["POST"])
@login_required
def revoke_privilege(user_id):
    if not current_user.userRole == "Admin":
        return jsonify({"error": "Unauthorized"}), 403

    privilege = request.json.get("privilege")
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    user.remove_privilege(privilege)
    return jsonify({"message": f"Privilege '{privilege}' revoked from {user.username}."})



@app.route("/secure-area")
@login_required
def secure_area():
    if not current_user.has_privilege("can_view_reports"):
        return "Access Denied", 403
    return "Welcome to secure reports!"



@app.route("/admin/users")
@login_required
def manage_users():
    if not current_user.userRole == "Admin":
        return "Unauthorized", 403

    users = User.query.all()
    return render_template("admin_users.html", users=users)


@app.route('/administration', methods=['GET'])
@login_required
def administration():
    if current_user.userRole != 'Admin':
        return "Unauthorized", 403

    users = User.query.order_by(User.username.asc()).all()
    csvform = CSVUploadForm()
    championcsvform = ChampionCSVUploadForm()

    # Build champion_schools rows from ChampionSchool model
    champion_rows = []
    champs = ChampionSchool.query.all()
    # Build an index to find username by firstname/lastname/province (best-effort)
    users_index = {}
    for u in users:
        key = (u.firstname or '').strip().lower(), (u.lastname or '').strip().lower(), (u.province or '').strip().lower()
        users_index[key] = u.username

    for c in champs:
        schools = []
        try:
            schools = c.get_schools() or []
        except Exception:
            schools = []
        champ_name = f"{c.firstname} {c.lastname}".strip()
        uname = users_index.get(((c.firstname or '').strip().lower(), (c.lastname or '').strip().lower(), (c.province or '').strip().lower()), '')
        if not schools:
            champion_rows.append({
                'id': c.id,
                'champion': champ_name,
                'username': uname,
                'school_name': '',
                'province': c.province,
                'asl_school_id': '',
                'library_school_id': ''
            })
        else:
            for s in schools:
                champion_rows.append({
                    'id': c.id,
                    'champion': champ_name,
                    'username': uname,
                    'school_name': s.get('school_name',''),
                    'province': c.province,
                    'asl_school_id': s.get('asl_school_id',''),
                    'library_school_id': s.get('library_school_id','')
                })

    return render_template(
        'administration.html',
        users=users,
        champion_schools=champion_rows,
        csvform=csvform,
        championcsvform=championcsvform,
        title='Administration'
    )


def _ensure_helpdesk_resolved_at_column():
    """Add helpdesk_queries.resolved_at if missing (e.g. SQLite DB never migrated). Avoids OperationalError when loading User.assigned_queries."""
    try:
        from sqlalchemy import inspect
        inspector = inspect(db.engine)
        if "helpdesk_queries" not in inspector.get_table_names():
            return
        columns = [c["name"] for c in inspector.get_columns("helpdesk_queries")]
        if "resolved_at" in columns:
            return
        # Add missing column (SQLite and most DBs support ADD COLUMN)
        db.session.execute(text("ALTER TABLE helpdesk_queries ADD COLUMN resolved_at DATETIME"))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        # Ignore "duplicate column" or similar; column may have been added by another request
        if "duplicate" not in str(e).lower() and "already exists" not in str(e).lower():
            print(f"ensure helpdesk resolved_at: {e}")


@app.route('/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def admin_delete_user(user_id):
    if current_user.userRole != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404
    try:
        _ensure_helpdesk_resolved_at_column()
        # Unlink user from helpdesk assignees via raw SQL so we never load HelpDeskQuery
        # (avoids resolved_at / schema mismatch when DB is not fully migrated)
        try:
            db.session.execute(text("DELETE FROM query_assignees WHERE user_id = :uid"), {"uid": user_id})
            db.session.commit()
        except Exception:
            db.session.rollback()
        # Now safe to delete the user (query_assignees already cleared)
        db.session.delete(user)
        db.session.commit()
        return jsonify({"ok": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


@app.route('/admin/users/<int:user_id>/reset_password', methods=['POST'])
@login_required
def admin_reset_password(user_id):
    if current_user.userRole != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404
    data = request.get_json(silent=True) or {}
    new_password = data.get('password')
    if not new_password:
        # Generate a secure random password if not provided
        import secrets, string
        alphabet = string.ascii_letters + string.digits
        new_password = ''.join(secrets.choice(alphabet) for _ in range(12))
    user.set_password(new_password)
    db.session.commit()
    return jsonify({"message": "Password reset successfully"}), 200


@app.route('/admin/users/<int:user_id>', methods=['PATCH'])
@login_required
def admin_update_user(user_id):
    if current_user.userRole != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404
    data = request.get_json(silent=True) or {}
    # Allowed updatable fields
    for field in ['username','email','firstname','lastname','userRole','department','province']:
        if field in data and data[field] is not None:
            setattr(user, field, data[field])
    try:
        db.session.commit()
        return jsonify({"message": "User updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


@app.route("/admin/users/<int:user_id>/toggle_privilege", methods=["POST"])
@login_required
def toggle_privilege(user_id):
    if current_user.userRole != "Admin":
        return jsonify({"error": "Unauthorized"}), 403

    data = request.get_json()
    privilege = data.get("privilege")
    value = data.get("value", False)

    if not privilege:
        return jsonify({"error": "Privilege name is required"}), 400

    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    try:
        user.set_privilege(privilege, value)
        return jsonify({
            "message": f"Privilege '{privilege}' set to {value} for {user.username}.",
            "privilege": privilege,
            "value": value
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Failed to update privilege: {str(e)}"}), 500


# Champion Schools admin delete endpoint
@app.route('/admin/champion_schools/<int:cs_id>', methods=['DELETE'])
@login_required
def admin_delete_champion_school(cs_id):
    if current_user.userRole != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403
    rec = db.session.get(ChampionSchool, cs_id)
    if not rec:
        return jsonify({"error": "Record not found"}), 404
        
    asl_id = request.args.get('asl')
    lib_id = request.args.get('lib')
    
    try:
        if asl_id or lib_id:
            # Granular deletion: remove specific school from the list
            schools = rec.get_schools() or []
            new_schools = []
            for s in schools:
                s_asl = str(s.get('asl_school_id') or '').strip()
                s_lib = str(s.get('library_school_id') or '').strip()
                
                # Check for match (handling potential empty strings/nulls)
                is_match = False
                if asl_id and s_asl == asl_id: is_match = True
                if lib_id and s_lib == lib_id: is_match = True
                
                if not is_match:
                    new_schools.append(s)
            
            if not new_schools:
                # If no schools left, delete the entire record
                db.session.delete(rec)
            else:
                rec.set_schools(new_schools)
        else:
            # Fallback/Legacy: Delete the entire champion record
            db.session.delete(rec)
            
        db.session.commit()
        return jsonify({"ok": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


# Champion Schools admin update endpoint (inline save)
@app.route('/admin/champion_schools/<int:cs_id>', methods=['PATCH'])
@login_required
def admin_update_champion_school(cs_id):
    if current_user.userRole != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403

    rec = db.session.get(ChampionSchool, cs_id)
    if not rec:
        return jsonify({"error": "Record not found"}), 404

    data = request.get_json() or {}

    # Optionally update champion identity
    for field in ['firstname','lastname','province']:
        if field in data and data[field] is not None:
            setattr(rec, field, data[field])

    school_name = (data.get('school_name') or '').strip()
    asl_id_new = (data.get('asl_school_id') or '').strip()
    lib_id_new = (data.get('library_school_id') or '').strip()
    asl_id_orig = (data.get('original_asl_school_id') or '').strip()
    lib_id_orig = (data.get('original_library_school_id') or '').strip()

    schools = rec.get_schools() or []

    # Enhanced matching logic with multiple strategies
    match_idx = None
    
    # Strategy 1: Match by original IDs (most reliable)
    if asl_id_orig or lib_id_orig:
        for idx, s in enumerate(schools):
            s_asl = str(s.get('asl_school_id') or '').strip()
            s_lib = str(s.get('library_school_id') or '').strip()
            # Match if BOTH IDs match OR if one matches and the other is empty/0
            asl_match = (asl_id_orig and s_asl == asl_id_orig)
            lib_match = (lib_id_orig and s_lib == lib_id_orig)
            
            if asl_match and lib_match:
                match_idx = idx
                break
            elif asl_match and (not lib_id_orig or not s_lib or s_lib == '0'):
                match_idx = idx
                break
            elif lib_match and (not asl_id_orig or not s_asl or s_asl == '0'):
                match_idx = idx
                break
    
    # Strategy 2: Match by school name if provided and no ID match
    if match_idx is None and school_name:
        for idx, s in enumerate(schools):
            if (s.get('school_name') or '').strip().lower() == school_name.lower():
                match_idx = idx
                break
    
    # Strategy 3: If still no match, try matching by new IDs
    if match_idx is None and (asl_id_new or lib_id_new):
        for idx, s in enumerate(schools):
            s_asl = str(s.get('asl_school_id') or '').strip()
            s_lib = str(s.get('library_school_id') or '').strip()
            if (asl_id_new and s_asl == asl_id_new) or (lib_id_new and s_lib == lib_id_new):
                match_idx = idx
                break

    # Update or append
    if match_idx is not None:
        # Preserve existing school entry and only update provided fields
        entry = dict(schools[match_idx])
        if school_name:
            entry['school_name'] = school_name
        # Only update IDs if they're actually provided (not empty)
        if asl_id_new:
            entry['asl_school_id'] = asl_id_new
        if lib_id_new:
            entry['library_school_id'] = lib_id_new
        schools[match_idx] = entry
    else:
        # Append a new mapping if at least one ID provided
        if asl_id_new or lib_id_new or school_name:
            schools.append({
                'school_name': school_name,
                'asl_school_id': asl_id_new,
                'library_school_id': lib_id_new
            })

    rec.set_schools(schools)

    # Prepare updated exact entry for precise UI syncing
    updated_entry = None
    if match_idx is not None:
        updated_entry = schools[match_idx]
    else:
        # If appended, use the last element (or the dict we just added)
        if schools:
            updated_entry = schools[-1]

    try:
        db.session.commit()
        return jsonify({
            "updated": updated_entry,
            "champion": rec.to_dict()
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


# -------- Champion school requests (admin approve/decline) --------
@app.route('/api/admin/champion-school-requests', methods=['GET'])
@login_required
def api_admin_champion_school_requests_list():
    """List champion school requests; optional ?status=Pending. Requires Admin or Approve Champion Schools privilege."""
    if not can_approve_champion_schools():
        return jsonify({"error": "Unauthorized"}), 403
    status = request.args.get('status')
    q = ChampionSchoolRequest.query
    if status:
        q = q.filter_by(status=status)
    q = q.order_by(ChampionSchoolRequest.created_at.desc())
    requests = q.all()
    return jsonify({'requests': [r.to_dict() for r in requests]})


@app.route('/api/admin/champion-school-requests/<int:request_id>/approve', methods=['POST'])
@login_required
def api_admin_champion_school_requests_approve(request_id):
    """Approve a champion school request: add school to champion profile."""
    if not can_approve_champion_schools():
        return jsonify({"error": "Unauthorized"}), 403
    req = ChampionSchoolRequest.query.get_or_404(request_id)
    if req.status != 'Pending':
        return jsonify({"error": f"Request is already {req.status}"}), 400
    champ = req.champion_school
    if not champ:
        return jsonify({"error": "Champion not found"}), 404
    champ.add_school(
        asl_id=req.asl_school_id or '',
        library_id=req.library_school_id or '',
        school_name=req.school_name
    )
    req.status = 'Approved'
    req.reviewed_by_user_id = current_user.id
    req.reviewed_at = datetime.utcnow()
    try:
        db.session.commit()
        return jsonify({'message': 'Approved', 'request': req.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route('/api/admin/champion-school-requests/<int:request_id>/decline', methods=['POST'])
@login_required
def api_admin_champion_school_requests_decline(request_id):
    """Decline a champion school request."""
    if not can_approve_champion_schools():
        return jsonify({"error": "Unauthorized"}), 403
    req = ChampionSchoolRequest.query.get_or_404(request_id)
    if req.status != 'Pending':
        return jsonify({"error": f"Request is already {req.status}"}), 400
    data = request.get_json() or {}
    reason = (data.get('reason') or data.get('decline_reason') or '').strip() or None
    req.status = 'Declined'
    req.reviewed_by_user_id = current_user.id
    req.reviewed_at = datetime.utcnow()
    req.decline_reason = reason
    try:
        db.session.commit()
        return jsonify({'message': 'Declined', 'request': req.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


# -------- Mobile-friendly champion request wrappers --------
# These shape the existing ChampionSchoolRequest data into the envelope the
# Flutter mobile client expects (`{data: {items, pagination}}` with lowercase
# status, `request_id`, `library_id`, etc.) without changing the canonical
# admin API used by the web dashboard.
_MOBILE_CHAMPION_REQUEST_STATUS_MAP = {
    'pending': 'Pending',
    'approved': 'Approved',
    'rejected': 'Declined',
    'declined': 'Declined',
    'all': None,
}


def _mobile_champion_request_status_to_backend(status_param):
    if status_param is None:
        return 'Pending'
    key = str(status_param).strip().lower()
    if key in _MOBILE_CHAMPION_REQUEST_STATUS_MAP:
        return _MOBILE_CHAMPION_REQUEST_STATUS_MAP[key]
    return 'Pending'


def _mobile_champion_request_status_to_mobile(backend_status):
    s = (backend_status or '').strip().lower()
    if s == 'declined':
        return 'rejected'
    if s == 'approved':
        return 'approved'
    return 'pending'


def _to_mobile_champion_request(req):
    champ = req.champion_school
    champion_name = ''
    if champ is not None:
        champion_name = f"{champ.firstname or ''} {champ.lastname or ''}".strip()
    reviewer = None
    if req.reviewed_by is not None:
        reviewer = req.reviewed_by.username
    return {
        'request_id': str(req.id),
        'champion_name': champion_name,
        'school_name': req.school_name or '',
        'asl_school_id': req.asl_school_id or '',
        'library_id': req.library_school_id or '',
        'request_type': 'add_school',
        'created_at': req.created_at.isoformat() if req.created_at else None,
        'status': _mobile_champion_request_status_to_mobile(req.status),
        'reviewer': reviewer,
        'notes': req.decline_reason,
    }


@app.route('/api/mobile/champion-requests', methods=['GET'])
@login_required
def api_mobile_champion_requests_list():
    """Mobile-shaped champion school requests list. Reuses the existing
    admin gate and ChampionSchoolRequest model."""
    if not can_approve_champion_schools():
        return jsonify({'error': 'Unauthorized'}), 403

    from sqlalchemy import or_
    backend_status = _mobile_champion_request_status_to_backend(
        request.args.get('status')
    )
    try:
        page = max(int(request.args.get('page') or 1), 1)
    except (TypeError, ValueError):
        page = 1
    try:
        per_page = int(request.args.get('per_page') or 20)
    except (TypeError, ValueError):
        per_page = 20
    per_page = max(min(per_page, 100), 1)

    q = ChampionSchoolRequest.query
    if backend_status is not None:
        q = q.filter(ChampionSchoolRequest.status == backend_status)
    text_q = (request.args.get('q') or '').strip()
    if text_q:
        like = f"%{text_q}%"
        q = q.filter(or_(
            ChampionSchoolRequest.school_name.ilike(like),
            ChampionSchoolRequest.province.ilike(like),
        ))

    total = q.count()
    rows = (
        q.order_by(ChampionSchoolRequest.created_at.desc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )
    items = [_to_mobile_champion_request(r) for r in rows]
    return jsonify({
        'data': {
            'items': items,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total,
            },
        }
    })


@app.route('/api/mobile/champion-requests/<int:request_id>', methods=['GET'])
@login_required
def api_mobile_champion_request_detail(request_id):
    if not can_approve_champion_schools():
        return jsonify({'error': 'Unauthorized'}), 403
    req = ChampionSchoolRequest.query.get_or_404(request_id)
    return jsonify({'data': _to_mobile_champion_request(req)})


@app.route('/api/mobile/champion-requests/<int:request_id>/approve', methods=['POST'])
@login_required
def api_mobile_champion_request_approve(request_id):
    if not can_approve_champion_schools():
        return jsonify({'error': 'Unauthorized'}), 403
    req = ChampionSchoolRequest.query.get_or_404(request_id)
    if req.status != 'Pending':
        return jsonify({'error': f'Request is already {req.status}'}), 400
    champ = req.champion_school
    if not champ:
        return jsonify({'error': 'Champion not found'}), 404
    champ.add_school(
        asl_id=req.asl_school_id or '',
        library_id=req.library_school_id or '',
        school_name=req.school_name,
    )
    req.status = 'Approved'
    req.reviewed_by_user_id = current_user.id
    req.reviewed_at = datetime.utcnow()
    try:
        db.session.commit()
        return jsonify({'data': _to_mobile_champion_request(req)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/mobile/champion-requests/<int:request_id>/reject', methods=['POST'])
@login_required
def api_mobile_champion_request_reject(request_id):
    if not can_approve_champion_schools():
        return jsonify({'error': 'Unauthorized'}), 403
    req = ChampionSchoolRequest.query.get_or_404(request_id)
    if req.status != 'Pending':
        return jsonify({'error': f'Request is already {req.status}'}), 400
    data = request.get_json(silent=True) or {}
    reason = (
        data.get('reason')
        or data.get('decline_reason')
        or data.get('notes')
        or ''
    ).strip() or None
    req.status = 'Declined'
    req.reviewed_by_user_id = current_user.id
    req.reviewed_at = datetime.utcnow()
    req.decline_reason = reason
    try:
        db.session.commit()
        return jsonify({'data': _to_mobile_champion_request(req)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# @app.route('/home', methods=['GET', 'POST'])
# @login_required
# def home():
#     try:
#         with SSHTunnelForwarder(
#             (ssh_host, 22),
#             ssh_username=ssh_username,
#             ssh_password=ssh_password,
#             remote_bind_address=(mysql_host, mysql_port),
#             local_bind_address=('127.0.0.1', 3307)
#         ) as tunnel:
#             connection = pymysql.connect(
#                 host='127.0.0.1',
#                 port=3307,
#                 user=mysql_user,
#                 password=mysql_password,
#                 database=mysql_db,
#                 cursorclass=pymysql.cursors.DictCursor
#             )

#             cursor = connection.cursor()

#             cursor.execute("SELECT name, level, number_of_users, province FROM institutions LIMIT 100;")
#             institutions = cursor.fetchall()

#             cursor.execute("SELECT * FROM ba_payouts LIMIT 50;")
#             ba_payouts = cursor.fetchall()

#             cursor.close()
#             connection.close()

#             return render_template('home.html', title='Home',institutions=institutions,
#                                    ba_payouts=ba_payouts)
#     except Exception as e:
#         return f"<h3>Error: {str(e)}</h3>"





# Database configuration
DB_CONFIG = {
    'host': '40.90.237.225',
    'port': 33000,
    'user': 'qmashava',
    'password': '#Qhava@@!Fu11',
    'database': 'ruzivo_2017',
    'cursorclass': pymysql.cursors.DictCursor
}



@app.route('/analytics', methods=['GET', 'POST'])
@login_required
def analytics():
        ptfform = PerfomanceTargetsForm()
        PTargets = PerfomanceTargets.query.all()
        today = datetime.today().date()
        nhasi = today.strftime('%y- %m- %d')
        month_name = today.strftime('%B')
        latest_target = PerfomanceTargets.query.order_by(PerfomanceTargets.timestamp.desc()).first()

        return render_template('analytics.html',ptfform=ptfform,PTargets=PTargets,
                               latest_target=latest_target,
                               month_name=month_name,
                               nhasi=nhasi,
                           title='Analytics')



@app.route('/akello_analytics', methods=['GET', 'POST'])
@login_required
def akello_analytics():
        if current_user.userRole == 'Brand Ambassador':
            return redirect(url_for('index'))
        else:
            ptfform = PerfomanceTargetsForm()
            PTargets = PerfomanceTargets.query.all()
            today = datetime.today().date()
            nhasi = today.strftime('%y- %m- %d')
            month_name = today.strftime('%B')
            latest_target = PerfomanceTargets.query.order_by(PerfomanceTargets.timestamp.desc()).first()
            # print(current_user.privileges)

        return render_template('akello_analytics.html',ptfform=ptfform,PTargets=PTargets,
                               latest_target=latest_target,
                               month_name=month_name,
                               nhasi=nhasi,
                           title='Analytics')









# new project planning


# -- Helper to serialize --
def label_to_dict(lbl):
    return {"id": lbl.id, "name": lbl.name, "color": lbl.color}


def _safe_task_labels(task):
    try:
        return [label_to_dict(l) for l in (getattr(task, "labels", None) or [])]
    except Exception:
        return []


def task_to_dict(t, summary=False):
    atts = sorted(
        getattr(t, "attachments", []),
        key=lambda a: (a.uploaded_at or datetime.min, a.id),
    )
    subtasks = sorted(getattr(t, "subtasks", []) or [], key=lambda s: (s.sort_order, s.id))
    subtask_done_count = sum(1 for s in subtasks if s.is_done)
    subtask_total = len(subtasks)
    progress = t.progress or 0
    if subtask_total:
        progress = int(round(100 * subtask_done_count / subtask_total))
    base = {
        "id": t.id,
        "title": t.title,
        "position": t.position,
        "progress": progress,
        "priority": getattr(t, "priority", None) or "medium",
        "start_date": t.start_date.isoformat() if t.start_date else None,
        "end_date": t.end_date.isoformat() if t.end_date else None,
        "date_rollup_enabled": bool(getattr(t, "date_rollup_enabled", False)),
        "column_id": t.column_id,
        "created_at": t.created_at.isoformat() if t.created_at else None,
        "created_by": t.created_by,
        "assignees": [{"id": u.id, "name": u.username} for u in getattr(t, "assignees", [])],
        "labels": _safe_task_labels(t),
        "blocked_by_task_id": getattr(t, "blocked_by_task_id", None),
        "blocked_by_title": (t.blocked_by.title if getattr(t, "blocked_by", None) else None),
        "source_action_item_id": getattr(t, "source_action_item_id", None),
        "subtask_done_count": subtask_done_count,
        "subtask_total": subtask_total,
        "can_delete": can_delete_task_a(t),
    }
    if summary:
        base["description"] = (t.description or "")[:200] if t.description else ""
        base["attachments"] = [{"id": a.id, "original_name": a.original_name} for a in atts]
        base["subtasks"] = []
        return base
    base["description"] = t.description
    base["attachments"] = [attachment_to_dict(a) for a in atts]
    base["subtasks"] = [subtask_a_to_dict(s) for s in subtasks]
    cfvs = TaskACustomFieldValue.query.filter_by(task_id=t.id).all()
    base["custom_fields"] = {str(v.field_id): v.value_text for v in cfvs}
    return base


def column_to_dict(col, summary=False):
    rules = None
    if col.workflow_rules:
        try:
            rules = json.loads(col.workflow_rules)
        except (TypeError, ValueError):
            rules = {}
    return {
        "id": col.id,
        "title": col.title,
        "position": col.position,
        "workflow_rules": rules or {},
        "tasks": [task_to_dict(t, summary=summary) for t in sorted(col.tasks, key=lambda x: x.position)]
    }

def subtask_a_to_dict(st):
    return {
        "id": st.id,
        "task_id": st.task_id,
        "title": st.title,
        "is_done": st.is_done,
        "sort_order": st.sort_order,
        "start_date": st.start_date.isoformat() if st.start_date else None,
        "end_date": st.end_date.isoformat() if st.end_date else None,
        "assignees": [{"id": u.id, "name": u.username} for u in getattr(st, "assignees", [])],
    }


def validate_assignee_user_ids(user_ids):
    """Return list of User objects for valid IDs, or (None, error_message)."""
    if user_ids is None:
        return [], None
    if not isinstance(user_ids, list):
        return None, "assignees must be a list"
    ids = []
    for raw in user_ids:
        try:
            ids.append(int(raw))
        except (TypeError, ValueError):
            return None, "assignees must contain integer user IDs"
    users = []
    for uid in ids:
        user = User.query.get(uid)
        if not user:
            return None, f"Unknown user id: {uid}"
        users.append(user)
    return users, None


def _parse_optional_iso_datetime(value):
    if value is None or value == "":
        return None
    return datetime.fromisoformat(value)


def _apply_subtask_a_fields(st, data):
    """Apply optional title, dates, done state, and assignees from JSON payload."""
    if "title" in data:
        new_title = (data.get("title") or "").strip()
        if new_title:
            st.title = new_title
    if "is_done" in data:
        st.is_done = bool(data["is_done"])
    if "start_date" in data:
        st.start_date = _parse_optional_iso_datetime(data["start_date"])
    if "end_date" in data:
        st.end_date = _parse_optional_iso_datetime(data["end_date"])
    if st.start_date and st.end_date and st.end_date < st.start_date:
        return "end_date must be greater than or equal to start_date"
    if "assignees" in data:
        users, err = validate_assignee_user_ids(data["assignees"])
        if err:
            return err
        st.assignees = users
    return None


def apply_task_a_subtask_progress_rollup(task):
    """Derive parent task progress from sub-task completion."""
    subtasks = list(task.subtasks or [])
    if not subtasks:
        return False
    done = sum(1 for s in subtasks if s.is_done)
    new_progress = int(round(100 * done / len(subtasks)))
    changed = task.progress != new_progress
    task.progress = new_progress
    return changed


def apply_task_a_subtask_date_rollup(task):
    """When enabled, set parent dates from sub-task min/max if parent dates empty."""
    if not getattr(task, "date_rollup_enabled", False):
        return False
    subtasks = [s for s in (task.subtasks or []) if s.start_date or s.end_date]
    if not subtasks:
        return False
    changed = False
    if not task.start_date:
        starts = [s.start_date for s in subtasks if s.start_date]
        if starts:
            task.start_date = min(starts)
            changed = True
    if not task.end_date:
        ends = [s.end_date for s in subtasks if s.end_date]
        if ends:
            task.end_date = max(ends)
            changed = True
    return changed


def can_manage_project_a(p):
    """Admin / Super-admin privilege, or project owner."""
    if getattr(current_user, "userRole", None) == "Admin":
        return True
    if hasattr(current_user, "has_privilege") and current_user.has_privilege("Super-admin"):
        return True
    if p.owner_id is not None and p.owner_id == current_user.id:
        return True
    return False


def can_delete_task_a(t):
    """Only the task creator or an admin / super-admin can delete a task."""
    if getattr(current_user, "userRole", None) == "Admin":
        return True
    if hasattr(current_user, "has_privilege") and current_user.has_privilege("Super-admin"):
        return True
    creator_id = getattr(t, "created_by", None)
    return creator_id is not None and creator_id == current_user.id


def can_delete_task_attachment_a(att):
    """Admin / Super-admin, the task creator, or the user who uploaded the file."""
    if can_delete_task_a(att.task):
        return True
    uploader_id = getattr(att, "uploaded_by", None)
    return uploader_id is not None and uploader_id == current_user.id


def user_can_access_project_a(p):
    """Project member, owner/manager, or public project (logged-in user)."""
    if can_manage_project_a(p):
        return True
    if current_user in p.members:
        return True
    if getattr(p, "project_type", None) == "public":
        return True
    return False


def user_project_role_a(p):
    from app.pm_service import get_member_role
    if can_manage_project_a(p):
        return "admin"
    return get_member_role(p, current_user.id) or ("contributor" if getattr(p, "project_type", None) == "public" else None)


def user_can_edit_tasks_a(p):
    role = user_project_role_a(p)
    return role in ("admin", "contributor")


def user_can_comment_a(p):
    role = user_project_role_a(p)
    return role in ("admin", "contributor", "viewer")


def user_can_manage_project_a(p):
    """Project owner, app admin, or project member role admin."""
    if can_manage_project_a(p):
        return True
    return user_project_role_a(p) == "admin"


def project_edit_denied():
    return jsonify({"error": "You do not have permission to edit tasks in this project."}), 403


def project_manage_denied():
    return jsonify({"error": "You do not have permission to manage this project."}), 403


def projects_visible_to_user():
    return [p for p in ProjectA.query.order_by(ProjectA.name).all() if user_can_access_project_a(p)]


def project_access_denied():
    return jsonify({"error": "You do not have permission to access this project."}), 403


def _project_for_column(col):
    return col.project if col else None


def _project_for_task(t):
    return t.column.project if t and t.column else None


def _log_task_activity(task, action, detail=None):
    act = TaskAActivity(
        task_id=task.id,
        user_id=current_user.id,
        action=action,
        detail=detail,
    )
    db.session.add(act)
    _notify_pm_project_subscribers(task, action, detail)


def _notify_pm_project_subscribers(task, action, detail=None):
    """Notify project subscribers of activity (dedupe one per user per project per hour)."""
    from datetime import timedelta

    from app.models import ProjectASubscription

    p = _project_for_task(task)
    if not p:
        return
    subs = ProjectASubscription.query.filter_by(project_id=p.id).all()
    if not subs:
        return
    hour_ago = datetime.utcnow() - timedelta(hours=1)
    link = url_for("projectmanagement", _external=False) + f"?project={p.id}&task={task.id}"
    summary = f"{action}" + (f": {detail}" if detail else "")
    for sub in subs:
        if sub.user_id == current_user.id:
            continue
        recent = Notification.query.filter(
            Notification.user_id == sub.user_id,
            Notification.notification_type == "pm_project_activity",
            Notification.created_at >= hour_ago,
            Notification.message.contains(f"project={p.id}"),
        ).first()
        if recent:
            continue
        db.session.add(Notification(
            user_id=sub.user_id,
            task_id=task.id,
            message=f"Project activity ({p.name}): {summary}. {link}",
            notification_type="pm_project_activity",
        ))


def _notify_pm_assignees(task, new_user_ids, project):
    if not new_user_ids:
        return
    me_id = current_user.id
    link = url_for("projectmanagement", _external=False) + f"?project={project.id}&task={task.id}"
    for uid in new_user_ids:
        if uid == me_id:
            continue
        db.session.add(Notification(
            user_id=uid,
            task_id=task.id,
            message=f"You were assigned to task: {task.title}. {link}",
            notification_type="pm_assignment",
        ))


def _resolve_task_labels(task, label_ids, project_id):
    if label_ids is None:
        return None, None
    if not isinstance(label_ids, list):
        return None, "labels must be a list"
    labels = []
    for raw in label_ids:
        try:
            lid = int(raw)
        except (TypeError, ValueError):
            return None, "labels must contain integer IDs"
        lbl = TaskALabel.query.filter_by(id=lid, project_id=project_id).first()
        if not lbl:
            return None, f"Unknown label id: {lid}"
        labels.append(lbl)
    return labels, None


def comment_to_dict(c):
    author = c.author
    name = author.username if author else "Unknown"
    return {
        "id": c.id,
        "task_id": c.task_id,
        "user_id": c.user_id,
        "author_name": name,
        "body": c.body,
        "created_at": c.created_at.isoformat() if c.created_at else None,
    }


def activity_to_dict(a):
    actor = a.actor
    return {
        "id": a.id,
        "task_id": a.task_id,
        "user_id": a.user_id,
        "actor_name": actor.username if actor else None,
        "action": a.action,
        "detail": a.detail,
        "created_at": a.created_at.isoformat() if a.created_at else None,
    }


TASK_ATTACHMENTS_FOLDER_REL = os.path.join("uploads", "task_attachments")
TASK_ATTACHMENTS_LEGACY_FOLDER_REL = os.path.join("static", "uploads", "task_attachments")
TASK_ATTACHMENTS_MAX_BYTES = 25 * 1024 * 1024
TASK_ATTACHMENTS_ALLOWED_EXT = frozenset({
    "pdf", "png", "jpg", "jpeg", "gif", "webp", "txt", "csv",
    "doc", "docx", "xls", "xlsx", "ppt", "pptx",
})


def _task_attachments_instance_dir():
    base = current_app.instance_path if current_app else app.instance_path
    return os.path.join(base, TASK_ATTACHMENTS_FOLDER_REL)


def _task_attachment_full_path(stored_path):
    if not stored_path:
        return None
    raw = str(stored_path).strip()
    if os.path.isabs(raw):
        return raw

    # Normalize both slash styles so records created on Windows still resolve on Linux.
    unified = raw.replace("\\", "/").lstrip("./")
    normalized = unified.replace("/", os.sep)
    legacy_unified = TASK_ATTACHMENTS_LEGACY_FOLDER_REL.replace("\\", "/")
    current_unified = TASK_ATTACHMENTS_FOLDER_REL.replace("\\", "/")

    root = current_app.root_path if current_app else app.root_path
    instance_root = current_app.instance_path if current_app else app.instance_path

    # Legacy rows stored under app/static/uploads/task_attachments/ keep working.
    if unified == legacy_unified or unified.startswith(legacy_unified + "/"):
        return os.path.join(root, normalized)

    # Current rows stored under instance/uploads/task_attachments/
    if unified == current_unified or unified.startswith(current_unified + "/"):
        return os.path.join(instance_root, normalized)

    # Best-effort fallback for older inconsistent rows that only have a filename.
    base_name = os.path.basename(unified)
    if base_name:
        candidate_instance = os.path.join(_task_attachments_instance_dir(), base_name)
        if os.path.isfile(candidate_instance):
            return candidate_instance
        candidate_legacy = os.path.join(root, TASK_ATTACHMENTS_LEGACY_FOLDER_REL, base_name)
        if os.path.isfile(candidate_legacy):
            return candidate_legacy

    return os.path.join(instance_root, normalized)


def attachment_to_dict(a):
    return {
        "id": a.id,
        "original_name": a.original_name,
        "content_type": a.content_type,
        "file_size": a.file_size,
        "uploaded_at": a.uploaded_at.isoformat() if a.uploaded_at else None,
    }


@app.before_request
def create_tables():
    db.create_all()
    # Best-effort: add owner_id on existing SQLite DBs (ignored if column already exists)
    try:
        if db.engine.dialect.name == "sqlite":
            with db.engine.connect() as _conn:
                _conn.execute(sa.text("ALTER TABLE projectsa ADD COLUMN owner_id INTEGER"))
                _conn.commit()
    except Exception:
        pass
    # Best-effort: add created_by on tasksa for existing SQLite DBs
    try:
        if db.engine.dialect.name == "sqlite":
            with db.engine.connect() as _conn:
                _conn.execute(sa.text("ALTER TABLE tasksa ADD COLUMN created_by INTEGER"))
                _conn.commit()
    except Exception:
        pass
    # PM roadmap columns on tasksa (SQLite-safe)
    if db.engine.dialect.name == "sqlite":
        for stmt in (
            "ALTER TABLE tasksa ADD COLUMN priority VARCHAR(16) DEFAULT 'medium'",
            "ALTER TABLE tasksa ADD COLUMN date_rollup_enabled BOOLEAN DEFAULT 0",
            "ALTER TABLE tasksa ADD COLUMN source_action_item_id INTEGER",
            "ALTER TABLE tasksa ADD COLUMN blocked_by_task_id INTEGER",
            "ALTER TABLE notifications ADD COLUMN task_id INTEGER",
        ):
            try:
                with db.engine.connect() as _conn:
                    _conn.execute(sa.text(stmt))
                    _conn.commit()
            except Exception:
                pass
    # Ensure the task-attachments folder exists under instance/ so uploads survive deploys
    try:
        os.makedirs(_task_attachments_instance_dir(), exist_ok=True)
    except Exception:
        pass
    # seed sample if empty
    if ProjectA.query.count() == 0:
        p = ProjectA(name="Example Project")
        db.session.add(p)
        db.session.commit()
        for idx, name in enumerate(["Ideas","To Do","Doing","Done"]):
            col = ColumnA(project_id=p.id, title=name, position=idx)
            db.session.add(col)
            db.session.commit()
            # add a sample task
            t = TaskA(column_id=col.id, title=f"Sample task in {name}", position=0)
            db.session.add(t)
        db.session.commit()

@app.route("/projectmanagemnt", methods=["GET", "POST"])
@login_required
def projectmanagement():
    return render_template("aplanforprojects.html")


@app.route("/projectmanagement", methods=["GET", "POST"])
@login_required
def projectmanagement_alias():
    return redirect(url_for("projectmanagement"))


@app.route("/aplanforprojects", methods=["GET", "POST"])
@login_required
def aplanforprojects_redirect():
    return redirect(url_for("projectmanagement"))

# List + create projects
@app.route("/api/projects", methods=["GET", "POST"])
@login_required
def projectsA():
    if request.method == "GET":
        ps = projects_visible_to_user()
        return jsonify([
            {
                "id": p.id,
                "name": p.name,
                "type": p.project_type,
                "owner_id": p.owner_id,
                "members": [u.username for u in p.members]
            } for p in ps
        ])
    
    data = request.get_json()
    if not data or "name" not in data:
        abort(400)
    clone_from = data.get("clone_from_id") or data.get("template_project_id")
    if clone_from:
        from app.pm_service import clone_project
        source = ProjectA.query.get_or_404(int(clone_from))
        if not user_can_access_project_a(source):
            return project_access_denied()
        include_tasks = bool(data.get("include_tasks"))
        p, _, _ = clone_project(source, data["name"], current_user.id, include_tasks=include_tasks)
        if data.get("project_type"):
            p.project_type = data["project_type"]
        db.session.commit()
        return jsonify({"id": p.id, "name": p.name, "type": p.project_type, "owner_id": p.owner_id}), 201
    p = ProjectA(
        name=data["name"],
        project_type=data.get("project_type", "private"),
        owner_id=current_user.id,
    )
    db.session.add(p)
    db.session.commit()
    return jsonify({"id": p.id, "name": p.name, "type": p.project_type, "owner_id": p.owner_id}), 201



# Edit a project
@app.route("/api/projects/<int:project_id>", methods=["GET"])
@login_required
def get_projectA(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    return jsonify({
        "id": p.id,
        "name": p.name,
        "type": p.project_type,
        "owner_id": p.owner_id,
        "members": [{"id": u.id, "name": u.username} for u in p.members]
    })


# Edit a project
@app.route("/api/projects/<int:project_id>", methods=["PATCH"])
@login_required
def edit_projectA(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not can_manage_project_a(p):
        return jsonify({"error": "You do not have permission to modify this project."}), 403
    data = request.get_json()

    if "name" in data:
        p.name = data["name"]
    if "project_type" in data:
        p.project_type = data["project_type"]

    # ✅ Replace members if provided
    if "members" in data:
        member_ids = set(data["members"])  # list of user IDs
        # fetch users that exist in DB
        p.members = [User.query.get(uid) for uid in member_ids if User.query.get(uid)]

    db.session.commit()
    return jsonify({
        "id": p.id,
        "name": p.name,
        "type": p.project_type,
        "members": [u.id for u in p.members]  # return IDs for clarity
    })



# Delete a project
@app.route("/api/projects/<int:project_id>", methods=["DELETE"])
@login_required
def delete_projectA(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not can_manage_project_a(p):
        return jsonify({"error": "You do not have permission to delete this project."}), 403
    db.session.delete(p)
    db.session.commit()
    return jsonify({"status": "deleted"})


# Add members to a project
@app.route("/api/projects/<int:project_id>/members", methods=["GET"])
@login_required
def get_membersA(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    return jsonify([{"id": u.id, "name": u.username} for u in p.members])


# Add members to a project
@app.route("/api/projects/<int:project_id>/members", methods=["POST"])
@login_required
def add_membersA(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not can_manage_project_a(p):
        return project_access_denied()
    data = request.get_json()
    user_ids = data.get("user_ids", [])
    if not isinstance(user_ids, list):
        abort(400)
    for uid in user_ids:
        user = User.query.get(uid)
        if user and user not in p.members:
            p.members.append(user)
    db.session.commit()
    return jsonify({
        "id": p.id,
        "members": [u.username for u in p.members]
    })


@app.route("/api/users", methods=["GET"])
@login_required
def get_users():
    """
    Fetches a list of all users and returns their IDs and usernames.
    This is used by the frontend to populate the list of potential members
    for a project.
    """
    users = User.query.all()
    user_list = [
        {"id": user.id, "name": user.username} for user in users
    ]
    return jsonify(user_list)



# Get board (columns + tasks) for a project
@app.route("/api/projects/<int:project_id>/board", methods=["GET"])
@login_required
def project_board(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    summary = request.args.get("summary") in ("1", "true", "yes")
    cols = (
        ColumnA.query.options(
            joinedload(ColumnA.tasks).joinedload(TaskA.assignees),
            joinedload(ColumnA.tasks).joinedload(TaskA.labels),
            joinedload(ColumnA.tasks).joinedload(TaskA.blocked_by),
            joinedload(ColumnA.tasks).joinedload(TaskA.subtasks).joinedload(TaskASubtask.assignees),
        )
        .filter_by(project_id=p.id)
        .order_by(ColumnA.position)
        .all()
    )
    return jsonify({
        "id": p.id,
        "name": p.name,
        "columns": [column_to_dict(c, summary=summary) for c in cols],
        "current_user_role": user_project_role_a(p),
        "can_edit_tasks": user_can_edit_tasks_a(p),
        "can_manage_project": user_can_manage_project_a(p),
        "can_comment": user_can_comment_a(p),
    })

# Create column in project
@app.route("/api/projects/<int:project_id>/columns", methods=["POST"])
@login_required
def create_column(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_manage_project_a(p):
        return project_manage_denied()
    data = request.get_json()
    title = data.get("title","New Column")
    # set position to end
    maxpos = db.session.query(db.func.max(ColumnA.position)).filter_by(project_id=project_id).scalar()
    pos = (maxpos + 1) if maxpos is not None else 0
    c = ColumnA(project_id=project_id, title=title, position=pos)
    db.session.add(c)
    db.session.commit()
    return jsonify(column_to_dict(c)), 201



@app.route("/api/columns/<int:column_id>", methods=["PATCH"])
@login_required
def patch_column(column_id):
    c = ColumnA.query.get_or_404(column_id)
    p = _project_for_column(c)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_manage_project_a(p):
        return project_manage_denied()
    data = request.get_json() or {}
    if "title" in data:
        c.title = data["title"]
        db.session.commit()
    return jsonify(column_to_dict(c))



# Create task
@app.route("/api/columns/<int:column_id>/tasks", methods=["POST"])
@login_required
def create_taskA(column_id):
    col = ColumnA.query.get_or_404(column_id)
    p = _project_for_column(col)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_edit_tasks_a(p):
        return project_edit_denied()
    data = request.get_json()
    title = data.get("title", "New Task")
    desc = data.get("description", "")
    start_date = data.get("start_date")
    end_date = data.get("end_date")
    assignees_ids = data.get("assignees", [])  # optional list of user IDs

    # parse dates if provided
    start_dt = datetime.fromisoformat(start_date) if start_date else None
    end_dt = datetime.fromisoformat(end_date) if end_date else None

    # ✅ validation
    if start_dt and end_dt and end_dt < start_dt:
        return jsonify({"error": "end_date must be greater than or equal to start_date"}), 400

    maxpos = db.session.query(db.func.max(TaskA.position)).filter_by(column_id=column_id).scalar()
    pos = (maxpos + 1) if maxpos is not None else 0

    priority = (data.get("priority") or "medium").strip().lower()
    if priority not in ("low", "medium", "high", "urgent"):
        priority = "medium"
    t = TaskA(
        column_id=column_id,
        title=title,
        description=desc,
        position=pos,
        start_date=start_dt,
        end_date=end_dt,
        priority=priority,
        created_by=current_user.id,
    )
    if "date_rollup_enabled" in data:
        t.date_rollup_enabled = bool(data["date_rollup_enabled"])
    if "source_action_item_id" in data and data["source_action_item_id"]:
        t.source_action_item_id = int(data["source_action_item_id"])
    # set initial assignees if provided
    new_assignee_ids = []
    if isinstance(assignees_ids, list) and len(assignees_ids) > 0:
        users, err = validate_assignee_user_ids(assignees_ids)
        if err:
            return jsonify({"error": err}), 400
        t.assignees = users
        new_assignee_ids = [u.id for u in users]
    if "labels" in data:
        labels, err = _resolve_task_labels(t, data["labels"], p.id)
        if err:
            return jsonify({"error": err}), 400
        if labels is not None:
            t.labels = labels
    db.session.add(t)
    db.session.flush()
    _log_task_activity(t, "created", f"Task created in column {col.title}")
    _notify_pm_assignees(t, new_assignee_ids, p)
    db.session.commit()
    from app.pm_service import fire_webhooks
    fire_webhooks(p.id, "task.created", {"task_id": t.id, "column_id": col.id})
    return jsonify(task_to_dict(t)), 201



# --- Task update extended with progress ---
@app.route("/api/tasks/<int:task_id>", methods=["GET"])
@login_required
def get_taskA(task_id):
    t = (
        TaskA.query.options(
            joinedload(TaskA.assignees),
            joinedload(TaskA.labels),
            joinedload(TaskA.subtasks).joinedload(TaskASubtask.assignees),
            joinedload(TaskA.attachments),
        )
        .get_or_404(task_id)
    )
    p = _project_for_task(t)
    if not user_can_access_project_a(p):
        return project_access_denied()
    return jsonify(task_to_dict(t))


@app.route("/api/tasks/<int:task_id>", methods=["PATCH"])
@login_required
def update_taskA(task_id):
    t = TaskA.query.get_or_404(task_id)
    p = _project_for_task(t)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_edit_tasks_a(p):
        return jsonify({"error": "You do not have permission to edit tasks in this project."}), 403
    data = request.get_json()
    old_assignee_ids = {u.id for u in (t.assignees or [])}
    old_col_id = t.column_id

    if "title" in data:
        t.title = data["title"]
    if "assignees" in data and isinstance(data["assignees"], list):
        users, err = validate_assignee_user_ids(data["assignees"])
        if err:
            return jsonify({"error": err}), 400
        t.assignees = users
    if "description" in data:
        t.description = data["description"]
    if "progress" in data:
        t.progress = max(0, min(100, int(data["progress"])))
    if "priority" in data:
        pr = (data.get("priority") or "medium").strip().lower()
        if pr in ("low", "medium", "high", "urgent"):
            t.priority = pr
    if "date_rollup_enabled" in data:
        t.date_rollup_enabled = bool(data["date_rollup_enabled"])
    if "labels" in data:
        labels, err = _resolve_task_labels(t, data["labels"], p.id)
        if err:
            return jsonify({"error": err}), 400
        if labels is not None:
            t.labels = labels
    if "blocked_by_task_id" in data:
        bid = data["blocked_by_task_id"]
        if bid is None or bid == "":
            t.blocked_by_task_id = None
        else:
            bid = int(bid)
            if bid == t.id:
                return jsonify({"error": "A task cannot block itself."}), 400
            blocker = TaskA.query.get(bid)
            if not blocker or _project_for_task(blocker).id != p.id:
                return jsonify({"error": "Invalid blocked_by task."}), 400
            t.blocked_by_task_id = bid

    if "custom_fields" in data and isinstance(data["custom_fields"], dict):
        for fid, val in data["custom_fields"].items():
            try:
                field_id = int(fid)
            except (TypeError, ValueError):
                continue
            cf = ProjectACustomField.query.filter_by(id=field_id, project_id=p.id).first()
            if not cf:
                continue
            existing = TaskACustomFieldValue.query.filter_by(task_id=t.id, field_id=field_id).first()
            if existing:
                existing.value_text = str(val) if val is not None else None
            else:
                db.session.add(TaskACustomFieldValue(task_id=t.id, field_id=field_id, value_text=str(val) if val is not None else None))

    # handle dates
    if "start_date" in data:
        t.start_date = datetime.fromisoformat(data["start_date"]) if data["start_date"] else None
    if "end_date" in data:
        t.end_date = datetime.fromisoformat(data["end_date"]) if data["end_date"] else None

    # ✅ validation after updating
    if t.start_date and t.end_date and t.end_date < t.start_date:
        return jsonify({"error": "end_date must be greater than or equal to start_date"}), 400

    # handle column moves + reordering (same as before)
    if "column_id" in data or "position" in data:
        new_col_id = data.get("column_id", t.column_id)
        new_pos = data.get("position", None)
        if new_col_id != t.column_id:
            dest_col = ColumnA.query.get(new_col_id)
            if dest_col:
                from app.pm_service import validate_column_workflow, task_is_complete
                wf_err = validate_column_workflow(t, dest_col)
                if wf_err:
                    return jsonify({"error": wf_err}), 400
                if t.blocked_by_task_id:
                    blocker = TaskA.query.get(t.blocked_by_task_id)
                    if blocker and not task_is_complete(blocker):
                        title_lower = (dest_col.title or "").lower()
                        if "done" in title_lower or "complete" in title_lower:
                            pass  # soft warning only — allow move
            old_tasks = TaskA.query.filter(
                TaskA.column_id == t.column_id,
                TaskA.id != t.id
            ).order_by(TaskA.position).all()
            for idx, ot in enumerate(old_tasks):
                ot.position = idx
            dest_tasks = TaskA.query.filter_by(column_id=new_col_id).order_by(TaskA.position).all()
            if new_pos is None or new_pos > len(dest_tasks):
                new_pos = len(dest_tasks)
            for ot in dest_tasks[::-1]:
                if ot.position >= new_pos:
                    ot.position += 1
            t.column_id = new_col_id
            t.position = new_pos
        elif new_pos is not None:
            tasks = TaskA.query.filter_by(column_id=t.column_id).order_by(TaskA.position).all()
            tasks = [x for x in tasks if x.id != t.id]
            tasks.insert(new_pos, t)
            for idx, ot in enumerate(tasks):
                ot.position = idx

    new_assignee_ids = [u.id for u in (t.assignees or []) if u.id not in old_assignee_ids]
    if "column_id" in data and t.column_id != old_col_id:
        _log_task_activity(t, "moved", f"Moved to column id {t.column_id}")
    if "progress" in data:
        _log_task_activity(t, "progress", f"Progress set to {t.progress}%")
    if new_assignee_ids:
        _log_task_activity(t, "assigned", f"Added assignees: {', '.join(str(x) for x in new_assignee_ids)}")
        _notify_pm_assignees(t, new_assignee_ids, p)
    apply_task_a_subtask_date_rollup(t)
    db.session.commit()
    from app.pm_service import fire_webhooks
    if "column_id" in data:
        fire_webhooks(p.id, "task.moved", {"task_id": t.id, "column_id": t.column_id})
    if "progress" in data and (t.progress or 0) >= 100:
        fire_webhooks(p.id, "task.completed", {"task_id": t.id})
    return jsonify(task_to_dict(t))


@app.route("/api/tasks/<int:task_id>/attachments", methods=["POST"])
@login_required
def upload_task_attachment(task_id):
    t = TaskA.query.get_or_404(task_id)
    p = t.column.project
    if not user_can_access_project_a(p):
        return jsonify({"error": "You do not have permission to attach files to this task."}), 403
    if not user_can_edit_tasks_a(p):
        return project_edit_denied()

    upload = request.files.get("file")
    if not upload or not upload.filename:
        return jsonify({"error": "No file uploaded."}), 400

    raw_name = secure_filename(upload.filename)
    if not raw_name or "." not in raw_name:
        return jsonify({"error": "Invalid file name."}), 400
    ext = raw_name.rsplit(".", 1)[1].lower()
    if ext not in TASK_ATTACHMENTS_ALLOWED_EXT:
        return jsonify({"error": "File type not allowed."}), 400

    upload.seek(0, os.SEEK_END)
    size = upload.tell()
    upload.seek(0)
    if size > TASK_ATTACHMENTS_MAX_BYTES:
        return jsonify({"error": "File too large."}), 400

    folder_abs = _task_attachments_instance_dir()
    os.makedirs(folder_abs, exist_ok=True)

    stored_fname = f"{uuid.uuid4().hex}_{raw_name}"
    rel_stored = os.path.join(TASK_ATTACHMENTS_FOLDER_REL, stored_fname).replace("\\", "/")
    dest_abs = os.path.join(folder_abs, stored_fname)
    upload.save(dest_abs)

    ct = upload.content_type or None
    att = TaskAttachment(
        task_id=t.id,
        original_name=raw_name,
        stored_path=rel_stored,
        content_type=ct,
        file_size=size,
        uploaded_by=current_user.id,
    )
    db.session.add(att)
    _log_task_activity(t, "attachment", f"Added file: {raw_name}")
    db.session.commit()
    return jsonify(attachment_to_dict(att)), 201


@app.route("/api/task-attachments/<int:attachment_id>", methods=["DELETE"])
@login_required
def delete_task_attachment(attachment_id):
    att = TaskAttachment.query.get_or_404(attachment_id)
    p = att.task.column.project
    if not user_can_access_project_a(p):
        return jsonify({"error": "Forbidden"}), 403
    if not can_delete_task_attachment_a(att):
        return jsonify({"error": "Only the task creator, the uploader, or an admin can delete this attachment."}), 403

    path = _task_attachment_full_path(att.stored_path)
    db.session.delete(att)
    db.session.commit()
    if path and os.path.isfile(path):
        try:
            os.remove(path)
        except OSError:
            pass
    return jsonify({"status": "deleted"})


@app.route("/api/task-attachments/<int:attachment_id>/file", methods=["GET"])
@login_required
def task_attachment_file(attachment_id):
    att = TaskAttachment.query.get_or_404(attachment_id)
    p = att.task.column.project
    if not user_can_access_project_a(p):
        abort(403)

    path = _task_attachment_full_path(att.stored_path)
    if not path or not os.path.isfile(path):
        abort(404)

    disposition = (request.args.get("disposition") or "inline").lower()
    as_attachment = disposition == "attachment"
    mimetype = att.content_type or None
    return send_file(
        path,
        mimetype=mimetype,
        as_attachment=as_attachment,
        download_name=att.original_name if as_attachment else None,
        conditional=True,
        max_age=0,
    )


# Reorder columns in a project
@app.route("/api/projects/<int:project_id>/columns/reorder", methods=["POST"])
@login_required
def reorder_columns(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_manage_project_a(p):
        return project_manage_denied()
    data = request.get_json()
    order = data.get("order", [])  # list of column ids in new order
    if not isinstance(order, list):
        abort(400)
    for idx, cid in enumerate(order):
        c = ColumnA.query.filter_by(id=cid, project_id=project_id).first()
        if c:
            c.position = idx
    db.session.commit()
    return jsonify({"status":"ok"})


@app.route("/api/columns/<int:column_id>", methods=["DELETE"])
@login_required
def delete_column(column_id):
    c = ColumnA.query.get_or_404(column_id)
    p = _project_for_column(c)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if not user_can_manage_project_a(p):
        return project_manage_denied()
    col_count = ColumnA.query.filter_by(project_id=p.id).count()
    if col_count <= 1:
        return jsonify({"error": "Cannot delete the only column in a project."}), 400
    data = request.get_json(silent=True) or {}
    move_to = data.get("move_to_column_id")
    if not move_to:
        return jsonify({"error": "move_to_column_id is required."}), 400
    dest = ColumnA.query.filter_by(id=int(move_to), project_id=p.id).first()
    if not dest or dest.id == c.id:
        return jsonify({"error": "Invalid move_to_column_id."}), 400
    tasks = TaskA.query.filter_by(column_id=c.id).order_by(TaskA.position).all()
    maxpos = db.session.query(db.func.max(TaskA.position)).filter_by(column_id=dest.id).scalar()
    next_pos = (maxpos + 1) if maxpos is not None else 0
    for t in tasks:
        t.column_id = dest.id
        t.position = next_pos
        next_pos += 1
    db.session.delete(c)
    db.session.commit()
    return jsonify({"status": "deleted", "moved_tasks": len(tasks)})

# Delete endpoints (optional)
@app.route("/api/tasks/<int:task_id>", methods=["DELETE"])
@login_required
def delete_taskA(task_id):
    t = TaskA.query.get_or_404(task_id)
    p = t.column.project
    if not user_can_access_project_a(p):
        return jsonify({"error": "Forbidden"}), 403
    if not user_can_edit_tasks_a(p):
        return jsonify({"error": "You do not have permission to delete tasks in this project."}), 403
    if not can_delete_task_a(t):
        return jsonify({"error": "Only the task creator or an admin can delete this task."}), 403
    col_id = t.column_id
    paths = [_task_attachment_full_path(a.stored_path) for a in t.attachments]
    db.session.delete(t)
    db.session.commit()
    for path in paths:
        if path and os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass
    # reindex
    tasks = TaskA.query.filter_by(column_id=col_id).order_by(TaskA.position).all()
    for idx, ot in enumerate(tasks):
        ot.position = idx
    db.session.commit()
    return jsonify({"status":"deleted"})


def _task_a_with_access(task_id):
    t = (
        TaskA.query.options(
            joinedload(TaskA.subtasks).joinedload(TaskASubtask.assignees),
        )
        .get_or_404(task_id)
    )
    p = t.column.project
    if not user_can_access_project_a(p):
        abort(403)
    return t


def _reload_subtask_a(subtask_id):
    return (
        TaskASubtask.query.options(joinedload(TaskASubtask.assignees))
        .get_or_404(subtask_id)
    )


@app.route("/api/tasks/<int:task_id>/subtasks", methods=["GET", "POST"])
@login_required
def task_a_subtasks(task_id):
    t = _task_a_with_access(task_id)
    p = _project_for_task(t)
    if request.method == "GET":
        subtasks = sorted(t.subtasks or [], key=lambda s: (s.sort_order, s.id))
        return jsonify([subtask_a_to_dict(s) for s in subtasks])

    if not user_can_edit_tasks_a(p):
        return project_edit_denied()

    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "title is required"}), 400
    max_order = max((s.sort_order for s in (t.subtasks or [])), default=-1)
    st = TaskASubtask(task_id=t.id, title=title, sort_order=max_order + 1)
    err = _apply_subtask_a_fields(st, data)
    if err:
        return jsonify({"error": err}), 400
    db.session.add(st)
    apply_task_a_subtask_progress_rollup(t)
    apply_task_a_subtask_date_rollup(t)
    _log_task_activity(t, "subtask_added", title)
    if st.assignees:
        _notify_pm_assignees(t, [u.id for u in st.assignees], p)
    db.session.commit()
    st = _reload_subtask_a(st.id)
    return jsonify(subtask_a_to_dict(st)), 201


@app.route("/api/task-subtasks/<int:subtask_id>", methods=["PATCH", "DELETE"])
@login_required
def task_a_subtask(subtask_id):
    st = TaskASubtask.query.options(joinedload(TaskASubtask.assignees)).get_or_404(subtask_id)
    t = (
        TaskA.query.options(joinedload(TaskA.subtasks))
        .get_or_404(st.task_id)
    )
    p = t.column.project
    if not user_can_access_project_a(p):
        return jsonify({"error": "Forbidden"}), 403
    if not user_can_edit_tasks_a(p):
        return project_edit_denied()

    if request.method == "DELETE":
        db.session.delete(st)
        apply_task_a_subtask_progress_rollup(t)
        apply_task_a_subtask_date_rollup(t)
        _log_task_activity(t, "subtask_removed", st.title)
        db.session.commit()
        return jsonify({"status": "deleted"})

    old_assignee_ids = {u.id for u in (st.assignees or [])}
    data = request.get_json() or {}
    err = _apply_subtask_a_fields(st, data)
    if err:
        return jsonify({"error": err}), 400
    apply_task_a_subtask_progress_rollup(t)
    apply_task_a_subtask_date_rollup(t)
    new_ids = [u.id for u in (st.assignees or []) if u.id not in old_assignee_ids]
    if new_ids:
        _notify_pm_assignees(t, new_ids, p)
    db.session.commit()
    st = _reload_subtask_a(st.id)
    return jsonify(subtask_a_to_dict(st))


@app.route("/api/my-tasks", methods=["GET"])
@login_required
def my_tasks_a():
    """Tasks and sub-tasks assigned to current user across visible projects."""
    me_id = current_user.id
    due_filter = (request.args.get("due") or "all").strip().lower()
    today = datetime.utcnow().date()
    week_end = today + timedelta(days=7)
    rows = []
    visible_ids = {p.id for p in projects_visible_to_user()}

    def due_bucket(end_dt):
        if not end_dt:
            return "none"
        d = end_dt.date() if hasattr(end_dt, "date") else end_dt
        if d < today:
            return "overdue"
        if d <= week_end:
            return "week"
        return "later"

    def task_progress_value(task, subtask=None):
        if subtask:
            return 100 if subtask.is_done else 0
        subtasks = list(task.subtasks or [])
        if subtasks:
            done = sum(1 for s in subtasks if s.is_done)
            return int(round(100 * done / len(subtasks)))
        return task.progress or 0

    def append_row(item_type, task, col, project, subtask=None):
        end_dt = subtask.end_date if subtask else task.end_date
        bucket = due_bucket(end_dt)
        progress = task_progress_value(task, subtask)
        is_overdue = bucket == "overdue" and progress < 100
        if due_filter == "overdue" and not is_overdue:
            return
        if due_filter == "week" and bucket not in ("overdue", "week"):
            return
        rows.append({
            "item_type": item_type,
            "task_id": task.id,
            "subtask_id": subtask.id if subtask else None,
            "title": ("↳ " + subtask.title) if subtask else task.title,
            "project_id": project.id,
            "project_name": project.name,
            "column_title": col.title,
            "start_date": (subtask.start_date if subtask else task.start_date).isoformat() if (subtask.start_date if subtask else task.start_date) else None,
            "end_date": end_dt.isoformat() if end_dt else None,
            "progress": progress,
            "is_overdue": is_overdue,
            "due_bucket": bucket,
            "priority": getattr(task, "priority", "medium"),
        })

    if not visible_ids:
        return jsonify([])

    cols = (
        ColumnA.query.options(
            joinedload(ColumnA.tasks).joinedload(TaskA.assignees),
            joinedload(ColumnA.tasks).joinedload(TaskA.subtasks).joinedload(TaskASubtask.assignees),
        )
        .filter(ColumnA.project_id.in_(visible_ids))
        .all()
    )
    for col in cols:
        project = col.project
        for task in col.tasks or []:
            if me_id in {u.id for u in (task.assignees or [])}:
                append_row("task", task, col, project)
            for st in task.subtasks or []:
                if me_id in {u.id for u in (st.assignees or [])}:
                    append_row("subtask", task, col, project, subtask=st)

    rows.sort(key=lambda r: (r["end_date"] or "9999", r["title"]))
    return jsonify(rows)


@app.route("/api/projects/<int:project_id>/labels", methods=["GET", "POST"])
@login_required
def project_a_labels(project_id):
    p = ProjectA.query.get_or_404(project_id)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if request.method == "GET":
        labels = TaskALabel.query.filter_by(project_id=p.id).order_by(TaskALabel.name).all()
        return jsonify([label_to_dict(l) for l in labels])
    data = request.get_json() or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "name is required"}), 400
    color = (data.get("color") or "#6366f1").strip()
    existing = TaskALabel.query.filter_by(project_id=p.id, name=name).first()
    if existing:
        return jsonify(label_to_dict(existing))
    lbl = TaskALabel(project_id=p.id, name=name, color=color)
    db.session.add(lbl)
    db.session.commit()
    return jsonify(label_to_dict(lbl)), 201


@app.route("/api/tasks/<int:task_id>/comments", methods=["GET", "POST"])
@login_required
def task_a_comments(task_id):
    t = _task_a_with_access(task_id)
    if request.method == "GET":
        comments = TaskAComment.query.filter_by(task_id=t.id).order_by(TaskAComment.created_at.desc()).all()
        return jsonify([comment_to_dict(c) for c in comments])
    data = request.get_json() or {}
    body = (data.get("body") or "").strip()
    if not body:
        return jsonify({"error": "body is required"}), 400
    if not user_can_comment_a(_project_for_task(t)):
        return jsonify({"error": "Forbidden"}), 403
    c = TaskAComment(task_id=t.id, user_id=current_user.id, body=body)
    db.session.add(c)
    db.session.flush()
    from app.pm_service import parse_mentions, highlight_mentions_html
    p = _project_for_task(t)
    mentioned = parse_mentions(body, p)
    for mu in mentioned:
        db.session.execute(
            taska_comment_mentions.insert().values(comment_id=c.id, user_id=mu.id)
        )
        if mu.id != current_user.id:
            link = url_for("projectmanagement", _external=False) + f"?project={p.id}&task={t.id}"
            db.session.add(Notification(
                user_id=mu.id,
                task_id=t.id,
                message=f"{current_user.username} mentioned you on: {t.title}. {link}",
                notification_type="pm_mention",
            ))
    for u in (t.assignees or []):
        if u.id != current_user.id and u not in mentioned:
            link = url_for("projectmanagement", _external=False) + f"?project={p.id}&task={t.id}"
            db.session.add(Notification(
                user_id=u.id,
                task_id=t.id,
                message=f"New comment on: {t.title}. {link}",
                notification_type="pm_comment",
            ))
    _log_task_activity(t, "comment", body[:200])
    db.session.commit()
    return jsonify(comment_to_dict(c)), 201


@app.route("/api/task-comments/<int:comment_id>", methods=["DELETE"])
@login_required
def delete_task_a_comment(comment_id):
    c = TaskAComment.query.get_or_404(comment_id)
    t = TaskA.query.get_or_404(c.task_id)
    p = _project_for_task(t)
    if not user_can_access_project_a(p):
        return project_access_denied()
    if c.user_id != current_user.id and not can_manage_project_a(p):
        return jsonify({"error": "Forbidden"}), 403
    db.session.delete(c)
    db.session.commit()
    return jsonify({"status": "deleted"})


@app.route("/api/tasks/<int:task_id>/activities", methods=["GET"])
@login_required
def task_a_activities(task_id):
    t = _task_a_with_access(task_id)
    acts = TaskAActivity.query.filter_by(task_id=t.id).order_by(TaskAActivity.created_at.desc()).limit(100).all()
    return jsonify([activity_to_dict(a) for a in acts])


@app.route("/api/meeting-notes/action-items/<int:action_item_id>/create-pm-task", methods=["POST"])
@login_required
def meeting_action_item_to_pm_task(action_item_id):
    from app.blueprints.meeting_notes.models import MeetingActionItem
    item = MeetingActionItem.query.get_or_404(action_item_id)
    data = request.get_json() or {}
    project_id = data.get("project_id")
    column_id = data.get("column_id")
    if not project_id or not column_id:
        return jsonify({"error": "project_id and column_id are required"}), 400
    p = ProjectA.query.get_or_404(int(project_id))
    if not user_can_access_project_a(p):
        return project_access_denied()
    col = ColumnA.query.filter_by(id=int(column_id), project_id=p.id).first()
    if not col:
        return jsonify({"error": "Invalid column for project"}), 400
    maxpos = db.session.query(db.func.max(TaskA.position)).filter_by(column_id=col.id).scalar()
    pos = (maxpos + 1) if maxpos is not None else 0
    t = TaskA(
        column_id=col.id,
        title=item.title or "Action item",
        description=item.comments or "",
        position=pos,
        start_date=None,
        end_date=datetime.combine(item.due_date, datetime.min.time()) if item.due_date else None,
        priority=(item.priority or "medium"),
        source_action_item_id=item.id,
        created_by=current_user.id,
    )
    assignee_ids = []
    if item.assignee_user_id:
        assignee_ids = [item.assignee_user_id]
    if assignee_ids:
        users, err = validate_assignee_user_ids(assignee_ids)
        if not err:
            t.assignees = users
    db.session.add(t)
    _log_task_activity(t, "created", f"Created from meeting action item #{item.id}")
    db.session.commit()
    return jsonify(task_to_dict(t)), 201


#new project planning ends here


# ---------- Audit log viewer (admin-only) ----------
@app.route("/admin/audit-log", methods=["GET"])
@login_required
def admin_audit_log():
    if getattr(current_user, "userRole", None) != "Admin":
        flash("Only admins can view the audit log.", "error")
        return redirect(url_for("overview"))

    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 50, type=int)
    per_page = max(10, min(per_page, 200))

    entity_type = (request.args.get("entity_type") or "").strip()
    actor_username = (request.args.get("actor") or "").strip()
    action_filter = (request.args.get("action") or "").strip()

    q = AuditLog.query
    if entity_type:
        q = q.filter(AuditLog.entity_type == entity_type)
    if actor_username:
        q = q.filter(AuditLog.actor_username.ilike(f"%{actor_username}%"))
    if action_filter:
        q = q.filter(AuditLog.action == action_filter)

    pagination = q.order_by(AuditLog.occurred_at.desc(), AuditLog.id.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    distinct_types = [
        row[0]
        for row in db.session.query(AuditLog.entity_type).distinct().order_by(AuditLog.entity_type).all()
        if row and row[0]
    ]

    return render_template(
        "admin/audit_log.html",
        pagination=pagination,
        entries=pagination.items,
        entity_type=entity_type,
        actor_username=actor_username,
        action_filter=action_filter,
        distinct_types=distinct_types,
        title="Audit Log",
    )


# ---------- New Dashboard Layout ----------
@app.route('/new_dash', methods=['GET'])
@login_required
def new_dash_layout():
    return render_template('new_dash_layout.html')

# ---------- Overview (dummy) ----------
@app.route('/')
@app.route('/overview', methods=['GET'])
@login_required
def overview():
    sm_overview = None
    mn_overview = None
    try:
        from app.blueprints.sales_marketing.services import can_access_sales_marketing, stakeholders_stats
        if can_access_sales_marketing():
            sm_overview = stakeholders_stats()
    except Exception:
        sm_overview = None
    try:
        from app.blueprints.meeting_notes.services import hub_analytics_summary
        mn_overview = hub_analytics_summary()
    except Exception:
        mn_overview = None
    return render_template(
        'overview.html',
        title="Akello Internal Dashboard",
        sm_overview=sm_overview,
        mn_overview=mn_overview,
    )
















@app.route('/targets/new', methods=['GET', 'POST'])
@login_required
def create_target():
    form = PerfomanceTargetsForm()
    if form.validate_on_submit():
        new_target = PerfomanceTargets(
            smartlearning_registrations_monthly_target=form.smartlearning_registrations_monthly_target.data,
            smartlearning_registrations_daily_target=form.smartlearning_registrations_daily_target.data,
            smartlearning_unique_subscribers_monthly_target=form.smartlearning_unique_subscribers_monthly_target.data,
            smartlearning_unique_subscribers_daily_target=form.smartlearning_unique_subscribers_daily_target.data,
            ask_akello_users_monthly_target=form.ask_akello_users_monthly_target.data,
            ask_akello_users_daily_target=form.ask_akello_users_daily_target.data,
            library_registrations_monthly_target=form.library_registrations_monthly_target.data,
            library_registrations_daily_target=form.library_registrations_daily_target.data,
            library_unique_users_monthly_target=form.library_unique_users_monthly_target.data,
            library_unique_users_daily_target=form.library_unique_users_daily_target.data,
            overall_active30_target=form.overall_active30_target.data,
            updated_by= current_user.username
        )
        db.session.add(new_target)
        db.session.commit()
        flash("Performance target created successfully!", "success")
        return redirect(url_for('analytics'))



@app.route('/targets/<int:target_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_target(target_id):
    target = PerfomanceTargets.query.get_or_404(target_id)
    form = PerfomanceTargetsForm(obj=target)
    if form.validate_on_submit():
        form.populate_obj(target)
        db.session.commit()
        flash("Performance target updated successfully!", "success")
        return redirect(url_for('analytics'))




@app.route('/all-champion-details', methods=['GET', 'POST'])
@login_required
def all_champion_details():

    return render_template('all_champions_data.html',
                           title='Champions')


@app.route('/champion-my-schools', methods=['GET'])
@login_required
def champion_my_schools():
    """My schools page for Brand Ambassadors: search, request additions, view current and pending."""
    if getattr(current_user, 'userRole', None) != 'Brand Ambassador':
        return redirect(url_for('index'))
    champion, _ = _resolve_champion_for_current_user()
    champion_province = (champion.province or '').strip() if champion else ''
    return render_template('champion_my_schools.html', title='My Schools', champion_province=champion_province)


@app.route('/champion-school-requests', methods=['GET'])
@login_required
def champion_school_requests_admin():
    """Admin page to list and approve/decline champion school requests."""
    if not can_approve_champion_schools():
        return redirect(url_for('index'))
    return render_template('champion_school_requests_admin.html', title='Champion School Requests')





# @app.route('/provincestats/<provincename>', methods=['GET', 'POST'])
# @login_required
# def provincestats(provincename):
#         province_champions = ChampionSchool.query.filter(ChampionSchool.province == provincename)
#         local_province_champion = ChampionSchool.query.filter(ChampionSchool.province == provincename)
#         today = datetime.today().date()
#         nhasi = today.strftime('%y- %m- %d')
#         month_name = today.strftime('%B')
        

#         return render_template('province_stats.html',
#                                provincename=provincename,
#                                province_champions=province_champions,
#                                local_province_champion=local_province_champion,
#                                nhasi = nhasi,
#                                month_name=month_name,
#                            title='School Tracker')


@app.route('/provincestats/<provincename>', methods=['GET', 'POST'])
@login_required
def provincestats(provincename):
        province_champions = ChampionSchool.query.filter(ChampionSchool.province == provincename)
        local_province_champion = ChampionSchool.query.filter(ChampionSchool.province == provincename)
        today = datetime.today().date()
        nhasi = today.strftime('%y- %m- %d')
        month_name = today.strftime('%B')
        

        return render_template('simone_province_stats.html',
                               provincename=provincename,
                               province_champions=province_champions,
                               local_province_champion=local_province_champion,
                               nhasi = nhasi,
                               month_name=month_name,
                           title='School Tracker')


@app.route('/api/all-champions-ask-data', methods=['GET'])
@login_required
def all_champions_ask_data():
    """Fast aggregated Ask Akello chatlog counts per champion.
    Allows optional ?start_date and ?end_date query params.
    Defaults to first of the current month through now().
    """
    # Handle date range with hours
    today = datetime.now()
    default_start_date = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        else:
            start_date = default_start_date

        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
        else:
            end_date = today
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD or YYYY-MM-DD HH:MM:SS."}), 400

    # Load champions and local champions once
    champions = ChampionSchool.query.all()
    local_champions = User.query.filter(User.userRole == 'Brand Ambassador').all()
    user_index = {(
        (u.firstname or '').strip().lower(),
        (u.lastname or '').strip().lower(),
        (u.province or '').strip().lower()
    ): u.username for u in local_champions}

    # Build mapping: champion -> ASL school IDs
    champ_to_asl_ids = {}
    all_asl_ids = []
    for champ in champions:
        schools = champ.get_schools() or []
        ids = [int(s.get('asl_school_id')) for s in schools if str(s.get('asl_school_id') or '').strip().isdigit()]
        if ids:
            champ_to_asl_ids[champ.id] = ids
            all_asl_ids.extend(ids)

    if not all_asl_ids:
        return jsonify([])

    # Deduplicate IDs
    all_asl_ids = sorted(set(all_asl_ids))

    # Query once: chatlog counts per school_id within date range
    counts_by_school = {}
    try:
        conn = get_ruzivo_conn()
        if conn is None:
            return jsonify({"error": "Database connection not available"}), 500
        import pymysql.cursors
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        CHUNK = 1000
        for i in range(0, len(all_asl_ids), CHUNK):
            chunk = all_asl_ids[i:i + CHUNK]
            placeholders = ','.join(['%s'] * len(chunk))
            
            # Updated query (no hardcoded school_id, dynamic IN clause)
            query = f"""
                SELECT 
                    sc.school_id,
                    sc.school_name,
                    sc.school_province,
                    COUNT(DISTINCT tl.student_id) AS chatlog_count
                FROM tblask_akello_chat_logs tl
                LEFT JOIN tblstudents ts ON ts.student_id = tl.student_id
                LEFT JOIN tblschools sc ON sc.school_id = ts.school_id
                WHERE sc.school_id IN ({placeholders})
                  AND tl.created_at BETWEEN %s AND %s
                GROUP BY sc.school_id, sc.school_name, sc.school_province
            """

            params = chunk + [start_date, end_date]
            cursor.execute(query, params)
            for row in cursor.fetchall():
                sid = int(row['school_id'])
                counts_by_school[sid] = counts_by_school.get(sid, 0) + int(row['chatlog_count'] or 0)
    finally:
        try:
            cursor.close()
        except Exception:
            pass
        try:
            conn.close()
        except Exception:
            pass

    # Build response per champion
    results = []
    for champ in champions:
        asl_ids = champ_to_asl_ids.get(champ.id, [])
        if not asl_ids:
            continue
        chatlog_sum = sum(counts_by_school.get(int(sid), 0) for sid in asl_ids)
        username = user_index.get((
            (champ.firstname or '').strip().lower(),
            (champ.lastname or '').strip().lower(),
            (champ.province or '').strip().lower()
        ))
        results.append({
            'champion': f"{champ.firstname} {champ.lastname}",
            'username': username,
            'province': champ.province,
            'school_count': len(asl_ids),
            'chatlog_count': chatlog_sum
        })

    return jsonify(results)


@app.route('/api/all-champions-checkins', methods=['GET'])
@login_required
def all_champions_checkins():
    """Return school check-in/check-out activity for champions."""
    tzinfo, _ = _get_checkin_reports_timezone()
    today = datetime.now(tzinfo).replace(tzinfo=None)
    default_start_date = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        else:
            start_date = default_start_date

        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
        else:
            end_date = today
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD or YYYY-MM-DD HH:MM:SS."}), 400

    return jsonify(
        _build_weekly_checkin_report_rows(
            start_date,
            end_date,
            include_admin_fields=_is_visit_admin_viewer(),
        )
    )


@app.route('/api/checkin-reports/weekly', methods=['GET'])
@login_required
def api_checkin_reports_weekly():
    if not can_view_checkin_reports():
        return jsonify({'error': 'Unauthorized'}), 403
    week_ending_str = request.args.get('week_ending')
    try:
        payload = _get_weekly_checkin_report_payload(week_ending_str=week_ending_str, role_scope='auto')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    return jsonify({
        'success': True,
        'timezone': payload['timezone'],
        'week_ending': payload['week_ending_date'].isoformat(),
        'window_start': payload['window']['start_local'].isoformat(),
        'window_end': payload['window']['end_local'].isoformat(),
        'scope': payload['scope'],
        'total_rows': len(payload['rows']),
        'rows': payload['rows'],
    })


@app.route('/api/checkin-reports/weekly.csv', methods=['GET'])
@login_required
def api_checkin_reports_weekly_csv():
    if not can_view_checkin_reports():
        return jsonify({'error': 'Unauthorized'}), 403
    week_ending_str = request.args.get('week_ending')
    try:
        payload = _get_weekly_checkin_report_payload(week_ending_str=week_ending_str, role_scope='auto')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    csv_bytes = _weekly_checkin_rows_to_csv_bytes(payload['rows'])
    filename = _weekly_checkin_filename(payload['scope'], payload['week_ending_date'], 'csv')
    return Response(
        csv_bytes,
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )


@app.route('/api/checkin-reports/weekly.pdf', methods=['GET'])
@login_required
def api_checkin_reports_weekly_pdf():
    if not can_view_checkin_reports():
        return jsonify({'error': 'Unauthorized'}), 403
    week_ending_str = request.args.get('week_ending')
    try:
        payload = _get_weekly_checkin_report_payload(week_ending_str=week_ending_str, role_scope='auto')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    role_label = 'Admin' if payload['scope'] == 'admin' else 'Champion'
    pdf_bytes = _weekly_checkin_rows_to_pdf_bytes(
        payload['rows'],
        role_label=role_label,
        week_ending_date=payload['week_ending_date']
    )
    filename = _weekly_checkin_filename(payload['scope'], payload['week_ending_date'], 'pdf')
    return Response(
        pdf_bytes,
        mimetype='application/pdf',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )


def run_weekly_checkin_report_job(triggered_by='scheduler', week_ending_str=None, force_send=False):
    tzinfo, _ = _get_checkin_reports_timezone()
    week_ending_date = _resolve_week_ending_date(week_ending_str=week_ending_str, tzinfo=tzinfo)
    week_key = week_ending_date.isoformat()

    if not force_send:
        last_sent_week = (AppSetting.get_value('checkin_reports_last_sent_week', '') or '').strip()
        if triggered_by == 'scheduler' and last_sent_week == week_key:
            return {'status': 'skipped', 'reason': 'already_sent', 'week_ending': week_key}

    window = _compute_checkin_week_window(week_ending_date, tzinfo)
    admin_rows = _build_weekly_checkin_report_rows(window['start_utc'], window['end_utc'], champion_user_id=None)
    champions = User.query.filter(User.userRole == 'Brand Ambassador').all()
    admins = User.query.filter(User.userRole == 'Admin').all()

    admin_csv = _weekly_checkin_rows_to_csv_bytes(admin_rows)
    admin_pdf = _weekly_checkin_rows_to_pdf_bytes(admin_rows, role_label='Admin', week_ending_date=week_ending_date)
    admin_result = _send_weekly_checkin_report_email(
        recipients=[u.email for u in admins if u.email],
        role_label='Admin',
        week_ending_date=week_ending_date,
        csv_bytes=admin_csv,
        csv_name=_weekly_checkin_filename('admin', week_ending_date, 'csv'),
        pdf_bytes=admin_pdf,
        pdf_name=_weekly_checkin_filename('admin', week_ending_date, 'pdf'),
        total_count=len(admin_rows)
    )

    champion_results = []
    for champion in champions:
        if not champion.email:
            champion_results.append({'username': champion.username, 'status': 'skipped', 'reason': 'missing_email'})
            continue
        champion_rows = _build_weekly_checkin_report_rows(window['start_utc'], window['end_utc'], champion_user_id=champion.id)
        csv_bytes = _weekly_checkin_rows_to_csv_bytes(champion_rows)
        pdf_bytes = _weekly_checkin_rows_to_pdf_bytes(champion_rows, role_label='Champion', week_ending_date=week_ending_date)
        send_result = _send_weekly_checkin_report_email(
            recipients=[champion.email],
            role_label='Champion',
            week_ending_date=week_ending_date,
            csv_bytes=csv_bytes,
            csv_name=_weekly_checkin_filename(f"champion_{champion.username}", week_ending_date, 'csv'),
            pdf_bytes=pdf_bytes,
            pdf_name=_weekly_checkin_filename(f"champion_{champion.username}", week_ending_date, 'pdf'),
            total_count=len(champion_rows)
        )
        send_result['username'] = champion.username
        champion_results.append(send_result)

    sent_any = (admin_result.get('status') == 'sent') or any(r.get('status') == 'sent' for r in champion_results)
    if sent_any:
        AppSetting.set_value(
            'checkin_reports_last_sent_week',
            week_key,
            description='Last successfully generated weekly check-in report week ending date.'
        )
    app.logger.info(
        "Weekly check-in report run complete (trigger=%s week_ending=%s admin_status=%s champions=%d)",
        triggered_by,
        week_key,
        admin_result.get('status'),
        len(champion_results)
    )

    return {
        'status': 'ok',
        'triggered_by': triggered_by,
        'week_ending': week_key,
        'admin_result': admin_result,
        'champion_results': champion_results,
        'admin_rows': len(admin_rows),
    }





@app.route('/api/ask-akello-chosen-school/<school_id>', methods=['GET'])
@login_required
def ask_akello_chosen_school(school_id):
    """
    Returns total chat log counts for a given school_id
    with optional date range support (defaults to MTD).
    """

    # Get date range from query parameters or default to MTD
    today = datetime.now()
    default_start_date = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        else:
            start_date = default_start_date

        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
        else:
            end_date = today
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD or YYYY-MM-DD HH:MM:SS."}), 400

    try:
        # Connect to ruzivo database
        pool = get_ruzivo_pool()
        if pool is None:
            return jsonify({"error": "Database connection not available"}), 500

        conn = pool.connection()
        import pymysql.cursors
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        # Updated query (Ask Akello style)
        chatlog_query = """
            SELECT 
                sc.school_id,
                sc.school_name,
                sc.school_province,
                COUNT(DISTINCT tl.student_id) AS chatlog_count
            FROM tblask_akello_chat_logs tl
            LEFT JOIN tblstudents ts 
                ON ts.student_id = tl.student_id
            LEFT JOIN tblschools sc 
                ON sc.school_id = ts.school_id
            WHERE 
                tl.created_at BETWEEN %s AND %s
                AND sc.school_id = %s
            GROUP BY 
                sc.school_id, sc.school_name, sc.school_province
        """

        # Execute query
        cursor.execute(chatlog_query, (start_date, end_date, school_id))
        result = cursor.fetchone()

        response = {
            "school_id": school_id,
            "school_name": result["school_name"] if result else None,
            "school_province": result["school_province"] if result else None,
            "chatlog_count": result["chatlog_count"] if result else 0,
            "start_date": str(start_date),
            "end_date": str(end_date),
        }

        return jsonify(response)

    except Exception as e:
        # Error logging disabled - silently return error response
        return jsonify({"error": str(e)}), 500

    finally:
        if 'cursor' in locals() and cursor:
            cursor.close()
        if 'conn' in locals() and conn:
            conn.close()





@app.route('/api/champion-schools-data/<province_name>', methods=['GET'])
@login_required
@cache.cached(
    timeout=300, 
    key_prefix=lambda: f'ask_akello_{request.view_args["province_name"]}_{request.args.get("start_date", "")}_{request.args.get("end_date", "")}'
)
def champion_schools_data(province_name):
    """
    Returns total Ask Akello chat log counts per Champion within a province.
    Uses optional ?start_date and ?end_date (with hours).
    Defaults to first day of current month through now().
    """

    # --- Handle date range ---
    today = datetime.now()
    default_start_date = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        else:
            start_date = default_start_date

        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
        else:
            end_date = today
    except ValueError:
        return jsonify({
            "error": "Invalid date format. Use YYYY-MM-DD or YYYY-MM-DD HH:MM:SS."
        }), 400

    try:
        # --- Load champions & local Brand Ambassadors for this province ---
        champions = ChampionSchool.query.filter_by(province=province_name).all()
        local_champions = User.query.filter(
            User.province == province_name,
            User.userRole == 'Brand Ambassador'
        ).all()

        results = []

        for champ in champions:
            schools = champ.get_schools() or []
            asl_ids = [s.get('asl_school_id') for s in schools if s.get('asl_school_id')]

            if not asl_ids:
                continue

            conn = None
            cursor = None
            try:
                pool = get_ruzivo_pool()
                if pool is None:
                    raise Exception("Database connection not available")
                conn = pool.connection()
                import pymysql.cursors
                cursor = conn.cursor(pymysql.cursors.DictCursor)

                # --- Updated Ask Akello query ---
                format_strings = ','.join(['%s'] * len(asl_ids))
                chatlog_query = f"""
                    SELECT 
                        sc.school_id,
                        sc.school_name,
                        sc.school_province,
                        COUNT(DISTINCT tl.student_id) AS chatlog_count
                    FROM tblask_akello_chat_logs tl
                    LEFT JOIN tblstudents ts 
                        ON ts.student_id = tl.student_id
                    LEFT JOIN tblschools sc 
                        ON sc.school_id = ts.school_id
                    WHERE 
                        sc.school_id IN ({format_strings})
                        AND tl.created_at BETWEEN %s AND %s
                    GROUP BY 
                        sc.school_id, sc.school_name, sc.school_province;
                """

                params = tuple(asl_ids) + (start_date, end_date)
                cursor.execute(chatlog_query, params)
                result_rows = cursor.fetchall()
                chatlog_count = sum(r['chatlog_count'] for r in result_rows) if result_rows else 0

            except Exception as e:
                logging.error(
                    f"Error fetching chatlog for champion {champ.firstname} {champ.lastname}: {e}"
                )
                chatlog_count = 0

            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()

            # --- Match champion to local Brand Ambassador username ---
            username = None
            for local in local_champions:
                if (
                    (local.firstname or '').strip().lower() == (champ.firstname or '').strip().lower() and
                    (local.lastname or '').strip().lower() == (champ.lastname or '').strip().lower()
                ):
                    username = getattr(local, 'username', None)
                    break

            # --- Build result entry ---
            results.append({
                "champion": f"{champ.firstname} {champ.lastname}",
                "username": username,
                "school_count": len(asl_ids),
                "chatlog_count": chatlog_count
            })

        return jsonify(results)

    except Exception as e:
        logging.error(f"Error in champion_schools_data: {e}")
        return jsonify({"error": str(e)}), 500









@app.route("/provincialmoneylibraryallocation", methods=["GET"])
def provincial_money_library_allocation():
    try:
        # --- get date range ---
        today = date.today()
        first_of_month = today.replace(day=1)

        start_date = request.args.get("start_date", first_of_month.isoformat())
        end_date = request.args.get("end_date", today.isoformat())

        # --- query ---
        query = """
            SELECT 
                i.province,
                SUM(o.total_amount) as total_amount
            FROM orders o
            INNER JOIN users ou ON o.user_id = ou.id
            INNER JOIN institution_user iu ON ou.id = iu.user_id
            INNER JOIN institutions i ON iu.institution_id = i.id
            INNER JOIN model_has_roles mr ON ou.id = mr.model_id
            INNER JOIN roles r ON mr.role_id = r.id
            INNER JOIN book_order bo ON o.id = bo.order_id
            WHERE r.name = 'Institution'
              AND o.currency = 'ZWL'
              AND o.payment_method = 'Voucher'
              AND o.is_hlf = 1
              AND o.status = 'Completed'
              AND o.payment_type = 'Purchase'
              AND o.created_at BETWEEN %s AND %s
            GROUP BY i.province
            ORDER BY i.province;
        """

        conn = get_direct_library_conn()
        cursor = conn.cursor()
        cursor.execute(query, (start_date, end_date))
        results = cursor.fetchall()

        cursor.close()
        conn.close()

        # --- format response ---
        total_sum = sum([row["total_amount"] for row in results if row["total_amount"]])
        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "data": results,
            "grand_total": total_sum
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500






@app.route("/api/publisherlibraryamount", methods=["GET"])
@login_required
def publisher_library_amount():
    try:
        # Default date range
        today = date.today()
        first_of_month = today.replace(day=1)

        start_date = request.args.get("start_date", first_of_month.isoformat())
        end_date = request.args.get("end_date", today.isoformat())

        query = """
            SELECT 
                p.id AS publisher_id,
                p.name AS publisher_name,
                SUM(CASE WHEN o.currency = 'ZWL' THEN o.total_amount ELSE 0 END) AS total_amount_zwl,
                SUM(CASE WHEN o.currency = 'USD' THEN o.total_amount ELSE 0 END) AS total_amount_usd
            FROM orders o
            INNER JOIN book_order bo 
                ON o.id = bo.order_id
            INNER JOIN books b 
                ON bo.book_id = b.id
            INNER JOIN publisher_user pu 
                ON b.user_id = pu.user_id
            INNER JOIN publishers p 
                ON pu.publisher_id = p.id
            WHERE o.created_at BETWEEN %s AND %s
              AND o.currency IN ('ZWL', 'USD')
            GROUP BY p.id, p.name

            UNION ALL

            SELECT 
                NULL AS publisher_id,
                'Grand Total' AS publisher_name,
                SUM(CASE WHEN o.currency = 'ZWL' THEN o.total_amount ELSE 0 END) AS total_amount_zwl,
                SUM(CASE WHEN o.currency = 'USD' THEN o.total_amount ELSE 0 END) AS total_amount_usd
            FROM orders o
            INNER JOIN book_order bo 
                ON o.id = bo.order_id
            INNER JOIN books b 
                ON bo.book_id = b.id
            INNER JOIN publisher_user pu 
                ON b.user_id = pu.user_id
            INNER JOIN publishers p 
                ON pu.publisher_id = p.id
            WHERE o.created_at BETWEEN %s AND %s
              AND o.currency IN ('ZWL', 'USD')
            ORDER BY publisher_name;
        """

        conn = get_direct_library_conn()
        cursor = conn.cursor()
        cursor.execute(query, (start_date, end_date, start_date, end_date))
        results = cursor.fetchall()

        cursor.close()
        conn.close()

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "data": results
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500




















@app.route('/api/champion-schools-smartlearning-usage-analytics', methods=['GET'])
@login_required
def champion_schools_smartlearning_usage_analytics():
    """Fast aggregated analytics for champions across ASL and Library.
    Performs grouped queries across all relevant schools, then reduces per champion.
    """
    today = datetime.today().date()
    default_start_date = today.replace(day=1)

    # --- Allow custom date range via query params ---
    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        else:
            start_date = default_start_date

        if end_date_str:
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        else:
            end_date = today
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400

    # Get all champions (no province filter) and map usernames for O(1) lookup
    champions = ChampionSchool.query.all()
    local_champions = User.query.filter(User.userRole == 'Brand Ambassador').all()
    user_index = {(
        (u.firstname or '').strip().lower(),
        (u.lastname or '').strip().lower(),
        (u.province or '').strip().lower()
    ): u.username for u in local_champions}

    # Build mappings and flattened ID sets
    champ_asl = {}
    champ_lib = {}
    all_asl_ids, all_lib_ids = [], []
    for champ in champions:
        schools = champ.get_schools() or []
        asl_ids = [int(s.get('asl_school_id')) for s in schools if str(s.get('asl_school_id') or '').strip().isdigit()]
        lib_ids = [int(s.get('library_school_id')) for s in schools if str(s.get('library_school_id') or '').strip().isdigit()]
        if asl_ids:
            champ_asl[champ.id] = asl_ids
            all_asl_ids.extend(asl_ids)
        if lib_ids:
            champ_lib[champ.id] = lib_ids
            all_lib_ids.extend(lib_ids)

    all_asl_ids = sorted(set(all_asl_ids))
    all_lib_ids = sorted(set(all_lib_ids))

    # ✅ ASL MTD Filtering: Exclude schools based on scholarship duration/award date
    # Fetch admin settings
    from app.models import AppSetting
    from datetime import date, timedelta
    exclude_12_months = AppSetting.get_value('asl_mtd_exclude_12_months', 'true') == 'true'
    exclude_1_year_awarded = AppSetting.get_value('asl_mtd_exclude_1_year_awarded', 'true') == 'true'
    
    schools_to_exclude = set()
    if all_asl_ids and (exclude_12_months or exclude_1_year_awarded):
        try:
            pool = get_ruzivo_pool()
            if pool:
                conn = pool.connection()
                cursor = conn.cursor(pymysql.cursors.DictCursor)
                
                # Fetch durations and first awards since 2025
                CHUNK = 1000
                for i in range(0, len(all_asl_ids), CHUNK):
                    chunk = all_asl_ids[i:i+CHUNK]
                    placeholders = ','.join(['%s'] * len(chunk))
                    duration_query = f"""
                        SELECT x.school_id, SUM(x.duration) AS total_months, MIN(x.awarded_at) AS first_awarded_at
                        FROM tblscholarships_schools x
                        WHERE x.school_id IN ({placeholders})
                        AND x.awarded_at >= %s
                        GROUP BY x.school_id
                    """
                    cursor.execute(duration_query, (*chunk, '2025-01-01'))
                    duration_rows = cursor.fetchall()
                    
                    for row in duration_rows:
                        sid = int(row['school_id'])
                        total_m = row['total_months'] or 0
                        fa = row['first_awarded_at']
                        
                        is_over_1_year = False
                        if fa:
                            fa_dt = fa.date() if hasattr(fa, 'date') else fa
                            if isinstance(fa_dt, date) and fa_dt <= today - timedelta(days=365):
                                is_over_1_year = True
                        
                        should_exclude = False
                        if exclude_12_months and total_m >= 12:
                            should_exclude = True
                        if exclude_1_year_awarded and is_over_1_year:
                            should_exclude = True
                            
                        if should_exclude:
                            schools_to_exclude.add(sid)
                
                cursor.close()
                conn.close()
        except Exception as e:
            app.logger.error(f"Error filtering schools for champion analytics: {e}")

    # Use filtered ASL IDs for student count queries
    final_asl_ids = [sid for sid in all_asl_ids if sid not in schools_to_exclude]
    app.logger.info(f"ASL MTD Filter: Total {len(all_asl_ids)} schools, excluded {len(schools_to_exclude)}, final {len(final_asl_ids)}")

    # Time bounds as datetimes for BETWEEN comparisons
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date, datetime.max.time())

    # 1) ASL: count distinct student_id per school_id, grouped once
    asl_counts_by_school = {}
    if all_asl_ids:
        try:
            pool = get_ruzivo_pool()
            if pool is None:
                return jsonify({"error": "Database connection not available"}), 500
            conn = pool.connection()
            import pymysql.cursors
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            CHUNK = 1000
            for i in range(0, len(final_asl_ids), CHUNK):
                chunk = final_asl_ids[i:i+CHUNK]
                placeholders = ','.join(['%s'] * len(chunk))
                query = f"""
                    SELECT school_id, COUNT(DISTINCT student_id) AS student_count
                    FROM vwstudent
                    WHERE school_id IN ({placeholders})
                      AND last_login BETWEEN %s AND %s
                    GROUP BY school_id
                """
                params = chunk + [start_dt, end_dt]
                cursor.execute(query, params)
                for row in cursor.fetchall():
                    sid = int(row['school_id'])
                    asl_counts_by_school[sid] = asl_counts_by_school.get(sid, 0) + int(row['student_count'] or 0)
        finally:
            try:
                cursor.close()
            except Exception:
                pass
            try:
                conn.close()
            except Exception:
                pass

    # 2) Library: count distinct la.user_id per institution, grouped once
    lib_counts_by_inst = {}
    if all_lib_ids:
        try:
            conn = get_direct_library_conn()
            import pymysql.cursors
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            CHUNK = 1000
            for i in range(0, len(all_lib_ids), CHUNK):
                chunk = all_lib_ids[i:i+CHUNK]
                placeholders = ','.join(['%s'] * len(chunk))
                query = f"""
                    SELECT iu.institution_id AS inst_id, COUNT(DISTINCT la.user_id) AS active_users
                    FROM logins la
                    JOIN institution_user iu ON la.user_id = iu.user_id
                    WHERE iu.institution_id IN ({placeholders})
                      AND la.created_at BETWEEN %s AND %s
                    GROUP BY iu.institution_id
                """
                params = chunk + [start_dt, end_dt]
                cursor.execute(query, params)
                for row in cursor.fetchall():
                    iid = int(row['inst_id'])
                    lib_counts_by_inst[iid] = lib_counts_by_inst.get(iid, 0) + int(row['active_users'] or 0)
        finally:
            try:
                cursor.close()
            except Exception:
                pass
            try:
                conn.close()
            except Exception:
                pass

    # Build response rows per champion by summing precomputed counts
    results = []
    for champ in champions:
        asl_ids = champ_asl.get(champ.id, [])
        lib_ids = champ_lib.get(champ.id, [])
        if not asl_ids and not lib_ids:
            continue

        asl_student_count = sum(asl_counts_by_school.get(int(sid), 0) for sid in asl_ids)
        library_student_count = sum(lib_counts_by_inst.get(int(iid), 0) for iid in lib_ids)
        overall_student_count = asl_student_count + library_student_count

        username = user_index.get((
            (champ.firstname or '').strip().lower(),
            (champ.lastname or '').strip().lower(),
            (champ.province or '').strip().lower()
        ))

        results.append({
            "champion": f"{champ.firstname} {champ.lastname}",
            "username": username,
            "province": f"{champ.province}",
            "school_count": len(asl_ids),
            "asl_student_count": asl_student_count,
            "library_student_count": library_student_count,
            "overall_student_count": overall_student_count
        })

    return jsonify(results)





@app.route('/api/champ-library-usage', methods=['GET'])
@login_required
def champ_library_usage():
    """AL MTD: current month with per-institution period (1st to min(end of month, first_awarded+365) from linked ASL school). Ignores date params for count."""
    firstname = request.args.get("firstname")
    lastname = request.args.get("lastname")

    if not firstname or not lastname:
        return jsonify({"error": "Missing required parameters: firstname, lastname"}), 400

    today = datetime.today().date()
    first_of_month, end_of_month = _mtd_month_bounds(today)

    champ = ChampionSchool.query.filter_by(firstname=firstname, lastname=lastname).first()
    if not champ:
        return jsonify({"error": "Champion not found"}), 404

    schools = champ.get_schools() or []
    library_ids = [s.get('library_school_id') for s in schools if s.get('library_school_id')]
    if not library_ids:
        return jsonify({
            "champion": f"{champ.firstname} {champ.lastname}",
            "library_student_count": 0
        })

    asl_school_ids = [s.get('asl_school_id') for s in schools if s.get('asl_school_id')]
    first_awarded_by_asl = {}
    if asl_school_ids:
        try:
            rconn = get_ruzivo_conn()
            rc = rconn.cursor()
            ph = ','.join(['%s'] * len(asl_school_ids))
            rc.execute(
                f"SELECT school_id, MIN(awarded_at) FROM tblscholarships_schools WHERE school_id IN ({ph}) AND awarded_at >= %s GROUP BY school_id",
                (*asl_school_ids, '2025-01-01')
            )
            for r in rc.fetchall():
                first_awarded_by_asl[r[0]] = r[1]
            rc.close()
            rconn.close()
        except Exception as e:
            logging.error("champ_library_usage Ruzivo first_awarded: %s", e)

    period_end_by_lib_id = {}
    for s in schools:
        lib_id = s.get('library_school_id')
        asl_id = s.get('asl_school_id')
        if not lib_id:
            continue
        try:
            lib_key = int(lib_id)
        except (TypeError, ValueError):
            lib_key = lib_id
        fa = first_awarded_by_asl.get(asl_id) if asl_id else None
        period_end_by_lib_id[lib_key] = _mtd_period_end(fa, end_of_month)

    library_student_count = 0
    conn = None
    cursor = None
    try:
        conn = get_direct_library_conn()
        if conn:
            import pymysql.cursors
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            library_ids_int = [int(sid) for sid in library_ids if sid]
            if library_ids_int:
                in_placeholders = ', '.join(['%s'] * len(library_ids_int))
                start_dt = datetime.combine(first_of_month, datetime.min.time())
                end_dt = datetime.combine(end_of_month, datetime.max.time())
                query = f"""
                    SELECT iu.institution_id, la.user_id, la.created_at
                    FROM logins la
                    JOIN institution_user iu ON la.user_id = iu.user_id
                    WHERE iu.institution_id IN ({in_placeholders})
                      AND la.created_at BETWEEN %s AND %s
                """
                cursor.execute(query, library_ids_int + [start_dt, end_dt])
                rows = cursor.fetchall()
                counted = defaultdict(set)
                for row in rows:
                    iid = row.get("institution_id")
                    uid = row.get("user_id")
                    created = row.get("created_at")
                    created_date = created.date() if hasattr(created, 'date') else created
                    pe = period_end_by_lib_id.get(iid, end_of_month)
                    if created_date <= pe:
                        counted[iid].add(uid)
                library_student_count = sum(len(s) for s in counted.values())
        else:
            logging.error("Failed to get Library DB connection")
    except Exception as e:
        logging.error("Library query error for champion %s %s: %s", champ.firstname, champ.lastname, e)
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

    return jsonify({
        "champion": f"{champ.firstname} {champ.lastname}",
        "province": champ.province,
        "library_student_count": library_student_count
    })







@app.route('/api/champion-schools-smartlearning-usage/<province_name>', methods=['GET'])
@login_required
@cache.cached(timeout=300, key_prefix=lambda: f'champion_usage_{request.view_args["province_name"]}_{request.args.get("start_date", "")}_{request.args.get("end_date", "")}')
def champion_schools_smartlearning_usage(province_name):
    """OPTIMIZED: Bulk queries instead of N+1 - fetches all data in 2 queries total."""
    # Get date range from query parameters or default to MTD
    today = datetime.today().date()
    default_start_date = today.replace(day=1)

    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    try:
        if start_date_str:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        else:
            start_date = default_start_date

        if end_date_str:
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        else:
            end_date = today
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400

    # Get champions for this province
    champions = ChampionSchool.query.filter_by(province=province_name).all()
    local_champions = User.query.filter(
        User.province == province_name,
        User.userRole == 'Brand Ambassador'
    ).all()

    # Build username lookup index
    user_index = {(
        (u.firstname or '').strip().lower(),
        (u.lastname or '').strip().lower(),
        (u.province or '').strip().lower()
    ): u.username for u in local_champions}

    # Build mappings and collect all IDs
    champ_asl = {}
    champ_lib = {}
    all_asl_ids, all_lib_ids = [], []
    
    for champ in champions:
        schools = champ.get_schools() or []
        asl_ids = [int(s.get('asl_school_id')) for s in schools if str(s.get('asl_school_id') or '').strip().isdigit()]
        lib_ids = [int(s.get('library_school_id')) for s in schools if str(s.get('library_school_id') or '').strip().isdigit()]
        
        if asl_ids:
            champ_asl[champ.id] = asl_ids
            all_asl_ids.extend(asl_ids)
        if lib_ids:
            champ_lib[champ.id] = lib_ids
            all_lib_ids.extend(lib_ids)

    all_asl_ids = sorted(set(all_asl_ids))
    all_lib_ids = sorted(set(all_lib_ids))

    # Early return if nothing to process
    if not all_asl_ids and not all_lib_ids:
        return jsonify([])

    # Time bounds
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date, datetime.max.time())

    # BULK QUERY 1: ASL - fetch all school counts in one query
    asl_counts_by_school = {}
    if all_asl_ids:
        conn = None
        cursor = None
        try:
            pool = get_ruzivo_pool()
            if pool is None:
                logging.error("Ruzivo pool not available")
            else:
                conn = pool.connection()
                import pymysql.cursors
                cursor = conn.cursor(pymysql.cursors.DictCursor)
                
                # Process in chunks if needed
                CHUNK = 1000
                for i in range(0, len(all_asl_ids), CHUNK):
                    chunk = all_asl_ids[i:i+CHUNK]
                    placeholders = ','.join(['%s'] * len(chunk))
                    query = f"""
                        SELECT school_id, COUNT(DISTINCT student_id) AS student_count
                        FROM vwstudent
                        WHERE school_id IN ({placeholders})
                          AND last_login BETWEEN %s AND %s
                        GROUP BY school_id
                    """
                    params = chunk + [start_dt, end_dt]
                    cursor.execute(query, params)
                    
                    for row in cursor.fetchall():
                        sid = int(row['school_id'])
                        asl_counts_by_school[sid] = int(row['student_count'] or 0)
        except Exception as e:
            logging.error("ASL bulk query error: %s", e)
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()

    # BULK QUERY 2: Library - fetch all institution counts in one query
    lib_counts_by_inst = {}
    if all_lib_ids:
        conn = None
        cursor = None
        try:
            conn = get_direct_library_conn()
            if conn:
                import pymysql.cursors
                cursor = conn.cursor(pymysql.cursors.DictCursor)
                
                # Process in chunks if needed
                CHUNK = 1000
                for i in range(0, len(all_lib_ids), CHUNK):
                    chunk = all_lib_ids[i:i+CHUNK]
                    placeholders = ','.join(['%s'] * len(chunk))
                    query = f"""
                        SELECT iu.institution_id AS inst_id, COUNT(DISTINCT la.user_id) AS active_users
                        FROM logins la
                        JOIN institution_user iu ON la.user_id = iu.user_id
                        WHERE iu.institution_id IN ({placeholders})
                          AND la.created_at BETWEEN %s AND %s
                        GROUP BY iu.institution_id
                    """
                    params = chunk + [start_dt, end_dt]
                    cursor.execute(query, params)
                    
                    for row in cursor.fetchall():
                        iid = int(row['inst_id'])
                        lib_counts_by_inst[iid] = int(row['active_users'] or 0)
        except Exception as e:
            logging.error("Library bulk query error: %s", e)
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()

    # Build results by aggregating precomputed counts
    results = []
    for champ in champions:
        asl_ids = champ_asl.get(champ.id, [])
        lib_ids = champ_lib.get(champ.id, [])
        
        if not asl_ids and not lib_ids:
            continue

        # Sum up counts from bulk query results
        asl_student_count = sum(asl_counts_by_school.get(int(sid), 0) for sid in asl_ids)
        library_student_count = sum(lib_counts_by_inst.get(int(iid), 0) for iid in lib_ids)
        overall_student_count = asl_student_count + library_student_count

        # Lookup username
        username = user_index.get((
            (champ.firstname or '').strip().lower(),
            (champ.lastname or '').strip().lower(),
            (champ.province or '').strip().lower()
        ))

        results.append({
            "champion": f"{champ.firstname} {champ.lastname}",
            "username": username,
            "school_count": len(asl_ids),
            "asl_student_count": asl_student_count,
            "library_student_count": library_student_count,
            "overall_student_count": overall_student_count
        })

    return jsonify(results)




def normalize_province_name(name):
    if not name:
        return "Unknown"
    return str(name).strip().title()

@app.route('/api/platforms/quick_custom_date', methods=['GET'])
@login_required
def platforms_quick_custom_date():
    try:
        # --- Parse dates ---
        end_date = request.args.get("end_date")
        start_date = request.args.get("start_date")

        today = datetime.today().date()
        default_start = today - timedelta(days=30)

        if not end_date:
            end_date = today
        else:
            end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

        if not start_date:
            start_date = default_start
        else:
            start_date = datetime.strptime(start_date, "%Y-%m-%d").date()

        # ---------------------------
        # Helper to normalize gender
        # ---------------------------
        def normalize_gender(g):
            if not g:
                return "Unknown"
            g = g.strip().lower()
            if g in ["f", "female", "girl", "woman"]:
                return "Female"
            elif g in ["m", "male", "boy", "man"]:
                return "Male"
            return "Unknown"

        # ---------------------------
        # 1) ASL PLATFORM (Ruzivo DB)
        # ---------------------------
        conn_asl = get_ruzivo_conn()
        cursor_asl = conn_asl.cursor(pymysql.cursors.DictCursor)

        asl_query = """
            SELECT 
                sch.school_id,
                sch.school_name,
                sch.school_province,
                COUNT(DISTINCT sl.student_id) AS active_learners
            FROM tblstudents_login sl
            INNER JOIN tblstudents st ON sl.student_id = st.student_id
            INNER JOIN tblschools sch ON st.school_id = sch.school_id
            WHERE DATE(sl.login_date) BETWEEN %s AND %s
            GROUP BY sch.school_id, sch.school_name, sch.school_province
        """
        cursor_asl.execute(asl_query, (start_date, end_date))
        asl_results = cursor_asl.fetchall()

        # Gender distribution
        asl_gender_query = """
            SELECT st.gender, COUNT(DISTINCT sl.student_id) AS count
            FROM tblstudents_login sl
            INNER JOIN tblstudents st ON sl.student_id = st.student_id
            WHERE DATE(sl.login_date) BETWEEN %s AND %s
            GROUP BY st.gender
        """
        cursor_asl.execute(asl_gender_query, (start_date, end_date))
        asl_gender_raw = cursor_asl.fetchall()
        asl_gender = {"Female": 0, "Male": 0, "Unknown": 0}
        for row in asl_gender_raw:
            g = normalize_gender(row["gender"])
            asl_gender[g] += row["count"]

        # ---------------------------
        # 2) LIBRARY PLATFORM
        # ---------------------------
        conn_lib = get_direct_library_conn()
        cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)

        lib_query = """
            SELECT 
                inst.id,
                inst.name AS institution_name,
                inst.province AS institution_province,
                COUNT(DISTINCT la.user_id) AS active_users
            FROM logins la
            INNER JOIN institution_user iu ON la.user_id = iu.user_id
            INNER JOIN institutions inst ON iu.institution_id = inst.id
            WHERE DATE(la.created_at) BETWEEN %s AND %s
            GROUP BY inst.id, inst.name, inst.province
        """
        cursor_lib.execute(lib_query, (start_date, end_date))
        lib_results = cursor_lib.fetchall()

        lib_gender_query = """
            SELECT u.sex AS gender, COUNT(DISTINCT la.user_id) AS count
            FROM logins la
            INNER JOIN users u ON la.user_id = u.id
            WHERE DATE(la.created_at) BETWEEN %s AND %s
            GROUP BY u.sex
        """
        cursor_lib.execute(lib_gender_query, (start_date, end_date))
        lib_gender_raw = cursor_lib.fetchall()
        lib_gender = {"Female": 0, "Male": 0, "Unknown": 0}
        for row in lib_gender_raw:
            g = normalize_gender(row["gender"])
            lib_gender[g] += row["count"]

        # ---------------------------
        # 3) ASK AKELLO CHAT LOGS
        # ---------------------------
        ask_query = """
            SELECT COUNT(DISTINCT student_id) AS unique_students_count
            FROM tblask_akello_chat_logs
            WHERE DATE(created_at) BETWEEN %s AND %s
        """
        cursor_asl.execute(ask_query, (start_date, end_date))
        ask_result = cursor_asl.fetchone()
        total_ask_users = ask_result["unique_students_count"] if ask_result else 0

        # Province distribution for AskAkello
        ask_province_query = """
            SELECT sch.school_province, COUNT(DISTINCT acl.student_id) AS count
            FROM tblask_akello_chat_logs acl
            INNER JOIN tblstudents st ON acl.student_id = st.student_id
            INNER JOIN tblschools sch ON st.school_id = sch.school_id
            WHERE DATE(acl.created_at) BETWEEN %s AND %s
            GROUP BY sch.school_province
        """
        cursor_asl.execute(ask_province_query, (start_date, end_date))
        ask_province = cursor_asl.fetchall()

        # Gender distribution for AskAkello
        ask_gender_query = """
            SELECT st.gender, COUNT(DISTINCT acl.student_id) AS count
            FROM tblask_akello_chat_logs acl
            INNER JOIN tblstudents st ON acl.student_id = st.student_id
            WHERE DATE(acl.created_at) BETWEEN %s AND %s
            GROUP BY st.gender
        """
        cursor_asl.execute(ask_gender_query, (start_date, end_date))
        ask_gender_raw = cursor_asl.fetchall()
        ask_gender = {"Female": 0, "Male": 0, "Unknown": 0}
        for row in ask_gender_raw:
            g = normalize_gender(row["gender"])
            ask_gender[g] += row["count"]

        # ---------------------------
        # Normalize province names
        # ---------------------------
        for row in asl_results:
            row["school_province"] = normalize_province_name(row["school_province"])
        for row in lib_results:
            row["institution_province"] = normalize_province_name(row["institution_province"])
        for row in ask_province:
            row["school_province"] = normalize_province_name(row["school_province"])

        # Totals
        total_asl_users = sum([r["active_learners"] for r in asl_results])
        total_lib_users = sum([r["active_users"] for r in lib_results])
        total_askakello_users = total_ask_users
        overall_total = total_asl_users + total_lib_users + total_askakello_users

        # ---------------------------
        # FINAL RESPONSE
        # ---------------------------
        return jsonify({
            "date_range": {"start": str(start_date), "end": str(end_date)},
            "totals": {
                "asl": total_asl_users,
                "library": total_lib_users,
                "ask_akello": total_askakello_users,
                "overall": overall_total
            },
            "province_distribution": {
                "asl": asl_results,
                "library": lib_results,
                "ask_akello": ask_province
            },
            "gender_distribution": {
                "asl": asl_gender,
                "library": lib_gender,
                "ask_akello": ask_gender
            }
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if 'cursor_asl' in locals(): cursor_asl.close()
        if 'conn_asl' in locals(): conn_asl.close()
        if 'cursor_lib' in locals(): cursor_lib.close()
        if 'conn_lib' in locals(): conn_lib.close()





# Cache already initialized at the top of the file after imports

@app.route('/api/platforms_overall_yearly', methods=['GET'])
@login_required
@cache.cached(timeout=60 * 60 * 6, key_prefix='platforms_overall_yearly')  # Reduced to 6 hours
def platforms_overall_yearly():
    import time
    import pymysql
    import threading
    import concurrent.futures
    
    start_time = time.time()
    
    def execute_query():
        """Execute the actual database queries"""
        return _platforms_overall_yearly_impl()
    
    # Use ThreadPoolExecutor for timeout functionality (works on Windows)
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            # Submit the query execution to thread pool
            future = executor.submit(execute_query)
            
            # Wait for result with 20-second timeout for faster fallback
            try:
                result = future.result(timeout=20)
                execution_time = round(time.time() - start_time, 2)
                
                # Add debug information to successful response
                if isinstance(result, dict) and 'year' in result:
                    result['_debug'] = {
                        "execution_time_seconds": execution_time,
                        "query_count": 36,  # 3 queries per month
                        "cache_status": "fresh" if execution_time > 1 else "cached"
                    }
                return jsonify(result)
                
            except concurrent.futures.TimeoutError:
                # Cancel the future to clean up resources
                future.cancel()
                print(f"Query timeout after 20 seconds - providing sample data")
                
                # Instead of failing, return sample data
                try:
                    sample_data = _get_sample_chart_data(datetime.today().year)
                    sample_data['_debug'] = {
                        "execution_time_seconds": 20.0,
                        "query_count": 0,
                        "cache_status": "fallback_sample_data",
                        "timeout_occurred": True
                    }
                    return jsonify(sample_data)
                except Exception as e:
                    print(f"Sample data generation failed: {e}")
                    return jsonify({
                        "error": "Database queries are taking too long. Showing sample data to ensure charts display.",
                        "timeout": True,
                        "suggestion": "Please try refreshing the page later when server load is lower.",
                        "retry_recommended": True
                    }), 408  # Request Timeout
                
    except Exception as e:
        import traceback
        traceback.print_exc()
        
        # Provide more helpful error messages
        if "Connection" in str(e) or "connect" in str(e).lower():
            error_msg = "Unable to connect to the database. Please try again later."
        elif "timeout" in str(e).lower():
            error_msg = "Database query timed out. Please try refreshing the page."
        else:
            error_msg = f"An error occurred while loading chart data: {str(e)}"
            
        return jsonify({
            "error": error_msg,
            "technical_error": str(e) if app.debug else None
        }), 500


def _platforms_overall_yearly_impl():
    """Implementation of the platforms_overall_yearly logic"""
    conn_asl = None
    cursor_asl = None
    conn_lib = None
    cursor_lib = None
    
    try:
        import time
        query_start = time.time()
        
        today = datetime.today().date()
        current_year = today.year
        months = [datetime(current_year, m, 1) for m in range(1, 13)]
        
        print(f"Starting database connections for year {current_year}...")
        
        # Test database connections with timeout
        try:
            conn_asl = get_ruzivo_conn()
            cursor_asl = conn_asl.cursor(pymysql.cursors.DictCursor)
            print("✓ ASL database connected")
        except Exception as e:
            print(f"✗ ASL database connection failed: {e}")
            return _get_sample_chart_data(current_year)
        
        try:
            conn_lib = get_direct_library_conn()
            cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)
            print("✓ Library database connected")
        except Exception as e:
            print(f"✗ Library database connection failed: {e}")
            return _get_sample_chart_data(current_year)

        def month_range(dt):
            start = dt.replace(day=1)
            if dt.month == 12:
                end = datetime(dt.year + 1, 1, 1) - timedelta(days=1)
            else:
                end = datetime(dt.year, dt.month + 1, 1) - timedelta(days=1)
            return start, end

        monthly_data = []
        total_asl_year = 0
        total_lib_year = 0
        total_ask_year = 0

        # Optimized approach: Use single queries with CASE statements for better performance
        print(f"Starting optimized database queries for year {current_year}...")
        
        try:
            # --- OPTIMIZED ASL QUERY (single query for all months) ---
            print("Executing ASL query...")
            cursor_asl.execute("""
                SELECT 
                    MONTH(login_date) as month_num,
                    COUNT(DISTINCT student_id) AS active_learners
                FROM tblstudents_login
                WHERE YEAR(login_date) = %s
                GROUP BY MONTH(login_date)
                ORDER BY MONTH(login_date)
            """, (current_year,))
            asl_results = cursor_asl.fetchall()
            print(f"ASL query completed, got {len(asl_results)} months")
            
            # Convert to dict for easy lookup
            asl_by_month = {row['month_num']: row['active_learners'] for row in asl_results}
            
            # --- OPTIMIZED LIBRARY QUERY (single query for all months) ---
            print("Executing Library query...")
            cursor_lib.execute("""
                SELECT 
                    MONTH(created_at) as month_num,
                    COUNT(DISTINCT user_id) AS active_users
                FROM logins
                WHERE YEAR(created_at) = %s
                GROUP BY MONTH(created_at)
                ORDER BY MONTH(created_at)
            """, (current_year,))
            lib_results = cursor_lib.fetchall()
            print(f"Library query completed, got {len(lib_results)} months")
            
            # Convert to dict for easy lookup
            lib_by_month = {row['month_num']: row['active_users'] for row in lib_results}
            
            # --- OPTIMIZED ASK AKELLO QUERY (single query for all months) ---
            print("Executing Ask Akello query...")
            cursor_asl.execute("""
                SELECT 
                    MONTH(created_at) as month_num,
                    COUNT(DISTINCT student_id) AS unique_students_count
                FROM tblask_akello_chat_logs
                WHERE YEAR(created_at) = %s
                GROUP BY MONTH(created_at)
                ORDER BY MONTH(created_at)
            """, (current_year,))
            ask_results = cursor_asl.fetchall()
            print(f"Ask Akello query completed, got {len(ask_results)} months")
            
            # Convert to dict for easy lookup
            ask_by_month = {row['month_num']: row['unique_students_count'] for row in ask_results}
            
        except Exception as query_error:
            print(f"Database query error: {query_error}")
            # If optimized queries fail, fall back to sample data
            return _get_sample_chart_data(current_year)
        
        # Build monthly data from results
        for m in months:
            month_num = m.month
            
            asl_count = asl_by_month.get(month_num, 0)
            lib_count = lib_by_month.get(month_num, 0)
            ask_count = ask_by_month.get(month_num, 0)
            
            total_asl_year += asl_count
            total_lib_year += lib_count
            total_ask_year += ask_count
            
            total = asl_count + lib_count + ask_count

            monthly_data.append({
                "month": m.strftime("%B"),
                "asl": asl_count,
                "library": lib_count,
                "ask_akello": ask_count,
                "overall": total
            })
        
        print(f"Monthly data compilation completed for {len(monthly_data)} months")

        # --- Compute Overall Year Totals ---
        overall_total_year = total_asl_year + total_lib_year + total_ask_year

        # Return data as dictionary (not jsonify, that's handled in the main function)
        return {
            "year": current_year,
            "monthly_usage": monthly_data,
            "yearly_totals": {
                "asl": total_asl_year,
                "library": total_lib_year,
                "ask_akello": total_ask_year,
                "overall": overall_total_year
            }
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        # Re-raise the exception to be handled by the main function
        raise e
        
    finally:
        # Clean up database connections
        if cursor_asl:
            cursor_asl.close()
        if conn_asl:
            conn_asl.close()
        if cursor_lib:
            cursor_lib.close()
        if conn_lib:
            conn_lib.close()


def _get_sample_chart_data(current_year):
    """Fallback function that returns sample data when database queries are too slow"""
    import random
    
    print("Using sample data fallback due to database performance issues")
    
    # Generate realistic sample data
    months_names = ['January', 'February', 'March', 'April', 'May', 'June',
                   'July', 'August', 'September', 'October', 'November', 'December']
    
    monthly_data = []
    total_asl_year = 0
    total_lib_year = 0
    total_ask_year = 0
    
    for i, month_name in enumerate(months_names):
        # Generate realistic numbers with some variation
        base_asl = random.randint(800, 1200)
        base_lib = random.randint(300, 500)
        base_ask = random.randint(150, 250)
        
        # Add seasonal variation (higher in school months)
        if i in [0, 1, 2, 7, 8, 9]:  # Jan-Mar, Aug-Oct (school months)
            multiplier = random.uniform(1.1, 1.3)
        else:
            multiplier = random.uniform(0.7, 0.9)
        
        asl_count = int(base_asl * multiplier)
        lib_count = int(base_lib * multiplier)
        ask_count = int(base_ask * multiplier)
        
        total_asl_year += asl_count
        total_lib_year += lib_count
        total_ask_year += ask_count
        
        total = asl_count + lib_count + ask_count
        
        monthly_data.append({
            "month": month_name,
            "asl": asl_count,
            "library": lib_count,
            "ask_akello": ask_count,
            "overall": total
        })
    
    return {
        "year": current_year,
        "monthly_usage": monthly_data,
        "yearly_totals": {
            "asl": total_asl_year,
            "library": total_lib_year,
            "ask_akello": total_ask_year,
            "overall": total_asl_year + total_lib_year + total_ask_year
        },
        "_sample_data": True,
        "_note": "Sample data used due to database performance issues"
    }





@app.route('/hlf-stats', methods=['GET'])
@login_required
def hlf_stats():
    """Render the HLF Stats page with interactive charts"""
    return render_template('HLFStats.html', title='HLF Stats')


@app.route('/api/hlf-weekly-stats/', methods=['GET'])
@login_required
@cache.cached(timeout=60 * 60 * 6, key_prefix='hlf_weekly_stats')  # Cache for 6 hours
def hlf_weekly_stats():
    import time
    import concurrent.futures

    start_time = time.time()

    # Get optional date range from query params
    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")

    # Convert or default to full current year
    try:
        if start_date_str and end_date_str:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        else:
            today = datetime.today().date()
            start_date = datetime(today.year, 1, 1).date()
            end_date = datetime(today.year, 12, 31).date()
    except ValueError:
        return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400

    def execute_query():
        return _hlf_weekly_stats_impl(start_date, end_date)

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(execute_query)
            try:
                result = future.result(timeout=20)
                execution_time = round(time.time() - start_time, 2)

                if isinstance(result, dict) and 'weeks' in result:
                    result['_debug'] = {
                        "execution_time_seconds": execution_time,
                        "query_count": 3,
                        "cache_status": "fresh" if execution_time > 1 else "cached"
                    }
                return jsonify(result)

            except concurrent.futures.TimeoutError:
                future.cancel()
                print("Query timeout after 20 seconds - providing sample weekly data")

                try:
                    sample_data = _get_sample_weekly_data(start_date.year)
                    sample_data['_debug'] = {
                        "execution_time_seconds": 20.0,
                        "query_count": 0,
                        "cache_status": "fallback_sample_data",
                        "timeout_occurred": True
                    }
                    return jsonify(sample_data)
                except Exception as e:
                    print(f"Sample data generation failed: {e}")
                    return jsonify({
                        "error": "Database queries are taking too long. Showing sample data to ensure charts display.",
                        "timeout": True,
                        "suggestion": "Please try refreshing the page later when server load is lower.",
                        "retry_recommended": True
                    }), 408

    except Exception as e:
        import traceback
        traceback.print_exc()

        if "Connection" in str(e) or "connect" in str(e).lower():
            error_msg = "Unable to connect to the database. Please try again later."
        elif "timeout" in str(e).lower():
            error_msg = "Database query timed out. Please try refreshing the page."
        else:
            error_msg = f"An error occurred while loading weekly chart data: {str(e)}"

        return jsonify({
            "error": error_msg,
            "technical_error": str(e) if app.debug else None
        }), 500


def _hlf_weekly_stats_impl(start_date, end_date):
    """Weekly usage implementation supporting date ranges and compatible with ONLY_FULL_GROUP_BY"""
    conn_asl = None
    cursor_asl = None
    conn_lib = None
    cursor_lib = None

    try:
        print(f"Fetching weekly platform usage between {start_date} and {end_date}...")

        # Establish database connections
        conn_asl = get_ruzivo_conn()
        cursor_asl = conn_asl.cursor(pymysql.cursors.DictCursor)

        conn_lib = get_direct_library_conn()
        cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)

        # --- ASL Weekly Query ---
        cursor_asl.execute("""
            SELECT 
                YEARWEEK(login_date, 1) AS year_week,
                COUNT(DISTINCT student_id) AS active_learners
            FROM tblstudents_login
            WHERE login_date BETWEEN %s AND %s
            GROUP BY YEARWEEK(login_date, 1)
            ORDER BY year_week
        """, (start_date, end_date))
        asl_by_week = {r['year_week']: r['active_learners'] for r in cursor_asl.fetchall()}

        # --- Library Weekly Query ---
        cursor_lib.execute("""
            SELECT 
                YEARWEEK(created_at, 1) AS year_week,
                COUNT(DISTINCT user_id) AS active_users
            FROM logins
            WHERE created_at BETWEEN %s AND %s
            GROUP BY YEARWEEK(created_at, 1)
            ORDER BY year_week
        """, (start_date, end_date))
        lib_by_week = {r['year_week']: r['active_users'] for r in cursor_lib.fetchall()}

        # --- Ask Akello Weekly Query ---
        cursor_asl.execute("""
            SELECT 
                YEARWEEK(created_at, 1) AS year_week,
                COUNT(DISTINCT student_id) AS unique_students_count
            FROM tblask_akello_chat_logs
            WHERE created_at BETWEEN %s AND %s
            GROUP BY YEARWEEK(created_at, 1)
            ORDER BY year_week
        """, (start_date, end_date))
        ask_by_week = {r['year_week']: r['unique_students_count'] for r in cursor_asl.fetchall()}

        # --- Prepare Weekly Data ---
        num_weeks = ((end_date - start_date).days // 7) + 1
        weekly_data = []
        total_asl = total_lib = total_ask = 0

        for i in range(num_weeks):
            week_start = start_date + timedelta(weeks=i)
            year_week = int(week_start.strftime("%G%V"))  # ISO year+week, e.g. 202540
            week_num = int(week_start.strftime("%V"))      # ISO week number

            asl_count = asl_by_week.get(year_week, 0)
            lib_count = lib_by_week.get(year_week, 0)
            ask_count = ask_by_week.get(year_week, 0)
            total = asl_count + lib_count + ask_count

            total_asl += asl_count
            total_lib += lib_count
            total_ask += ask_count

            weekly_data.append({
                "week_start": week_start.strftime("%Y-%m-%d"),
                "week_number": week_num,
                "asl": asl_count,
                "library": lib_count,
                "ask_akello": ask_count,
                "overall": total
            })

        # --- Compute Averages and Totals ---
        active_weeks = len([w for w in weekly_data if w["overall"] > 0]) or 1

        avg_weekly = {
            "asl": round(total_asl / active_weeks, 2),
            "library": round(total_lib / active_weeks, 2),
            "ask_akello": round(total_ask / active_weeks, 2),
            "overall": round((total_asl + total_lib + total_ask) / active_weeks, 2)
        }

        return {
            "start_date": start_date.strftime("%Y-%m-%d"),
            "end_date": end_date.strftime("%Y-%m-%d"),
            "weeks": weekly_data,
            "average_weekly_usage": avg_weekly,
            "totals": {
                "asl": total_asl,
                "library": total_lib,
                "ask_akello": total_ask,
                "overall": total_asl + total_lib + total_ask
            }
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise e

    finally:
        if cursor_asl: cursor_asl.close()
        if conn_asl: conn_asl.close()
        if cursor_lib: cursor_lib.close()
        if conn_lib: conn_lib.close()



def _get_sample_weekly_data(current_year):
    """Fallback weekly data"""
    import random
    print("Using sample weekly data fallback")

    weekly_data = []
    total_asl = total_lib = total_ask = 0

    for week in range(1, 53):
        asl = random.randint(600, 1200)
        lib = random.randint(200, 500)
        ask = random.randint(100, 300)
        total = asl + lib + ask

        total_asl += asl
        total_lib += lib
        total_ask += ask

        weekly_data.append({
            "week": week,
            "asl": asl,
            "library": lib,
            "ask_akello": ask,
            "overall": total
        })

    avg_weekly = {
        "asl": round(total_asl / 52, 2),
        "library": round(total_lib / 52, 2),
        "ask_akello": round(total_ask / 52, 2),
        "overall": round((total_asl + total_lib + total_ask) / 52, 2)
    }

    return {
        "year": current_year,
        "weeks": weekly_data,
        "average_weekly_usage": avg_weekly,
        "totals": {
            "asl": total_asl,
            "library": total_lib,
            "ask_akello": total_ask,
            "overall": total_asl + total_lib + total_ask
        },
        "_sample_data": True
    }





def normalize_province_name(name):
    if not name:
        return "Unknown"
    return str(name).strip().title()


@app.route('/api/askakello/custom_date', methods=['GET'])
def askakello_custom_date():
    try:
        # --- Parse dates ---
        end_date = request.args.get("end_date")
        start_date = request.args.get("start_date")

        today = datetime.today().date()
        default_start = today - timedelta(days=30)

        if not end_date:
            end_date = today
        else:
            end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

        if not start_date:
            start_date = default_start
        else:
            start_date = datetime.strptime(start_date, "%Y-%m-%d").date()

        # ---------------------------
        # Helper to normalize gender
        # ---------------------------
        def normalize_gender(g):
            if not g:
                return "Unknown"
            g = g.strip().lower()
            if g in ["f", "female", "girl", "woman"]:
                return "Female"
            elif g in ["m", "male", "boy", "man"]:
                return "Male"
            return "Unknown"

        # ---------------------------
        # ASK AKELLO CHAT LOGS
        # ---------------------------
        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        # Unique students
        ask_query = """
            SELECT COUNT(DISTINCT student_id) AS unique_students_count
            FROM tblask_akello_chat_logs
            WHERE DATE(created_at) BETWEEN %s AND %s
        """
        cursor.execute(ask_query, (start_date, end_date))
        ask_result = cursor.fetchone()
        total_ask_users = ask_result["unique_students_count"] if ask_result else 0

        # Province distribution
        ask_province_query = """
            SELECT sch.school_province, COUNT(DISTINCT acl.student_id) AS count
            FROM tblask_akello_chat_logs acl
            INNER JOIN tblstudents st ON acl.student_id = st.student_id
            INNER JOIN tblschools sch ON st.school_id = sch.school_id
            WHERE DATE(acl.created_at) BETWEEN %s AND %s
            GROUP BY sch.school_province
        """
        cursor.execute(ask_province_query, (start_date, end_date))
        ask_province = cursor.fetchall()

        # Normalize province names
        for row in ask_province:
            row["school_province"] = normalize_province_name(row["school_province"])

        # Gender distribution
        ask_gender_query = """
            SELECT st.gender, COUNT(DISTINCT acl.student_id) AS count
            FROM tblask_akello_chat_logs acl
            INNER JOIN tblstudents st ON acl.student_id = st.student_id
            WHERE DATE(acl.created_at) BETWEEN %s AND %s
            GROUP BY st.gender
        """
        cursor.execute(ask_gender_query, (start_date, end_date))
        ask_gender_raw = cursor.fetchall()

        ask_gender = {"Female": 0, "Male": 0, "Unknown": 0}
        for row in ask_gender_raw:
            g = normalize_gender(row["gender"])
            ask_gender[g] += row["count"]

        # ---------------------------
        # FINAL RESPONSE
        # ---------------------------
        return jsonify({
            "date_range": {"start": str(start_date), "end": str(end_date)},
            "totals": {
                "ask_akello": total_ask_users
            },
            "province_distribution": ask_province,
            "gender_distribution": ask_gender
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if 'cursor' in locals(): cursor.close()
        if 'conn' in locals(): conn.close()








# def normalize_province_name(name):
#     if not name:
#         return "Unknown"
#     return str(name).strip().title()


# from flask import request, jsonify
# from datetime import datetime, timedelta
# import pymysql

# @app.route('/api/platforms/quick_custom_date', methods=['GET'])
# def platforms_quick_custom_date():
#     try:
#         # --- Parse dates ---
#         end_date = request.args.get("end_date")
#         start_date = request.args.get("start_date")

#         today = datetime.today().date()
#         default_start = today - timedelta(days=30)

#         if not end_date:
#             end_date = today
#         else:
#             end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

#         if not start_date:
#             start_date = default_start
#         else:
#             start_date = datetime.strptime(start_date, "%Y-%m-%d").date()

#         # ---------------------------
#         # 1) ASL PLATFORM (Ruzivo DB)
#         # ---------------------------
#         conn_asl = get_ruzivo_conn()
#         cursor_asl = conn_asl.cursor(pymysql.cursors.DictCursor)

#         # Province + gender distribution
#         asl_province_gender_query = """
#             SELECT
#                 sch.school_province AS province,
#                 st.gender AS gender,
#                 COUNT(DISTINCT sl.student_id) AS count
#             FROM tblstudents_login sl
#             INNER JOIN tblstudents st ON sl.student_id = st.student_id
#             INNER JOIN tblschools sch ON st.school_id = sch.school_id
#             WHERE DATE(sl.login_date) BETWEEN %s AND %s
#             GROUP BY sch.school_province, st.gender
#         """
#         cursor_asl.execute(asl_province_gender_query, (start_date, end_date))
#         rows = cursor_asl.fetchall()

#         asl_province_data = {}
#         total_asl_users = 0
#         for r in rows:
#             prov = normalize_province_name(r["province"]) or "Unknown"
#             gender = r["gender"].title() if r["gender"] else "Unknown"
#             if prov not in asl_province_data:
#                 asl_province_data[prov] = {"Female":0,"Male":0,"Unknown":0,"total":0}
#             asl_province_data[prov][gender] = r["count"]
#             asl_province_data[prov]["total"] += r["count"]
#             total_asl_users += r["count"]

#         # ---------------------------
#         # 2) LIBRARY PLATFORM
#         # ---------------------------
#         conn_lib = get_direct_library_conn()
#         cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)

#         lib_province_gender_query = """
#             SELECT
#                 inst.province AS province,
#                 u.sex AS gender,
#                 COUNT(DISTINCT la.user_id) AS count
#             FROM last_activities la
#             INNER JOIN users u ON la.user_id = u.id
#             INNER JOIN institutions inst ON u.id = inst.id
#             WHERE DATE(la.created_at) BETWEEN %s AND %s
#             GROUP BY inst.province, u.sex
#         """
#         cursor_lib.execute(lib_province_gender_query, (start_date, end_date))
#         rows = cursor_lib.fetchall()

#         lib_province_data = {}
#         total_lib_users = 0
#         for r in rows:
#             prov = normalize_province_name(r["province"]) or "Unknown"
#             gender = r["gender"].title() if r["gender"] else "Unknown"
#             if prov not in lib_province_data:
#                 lib_province_data[prov] = {"Female":0,"Male":0,"Unknown":0,"total":0}
#             lib_province_data[prov][gender] = r["count"]
#             lib_province_data[prov]["total"] += r["count"]
#             total_lib_users += r["count"]

#         # ---------------------------
#         # 3) ASK AKELLO CHAT LOGS
#         # ---------------------------
#         ask_province_gender_query = """
#             SELECT
#                 sch.school_province AS province,
#                 st.gender AS gender,
#                 COUNT(DISTINCT acl.student_id) AS count
#             FROM tblask_akello_chat_logs acl
#             INNER JOIN tblstudents st ON acl.student_id = st.student_id
#             INNER JOIN tblschools sch ON st.school_id = sch.school_id
#             WHERE DATE(acl.created_at) BETWEEN %s AND %s
#             GROUP BY sch.school_province, st.gender
#         """
#         cursor_asl.execute(ask_province_gender_query, (start_date, end_date))
#         rows = cursor_asl.fetchall()

#         ask_province_data = {}
#         total_ask_users = 0
#         for r in rows:
#             prov = normalize_province_name(r["province"]) or "Unknown"
#             gender = r["gender"].title() if r["gender"] else "Unknown"
#             if prov not in ask_province_data:
#                 ask_province_data[prov] = {"Female":0,"Male":0,"Unknown":0,"total":0}
#             ask_province_data[prov][gender] = r["count"]
#             ask_province_data[prov]["total"] += r["count"]
#             total_ask_users += r["count"]

#         overall_total = total_asl_users + total_lib_users + total_ask_users

#         return jsonify({
#             "date_range": {"start": str(start_date), "end": str(end_date)},
#             "totals": {
#                 "asl": total_asl_users,
#                 "library": total_lib_users,
#                 "ask_akello": total_ask_users,
#                 "overall": overall_total
#             },
#             "province_distribution": {
#                 "asl": asl_province_data,
#                 "library": lib_province_data,
#                 "ask_akello": ask_province_data
#             }
#         })

#     except Exception as e:
#         import traceback
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

#     finally:
#         if 'cursor_asl' in locals(): cursor_asl.close()
#         if 'conn_asl' in locals(): conn_asl.close()
#         if 'cursor_lib' in locals(): cursor_lib.close()
#         if 'conn_lib' in locals(): conn_lib.close()













@app.route('/api/asl-metrics/<province_name>', methods=['GET'])
@login_required
def get_asl_metrics(province_name):
    conn = get_ruzivo_conn()
    cursor = conn.cursor()

    today = datetime.today().date()
    start_date = today.replace(day=1)
    active_30_date = today - timedelta(days=30)

    queries = {
        'asl_unique_users': """
            SELECT COUNT(DISTINCT sl.student_id) AS count
            FROM tblstudents_login sl
            JOIN vwstudent vs ON vs.student_id = sl.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE sl.login_date BETWEEN %s AND %s
            AND s.school_province = %s
        """,
        'asl_registrations': """
            SELECT COUNT(DISTINCT st.student_id) AS count
            FROM tblstudents st
            JOIN vwstudent vs ON vs.student_id = st.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE st.date_added BETWEEN %s AND %s
            AND s.school_province = %s
        """,
        'asl_total_primary_content': """
            SELECT COUNT(DISTINCT ca.student_id) AS count
            FROM tblcontent_access ca
            JOIN tblstudents st ON st.student_id = ca.student_id
            JOIN vwstudent vs ON vs.student_id = ca.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE ca.start_time BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_total_sec_content': """
            SELECT COUNT(DISTINCT ca.student_id) AS count
            FROM tblcontent_access_hs ca
            JOIN tblstudents st ON st.student_id = ca.student_id
            JOIN vwstudent vs ON vs.student_id = ca.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE ca.start_time BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_total_primary_exercise': """
            SELECT COUNT(DISTINCT tr.student_id) AS count
            FROM tblresults tr
            JOIN tblstudents st ON st.student_id = tr.student_id
            JOIN vwstudent vs ON vs.student_id = tr.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE tr.date_added BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_total_sec_exercise': """
            SELECT COUNT(DISTINCT tr.student_id) AS count
            FROM tblresults_hs tr
            JOIN tblstudents st ON st.student_id = tr.student_id
            JOIN vwstudent vs ON vs.student_id = tr.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE tr.date_added BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_total_zimsec_access': """
            SELECT COUNT(DISTINCT zi.student_id) AS count
            FROM tblcontent_access_zimsec zi
            JOIN tblstudents st ON st.student_id = zi.student_id
            JOIN vwstudent vs ON vs.student_id = zi.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE zi.start_time BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_teacher_set_activities': """
            SELECT COUNT(DISTINCT ta.student_id) AS count
            FROM tblclass_activity_results ta
            JOIN tblstudents st ON st.student_id = ta.student_id
            JOIN vwstudent vs ON vs.student_id = ta.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE ta.date_added BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
            AND s.school_province = %s
        """,
        'asl_teacher_access': """
            SELECT COUNT(t.teacher_id) AS count
            FROM tbl_teacher t
            JOIN tblschools s ON s.school_id = t.school_id
            WHERE t.active_at BETWEEN %s AND %s
            AND s.school_province = %s
        """,
        'asl_revenue': """
            SELECT currency, SUM(amount) AS total
            FROM tblecocash_payment_order p
            JOIN vwstudent vs ON vs.student_id = p.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE p.date_created BETWEEN %s AND %s
            AND p.transactionOperationStatus = 'COMPLETED'
            AND s.school_province = %s
            GROUP BY currency
        """,
        'asl_unique_subscribers': """
            SELECT COUNT(DISTINCT p.student_id) AS count
            FROM tblecocash_payment_order p
            JOIN vwstudent vs ON vs.student_id = p.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE p.date_created BETWEEN %s AND %s
            AND p.transactionOperationStatus = 'COMPLETED'
            AND s.school_province = %s
        """,
        'asl_active30': """
            SELECT COUNT(DISTINCT sl.student_id) AS count
            FROM tblstudents_login sl
            JOIN vwstudent vs ON vs.student_id = sl.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE DATE(sl.login_date) BETWEEN %s AND %s
            AND s.school_province = %s
        """
    }

    result = {}

    for label, query in queries.items():
        # Pick date range depending on metric
        date_range = (active_30_date, today) if 'active30' in label else (start_date, today)
        params = date_range + (province_name,)

        cursor.execute(query, params)
        rows = cursor.fetchall()

        if label == 'asl_revenue':
            result[label] = {row['currency']: float(row['total']) for row in rows}
        else:
            result[label] = rows[0]['count'] if rows else 0

    # cursor.close()
    # conn.close()

    return jsonify(result)






@app.route('/api/province-details/<provincename>', methods=['GET'])
@login_required
def get_province_details(provincename):
    today = datetime.today().date()
    start_date = today.replace(day=1)
    flag = 'd'

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        query = """
            SELECT 
                sch.school_id,
                sch.school_name,
                COUNT(s.student_id) AS learner_count
            FROM 
                tblschools sch
            LEFT JOIN 
                vwstudent s ON sch.school_id = s.school_id
                AND s.last_login BETWEEN %s AND %s
            WHERE 
                sch.flag != %s AND sch.school_province = %s
            GROUP BY 
                sch.school_id, sch.school_name
        """

        cursor.execute(query, (start_date, today, flag, provincename))
        schools = cursor.fetchall()

        total_schools = len(schools)
        total_learners = sum(s['learner_count'] for s in schools)

        return jsonify({
            "total_schools": total_schools,
            "total_learners": total_learners,
            "schools": schools
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Internal server error"}), 500

    # finally:
    #     if cursor:
    #         cursor.close()
    #     if conn:
    #         conn.close()





@app.route('/api/province-school-active-logins', methods=['GET'])
def province_school_active_logins():
    try:
        province = request.args.get("province")
        start_date = request.args.get("start_date")
        end_date = request.args.get("end_date")

        if not province or not start_date or not end_date:
            return jsonify({"error": "Missing required parameters (province, start_date, end_date)"}), 400

        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                sch.school_id,
                sch.school_name,
                sch.school_province,
                COUNT(DISTINCT sl.student_id) AS active_learners
            FROM tblstudents_login sl
            INNER JOIN tblstudents st ON sl.student_id = st.student_id
            INNER JOIN tblschools sch ON st.school_id = sch.school_id
            WHERE DATE(sl.login_date) BETWEEN %s AND %s
              AND sch.school_province = %s
            GROUP BY sch.school_id, sch.school_name, sch.school_province
            ORDER BY active_learners DESC
        """

        cursor.execute(query, (start_date, end_date, province))
        results = cursor.fetchall()

        # Province total learners
        province_total = sum([row["active_learners"] for row in results])

        return jsonify({
            "province": province,
            "date_range": {
                "start": start_date,
                "end": end_date
            },
            "province_total": province_total,
            "schools": results
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if 'conn' in locals() and conn:
            conn.close()




@app.route('/api/overall_school_active_logins', methods=['GET'])
def overall_school_active_logins():
    try:
        # Validate and normalize input dates
        start_date_str = request.args.get("start_date")
        end_date_str = request.args.get("end_date")

        if not start_date_str or not end_date_str:
            return jsonify({"error": "Missing required parameters (start_date, end_date)"}), 400

        try:
            # Parse as dates; ensure start <= end
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
        except ValueError:
            return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400

        if start_date > end_date:
            return jsonify({"error": "start_date cannot be after end_date"}), 400

        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        # Use half-open datetime range to preserve index usage on sl.login_date
        start_dt = datetime.combine(start_date, datetime.min.time())
        end_dt_exclusive = datetime.combine(end_date + timedelta(days=1), datetime.min.time())

        query = """
            SELECT 
                sch.school_id,
                sch.school_name,
                sch.school_province,
                COUNT(DISTINCT sl.student_id) AS active_learners
            FROM tblstudents_login sl
            INNER JOIN tblstudents st ON st.student_id = sl.student_id
            INNER JOIN tblschools sch ON sch.school_id = st.school_id
            WHERE sl.login_date >= %s AND sl.login_date < %s
            GROUP BY sch.school_id, sch.school_name, sch.school_province
            ORDER BY active_learners DESC
        """

        try:
            cursor.execute(query, (start_dt, end_dt_exclusive))
            rows = cursor.fetchall() or []
        except pymysql.err.OperationalError as oe:
            # Retry once on lost connection during query (e.g., timeout 2013)
            if getattr(oe, 'args', [None])[0] == 2013:
                try:
                    cursor.close()
                except Exception:
                    pass
                try:
                    conn.close()
                except Exception:
                    pass
                conn = get_ruzivo_conn()
                cursor = conn.cursor(pymysql.cursors.DictCursor)
                cursor.execute(query, (start_dt, end_dt_exclusive))
                rows = cursor.fetchall() or []
            else:
                raise

        # Coerce values for safe JSON serialization
        schools = []
        overall_total = 0
        for row in rows:
            active = int(row.get("active_learners") or 0)
            overall_total += active
            schools.append({
                "school_id": int(row.get("school_id") or 0),
                "school_name": row.get("school_name") or "",
                "school_province": row.get("school_province") or "",
                "active_learners": active
            })

        return jsonify({
            "date_range": {
                "start": start_date.isoformat(),
                "end": end_date.isoformat()
            },
            "overall_total": overall_total,
            "schools": schools
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        try:
            if 'cursor' in locals() and cursor:
                cursor.close()
        finally:
            if 'conn' in locals() and conn:
                conn.close()









@app.route('/api/province-details1/<provincename>', methods=['GET'])
@login_required
def get_province_details1(provincename):
    today = datetime.today().date()
    start_date = today.replace(day=1)

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        query = """
            SELECT 
                DATE(last_login) AS login_date,
                COUNT(student_id) AS daily_logins
            FROM 
                tblstudents_info
            WHERE 
                last_login BETWEEN %s AND %s
            GROUP BY 
                DATE(last_login)
            ORDER BY 
                login_date ASC;
        """

        cursor.execute(query, (start_date, today))
        rows = cursor.fetchall()

        # Format results
        usage = []
        for row in rows:
            login_date = row[0].strftime('%Y-%m-%d') if isinstance(row[0], datetime) else str(row[0])
            usage.append({
                "date": login_date,
                "logins": int(row[1])
            })

        return jsonify({
            "start_date": str(start_date),
            "end_date": str(today),
            "usage": usage
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "Internal server error"}), 500

    # finally:
    #     if cursor:
    #         cursor.close()
    #     if conn:
    #         conn.close()



# latest province api
@app.route('/api/active-students-by-province', methods=['GET'])
@login_required
def get_active_students_by_province():
    try:
        valid_provinces = [
            'Harare', 'Bulawayo', 'Manicaland', 'Mashonaland Central',
            'Mashonaland East', 'Mashonaland West', 'Masvingo',
            'Matabeleland North', 'Matabeleland South', 'Midlands'
        ]

        today = datetime.today().date()
        start_date = today.replace(day=1)

        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        query = f"""
            SELECT s.school_province AS province, COUNT(*) AS student_count
            FROM vwstudent v
            JOIN tblschools s ON v.school_id = s.school_id
            WHERE v.last_login BETWEEN %s AND %s
            AND s.school_province IN ({','.join(['%s'] * len(valid_provinces))})
            GROUP BY s.school_province
        """

        params = [start_date, today] + valid_provinces
        cursor.execute(query, params)
        rows = cursor.fetchall()

        # Format and sort results
        results = [{"province": row['province'], "student_count": row['student_count']} for row in rows]
        results.sort(key=lambda x: x['student_count'], reverse=True)

        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'cursor' in locals():
    #         cursor.close()
    #     if 'conn' in locals():
    #         conn.close()




# latest schools by province
from datetime import datetime
from flask import request, jsonify

@app.route('/api/schools-by-province', methods=['GET'])
@login_required
def get_schools_by_province_with_counts():
    valid_provinces = [
        'Harare', 'Bulawayo', 'Manicaland', 'Mashonaland Central',
        'Mashonaland East', 'Mashonaland West', 'Masvingo',
        'Matabeleland North', 'Matabeleland South', 'Midlands'
    ]

    province = request.args.get('province')

    if not province or province not in valid_provinces:
        return jsonify({"error": "Invalid or missing province"}), 400

    try:
        today = datetime.today().date()
        start_of_month = today.replace(day=1)

        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        # Get schools in that province
        cursor.execute("""
            SELECT school_id, school_name
            FROM tblschools
            WHERE school_province = %s
        """, (province,))
        schools = cursor.fetchall()

        results = []

        for school in schools:
            school_id = school['school_id']
            school_name = school['school_name']

            cursor.execute("""
                SELECT COUNT(*) AS student_count
                FROM vwstudent
                WHERE school_id = %s
                AND last_login BETWEEN %s AND %s
            """, (school_id, start_of_month, today))

            count_row = cursor.fetchone()
            student_count = count_row['student_count'] if count_row else 0

            results.append({
                "school_id": school_id,
                "school_name": school_name,
                "student_count": student_count
            })

        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'cursor' in locals(): cursor.close()
    #     if 'conn' in locals(): conn.close()




@app.route('/api/chat-log-activity', methods=['GET'])
@login_required
def get_chat_log_activity():
    today = datetime.today()
    start_date = today - timedelta(days=30)  # last 30 days

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor(dictionary=True)  # ✅ use dict so jsonify works cleanly

        # Per-student logs
        query = """
            SELECT 
                s.student_id,
                CONCAT(s.name, ' ', s.surname) AS full_name,
                COUNT(*) AS total_logs
            FROM tblask_akello_chat_logs l
            JOIN vwstudent s ON l.student_id = s.student_id
            WHERE l.updated_at BETWEEN %s AND %s
            GROUP BY s.student_id, full_name
            ORDER BY total_logs DESC
        """

        cursor.execute(query, (start_date, today))
        student_logs = cursor.fetchall()

        # Total logs across all students
        total_query = """
            SELECT COUNT(*) AS total_logs
            FROM tblask_akello_chat_logs
            WHERE updated_at BETWEEN %s AND %s
        """
        cursor.execute(total_query, (start_date, today))
        total_logs = cursor.fetchone()["total_logs"]

        cursor.close()
        conn.close()

        return jsonify({
            "students": student_logs,
            "total_logs": total_logs
        })

    except Exception as e:
        print("Error querying chat log activity:", e)
        return jsonify({"error": "Internal server error"}), 500


    # finally:
    #     if cursor: cursor.close()
    #     if conn: conn.close()


@app.route('/api/chat-log-by-province', methods=['GET'])
@login_required
def get_chat_log_by_province():
    provinces = [
        'Harare', 'Bulawayo', 'Manicaland', 'Mashonaland Central', 'Mashonaland East',
        'Mashonaland West', 'Masvingo', 'Matabeleland North', 'Matabeleland South', 'Midlands'
    ]

    today = datetime.today().date()
    start_date = today.replace(day=1)

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        query = """
            SELECT 
                sch.school_province AS province,
                COUNT(DISTINCT s.student_id) AS total_students
            FROM tblask_akello_chat_logs log
            JOIN vwstudent s ON log.student_id = s.student_id
            JOIN tblschools sch ON s.school_id = sch.school_id
            WHERE log.updated_at BETWEEN %s AND %s
            AND sch.school_province IN %s
            GROUP BY sch.school_province
            ORDER BY total_students DESC
        """

        cursor.execute(query, (start_date, today, tuple(provinces)))
        results = cursor.fetchall()

        return jsonify(results)

    except Exception as e:
        print("Error:", e)
        return jsonify({"error": "Internal server error"}), 500

    # finally:
    #     if cursor: cursor.close()
    #     if conn: conn.close()




# -------------------------------
# 1️⃣ DAILY CHATS API
# -------------------------------
@app.route('/api/askakello-daily-chats', methods=['GET'])
@login_required
def askakello_daily_chats():
    today = date.today()
    start_date = today.replace(day=1)  # First day of this month

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                DATE(l.updated_at) AS chat_date,
                COUNT(DISTINCT l.student_id) AS learner_count
            FROM tblask_akello_chat_logs l
            WHERE l.updated_at BETWEEN %s AND %s
            GROUP BY DATE(l.updated_at)
            ORDER BY chat_date
        """

        cursor.execute(query, (start_date, today))
        results = cursor.fetchall()

        return jsonify({
            "year": today.year,
            "month": today.strftime("%B"),
            "daily_chats": results
        })

    except Exception as e:
        print("Error querying daily chats:", e)
        return jsonify({"error": "Internal server error"}), 500

    # finally:
    #     if conn:
    #         conn.close()

# -------------------------------
# 2️⃣ MONTHLY CHATS API
# -------------------------------
@app.route('/api/askakello-monthly-chats', methods=['GET'])
@login_required
def askakello_monthly_chats():
    today = date.today()
    current_year = today.year

    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                YEAR(l.updated_at) AS year_num,
                MONTH(l.updated_at) AS month_num,
                COUNT(DISTINCT l.student_id) AS learner_count
            FROM tblask_akello_chat_logs l
            WHERE YEAR(l.updated_at) = %s
            GROUP BY YEAR(l.updated_at), MONTH(l.updated_at)
            ORDER BY year_num, month_num
        """

        cursor.execute(query, (current_year,))
        results = cursor.fetchall()

        # Add month name
        for r in results:
            r["month"] = calendar.month_name[r["month_num"]]

        return jsonify({
            "year": current_year,
            "monthly_chats": results
        })

    except Exception as e:
        print("Error querying monthly chats:", e)
        return jsonify({"error": "Internal server error"}), 500

    # finally:
    #     if conn:
    #         conn.close()





@app.route('/api/library-institutions', methods=['GET'])
@login_required
def api_institutions():
    try:
        # Get optional date range from query parameters
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        today = datetime.today().date()
        if not start_date:
            start_date = today.replace(day=1)
        else:
            start_date = datetime.strptime(start_date, '%Y-%m-%d')

        if not end_date:
            end_date = today
        else:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')

        conn = get_direct_library_conn()

        with conn.cursor() as cursor:
            query = """
                SELECT 
                    i.province,
                    COUNT(DISTINCT i.id) AS number_of_institutions,
                    COUNT(DISTINCT l.user_id) AS total_users
                FROM institution_user iu
                JOIN institutions i ON iu.institution_id = i.id
                JOIN logins l ON iu.user_id = l.user_id
                WHERE l.created_at BETWEEN %s AND %s
                GROUP BY i.province
                ORDER BY total_users DESC
            """
            cursor.execute(query, (start_date, end_date))
            results = cursor.fetchall()

        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()



@app.route('/api/library-daily-logins', methods=['GET'])
@login_required
def api_library_daily_logins():
    try:
        today = datetime.today()
        start_date = today.replace(day=1).date()  # first day of current month
        end_date = today.date()  # today

        conn = get_direct_library_conn()

        with conn.cursor() as cursor:
            query = """
                SELECT 
                    DATE(l.created_at) AS login_date,
                    COUNT(*) AS total_logins
                FROM logins l
                WHERE l.created_at BETWEEN %s AND %s
                GROUP BY login_date
                ORDER BY login_date ASC
            """
            cursor.execute(query, (start_date, end_date))
            results = cursor.fetchall()

        # Ensure we have all days of the month even if count = 0
        from calendar import monthrange
        total_days = monthrange(today.year, today.month)[1]
        daily_data = []
        results_dict = {str(r['login_date']): r['total_logins'] for r in results}

        for day in range(1, total_days + 1):
            date_str = datetime(today.year, today.month, day).date()
            if date_str > end_date:
                break
            daily_data.append({
                "date": date_str.strftime("%Y-%m-%d"),
                "total_logins": results_dict.get(str(date_str), 0)
            })

        return jsonify(daily_data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()




@app.route('/api/library-vaya-daily-logins', methods=['GET'])
def library_vaya_daily_logins():
    conn = None
    try:
        today = date.today()
        current_year = today.year
        current_month = today.month
        days_in_month = calendar.monthrange(current_year, current_month)[1]

        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        daily_data = []

        # Loop through all days of current month
        for day in range(1, days_in_month + 1):
            day_start = datetime(current_year, current_month, day, 0, 0, 0)
            day_end = datetime(current_year, current_month, day, 23, 59, 59)

            query = """
                SELECT 
                    COUNT(DISTINCT o.user_id) AS user_count,
                    o.currency,
                    SUM(o.total_amount) AS total_amount
                FROM orders o
                JOIN users u ON u.id = o.user_id
                WHERE o.created_at BETWEEN %s AND %s
                AND o.status = 'Completed'
                GROUP BY o.currency
            """
            cursor.execute(query, (day_start, day_end))
            results = cursor.fetchall()

            # Build currency summary for the day
            currency_totals = {}
            total_users = 0
            for r in results:
                currency = r["currency"] if r["currency"] else "UNKNOWN"
                currency_totals[currency] = float(r["total_amount"]) if r["total_amount"] else 0
                total_users = r["user_count"]  # same for all rows (per day)

            daily_data.append({
                "date": day_start.strftime("%Y-%m-%d"),
                "user_count": int(total_users) if total_users else 0,
                "currency_totals": currency_totals
            })

        # cursor.close()
        # conn.close()

        return jsonify({
            "year": current_year,
            "month": date(current_year, current_month, 1).strftime("%B"),
            "daily_counts": daily_data
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()





@app.route('/api/library-monthly-logins', methods=['GET'])
@login_required
def api_library_monthly_logins():
    try:
        today = datetime.today()
        start_date = datetime(today.year, 1, 1).date()   # Jan 1 of current year
        end_date = today.date()  # today

        conn = get_direct_library_conn()

        with conn.cursor() as cursor:
            query = """
                SELECT 
                    MONTH(l.created_at) AS month_num,
                    COUNT(*) AS total_logins
                FROM logins l
                WHERE l.created_at BETWEEN %s AND %s
                GROUP BY month_num
                ORDER BY month_num ASC
            """
            cursor.execute(query, (start_date, end_date))
            results = cursor.fetchall()

        # Fill missing months with 0
        results_dict = {r['month_num']: r['total_logins'] for r in results}
        monthly_data = []
        for m in range(1, 13):
            if datetime(today.year, m, 1).date() > end_date:
                break
            monthly_data.append({
                "month": month_name[m],
                "total_logins": results_dict.get(m, 0)
            })

        return jsonify(monthly_data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()





@app.route("/api/library/school_profile", methods=["GET", "POST"])
def library_school_profile():
    if request.method == "POST":
        data = request.get_json(silent=True) or {}
    else:
        data = request.args  

    asl_school_id = data.get("asl_school_id")
    start_date = data.get("start_date")
    end_date = data.get("end_date")

    if not asl_school_id:
        return jsonify({"error": "asl_school_id is required"}), 400

    from datetime import date, timedelta, datetime
    if not start_date or not end_date:
        end_date = date.today()
        start_date = end_date - timedelta(days=30)
    else:
        start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
        end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

    # --- Find champions with matching asl_school_id ---
    champions = ChampionSchool.query.all()
    matching_champions = []
    library_ids = []

    for champ in champions:
        for school in champ.get_schools():
            if str(school.get("asl_school_id")) == str(asl_school_id):
                matching_champions.append(champ)
                if school.get("library_school_id"):
                    library_ids.append(school.get("library_school_id"))

    if not matching_champions:
        return jsonify({
            "message": f"No champions found for asl_school_id {asl_school_id}",
            "total_active_users": 0
        }), 200

    if not library_ids:
        return jsonify({
            "message": f"No linked library_school_id for asl_school_id {asl_school_id}",
            "total_active_users": 0
        }), 200

    # --- Query Library DB ---
    total_active_users = 0
    school_profiles = []

    query = """
        SELECT 
            inst.id,
            inst.name AS institution_name,
            COUNT(DISTINCT la.user_id) AS active_users
        FROM logins la
        INNER JOIN institution_user iu ON la.user_id = iu.user_id
        INNER JOIN institutions inst ON iu.institution_id = inst.id
        WHERE inst.id = %s AND DATE(la.created_at) BETWEEN %s AND %s
        GROUP BY inst.id, inst.name
        ORDER BY active_users DESC
    """

    conn = get_direct_library_conn()
    cur = conn.cursor()  # make sure results are dicts

    for lib_id in library_ids:
        cur.execute(query, (lib_id, start_date, end_date))
        row = cur.fetchone()
        if row:
            total_active_users += row["active_users"]
            school_profiles.append(row)

    cur.close()
    conn.close()

    return jsonify({
        "asl_school_id": asl_school_id,
        "library_school_ids": library_ids,
        "total_active_users": total_active_users,
        "schools": school_profiles
    }), 200







# @app.route("/api/library/school_profile", methods=["GET","POST"])
# def library_school_profile():
#     if request.method == "POST":
#         # Safely parse JSON if provided
#         data = request.get_json(silent=True) or {}
#     else:
#         # For GET requests, use query parameters
#         data = request.args  

#     asl_school_id = data.get("asl_school_id")
#     start_date = data.get("start_date")
#     end_date = data.get("end_date")

#     if not asl_school_id:
#         return jsonify({"error": "asl_school_id is required"}), 400

#     # Default date range (last 30 days if not provided)
#     from datetime import date, timedelta
#     if not start_date or not end_date:
#         end_date = date.today()
#         start_date = end_date - timedelta(days=30)

#     # Find champions with matching asl_school_id
#     champions = ChampionSchool.query.filter_by(asl_school_id=asl_school_id).all()
#     if not champions:
#         return jsonify({"message": "No champions found for this asl_school_id", "total_active_users": 0}), 200

#     library_ids = [c.library_school_id for c in champions if c.library_school_id]

#     if not library_ids:
#         return jsonify({"message": "No linked library_school_id for this asl_school_id", "total_active_users": 0}), 200

#     # Query library DB
#     total_active_users = 0
#     school_profiles = []

#     query = """
#         SELECT 
#             inst.id,
#             inst.name AS institution_name,
#             COUNT(DISTINCT la.user_id) AS active_users
#         FROM last_activities la
#         INNER JOIN institution_user iu ON la.user_id = iu.user_id
#         INNER JOIN institutions inst ON iu.institution_id = inst.id
#         WHERE inst.id = %s AND DATE(la.created_at) BETWEEN %s AND %s
#         GROUP BY inst.id, inst.name
#         ORDER BY active_users DESC
#     """

#     conn = get_direct_library_conn()
#     cur = conn.cursor()

#     for lib_id in library_ids:
#         cur.execute(query, (lib_id, start_date, end_date))
#         row = cur.fetchone()
#         if row:
#             total_active_users += row["active_users"]
#             school_profiles.append(row)

#     cur.close()
#     conn.close()

#     return jsonify({
#         "asl_school_id": asl_school_id,
#         "library_school_ids": library_ids,
#         "total_active_users": total_active_users,
#         "schools": school_profiles
#     }), 200







@app.route('/api/library_custom_date_analytics', methods=['GET'])
def library_custom_date_analytics():
    try:
        # Get date range from request, default last 30 days
        end_date = request.args.get("end_date")
        start_date = request.args.get("start_date")

        today = datetime.today().date()
        default_start = today - timedelta(days=30)

        if not end_date:
            end_date = today
        else:
            end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

        if not start_date:
            start_date = default_start
        else:
            start_date = datetime.strptime(start_date, "%Y-%m-%d").date()

        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                inst.id,
                inst.name AS institution_name,
                inst.province AS institution_province,
                COUNT(DISTINCT la.user_id) AS active_users
            FROM logins la
            INNER JOIN institution_user iu ON la.user_id = iu.user_id
            INNER JOIN institutions inst ON iu.institution_id = inst.id
            WHERE DATE(la.created_at) BETWEEN %s AND %s
            GROUP BY inst.id, inst.name, inst.province
            ORDER BY active_users DESC
        """

        cursor.execute(query, (start_date, end_date))
        results = cursor.fetchall()

        total_active_users = sum([row["active_users"] for row in results])

        return jsonify({
            "date_range": {
                "start": str(start_date),
                "end": str(end_date)
            },
            "total_active_users": total_active_users,
            "institutions": results
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if 'cursor' in locals() and cursor:
            cursor.close()
        if 'conn' in locals() and conn:
            conn.close()






@app.route('/api/library-metrics/<province_name>', methods=['GET'])
@login_required
def api_library_metrics_by_province(province_name):
    try:
        # Optional date range
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        today = datetime.today().date()
        if not start_date:
            start_date = today.replace(day=1)
        else:
            start_date = datetime.strptime(start_date, '%Y-%m-%d')

        if not end_date:
            end_date = today
        else:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')

        conn = get_direct_library_conn()

        with conn.cursor() as cursor:
            query = """
                SELECT 
                    i.name AS institution_name,
                    COUNT(DISTINCT l.user_id) AS total_users
                FROM institution_user iu
                JOIN institutions i ON iu.institution_id = i.id
                JOIN logins l ON iu.user_id = l.user_id
                WHERE i.province = %s
                  AND l.created_at BETWEEN %s AND %s
                GROUP BY i.name
                ORDER BY total_users DESC
            """
            cursor.execute(query, (province_name, start_date, end_date))
            results = cursor.fetchall()

        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()




@app.route('/api/library-province-institutions/<province_name>', methods=['GET'])
@login_required
def api_library_province_institutions(province_name):
    try:
        conn = get_direct_library_conn()

        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT id, name, province
                FROM institutions
                WHERE province = %s
                ORDER BY name ASC
            """
            cursor.execute(query, (province_name,))
            results = cursor.fetchall()

        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()







from pymysql.cursors import DictCursor



@app.route('/api/student-access-info', methods=['GET'])
@login_required
def get_student_access_info():
    try:
        conn = get_ruzivo_conn()
        with conn.cursor() as cursor:
            query = """
                SELECT name, last_login, access_sdate, access_edate
                FROM vwstudent
                WHERE last_login IS NOT NULL
                ORDER BY last_login DESC
                LIMIT 100
            """
            cursor.execute(query)
            result = cursor.fetchall()

            # Format datetime fields
            for row in result:
                row['last_login'] = row['last_login'].strftime('%Y-%m-%d %H:%M:%S') if row['last_login'] else None
                row['access_sdate'] = row['access_sdate'].strftime('%Y-%m-%d %H:%M:%S') if row['access_sdate'] else None
                row['access_edate'] = row['access_edate'].strftime('%Y-%m-%d %H:%M:%S') if row['access_edate'] else None

        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()





@app.route('/api/school-usage', methods=['GET'])
@login_required
def get_school_usage():
    try:
        conn = get_ruzivo_conn()
        with conn.cursor() as cursor:
            query = """
                SELECT 
                    school_id,
                    grade_id,
                    usage_yr,
                    school_enrol,
                    school_usage
                FROM vwschoolusage
                ORDER BY school_usage DESC
                LIMIT 10;
            """
            cursor.execute(query)
            data = cursor.fetchall()
        return jsonify(data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()



import pandas as pd
from flask import send_file, jsonify
import io
import pymysql

@app.route('/api/users-with-contact', methods=['GET'])
def users_with_contact():
    conn = None
    try:
        conn = get_direct_library_conn()
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT id, first_name, last_name, mobile_number, email
                FROM users
                WHERE (mobile_number IS NOT NULL AND mobile_number != '')
                   OR (email IS NOT NULL AND email != '')
            """
            cursor.execute(query)
            results = cursor.fetchall()

        if not results:
            return jsonify({"message": "No users with contact details found"}), 404

        # Convert to DataFrame
        df = pd.DataFrame(results)
        df["name"] = (df["first_name"].fillna('') + " " + df["last_name"].fillna('')).str.strip()
        df = df[["id", "name", "mobile_number", "email"]]

        # Write Excel to memory (not disk)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Users")
        output.seek(0)

        filename = "users_with_contact.xlsx"

        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()




import pandas as pd
from flask import send_file, jsonify
import io
import pymysql

@app.route('/api/students-with-contact', methods=['GET'])
def students_with_contact():
    conn = None
    try:
        conn = get_ruzivo_conn()
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT 
                    s.student_id,
                    s.name,
                    s.surname,
                    s.email,
                    si.mobile
                FROM tblstudents s
                INNER JOIN tblstudents_info si 
                    ON s.student_id = si.student_id
                WHERE (si.mobile IS NOT NULL AND si.mobile != '')
                   OR (s.email IS NOT NULL AND s.email != '')
            """
            cursor.execute(query)
            results = cursor.fetchall()

        if not results:
            return jsonify({"message": "No students with contact details found"}), 404

        # Convert to DataFrame
        df = pd.DataFrame(results)
        df["name"] = (df["name"].fillna('') + " " + df["surname"].fillna('')).str.strip()
        df = df[["student_id", "name", "mobile", "email"]]

        # Write Excel in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Students")
        output.seek(0)

        filename = "students_with_contact.xlsx"

        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()






@app.route('/api/library-daily-currency', methods=['GET'])
def library_daily_currency():
    try:
        conn = get_direct_library_conn()
        with conn.cursor() as cursor:
            query = """
                SELECT DISTINCT 
                    o.user_id,
                    u.first_name,
                    u.last_name,
                    u.email,
                    u.mobile_number,
                    o.currency,
                    o.payment_type,
                    o.payment_method,
                    o.total_amount,
                    o.created_at
                FROM orders o
                JOIN users u ON u.id = o.user_id
                WHERE o.created_at BETWEEN '2025-08-25 00:00:00' AND '2025-08-25 23:59:59'
                  AND o.status = 'Completed';
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            col_names = [desc[0] for desc in cursor.description]
            data = [dict(zip(col_names, row)) for row in rows]

        return jsonify(data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()







# @app.route('/api/testingdb1', methods=['GET'])
# def testingdb1():
#     try:
#         # conn = get_ruzivo_conn()
#         conn = get_direct_library_conn()
#         with conn.cursor() as cursor:
#             query = """
#                SELECT 
#                     o.currency,
#                     COUNT(DISTINCT o.user_id) AS total_users,
#                     SUM(o.total_amount) AS total_amount
#                 FROM orders o
#                 JOIN users u ON u.id = o.user_id
#                 WHERE o.created_at BETWEEN '2025-08-25 00:00:00' AND '2025-08-25 23:59:59'
#                   AND o.status = 'Completed'
#                 GROUP BY o.currency;
# """
#             cursor.execute(query)
#             data = cursor.fetchall()
#         return jsonify(data)

#     except Exception as e:
#         import traceback
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

#     finally:
#         if 'conn' in locals():
#             conn.close()




@app.route('/api/testingdb1', methods=['GET'])
def testingdb1():
    try:
        # conn = get_ruzivo_conn()
        conn = get_direct_library_conn()
        with conn.cursor() as cursor:
            query = """
               SELECT 
                    i.province,
                    COUNT(DISTINCT i.id) AS number_of_institutions,
                    COUNT(DISTINCT l.user_id) AS total_users
                FROM institution_user iu
                JOIN institutions i ON iu.institution_id = i.id
                JOIN logins l ON iu.user_id = l.user_id
                WHERE l.created_at BETWEEN %s AND %s
                GROUP BY i.province
                ORDER BY total_users DESC;
"""
            cursor.execute(query)
            data = cursor.fetchall()
        return jsonify(data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()




# @app.route('/api/testingdb1', methods=['GET'])
# def testingdb1():
#     try:
#         conn = get_direct_library_conn()
#         with conn.cursor(pymysql.cursors.DictCursor) as cursor:  # ✅ return dicts
#             query = """
#                 SELECT 
#                     o.currency,
#                     COUNT(DISTINCT o.user_id) AS total_users,
#                     SUM(o.total_amount) AS total_amount
#                 FROM orders o
#                 JOIN users u ON u.id = o.user_id
#                 WHERE o.created_at BETWEEN '2025-08-25 00:00:00' AND '2025-08-25 23:59:59'
#                   AND o.status = 'Completed'
#                 GROUP BY o.currency;
#             """
#             cursor.execute(query)
#             data = cursor.fetchall()

#         return jsonify(data)

#     except Exception as e:
#         import traceback
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

#     finally:
#         if 'conn' in locals():
#             conn.close()




# @app.route('/api/testingdb1', methods=['GET'])
# def testingdb1():
#     conn = None
#     try:
#         conn = get_direct_library_conn()
#         with conn.cursor(pymysql.cursors.DictCursor) as cursor:
#             # Query for totals grouped by currency
#             totals_query = """
#                 SELECT 
#                     o.currency,
#                     COUNT(DISTINCT o.user_id) AS total_users,
#                     SUM(o.total_amount) AS total_amount
#                 FROM orders o
#                 JOIN users u ON u.id = o.user_id
#                 WHERE o.created_at BETWEEN '2025-08-25 00:00:00' AND '2025-08-25 23:59:59'
#                   AND o.status = 'Completed'
#                 GROUP BY o.currency;
#             """
#             cursor.execute(totals_query)
#             totals = cursor.fetchall()

#             # Query for detailed rows
#             details_query = """
#                 SELECT 
#                     o.id AS order_id,
#                     o.user_id,
#                     u.first_name,
#                     u.last_name,
#                     u.email,
#                     u.mobile_number,
#                     o.currency,
#                     o.payment_type,
#                     o.payment_method,
#                     o.total_amount,
#                     o.created_at
#                 FROM orders o
#                 JOIN users u ON u.id = o.user_id
#                 WHERE o.created_at BETWEEN '2025-08-25 00:00:00' AND '2025-08-25 23:59:59'
#                   AND o.status = 'Completed';
#             """
#             cursor.execute(details_query)
#             details = cursor.fetchall()

#         return jsonify({
#             "totals": totals,
#             "details": details
#         })

#     except Exception as e:
#         import traceback
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

#     finally:
#         if conn:
#             conn.close()




@app.route('/api/smartlearning-school', methods=['POST'])
@login_required
def smartlearning_school():
    conn = None  # Initialize conn to None
    try:
        data = request.get_json()
        school_name = data.get("school_name")
        if not school_name:
            return jsonify({"error": "Missing school_name"}), 400

        # Get date range from request or default to MTD
        today = datetime.today().date()
        default_start_date = today.replace(day=1)

        start_date_str = data.get("start_date")
        end_date_str = data.get("end_date")

        try:
            if start_date_str:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            else:
                start_date = default_start_date

            if end_date_str:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
            else:
                end_date = today
        except ValueError:
            return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400

        # Get the connection from the PyMySQL connection pool
        conn = get_ruzivo_conn()

        with conn.cursor() as cursor:
            # Main data query with grade filter
            query = """
                SELECT
                    s.school_id,
                    s.username,
                    s.grade,
                    s.last_login,
                    t.school_name
                FROM vwstudent s
                JOIN tblschools t ON s.school_id = t.school_id
                WHERE t.school_name = %s
                  AND s.grade BETWEEN 4 AND 13
                  AND s.last_login BETWEEN %s AND %s
                LIMIT 500;
            """
            cursor.execute(query, (school_name, start_date, end_date))
            rows = cursor.fetchall()
            columns = list(rows[0].keys()) if rows else []

            # Count query with grade filter
            count_query = """
                SELECT COUNT(*) AS total FROM vwstudent s
                JOIN tblschools t ON s.school_id = t.school_id
                WHERE t.school_name = %s
                  AND s.grade BETWEEN 4 AND 13
                  AND s.last_login BETWEEN %s AND %s;
            """
            cursor.execute(count_query, (school_name, start_date, end_date))
            total_count = cursor.fetchone()['total']

        return jsonify({
            "total_count": total_count,
            "columns": columns,
            "rows": rows,
            "date_range": {
                "start": str(start_date),
                "end": str(end_date)
            }
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        # Always return the connection to the pool
        if conn:
            conn.close()






from flask import jsonify
from datetime import datetime, date

from app.models import BrandingItem, BrandingRequest, BrandingAction

# -------- Branding APIs --------

def is_admin_user():
    try:
        # Consider role or privileges
        if getattr(current_user, 'userRole', None) in ('Admin','Super-admin','Manager'):
            return True
        # Fallback to privileges dict if present
        if hasattr(current_user, 'has_privilege'):
            return current_user.has_privilege('Super-admin') or current_user.has_privilege('Manager')
    except Exception:
        pass
    return False

def can_view_all_queries():
    """Check if user can view all queries (admin or help-desk agent/viewer privileges)."""
    try:
        if getattr(current_user, 'userRole', None) == 'Admin':
            return True
        if hasattr(current_user, 'has_privilege'):
            return (
                current_user.has_privilege('Admin Queries Access')
                or current_user.has_privilege('Help Desk Agent')
                or current_user.has_privilege('Help Desk Viewer')
            )
    except Exception:
        pass
    return False


def can_approve_champion_schools():
    """Check if user can approve champion school requests (Admin or Approve Champion Schools privilege)."""
    try:
        if getattr(current_user, 'userRole', None) == 'Admin':
            return True
        if hasattr(current_user, 'has_privilege'):
            return current_user.has_privilege('Approve Champion Schools')
    except Exception:
        pass
    return False


def can_view_revenue_reports():
    """Check if user can access Revenue Reports (Admin or Revenue Reports privilege)."""
    try:
        if getattr(current_user, 'userRole', None) == 'Admin':
            return True
        if hasattr(current_user, 'has_privilege'):
            return current_user.has_privilege('Revenue Reports')
    except Exception:
        pass
    return False


def can_view_checkin_reports():
    """Check if user can access weekly check-in reports.

    Accepts the canonical Admin / Brand Ambassador role names, plus common
    SuperAdmin / Manager variants used elsewhere in the app. Case- and
    whitespace-insensitive so mobile clients that lowercase the role name do
    not get spurious 403s.
    """
    try:
        role = (getattr(current_user, 'userRole', '') or '').strip().lower()
        admin_roles = {
            'admin',
            'super admin',
            'superadmin',
            'super-admin',
            'manager',
        }
        if role in admin_roles:
            return True
        if role == 'brand ambassador':
            return True
        if hasattr(current_user, 'has_privilege'):
            try:
                if current_user.has_privilege('Check-in Reports'):
                    return True
            except Exception:
                pass
    except Exception:
        pass
    return False


def _get_checkin_reports_timezone():
    tz_name = (AppSetting.get_value('checkin_reports_timezone', 'Africa/Harare') or 'Africa/Harare').strip()
    try:
        return ZoneInfo(tz_name), tz_name
    except Exception:
        app.logger.warning("Invalid checkin reports timezone '%s'. Falling back to UTC.", tz_name)
        return timezone.utc, 'UTC'


def _resolve_week_ending_date(week_ending_str=None, tzinfo=None):
    tzinfo = tzinfo or timezone.utc
    now_local = datetime.now(tzinfo)
    if week_ending_str:
        try:
            week_end_date = datetime.strptime(week_ending_str, '%Y-%m-%d').date()
        except ValueError:
            raise ValueError("week_ending must use YYYY-MM-DD format")
    else:
        days_since_friday = (now_local.weekday() - 4) % 7
        week_end_date = (now_local - timedelta(days=days_since_friday)).date()
    if week_end_date.weekday() != 4:
        raise ValueError("week_ending must be a Friday date")
    return week_end_date


def _compute_checkin_week_window(week_ending_date, tzinfo):
    start_date = week_ending_date - timedelta(days=6)
    start_local = datetime.combine(start_date, datetime.min.time(), tzinfo)
    end_local = datetime.combine(week_ending_date, datetime.max.time(), tzinfo)
    start_utc = start_local.astimezone(timezone.utc).replace(tzinfo=None)
    end_utc = end_local.astimezone(timezone.utc).replace(tzinfo=None)
    return {
        'start_local': start_local,
        'end_local': end_local,
        'start_utc': start_utc,
        'end_utc': end_utc,
    }


def _build_weekly_checkin_report_rows(
    start_utc,
    end_utc,
    champion_user_id=None,
    *,
    include_admin_fields=False,
):
    logs_query = SchoolVisitLog.query.filter(
        SchoolVisitLog.checkin_at >= start_utc,
        SchoolVisitLog.checkin_at <= end_utc
    )
    if champion_user_id is not None:
        logs_query = logs_query.filter(SchoolVisitLog.user_id == champion_user_id)
    logs = logs_query.order_by(SchoolVisitLog.checkin_at.desc()).all()

    user_lookup = {
        u.username: u
        for u in User.query.filter(User.userRole == 'Brand Ambassador').all()
    }
    rows = []
    for log in logs:
        champion_user = user_lookup.get(log.username)
        champion_name = log.username
        champion_province = None
        champion_email = None
        if champion_user:
            champion_name = f"{champion_user.firstname} {champion_user.lastname}".strip() or champion_user.username
            champion_province = champion_user.province
            champion_email = champion_user.email

        rows.append(
            _school_visit_log_report_row(
                log,
                champion_name=champion_name,
                champion_province=champion_province,
                champion_email=champion_email,
                include_admin_fields=include_admin_fields,
            )
        )
    return rows


def _weekly_checkin_report_headers():
    return [
        "champion",
        "username",
        "province",
        "school_name",
        "status",
        "checkin_at",
        "checkout_at",
        "duration_minutes",
        "asl_school_id",
        "library_id",
        "location_source",
        "location_text",
        "checkin_latitude",
        "checkin_longitude",
        "checkout_location_source",
        "checkout_location_text",
        "checkout_latitude",
        "checkout_longitude",
        "checkin_device_install_id",
        "checkout_device_install_id",
        "visit_grades",
        "teacher_name",
        "teacher_contact",
    ]


def _weekly_checkin_rows_to_csv_bytes(rows):
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=_weekly_checkin_report_headers())
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k) for k in _weekly_checkin_report_headers()})
    return output.getvalue().encode('utf-8')


def _weekly_checkin_rows_to_pdf_bytes(rows, *, role_label, week_ending_date):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=20, rightMargin=20, topMargin=20, bottomMargin=20)
    styles = getSampleStyleSheet()
    title = Paragraph(
        f"Weekly Check-in Report ({role_label}) - Week ending {week_ending_date.isoformat()}",
        styles["Heading3"]
    )
    subtitle = Paragraph(f"Total records: {len(rows)}", styles["BodyText"])

    table_rows = [["Champion", "Province", "School", "Status", "Check-in", "Check-out", "Duration", "Location"]]
    for row in rows:
        location_display = row.get("location_text") or ""
        if row.get("checkin_latitude") is not None and row.get("checkin_longitude") is not None:
            location_display = f"{location_display} ({row.get('checkin_latitude')}, {row.get('checkin_longitude')})".strip()
        table_rows.append([
            str(row.get("champion") or "-"),
            str(row.get("province") or "-"),
            str(row.get("school_name") or "-"),
            str(row.get("status") or "-"),
            str(row.get("checkin_at") or "-"),
            str(row.get("checkout_at") or "-"),
            str(row.get("duration_minutes") if row.get("duration_minutes") is not None else "-"),
            location_display or "-",
        ])

    table = Table(table_rows, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f4f4f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
    ]))

    doc.build([title, Spacer(1, 8), subtitle, Spacer(1, 10), table])
    return buffer.getvalue()


def _weekly_checkin_filename(role_scope, week_ending_date, ext):
    return f"checkins_weekly_{role_scope}_{week_ending_date.isoformat()}.{ext}"


def _send_email_message(message):
    if app.config.get('MAIL_SUPPRESS_SEND'):
        app.logger.info("[MAIL_SUPPRESS_SEND] Email skipped for recipients=%s", message.get_all('To', []))
        return True

    smtp_host = app.config.get('MAIL_SERVER', 'smtp.gmail.com')
    smtp_port = app.config.get('MAIL_PORT', 587)
    use_tls = app.config.get('MAIL_USE_TLS', True)
    username = app.config.get('MAIL_USERNAME')
    password = app.config.get('MAIL_PASSWORD')

    if not (username and password):
        app.logger.warning('Email not sent: MAIL_USERNAME/PASSWORD not fully configured')
        return False

    with smtplib.SMTP(smtp_host, smtp_port) as server:
        if use_tls:
            server.starttls()
        server.login(username, password)
        server.send_message(message)
    return True


def _send_weekly_checkin_report_email(*, recipients, role_label, week_ending_date, csv_bytes, csv_name, pdf_bytes, pdf_name, total_count):
    recipients = [r for r in (recipients or []) if r]
    if not recipients:
        return {'status': 'skipped', 'reason': 'no_recipients'}

    sender = app.config.get('MAIL_DEFAULT_SENDER', app.config.get('MAIL_USERNAME'))
    if not sender:
        app.logger.warning('Weekly check-in report email skipped: sender not configured')
        return {'status': 'skipped', 'reason': 'sender_not_configured'}

    msg = EmailMessage()
    msg['Subject'] = f"Weekly Check-in Report ({role_label}) - {week_ending_date.isoformat()}"
    msg['From'] = sender
    msg['To'] = ', '.join(recipients)
    msg.set_content(
        f"Attached are the weekly check-in reports ({role_label}) for week ending {week_ending_date.isoformat()}.\n"
        f"Total records: {total_count}."
    )
    msg.add_alternative(
        f"<p>Attached are the weekly check-in reports (<strong>{role_label}</strong>) for week ending "
        f"<strong>{week_ending_date.isoformat()}</strong>.</p><p>Total records: <strong>{total_count}</strong>.</p>",
        subtype='html'
    )
    msg.add_attachment(csv_bytes, maintype='text', subtype='csv', filename=csv_name)
    msg.add_attachment(pdf_bytes, maintype='application', subtype='pdf', filename=pdf_name)

    sent = _send_email_message(msg)
    return {'status': 'sent' if sent else 'failed', 'recipients': recipients}


def _get_weekly_checkin_report_payload(week_ending_str=None, role_scope='auto'):
    tzinfo, tz_name = _get_checkin_reports_timezone()
    week_ending_date = _resolve_week_ending_date(week_ending_str, tzinfo=tzinfo)
    window = _compute_checkin_week_window(week_ending_date, tzinfo)

    if role_scope == 'auto':
        role = (getattr(current_user, 'userRole', '') or '').strip().lower()
        admin_roles = {
            'admin',
            'super admin',
            'superadmin',
            'super-admin',
            'manager',
        }
        scope = 'admin' if role in admin_roles else 'champion'
    else:
        scope = role_scope

    champion_user_id = current_user.id if scope == 'champion' else None
    rows = _build_weekly_checkin_report_rows(
        window['start_utc'],
        window['end_utc'],
        champion_user_id=champion_user_id,
        include_admin_fields=(scope == 'admin'),
    )

    return {
        'rows': rows,
        'scope': scope,
        'week_ending_date': week_ending_date,
        'timezone': tz_name,
        'window': window,
    }


def get_champion_school_approvers():
    """Return list of User ids who can approve champion school requests (Admin role; privilege holders can still approve via UI)."""
    return [u.id for u in User.query.filter(User.userRole == 'Admin').all()]


@app.route('/api/branding/inventory', methods=['GET'])
@login_required
def api_branding_inventory():
    items = BrandingItem.query.order_by(BrandingItem.platform.asc(), BrandingItem.item_type.asc()).all()
    return jsonify([i.to_dict() for i in items])

@app.route('/api/branding/requests', methods=['GET'])
@login_required
def api_branding_requests():
    reqs = BrandingRequest.query.order_by(BrandingRequest.created_at.desc()).all()
    # include actions history for clarity
    out = []
    for r in reqs:
        row = r.to_dict()
        acts = BrandingAction.query.filter_by(request_id=r.id).order_by(BrandingAction.created_at.asc()).all()
        row['actions'] = [a.to_dict() for a in acts]
        out.append(row)
    return jsonify(out)

@app.route('/api/branding/request', methods=['POST'])
@login_required
def api_branding_request_create():
    data = request.get_json() or {}
    platform = (data.get('platform') or '').strip()
    item_type = (data.get('item_type') or '').strip()
    quantity = int(data.get('quantity') or 1)
    event_name = (data.get('event_name') or '').strip()
    checkout_date = data.get('checkout_date')
    return_date = data.get('return_date')

    if not platform or not item_type or quantity <= 0:
        return jsonify({'error': 'platform, item_type and positive quantity are required'}), 400

    try:
        co_dt = datetime.strptime(checkout_date, '%Y-%m-%d').date() if checkout_date else None
        rt_dt = datetime.strptime(return_date, '%Y-%m-%d').date() if return_date else None
    except Exception:
        return jsonify({'error': 'Invalid date format, expected YYYY-MM-DD'}), 400

    req = BrandingRequest(
        requester_username=current_user.username,
        event_name=event_name,
        platform=platform,
        item_type=item_type,
        quantity=quantity,
        checkout_date=co_dt,
        return_date=rt_dt,
        status='Pending'
    )
    db.session.add(req)
    db.session.commit()
    return jsonify(req.to_dict()), 201

@app.route('/api/branding/requests/<int:req_id>/approve', methods=['POST'])
@login_required
def api_branding_request_approve(req_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    req = db.session.get(BrandingRequest, req_id)
    if not req:
        return jsonify({'error': 'Request not found'}), 404
    if req.status == 'Approved':
        return jsonify({'message': 'Already approved', 'request': req.to_dict()}), 200

    item = BrandingItem.query.filter_by(platform=req.platform, item_type=req.item_type).first()
    if not item:
        return jsonify({'error': 'Inventory item not found for this platform/type'}), 404
    if item.quantity_available < req.quantity:
        return jsonify({'error': 'Insufficient quantity available', 'available': item.quantity_available}), 400

    item.quantity_available -= req.quantity
    req.status = 'Approved'
    db.session.add(BrandingAction(request_id=req.id, action='approve', actor_username=current_user.username))
    db.session.commit()
    return jsonify({'request': req.to_dict(), 'inventory': item.to_dict()}), 200

@app.route('/api/branding/requests/<int:req_id>/decline', methods=['POST'])
@login_required
def api_branding_request_decline(req_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    req = db.session.get(BrandingRequest, req_id)
    if not req:
        return jsonify({'error': 'Request not found'}), 404
    data = request.get_json() or {}
    comment = (data.get('comment') or '').strip() or None
    req.status = 'Declined'
    db.session.add(BrandingAction(request_id=req.id, action='decline', actor_username=current_user.username, comment=comment))
    db.session.commit()
    return jsonify({'request': req.to_dict()}), 200

@app.route('/api/branding/requests/<int:req_id>/mark_returned', methods=['POST'])
@login_required
def api_branding_request_mark_returned(req_id):
    req = db.session.get(BrandingRequest, req_id)
    if not req:
        return jsonify({'error': 'Request not found'}), 404
    if req.requester_username != current_user.username and current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    if req.status != 'Approved':
        return jsonify({'error': 'Only approved requests can be marked returned'}), 400
    req.status = 'Return Pending'
    db.session.add(BrandingAction(request_id=req.id, action='mark_returned', actor_username=current_user.username))
    db.session.commit()
    return jsonify({'request': req.to_dict()}), 200

@app.route('/api/branding/requests/<int:req_id>/ack_return', methods=['POST'])
@login_required
def api_branding_request_ack_return(req_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    req = db.session.get(BrandingRequest, req_id)
    if not req:
        return jsonify({'error': 'Request not found'}), 404
    if req.status != 'Return Pending':
        return jsonify({'error': 'Request is not pending return'}), 400
    item = BrandingItem.query.filter_by(platform=req.platform, item_type=req.item_type).first()
    if not item:
        # Create inventory entry if missing
        item = BrandingItem(platform=req.platform, item_type=req.item_type, quantity_available=0)
        db.session.add(item)
        db.session.flush()
    item.quantity_available += req.quantity
    req.status = 'Returned'
    db.session.add(BrandingAction(request_id=req.id, action='ack_return', actor_username=current_user.username))
    db.session.commit()
    return jsonify({'request': req.to_dict(), 'inventory': item.to_dict()}), 200

@app.route('/api/branding/requests/<int:req_id>/decline_return', methods=['POST'])
@login_required
def api_branding_request_decline_return(req_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    req = db.session.get(BrandingRequest, req_id)
    if not req:
        return jsonify({'error': 'Request not found'}), 404
    if req.status != 'Return Pending':
        return jsonify({'error': 'Request is not pending return'}), 400
    data = request.get_json() or {}
    comment = (data.get('comment') or '').strip() or None
    # Decline return: keep request Approved (still out with user)
    req.status = 'Approved'
    db.session.add(BrandingAction(request_id=req.id, action='decline_return', actor_username=current_user.username, comment=comment))
    db.session.commit()
    return jsonify({'request': req.to_dict()}), 200

# -------- End Branding APIs --------

# Inventory management (Admin)
@app.route('/api/branding/inventory', methods=['POST'])
@login_required
def api_branding_inventory_upsert():
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    data = request.get_json() or {}
    platform = (data.get('platform') or '').strip()
    item_type = (data.get('item_type') or '').strip()
    try:
        qty = int(data.get('quantity_available'))
    except Exception:
        return jsonify({'error': 'quantity_available must be an integer'}), 400
    if not platform or not item_type:
        return jsonify({'error': 'platform and item_type are required'}), 400
    item = BrandingItem.query.filter_by(platform=platform, item_type=item_type).first()
    if item:
        item.quantity_available = qty
    else:
        item = BrandingItem(platform=platform, item_type=item_type, quantity_available=qty)
        db.session.add(item)
    db.session.commit()
    return jsonify(item.to_dict()), 200

@app.route('/api/branding/inventory/<int:item_id>', methods=['PATCH'])
@login_required
def api_branding_inventory_update(item_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    item = db.session.get(BrandingItem, item_id)
    if not item:
        return jsonify({'error': 'Not found'}), 404
    data = request.get_json() or {}
    if 'platform' in data and data['platform']:
        item.platform = data['platform']
    if 'item_type' in data and data['item_type']:
        item.item_type = data['item_type']
    if 'quantity_available' in data and data['quantity_available'] is not None:
        try:
            item.quantity_available = int(data['quantity_available'])
        except Exception:
            return jsonify({'error': 'quantity_available must be integer'}), 400
    db.session.commit()
    return jsonify(item.to_dict())

@app.route('/api/branding/inventory/<int:item_id>', methods=['DELETE'])
@login_required
def api_branding_inventory_delete(item_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    item = db.session.get(BrandingItem, item_id)
    if not item:
        return jsonify({'error': 'Not found'}), 404
    db.session.delete(item)
    db.session.commit()
    return jsonify({'ok': True})

# Bulk requests: create multiple BrandingRequest rows from one submission
@app.route('/api/branding/requests/bulk', methods=['POST'])
@login_required
def api_branding_requests_bulk():
    data = request.get_json() or {}
    items = data.get('items') or []
    event_name = (data.get('event_name') or '').strip()
    checkout_date = data.get('checkout_date')
    return_date = data.get('return_date')
    try:
        co_dt = datetime.strptime(checkout_date, '%Y-%m-%d').date() if checkout_date else None
        rt_dt = datetime.strptime(return_date, '%Y-%m-%d').date() if return_date else None
    except Exception:
        return jsonify({'error': 'Invalid date format'}), 400
    created = []
    for it in items:
        platform = (it.get('platform') or '').strip()
        item_type = (it.get('item_type') or '').strip()
        try:
            qty = int(it.get('quantity') or 0)
        except Exception:
            qty = 0
        if not platform or not item_type or qty <= 0:
            continue
        req = BrandingRequest(
            requester_username=current_user.username,
            event_name=event_name,
            platform=platform,
            item_type=item_type,
            quantity=qty,
            checkout_date=co_dt,
            return_date=rt_dt,
            status='Pending'
        )
        db.session.add(req)
        created.append(req)
    if not created:
        return jsonify({'error': 'No valid items to create'}), 400
    db.session.commit()
    return jsonify([r.to_dict() for r in created]), 201

# -------- Collateral Management APIs --------

@app.route('/api/collateral/items', methods=['GET'])
@login_required
def api_collateral_items_list():
    items = CollateralItems.query.order_by(CollateralItems.created_at.desc()).all()
    return jsonify({'items': [item.to_dict() for item in items]}), 200

@app.route('/api/collateral/items', methods=['POST'])
@login_required
def api_collateral_items_create():
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    data = request.get_json() or {}
    collateral_name = (data.get('collateral_name') or '').strip()
    status = (data.get('status') or 'available').strip()
    
    if not collateral_name:
        return jsonify({'error': 'collateral_name is required'}), 400
    if status not in ['available', 'unavailable']:
        return jsonify({'error': 'status must be available or unavailable'}), 400
    
    item = CollateralItems(
        collateral_name=collateral_name,
        status=status,
        added_by=current_user.username
    )
    db.session.add(item)
    db.session.commit()
    return jsonify(item.to_dict()), 201

@app.route('/api/collateral/items/<int:item_id>/status', methods=['PUT'])
@login_required
def api_collateral_items_toggle_status(item_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    item = db.session.get(CollateralItems, item_id)
    if not item:
        return jsonify({'error': 'Item not found'}), 404
    data = request.get_json() or {}
    new_status = (data.get('status') or '').strip()
    if new_status not in ['available', 'unavailable']:
        return jsonify({'error': 'status must be available or unavailable'}), 400
    item.status = new_status
    db.session.commit()
    return jsonify(item.to_dict()), 200

@app.route('/api/collateral/items/<int:item_id>', methods=['DELETE'])
@login_required
def api_collateral_items_delete(item_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    item = db.session.get(CollateralItems, item_id)
    if not item:
        return jsonify({'error': 'Item not found'}), 404
    db.session.delete(item)
    db.session.commit()
    return jsonify({'ok': True}), 200

# Collateral Request APIs
@app.route('/api/collateral/requests', methods=['GET'])
@login_required
def api_collateral_requests_list():
    # Admins see all requests, users see only their own
    if is_admin_user():
        requests = CollateralRequest.query.order_by(CollateralRequest.created_at.desc()).all()
    else:
        requests = CollateralRequest.query.filter_by(requester_username=current_user.username).order_by(CollateralRequest.created_at.desc()).all()
    return jsonify({'requests': [req.to_dict() for req in requests]}), 200

@app.route('/api/collateral/requests', methods=['POST'])
@login_required
def api_collateral_requests_create():
    data = request.get_json() or {}
    collateral_item_id = data.get('collateral_item_id')
    event_details = (data.get('event_details') or '').strip()
    needed_by_date = data.get('needed_by_date')
    
    if not collateral_item_id or not event_details or not needed_by_date:
        return jsonify({'error': 'collateral_item_id, event_details, and needed_by_date are required'}), 400
    
    # Verify collateral item exists and is available
    item = db.session.get(CollateralItems, collateral_item_id)
    if not item:
        return jsonify({'error': 'Collateral item not found'}), 404
    if item.status != 'available':
        return jsonify({'error': 'Collateral item is not available'}), 400
    
    # Parse date
    try:
        needed_date = datetime.strptime(needed_by_date, '%Y-%m-%d').date()
    except Exception:
        return jsonify({'error': 'Invalid date format. Use YYYY-MM-DD'}), 400
    
    # Create request
    collateral_request = CollateralRequest(
        collateral_item_id=collateral_item_id,
        requester_username=current_user.username,
        event_details=event_details,
        needed_by_date=needed_date,
        status='Pending'
    )
    db.session.add(collateral_request)
    db.session.commit()
    return jsonify(collateral_request.to_dict()), 201

@app.route('/api/collateral/requests/<int:request_id>/approve', methods=['POST'])
@login_required
def api_collateral_requests_approve(request_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    
    collateral_request = db.session.get(CollateralRequest, request_id)
    if not collateral_request:
        return jsonify({'error': 'Request not found'}), 404
    if collateral_request.status != 'Pending':
        return jsonify({'error': 'Request is not pending'}), 400
    
    collateral_request.status = 'Approved'
    collateral_request.approved_by = current_user.username
    db.session.commit()
    return jsonify(collateral_request.to_dict()), 200

@app.route('/api/collateral/requests/<int:request_id>/decline', methods=['POST'])
@login_required
def api_collateral_requests_decline(request_id):
    if not is_admin_user():
        return jsonify({'error': 'Unauthorized'}), 403
    
    collateral_request = db.session.get(CollateralRequest, request_id)
    if not collateral_request:
        return jsonify({'error': 'Request not found'}), 404
    if collateral_request.status != 'Pending':
        return jsonify({'error': 'Request is not pending'}), 400
    
    data = request.get_json() or {}
    decline_reason = (data.get('decline_reason') or '').strip()
    if not decline_reason:
        return jsonify({'error': 'decline_reason is required'}), 400
    
    collateral_request.status = 'Declined'
    collateral_request.approved_by = current_user.username
    collateral_request.decline_reason = decline_reason
    db.session.commit()
    return jsonify(collateral_request.to_dict()), 200

# -------- End Collateral APIs --------

@app.route('/api/province-trainers', methods=['GET'])
def api_province_trainers():
    try:
        # Get today's date and first of the month
        today = datetime.today().date()
        start_date = today.replace(day=1)

        conn = get_ruzivo_conn()
        with conn.cursor() as cursor:
            # Main data query
            query = """
                SELECT 
                    s.school_id,
                    s.username,
                    s.grade,
                    s.last_login,
                    t.school_name
                FROM vwstudent s
                JOIN tblschools t ON s.school_id = t.school_id
                WHERE t.school_name = %s
                  AND s.last_login BETWEEN %s AND %s
                LIMIT 10;
            """
            cursor.execute(query, (
                'MUFAKOSE NO 2 HIGH SCHOOL',
                start_date,
                today
            ))
            columns = [col[0] for col in cursor.description]
            rows = cursor.fetchall()

            # Count query (same filters)
            count_query = """
                SELECT COUNT(*) AS total FROM vwstudent s
                JOIN tblschools t ON s.school_id = t.school_id
                WHERE t.school_name = %s
                  AND s.last_login BETWEEN %s AND %s;
            """
            cursor.execute(count_query, (
                'MUFAKOSE NO 2 HIGH SCHOOL',
                start_date,
                today
            ))
            total_count = cursor.fetchone()['total']

        return jsonify({
            "total_count": total_count,
            "columns": columns,
            "rows": rows,
            "date_range": {
                "start": str(start_date),
                "end": str(today)
            }
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()





@app.route('/smartlearning-metrics-update',  methods=['GET','POST'])
@login_required
def smartlearning_metrics_update():
    try:
        # Read form data instead of JSON
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')

        today = date.today()
        start_date = start_date or today.replace(day=1).strftime('%Y-%m-%d')
        end_date = end_date or today.strftime('%Y-%m-%d')
        active_30_date = (datetime.strptime(end_date, '%Y-%m-%d') - timedelta(days=30)).strftime('%Y-%m-%d')



        queries = {
            "asl_unique_users": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents_login
                WHERE login_date BETWEEN DATE('{start_date}') AND DATE('{end_date}')
            """,
            "asl_registrations": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents
                WHERE date_added BETWEEN DATE('{start_date}') AND DATE('{end_date}')
            """,
            "asl_total_primary_content": f"""
                SELECT COUNT(DISTINCT ca.student_id) AS value
                FROM tblcontent_access ca
                JOIN tblstudents ts ON ts.student_id = ca.student_id
                WHERE ca.start_time BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_sec_content": f"""
                SELECT COUNT(DISTINCT ca.student_id) AS value
                FROM tblcontent_access_hs ca
                JOIN tblstudents ts ON ts.student_id = ca.student_id
                WHERE ca.start_time BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_primary_exercise": f"""
                SELECT COUNT(DISTINCT tr.student_id) AS value
                FROM tblresults tr
                JOIN tblstudents ts ON ts.student_id = tr.student_id
                WHERE tr.date_added BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_sec_exercise": f"""
                SELECT COUNT(DISTINCT tr.student_id) AS value
                FROM tblresults_hs tr
                JOIN tblstudents ts ON ts.student_id = tr.student_id
                WHERE tr.date_added BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_zimsec_access": f"""
                SELECT COUNT(DISTINCT zi.student_id) AS value
                FROM tblcontent_access_zimsec zi
                JOIN tblstudents ts ON ts.student_id = zi.student_id
                WHERE zi.start_time BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_teacher_set_activities": f"""
                SELECT COUNT(DISTINCT ta.student_id) AS value
                FROM tblclass_activity_results ta
                JOIN tblstudents ts ON ts.student_id = ta.student_id
                WHERE ta.date_added BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_teacher_access": f"""
                SELECT COUNT(teacher_id) AS value
                FROM tbl_teacher
                WHERE active_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
            """,
            "asl_revenue": f"""
                SELECT currency, SUM(amount) AS value
                FROM tblecocash_payment_order
                WHERE date_created BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND transactionOperationStatus = 'COMPLETED'
                GROUP BY currency
            """,
            "asl_unique_subscribers": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblecocash_payment_order
                WHERE date_created BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND transactionOperationStatus = 'COMPLETED'
            """,
            "asl_active30": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents_login
                WHERE login_date BETWEEN DATE('{active_30_date}') AND DATE('{end_date}')
            """
        }

        results = {}
        conn = get_ruzivo_conn()

        with conn.cursor() as cursor:
            for label, query in queries.items():
                cursor.execute(query)
                rows = cursor.fetchall()

                if "Revenue" in label:
                    results[label] = ', '.join(
                        f"{row['currency']}: {row['value']}" for row in rows
                    ) if rows else "N/A"
                else:
                    results[label] = rows[0]['value'] if rows else 0


        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()


# @app.route('/asl_active30', methods=['GET', 'POST'])
# def asl_active30():
#     try:
#         # Date range: last 30 days, inclusive of "today"
#         today = date.today()
#         end_date = today.strftime('%Y-%m-%d')                  # inclusive in response
#         active_30_date = (today - timedelta(days=30)).strftime('%Y-%m-%d')
#         # for SQL range: >= start and < (end + 1 day) to include the full end_date
#         end_exclusive = (today + timedelta(days=1)).strftime('%Y-%m-%d')

#         conn = get_ruzivo_conn()

#         # 0) Canonical list of provinces (ensures we show provinces with 0)
#         sql_provinces = """
#             SELECT DISTINCT TRIM(school_province) AS school_province
#             FROM tblstudent_school
#             WHERE school_province IS NOT NULL AND TRIM(school_province) <> ''
#             ORDER BY school_province
#         """
#         with conn.cursor() as cur:
#             cur.execute(sql_provinces)
#             provinces = [row['school_province'] for row in cur.fetchall()]

#         # 1) Active students (distinct) in window
#         # We'll LEFT JOIN this to students & schools so we don't lose anyone
#         # Normalize gender and province in SQL
#         sql_breakdown = """
#             WITH active_students AS (
#                 SELECT DISTINCT sl.student_id
#                 FROM tblstudents_login sl
#                 WHERE sl.login_date >= %s AND sl.login_date < %s
#             )
#             SELECT
#                 CASE
#                     WHEN UPPER(TRIM(s.gender)) IN ('MALE','M') THEN 'Male'
#                     WHEN UPPER(TRIM(s.gender)) IN ('FEMALE','F') THEN 'Female'
#                     ELSE 'Unknown'
#                 END AS gender,
#                 COALESCE(NULLIF(TRIM(ss.school_province), ''), 'Unknown') AS school_province,
#                 COUNT(DISTINCT a.student_id) AS active_count
#             FROM active_students a
#             LEFT JOIN tblstudents s ON a.student_id = s.student_id
#             LEFT JOIN tblstudent_school ss ON s.school_id = ss.student_school_id
#             GROUP BY gender, school_province
#         """

#         with conn.cursor() as cur:
#             cur.execute(sql_breakdown, (active_30_date, end_exclusive))
#             breakdown_raw = cur.fetchall()

#         # 2) Per-province totals (still from the same active_students set)
#         sql_province_totals = """
#             WITH active_students AS (
#                 SELECT DISTINCT sl.student_id
#                 FROM tblstudents_login sl
#                 WHERE sl.login_date >= %s AND sl.login_date < %s
#             )
#             SELECT
#                 COALESCE(NULLIF(TRIM(ss.school_province), ''), 'Unknown') AS school_province,
#                 COUNT(DISTINCT a.student_id) AS total_active
#             FROM active_students a
#             LEFT JOIN tblstudents s ON a.student_id = s.student_id
#             LEFT JOIN tblstudent_school ss ON s.school_id = ss.student_school_id
#             GROUP BY school_province
#         """
#         with conn.cursor() as cur:
#             cur.execute(sql_province_totals, (active_30_date, end_exclusive))
#             province_totals_raw = cur.fetchall()

#         # 3) Overall distinct total (matches above set)
#         sql_total = """
#             SELECT COUNT(DISTINCT sl.student_id) AS total_active
#             FROM tblstudents_login sl
#             WHERE sl.login_date >= %s AND sl.login_date < %s
#         """
#         with conn.cursor() as cur:
#             cur.execute(sql_total, (active_30_date, end_exclusive))
#             total_active = cur.fetchone()['total_active']

#         conn.close()

#         # Normalize genders set
#         genders = ['Male', 'Female', 'Unknown']

#         # Ensure 'Unknown' province is included if present in raw
#         raw_provinces = {r['school_province'] for r in breakdown_raw} | {r['school_province'] for r in province_totals_raw}
#         if 'Unknown' in raw_provinces and 'Unknown' not in provinces:
#             provinces.append('Unknown')

#         # Build fast lookup maps
#         bmap = {(r['school_province'], r['gender']): r['active_count'] for r in breakdown_raw}
#         tmap = {r['school_province']: r['total_active'] for r in province_totals_raw}

#         # Fill zeros for missing province/gender combos
#         breakdown = []
#         for prov in provinces:
#             for g in genders:
#                 breakdown.append({
#                     "school_province": prov,
#                     "gender": g,
#                     "active_count": int(bmap.get((prov, g), 0))
#                 })

#         # Province totals (ensure zeros for missing)
#         province_totals = [{"school_province": prov, "total_active": int(tmap.get(prov, 0))} for prov in provinces]

#         return jsonify({
#             "active_30_date": active_30_date,
#             "end_date": end_date,
#             "total_active": int(total_active),
#             "genders": genders,
#             "provinces": provinces,
#             "province_totals": province_totals,
#             "breakdown": breakdown
#         })

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

from datetime import date, timedelta
from flask import jsonify

def normalize_province(name):
    """Normalize province names consistently."""
    if not name:
        return 'Unknown'
    return name.replace('_', ' ').strip().title()

@app.route('/asl_active30', methods=['GET'])
@login_required
def asl_active30():
    conn = None
    try:
        today = date.today()
        active_30_date = (today - timedelta(days=30)).strftime('%Y-%m-%d')
        end_exclusive = (today + timedelta(days=1)).strftime('%Y-%m-%d')

        conn = get_ruzivo_conn()

        # --- Step 1: Get total_active directly (raw) ---
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COUNT(DISTINCT l.student_id) AS total_active
                FROM tblstudents_login l
                WHERE l.login_date >= %s AND l.login_date < %s
            """, (active_30_date, end_exclusive))
            total_active = cur.fetchone()['total_active'] or 0

        if total_active == 0:
            return jsonify({
                "active_30_date": active_30_date,
                "end_date": today.strftime('%Y-%m-%d'),
                "total_active": 0,
                "gender_totals": {},
                "province_totals": [],
                "breakdown": []
            })

        # --- Step 2: Define active set query for re-use ---
        active_set_sql = """
            SELECT DISTINCT l.student_id
            FROM tblstudents_login l
            WHERE l.login_date >= %s AND l.login_date < %s
        """

        # --- Step 3: Execute breakdown queries ---
        with conn.cursor() as cur:
            # Gender totals
            cur.execute(f"""
                SELECT
                    CASE
                        WHEN UPPER(TRIM(v.gender)) IN ('MALE','M') THEN 'Male'
                        WHEN UPPER(TRIM(v.gender)) IN ('FEMALE','F') THEN 'Female'
                        ELSE 'Unknown'
                    END AS gender,
                    COUNT(*) AS active_count
                FROM ({active_set_sql}) a
                LEFT JOIN vwstudentschools v ON v.student_id = a.student_id
                GROUP BY gender
            """, (active_30_date, end_exclusive))
            gender_rows = cur.fetchall()

            # Province totals
            cur.execute(f"""
                SELECT
                    COALESCE(TRIM(s.school_province), 'Unknown') AS school_province,
                    COUNT(*) AS province_total_active
                FROM ({active_set_sql}) a
                LEFT JOIN vwstudentschools v ON v.student_id = a.student_id
                LEFT JOIN tblschools s ON s.school_id = v.school_id
                GROUP BY s.school_province
            """, (active_30_date, end_exclusive))
            province_rows = cur.fetchall()

            # Province + Gender breakdown
            cur.execute(f"""
                SELECT
                    CASE
                        WHEN UPPER(TRIM(v.gender)) IN ('MALE','M') THEN 'Male'
                        WHEN UPPER(TRIM(v.gender)) IN ('FEMALE','F') THEN 'Female'
                        ELSE 'Unknown'
                    END AS gender,
                    COALESCE(TRIM(s.school_province), 'Unknown') AS school_province,
                    COUNT(*) AS active_count
                FROM ({active_set_sql}) a
                LEFT JOIN vwstudentschools v ON v.student_id = a.student_id
                LEFT JOIN tblschools s ON s.school_id = v.school_id
                GROUP BY gender, s.school_province
            """, (active_30_date, end_exclusive))
            breakdown_rows = cur.fetchall()

        # --- Step 4: Normalize + combine ---
        def normalize_province(name):
            if not name:
                return 'Unknown'
            cleaned = name.replace('_', ' ').replace(';', '').strip().title()
            # Handle special cases
            if cleaned.lower() in ['bulawayo', 'bulawyo', 'bulawyoa', 'bulawyao', 'bulawayo']:
                return 'Bulawayo'
            return cleaned

        gender_totals = {}
        for r in gender_rows:
            g = r['gender'] or 'Unknown'
            count = int(r['active_count'] or 0)
            if count > 0:
                gender_totals[g] = gender_totals.get(g, 0) + count

        province_totals_map = {}
        for r in province_rows:
            prov = normalize_province(r['school_province'])
            count = int(r['province_total_active'] or 0)
            if count > 0:
                province_totals_map[prov] = province_totals_map.get(prov, 0) + count
        province_totals = [{"school_province": p, "province_total_active": c}
                           for p, c in sorted(province_totals_map.items())]

        breakdown_map = {}
        for r in breakdown_rows:
            prov = normalize_province(r['school_province'])
            g = r['gender'] or 'Unknown'
            count = int(r['active_count'] or 0)
            if count > 0:
                key = (prov, g)
                breakdown_map[key] = breakdown_map.get(key, 0) + count
        breakdown = [{"school_province": prov, "gender": g, "active_count": c}
                     for (prov, g), c in breakdown_map.items()]

        # --- Step 5: Consistency checks ---
        sum_genders = sum(gender_totals.values())
        sum_provinces = sum(p['province_total_active'] for p in province_totals)

        # Adjust if mismatch
        if sum_genders != total_active:
            gender_totals['Unknown'] = gender_totals.get('Unknown', 0) + (total_active - sum_genders)
        if sum_provinces != total_active:
            province_totals_map['Unknown'] = province_totals_map.get('Unknown', 0) + (total_active - sum_provinces)
            province_totals = [{"school_province": p, "province_total_active": c}
                               for p, c in sorted(province_totals_map.items())]

        return jsonify({
            "active_30_date": active_30_date,
            "end_date": today.strftime('%Y-%m-%d'),
            "total_active": total_active,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    # finally:
    #     if conn:
    #         conn.close()






from flask import request, jsonify
from datetime import date, timedelta
import pymysql

# @app.route('/api/asl-active90', methods=['GET','POST'])
# def asl_active90():
#     conn = None
#     try:
#         # Read custom range from form
#         # start_date = request.form.get("start_date")
#         # end_date = request.form.get("end_date")

#         start_date = date(2025, 7, 1)
#         end_date = date(2025, 7, 24)

#         if not start_date or not end_date:
#             today = date.today()
#             end_date = today
#             start_date = today - timedelta(days=90)
#         else:
#             # start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
#             # end_date = datetime.strptime(end_date, "%Y-%m-%d").date()
#             start_date = date(2025, 7, 1)
#             end_date = date(2025, 7, 24)

#         conn = get_ruzivo_conn()
#         query = """
#             SELECT COUNT(DISTINCT student_id) AS asl_active30
#             FROM tblstudents_login
#             WHERE DATE(login_date) BETWEEN %s AND %s
#         """

#         df = pd.read_sql_query(query, conn, params=[start_date, end_date])
#         conn.close()

#         # Get result
#         asl_count = int(df["asl_active30"].iloc[0])

#         # If user wants CSV/Excel download
#         if "download" in request.form:
#             output_format = request.form.get("download")  # "csv" or "excel"
#             output = io.BytesIO()

#             if output_format == "csv":
#                 df.to_csv(output, index=False)
#                 output.seek(0)
#                 return send_file(
#                     output,
#                     as_attachment=True,
#                     download_name=f"asl_active30_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
#                     mimetype="text/csv"
#                 )
#             else:  # default to Excel
#                 with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
#                     df.to_excel(writer, index=False, sheet_name="ASLActive30")
#                 output.seek(0)
#                 return send_file(
#                     output,
#                     as_attachment=True,
#                     download_name=f"asl_active30_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
#                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#                 )

#         return jsonify({
#             "start_date": str(start_date),
#             "end_date": str(end_date),
#             "asl_active30": asl_count
#         })

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500



@app.route('/custom-active-date', methods=['GET', 'POST'])
def custom_active_date():

    return render_template(
        'customactivedate.html',
        title='Analytics'
    )



@app.route('/api/asl-active90', methods=['GET'])
def asl_active90():
    conn = None
    try:
        today = date.today()
        current_year = today.year
        current_month = today.month

        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        monthly_data = []

        for month in range(1, current_month + 1):
            start_date = date(current_year, month, 1)
            end_day = calendar.monthrange(current_year, month)[1]
            end_date = date(current_year, month, end_day)

            query = """
                SELECT COUNT(DISTINCT student_id) AS asl_active
                FROM tblstudents_login
                WHERE DATE(login_date) BETWEEN %s AND %s
            """
            cursor.execute(query, (start_date, end_date))
            result = cursor.fetchone()

            monthly_data.append({
                "month": start_date.strftime("%B"),
                "year": current_year,
                "asl_active": int(result["asl_active"]) if result and result["asl_active"] else 0
            })

        cursor.close()
        conn.close()

        return jsonify({
            "year": current_year,
            "monthly_logins": monthly_data
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()




@app.route('/api/asl-filtered-count', methods=['GET'])
def asl_filtered_count():
    conn = None
    try:
        conn = get_ruzivo_conn()
        with conn.cursor() as cursor:
            query = """
                SELECT COUNT(DISTINCT sl.student_id) AS filtered_count
                FROM tblstudents_login sl
                WHERE sl.student_id NOT IN (
                    SELECT DISTINCT tp.student_id
                    FROM tblecocash_payment_order tp
                    WHERE transactionOperationStatus = 'COMPLETED'
                )
                AND sl.student_id NOT IN (
                    SELECT DISTINCT sl2.student_id
                    FROM tblstudents_login sl2
                    WHERE sl2.login_date BETWEEN '2025-06-01 00:00:00' AND '2025-08-30 23:59:59'
                )
                AND sl.login_date BETWEEN '2025-07-1 00:00:00' AND '2025-08-24 23:59:59';
            """
            cursor.execute(query)
            result = cursor.fetchone()

        return jsonify({
            "filtered_count": int(result[0]) if result and result[0] else 0
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()





import logging

# Province normalization helper (reuse same one as active30)
def normalize_province(name):
    if not name:
        return 'Unknown'

    cleaned = name.replace('_', ' ').replace(';', '').strip().title()

    typo_map = {
        "Bulawyo": "Bulawayo",
        "Bulawyoa": "Bulawayo",
        "Bulawyao": "Bulawayo",
        "Bul;awayo": "Bulawayo",
    }

    if cleaned in typo_map:
        corrected = typo_map[cleaned]
        logging.info(f"Province name corrected: '{cleaned}' → '{corrected}'")
        return corrected

    return cleaned


@app.route('/asl_registrations', methods=['GET', 'POST'])
@login_required
def asl_registrations():
    try:
        today = date.today()
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')

        start_date = start_date or today.replace(day=1).strftime('%Y-%m-%d')
        end_date = end_date or today.strftime('%Y-%m-%d')

        conn = get_ruzivo_conn()

        # --- Step 1: Get total registrations exactly like base query ---
        sql_total = """
            SELECT COUNT(DISTINCT student_id) AS value
            FROM tblstudents
            WHERE DATE(date_added) BETWEEN %s AND %s
        """
        with conn.cursor() as cur:
            cur.execute(sql_total, (start_date, end_date))
            total_registrations = cur.fetchone()['value']

        # --- Step 2: Get breakdown by gender & province ---
        sql_breakdown = """
            SELECT
                t.student_id,
                CASE
                    WHEN UPPER(TRIM(v.gender)) IN ('MALE','M') THEN 'Male'
                    WHEN UPPER(TRIM(v.gender)) IN ('FEMALE','F') THEN 'Female'
                    ELSE 'Unknown'
                END AS gender,
                v.province AS school_province
            FROM tblstudents t
            LEFT JOIN vwstudentschools v ON t.student_id = v.student_id
            WHERE DATE(t.date_added) BETWEEN %s AND %s
        """
        with conn.cursor() as cur:
            cur.execute(sql_breakdown, (start_date, end_date))
            rows = cur.fetchall()

        # conn.close()

        # --- Step 3: Normalize & Aggregate ---
        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = []

        seen = set()
        for r in rows:
            sid = r['student_id']
            if sid in seen:   # enforce DISTINCT student_id
                continue
            seen.add(sid)

            g = r['gender'] or "Unknown"
            prov = normalize_province(r['school_province'])

            gender_totals[g] = gender_totals.get(g, 0) + 1
            province_totals[prov] = province_totals.get(prov, 0) + 1

            breakdown.append({"school_province": prov, "gender": g, "count": 1})

        # --- Step 4: Consistency check ---
        gender_sum = sum(gender_totals.values())
        province_sum = sum(province_totals.values())

        if gender_sum != total_registrations:
            logging.warning(f"Gender totals mismatch ({gender_sum} vs {total_registrations}), correcting.")
            total_registrations = gender_sum

        if province_sum != total_registrations:
            logging.warning(f"Province totals mismatch ({province_sum} vs {total_registrations}), correcting.")
            total_registrations = province_sum

        province_totals_list = [{"school_province": k, "total": v} for k, v in province_totals.items()]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_registrations": total_registrations,
            "gender_totals": gender_totals,
            "province_totals": province_totals_list,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500




def normalize_province(name):
    """Clean up province names (case, underscores, typos)."""
    if not name:
        return "Unknown"
    cleaned = name.replace("_", " ").replace(";", "").strip().title()
    if cleaned.lower().startswith("bul"):
        return "Bulawayo"
    return cleaned

@app.route("/api/asl_total_primary_content", methods=["GET", "POST"])
@login_required
def asl_total_primary_content():
    try:
        today = date.today()
        start_date = request.form.get("start_date") or today.replace(day=1).strftime("%Y-%m-%d")
        end_date = request.form.get("end_date") or today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Total distinct students
        sql_total = """
            SELECT COUNT(DISTINCT ca.student_id) AS total
            FROM tblcontent_access ca
            JOIN tblstudents st ON st.student_id = ca.student_id
            JOIN vwstudent v ON v.student_id = ca.student_id
            JOIN tblschools s ON s.school_id = v.school_id
            WHERE ca.start_time BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
        """
        with conn.cursor() as cur:
            cur.execute(sql_total, (start_date, end_date))
            total_students = cur.fetchone()["total"]

        # 2) Gender + province breakdown
        sql_breakdown = """
            SELECT v.student_id, v.gender, s.school_province
            FROM tblcontent_access ca
            JOIN tblstudents st ON st.student_id = ca.student_id
            JOIN vwstudent v ON v.student_id = ca.student_id
            JOIN tblschools s ON s.school_id = v.school_id
            WHERE ca.start_time BETWEEN %s AND %s
            AND st.student_type != 'STAFF'
        """
        with conn.cursor() as cur:
            cur.execute(sql_breakdown, (start_date, end_date))
            rows = cur.fetchall()

        # conn.close()

        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = {}

        counted_students = set()  # to avoid double-counting

        for row in rows:
            student_id = row["student_id"]
            if student_id in counted_students:
                continue
            counted_students.add(student_id)

            # Normalize gender
            gender = row["gender"].strip().upper() if row["gender"] else "UNKNOWN"
            if gender == "M":
                gender = "Male"
            elif gender == "F":
                gender = "Female"
            elif gender not in ["Male", "Female"]:
                gender = "Unknown"

            # Normalize province
            province = normalize_province(row["school_province"])

            # Aggregate counts
            gender_totals[gender] += 1
            province_totals[province] = province_totals.get(province, 0) + 1
            breakdown[(province, gender)] = breakdown.get((province, gender), 0) + 1

        # Build breakdown list
        breakdown_list = [{"school_province": p, "gender": g, "count": c} for (p, g), c in breakdown.items()]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_students": total_students,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown_list
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500





def normalize_gender(g):
    if not g:
        return "Unknown"
    g = g.strip().upper()
    if g in ["M", "MALE"]:
        return "Male"
    if g in ["F", "FEMALE"]:
        return "Female"
    return "Unknown"

def normalize_province(name):
    if not name:
        return "Unknown"
    cleaned = name.replace("_", " ").replace(";", "").strip().title()
    if cleaned.lower().startswith("bul"):
        return "Bulawayo"
    return cleaned

@app.route("/api/asl_teacher_set_activities", methods=["GET", "POST"])
@login_required
def asl_teacher_set_activities():
    try:
        today = date.today()
        start_date = request.form.get("start_date") or today.replace(day=1).strftime("%Y-%m-%d")
        end_date = request.form.get("end_date") or today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Get distinct student_ids
        sql_students = """
            SELECT DISTINCT ta.student_id
            FROM tblclass_activity_results ta
            JOIN tblstudents ts ON ts.student_id = ta.student_id
            WHERE ta.date_added BETWEEN %s AND %s
              AND ts.student_type != 'STAFF'
        """
        with conn.cursor() as cur:
            cur.execute(sql_students, (start_date, end_date))
            students = [row["student_id"] for row in cur.fetchall()]

        if not students:
            conn.close()
            return jsonify({
                "start_date": start_date,
                "end_date": end_date,
                "total_teacher_set_activities": 0,
                "gender_totals": {},
                "province_totals": {},
                "breakdown": []
            })

        # 2) Build dynamic IN clause
        placeholders = ",".join(["%s"] * len(students))
        sql_details = f"""
            SELECT v.student_id, v.gender, s.school_province
            FROM vwstudentschools v
            JOIN tblschools s ON v.school_id = s.school_id
            WHERE v.student_id IN ({placeholders})
        """

        with conn.cursor() as cur:
            cur.execute(sql_details, tuple(students))
            details = cur.fetchall()

        # conn.close()

        # 3) Aggregate counts
        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = {}

        for row in details:
            gender = normalize_gender(row["gender"])
            province = normalize_province(row["school_province"])

            gender_totals[gender] += 1
            province_totals[province] = province_totals.get(province, 0) + 1

            key = (province, gender)
            breakdown[key] = breakdown.get(key, 0) + 1

        total_teacher_set_activities = len(set(students))

        # ✅ Ensure consistency
        # if sum(gender_totals.values()) != total_teacher_set_activities:
        #     raise ValueError("Gender totals mismatch with overall total.")
        # if sum(province_totals.values()) != total_teacher_set_activities:
        #     raise ValueError("Province totals mismatch with overall total.")

        breakdown_list = []
        for (prov, gender), count in breakdown.items():
            breakdown_list.append({
                "school_province": prov,
                "gender": gender,
                "count": count
            })

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_teacher_set_activities": total_teacher_set_activities,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown_list
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500





def normalize_gender(gender):
    """Normalize gender values."""
    if not gender:
        return "Unknown"
    g = gender.strip().lower()
    if g in ["male", "m"]:
        return "Male"
    if g in ["female", "f"]:
        return "Female"
    return "Unknown"

def normalize_province(name):
    """Clean up province names (case, underscores, typos)."""
    if not name:
        return "Unknown"
    cleaned = name.replace("_", " ").replace(";", "").strip().title()
    if cleaned.lower().startswith("bul"):
        return "Bulawayo"
    return cleaned


@app.route("/api/asl_teacher_access", methods=["GET", "POST"])
@login_required
def asl_teacher_access():
    try:
        today = date.today()
        start_date = request.form.get("start_date") or today.replace(day=1).strftime("%Y-%m-%d")
        end_date   = request.form.get("end_date")   or today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # Pull all teachers active in range with their gender and province
        sql = """
            SELECT t.teacher_id, t.gender, s.school_province
            FROM tbl_teacher t
            JOIN tblschools s ON s.school_id = t.school_id
            WHERE t.active_at BETWEEN %s AND %s
        """
        with conn.cursor(pymysql.cursors.DictCursor) as cur:
            cur.execute(sql, (start_date, end_date))
            rows = cur.fetchall()

        # conn.close()

        if not rows:
            return jsonify({
                "start_date": start_date,
                "end_date": end_date,
                "total_teachers": 0,
                "gender_totals": {},
                "province_totals": {},
                "breakdown": []
            })

        gender_totals  = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = {}

        for r in rows:
            g = normalize_gender(r.get("gender"))
            p = normalize_province(r.get("school_province"))

            gender_totals[g] += 1
            province_totals[p] = province_totals.get(p, 0) + 1
            breakdown[(p, g)] = breakdown.get((p, g), 0) + 1

        total_teachers = len(rows)

        # Consistency (they should always match since all counts derive from `rows`)
        # If you prefer hard enforcement, uncomment the raises.
        # if sum(gender_totals.values()) != total_teachers:
        #     raise ValueError("Gender totals mismatch with overall total.")
        # if sum(province_totals.values()) != total_teachers:
        #     raise ValueError("Province totals mismatch with overall total.")

        breakdown_list = [
            {"school_province": prov, "gender": gen, "count": cnt}
            for (prov, gen), cnt in breakdown.items()
        ]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_teachers": total_teachers,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown_list
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500




def normalize_province(name):
    """Clean up province names (case, underscores, typos)."""
    if not name:
        return "Unknown"
    cleaned = name.replace("_", " ").replace(";", "").strip().title()
    # Handle specific typos
    if cleaned.lower().startswith("bul"):
        return "Bulawayo"
    return cleaned

@app.route("/api/asl_unique_users", methods=["GET", "POST"])
@login_required
def asl_unique_users():
    try:
        today = date.today()
        start_date = request.form.get("start_date") or today.replace(day=1).strftime("%Y-%m-%d")
        end_date = request.form.get("end_date") or today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Get distinct users who logged in within range
        sql_users = """
            SELECT DISTINCT student_id
            FROM tblstudents_login
            WHERE login_date BETWEEN %s AND %s
        """
        with conn.cursor() as cur:
            cur.execute(sql_users, (start_date, end_date))
            users = [row["student_id"] for row in cur.fetchall()]

        if not users:
            return jsonify({
                "start_date": start_date,
                "end_date": end_date,
                "total_unique_users": 0,
                "gender_totals": {},
                "province_totals": {},
                "breakdown": []
            })

        # 2) Get gender + school_id from vwstudentschools
        sql_details = """
            SELECT v.student_id, v.gender, v.school_id, s.school_province
            FROM vwstudentschools v
            JOIN tblschools s ON v.school_id = s.school_id
            WHERE v.student_id IN %s
        """
        with conn.cursor() as cur:
            cur.execute(sql_details, (tuple(users),))
            details = cur.fetchall()

        # conn.close()

        # 3) Aggregate by gender & province
        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = {}

        for row in details:
            gender = row["gender"].strip().capitalize() if row["gender"] else "Unknown"
            if gender not in ["Male", "Female"]:
                gender = "Unknown"

            province = normalize_province(row["school_province"])

            # Count by gender
            gender_totals[gender] += 1

            # Count by province
            if province not in province_totals:
                province_totals[province] = 0
            province_totals[province] += 1

            # Breakdown province × gender
            key = (province, gender)
            breakdown[key] = breakdown.get(key, 0) + 1

        total_unique_users = len(users)

        # ✅ Consistency checks
        # if sum(gender_totals.values()) != total_unique_users:
        #     raise ValueError("Gender totals mismatch with overall total.")
        # if sum(province_totals.values()) != total_unique_users:
        #     raise ValueError("Province totals mismatch with overall total.")

        # 4) Build breakdown list for frontend
        breakdown_list = []
        for (prov, gender), count in breakdown.items():
            breakdown_list.append({
                "school_province": prov,
                "gender": gender,
                "count": count
            })

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_unique_users": total_unique_users,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown_list
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route('/api/asl_unique_subscribers', methods=['GET', 'POST'])
@login_required
def asl_unique_subscribers():
    try:
        today = date.today()
        # default: first of month to today
        start_date = request.form.get('start_date') or today.replace(day=1).strftime('%Y-%m-%d')
        end_date = request.form.get('end_date') or today.strftime('%Y-%m-%d')

        conn = get_ruzivo_conn()

        # --- 1) Total unique subscribers directly from payments ---
        sql_total = """
            SELECT COUNT(DISTINCT p.student_id) AS total_subscribers
            FROM tblecocash_payment_order p
            WHERE p.date_created BETWEEN %s AND %s
              AND p.transactionOperationStatus = 'COMPLETED'
        """
        with conn.cursor() as cur:
            cur.execute(sql_total, (start_date, end_date))
            total_row = cur.fetchone()
            total_subscribers = total_row['total_subscribers'] if total_row else 0

        # --- 2) Breakdown with LEFT JOIN so we don't drop unmatched students ---
        sql_breakdown = """
            SELECT
                CASE
                    WHEN UPPER(TRIM(v.gender)) IN ('MALE','M') THEN 'Male'
                    WHEN UPPER(TRIM(v.gender)) IN ('FEMALE','F') THEN 'Female'
                    ELSE 'Unknown'
                END AS gender,
                TRIM(s.school_province) AS school_province,
                COUNT(DISTINCT p.student_id) AS cnt
            FROM tblecocash_payment_order p
            LEFT JOIN vwstudentschools v ON p.student_id = v.student_id
            LEFT JOIN tblschools s ON v.school_id = s.school_id
            WHERE p.date_created BETWEEN %s AND %s
              AND p.transactionOperationStatus = 'COMPLETED'
            GROUP BY gender, s.school_province
        """
        with conn.cursor() as cur:
            cur.execute(sql_breakdown, (start_date, end_date))
            breakdown_raw = cur.fetchall()

        # conn.close()

        # --- 3) Normalize & aggregate ---
        def normalize_province(name):
            if not name: return "Unknown"
            name = name.replace("_", " ").replace(";", "").strip().title()
            if name.lower() in ["bulawayo", "bul;awayo", "bul awayo"]:
                return "Bulawayo"
            return name

        breakdown_agg = {}
        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}

        for row in breakdown_raw:
            prov = normalize_province(row['school_province'])
            g = row['gender']
            c = int(row['cnt'])

            if prov not in breakdown_agg:
                breakdown_agg[prov] = {"Male": 0, "Female": 0, "Unknown": 0}

            breakdown_agg[prov][g] += c
            gender_totals[g] += c
            province_totals[prov] = province_totals.get(prov, 0) + c

        # --- 4) Guarantee consistency: match raw total query ---
        if sum(gender_totals.values()) != total_subscribers:
            total_subscribers = sum(gender_totals.values())
        if sum(province_totals.values()) != total_subscribers:
            total_subscribers = sum(province_totals.values())

        # --- 5) Format breakdown ---
        breakdown = []
        for prov, gdict in breakdown_agg.items():
            for g, val in gdict.items():
                if val > 0:
                    breakdown.append({
                        "school_province": prov,
                        "gender": g,
                        "count": val
                    })

        province_totals_list = [
            {"school_province": prov, "total_count": val}
            for prov, val in province_totals.items() if val > 0
        ]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_subscribers": total_subscribers,
            "gender_totals": gender_totals,
            "province_totals": province_totals_list,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500



def normalize_province(name):
    """
    Normalize province names:
    - Replace underscores, semicolons with spaces
    - Strip whitespace
    - Title-case
    - Map known variants to canonical form
    """
    if not name:
        return "Unknown"
    clean = name.replace("_", " ").replace(";", "").strip().title()
    mappings = {
        "Bulawayo": ["Bulawayo", "Bul;Awayo", "Bul Awayo"],
    }
    for canonical, variants in mappings.items():
        if clean in variants:
            return canonical
    return clean

@app.route("/api/asl_total_zimsec_access", methods=["GET"])
@login_required
def asl_total_zimsec_access():
    try:
        today = date.today()
        start_date = today.replace(day=1).strftime("%Y-%m-%d")  # first of month
        end_date = today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Baseline total + student list
        sql_students = """
            SELECT DISTINCT zi.student_id
            FROM tblcontent_access_zimsec zi
            JOIN tblstudents ts ON ts.student_id = zi.student_id
            WHERE zi.start_time BETWEEN %s AND %s
              AND ts.student_type != 'STAFF'
        """
        with conn.cursor() as cur:
            cur.execute(sql_students, (start_date, end_date))
            student_ids = [row["student_id"] for row in cur.fetchall()]
        total_access = len(student_ids)

        if total_access == 0:
            return jsonify({
                "start_date": start_date,
                "end_date": end_date,
                "total_access": 0,
                "genders": ["Male", "Female", "Unknown"],
                "gender_totals": {"Male": 0, "Female": 0, "Unknown": 0},
                "province_totals": [],
                "breakdown": []
            })

        # 2) Enrich with gender + province (LEFT JOIN so nothing is dropped)
        sql_breakdown = f"""
            SELECT
                ts.student_id,
                CASE
                    WHEN UPPER(TRIM(vs.gender)) IN ('MALE','M') THEN 'Male'
                    WHEN UPPER(TRIM(vs.gender)) IN ('FEMALE','F') THEN 'Female'
                    ELSE 'Unknown'
                END AS gender,
                TRIM(s.school_province) AS province
            FROM tblstudents ts
            LEFT JOIN vwstudentschools vs ON ts.student_id = vs.student_id
            LEFT JOIN tblschools s ON vs.school_id = s.school_id
            WHERE ts.student_id IN ({",".join(["%s"] * len(student_ids))})
        """
        with conn.cursor() as cur:
            cur.execute(sql_breakdown, student_ids)
            enriched = cur.fetchall()

        # conn.close()

        # 3) Normalize + aggregate
        def normalize_province(name):
            if not name:
                return "Unknown"
            clean = name.replace("_", " ").replace(";", "").strip().title()
            mappings = {"Bulawayo": ["Bulawayo", "Bul;Awayo", "Bul Awayo"]}
            for canonical, variants in mappings.items():
                if clean in variants:
                    return canonical
            return clean

        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown_agg = {}

        for row in enriched:
            gender = row["gender"] or "Unknown"
            prov = normalize_province(row["province"])
            gender_totals[gender] = gender_totals.get(gender, 0) + 1
            province_totals[prov] = province_totals.get(prov, 0) + 1
            breakdown_agg[(prov, gender)] = breakdown_agg.get((prov, gender), 0) + 1

        breakdown = [
            {"school_province": prov, "gender": g, "count": cnt}
            for (prov, g), cnt in breakdown_agg.items()
            if cnt > 0
        ]

        province_totals_list = [
            {"school_province": prov, "total": cnt}
            for prov, cnt in province_totals.items()
            if cnt > 0
        ]

        # ✅ Consistency check
        if sum(gender_totals.values()) != total_access:
            total_access = sum(gender_totals.values())
        if sum([p["total"] for p in province_totals_list]) != total_access:
            total_access = sum([p["total"] for p in province_totals_list])

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_access": total_access,
            "genders": ["Male", "Female", "Unknown"],
            "gender_totals": gender_totals,
            "province_totals": province_totals_list,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route("/api/asl_total_sec_exercise", methods=["GET"])
@login_required
def asl_total_sec_exercise():
    try:
        today = date.today()
        start_date = today.replace(day=1).strftime("%Y-%m-%d")
        end_date = today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Query all distinct students with province and gender
        sql = """
            SELECT DISTINCT tr.student_id,
                CASE
                    WHEN UPPER(TRIM(vs.gender)) IN ('MALE','M') THEN 'Male'
                    WHEN UPPER(TRIM(vs.gender)) IN ('FEMALE','F') THEN 'Female'
                    ELSE 'Unknown'
                END AS gender,
                TRIM(s.school_province) AS province
            FROM tblresults_hs tr
            LEFT JOIN vwstudent vs ON vs.student_id = tr.student_id
            LEFT JOIN tblstudents st ON st.student_id = tr.student_id
            LEFT JOIN tblschools s ON s.school_id = vs.school_id
            WHERE tr.date_added BETWEEN %s AND %s
              AND (st.student_type IS NULL OR st.student_type != 'STAFF')
        """
        with conn.cursor() as cur:
            cur.execute(sql, (start_date, end_date))
            rows = cur.fetchall()

        # conn.close()

        # 2) Normalize province names
        def normalize_province(name):
            if not name:
                return "Unknown"
            clean = name.replace("_", " ").replace(";", "").strip().title()
            mappings = {"Bulawayo": ["Bulawayo", "Bul;Awayo", "Bul Awayo"]}
            for canonical, variants in mappings.items():
                if clean in variants:
                    return canonical
            return clean

        # 3) Aggregate totals
        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown_agg = {}

        for r in rows:
            gender = r["gender"] or "Unknown"
            prov = normalize_province(r["province"])
            gender_totals[gender] += 1
            province_totals[prov] = province_totals.get(prov, 0) + 1
            breakdown_agg[(prov, gender)] = breakdown_agg.get((prov, gender), 0) + 1

        total_exercises = len({r["student_id"] for r in rows})  # match main query

        breakdown = [
            {"school_province": prov, "gender": g, "count": cnt}
            for (prov, g), cnt in breakdown_agg.items()
            if cnt > 0
        ]

        province_totals_list = [
            {"school_province": prov, "total": cnt}
            for prov, cnt in province_totals.items()
            if cnt > 0
        ]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_exercises": total_exercises,
            "genders": ["Male", "Female", "Unknown"],
            "gender_totals": gender_totals,
            "province_totals": province_totals_list,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500





def normalize_province(name):
    if not name:
        return 'Unknown'
    # Handle Bulawayo variants
    if name.upper().replace(";", "").replace("_", "").replace(" ", "") == "BULAWAYO":
        return "Bulawayo"
    return name.replace('_', ' ').strip().title()

def normalize_gender(g):
    if not g:
        return 'Unknown'
    g_clean = g.strip().upper()
    if g_clean in ['M', 'MALE']:
        return 'Male'
    elif g_clean in ['F', 'FEMALE']:
        return 'Female'
    return 'Unknown'

@app.route('/api/asl_total_primary_exercise', methods=['GET'])
@login_required
def asl_total_primary_exercise():
    try:
        today = date.today()
        start_date = today.replace(day=1).strftime('%Y-%m-%d')  # first day of current month
        end_date = today.strftime('%Y-%m-%d')

        conn = get_ruzivo_conn()

        # Count distinct student_id with proper joins
        sql = """
            SELECT tr.student_id, vs.gender, s.school_province
            FROM tblresults tr
            JOIN tblstudents st ON st.student_id = tr.student_id
            JOIN vwstudent vs ON vs.student_id = tr.student_id
            JOIN tblschools s ON s.school_id = vs.school_id
            WHERE tr.date_added BETWEEN %s AND %s
              AND st.student_type != 'STAFF'
            GROUP BY tr.student_id, vs.gender, s.school_province
        """
        with conn.cursor() as cur:
            cur.execute(sql, (start_date, end_date))
            rows = cur.fetchall()

        # conn.close()

        breakdown_agg = {}
        province_totals_agg = {}
        total_exercises = 0

        for r in rows:
            gender = normalize_gender(r['gender'])
            prov = normalize_province(r['school_province'])

            breakdown_agg[(prov, gender)] = breakdown_agg.get((prov, gender), 0) + 1
            province_totals_agg[prov] = province_totals_agg.get(prov, 0) + 1
            total_exercises += 1

        breakdown = [{"school_province": p, "gender": g, "count": c} 
                     for (p, g), c in breakdown_agg.items() if c > 0]

        province_totals = [{"school_province": p, "count": c} 
                           for p, c in province_totals_agg.items() if c > 0]

        # Ensure consistency
        assert sum(d['count'] for d in breakdown) == total_exercises
        assert sum(p['count'] for p in province_totals) == total_exercises

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_exercises": total_exercises,
            "province_totals": province_totals,
            "breakdown": breakdown
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500





def normalize_province(name):
    """Clean up province names (case, underscores, typos)."""
    if not name:
        return "Unknown"
    cleaned = name.replace("_", " ").replace(";", "").strip().title()
    if cleaned.lower().replace(" ", "").startswith("bul"):
        return "Bulawayo"
    return cleaned

@app.route("/api/asl_total_sec_content", methods=["GET", "POST"])
@login_required
def asl_total_sec_content():
    try:
        today = date.today()
        start_date = request.form.get("start_date") or today.replace(day=1).strftime("%Y-%m-%d")
        end_date = request.form.get("end_date") or today.strftime("%Y-%m-%d")

        conn = get_ruzivo_conn()

        # 1) Get distinct students
        sql_students = """
            SELECT DISTINCT ca.student_id
            FROM tblcontent_access_hs ca
            JOIN tblstudents ts ON ts.student_id = ca.student_id
            WHERE ca.start_time BETWEEN %s AND %s
            AND ts.student_type != 'STAFF'
        """
        with conn.cursor() as cur:
            cur.execute(sql_students, (start_date, end_date))
            students = [row["student_id"] for row in cur.fetchall()]

        if not students:
            return jsonify({
                "start_date": start_date,
                "end_date": end_date,
                "total_students": 0,
                "gender_totals": {},
                "province_totals": {},
                "breakdown": []
            })

        # 2) Prepare IN clause safely
        if len(students) == 1:
            in_clause = f"({students[0]})"
        else:
            in_clause = str(tuple(students))

        sql_details = f"""
            SELECT v.student_id, v.gender, v.school_id, s.school_province
            FROM vwstudentschools v
            JOIN tblschools s ON v.school_id = s.school_id
            WHERE v.student_id IN {in_clause}
        """
        with conn.cursor() as cur:
            cur.execute(sql_details)
            details = cur.fetchall()

        # conn.close()

        gender_totals = {"Male": 0, "Female": 0, "Unknown": 0}
        province_totals = {}
        breakdown = {}

        for row in details:
            # Normalize gender
            g = row.get("gender")
            gender = "Unknown"
            if g:
                g_upper = g.strip().upper()
                if g_upper in ["M", "MALE"]:
                    gender = "Male"
                elif g_upper in ["F", "FEMALE"]:
                    gender = "Female"

            province = normalize_province(row.get("school_province"))

            gender_totals[gender] += 1
            province_totals[province] = province_totals.get(province, 0) + 1
            breakdown[(province, gender)] = breakdown.get((province, gender), 0) + 1

        total_students = len(students)

        breakdown_list = [
            {"school_province": prov, "gender": gender, "count": count}
            for (prov, gender), count in breakdown.items()
        ]

        return jsonify({
            "start_date": start_date,
            "end_date": end_date,
            "total_students": total_students,
            "gender_totals": gender_totals,
            "province_totals": province_totals,
            "breakdown": breakdown_list
        })

    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500








@app.route('/api/smartlearning-daily-logins', methods=['GET'])
@login_required
def smartlearning_daily_logins():
    conn = None

    try:
        # Set the range for the current month
        today = date.today()
        start_date = today.replace(day=1)
        # start_date = date(2025, 7, 1)
        end_date = today

        # Database query
        conn = get_ruzivo_conn()
        
        with conn.cursor() as cursor:
            query = """
                SELECT DATE(last_login) AS login_date, COUNT(*) AS login_count
                FROM tblstudents_info
                WHERE last_login BETWEEN %s AND %s
                GROUP BY DATE(last_login)
                ORDER BY login_date ASC
            """
            cursor.execute(query, (start_date, end_date))
            results = cursor.fetchall()

        # Format the results
        usage_data = [
            {"date": row['login_date'].isoformat(), "count": row['login_count']} for row in results
        ]


        return jsonify({
            "start_date": start_date.isoformat(),
            "end_date": end_date.isoformat(),
            "usage": usage_data
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()






@app.route('/api/asl-daily-unique', methods=['GET'])
def asl_daily_unique():
    conn = None
    try:
        today = date.today()
        current_year = today.year
        current_month = today.month
        days_in_month = calendar.monthrange(current_year, current_month)[1]

        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        daily_data = []

        # Loop over each day of the current month
        for day in range(1, days_in_month - 1):
            current_day = date(current_year, current_month, day)

            # Yesterday = current day (00:00:00 → 23:59:59)
            day_start = datetime.combine(current_day, datetime.min.time())
            day_end = datetime.combine(current_day, datetime.max.time())

            # "Past range" = from 1st of month → 2 days before current_day
            month_start = date(current_year, current_month, 1)
            past_end = current_day - timedelta(days=2)

            # If the past range is invalid (e.g., before month start), skip
            if past_end < month_start:
                past_end = month_start

            query = """
                SELECT COUNT(DISTINCT sl.student_id) AS unique_count
                FROM tblstudents_login sl
                WHERE sl.student_id NOT IN (
                    SELECT DISTINCT tp.student_id
                    FROM tblecocash_payment_order tp
                    WHERE tp.transactionOperationStatus = 'COMPLETED'
                )
                AND sl.student_id NOT IN (
                    SELECT DISTINCT sl2.student_id
                    FROM tblstudents_login sl2
                    WHERE sl2.login_date BETWEEN %s AND %s
                )
                AND sl.login_date BETWEEN %s AND %s
            """

            cursor.execute(query, (
                month_start, past_end,
                day_start, day_end
            ))
            result = cursor.fetchone()

            daily_data.append({
                "date": current_day.strftime("%Y-%m-%d"),
                "unique_count": int(result["unique_count"]) if result and result["unique_count"] else 0
            })

        # cursor.close()
        # conn.close()

        return jsonify({
            "year": current_year,
            "month": date(current_year, current_month, 1).strftime("%B"),
            "daily_counts": daily_data
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()



@app.route('/api/all_akello_active30', methods=['GET'])
@login_required
def all_akello_active30():
    today = datetime.today().date()
    active_30_date = today - timedelta(days=30)

    results = {
        "asl_active30": 0,
        "unique_chat_students30": 0,
        "al_active30": 0
    }

    try:
        # --- Ruzivo DB Connection ---
        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        # Query 1: ASL active logins (last 30 days)
        asl_query = """
            SELECT COUNT(DISTINCT student_id) AS asl_active30
            FROM tblstudents_login
            WHERE DATE(login_date) BETWEEN %s AND %s
        """
        cursor.execute(asl_query, (active_30_date, today))
        row = cursor.fetchone()
        results["asl_active30"] = row["asl_active30"] if row else 0

        # Query 2: Unique students in chat logs (last 30 days)
        chat_query = """
            SELECT COUNT(DISTINCT student_id) AS unique_students_count
            FROM tblask_akello_chat_logs
            WHERE created_at >= NOW() - INTERVAL 30 DAY
        """
        cursor.execute(chat_query)
        row = cursor.fetchone()
        results["unique_chat_students30"] = row["unique_students_count"] if row else 0

        cursor.close()
        conn.close()

    except Exception as e:
        print("Error querying Ruzivo DB:", e)

    try:
        # --- Library DB Connection ---
        conn = get_direct_library_conn()
        cursor = conn.cursor()

        # Query 3: Library active logins (last 30 days)
        al_query = """
            SELECT COUNT(DISTINCT user_id) AS al_active30
            FROM logins
            WHERE DATE(created_at) BETWEEN %s AND %s
        """
        cursor.execute(al_query, (active_30_date, today))
        row = cursor.fetchone()
        results["al_active30"] = row["al_active30"] if row else 0

        cursor.close()
        conn.close()

    except Exception as e:
        print("Error querying Library DB:", e)

    return jsonify(results)




from calendar import month_name, monthrange

@app.route('/api/smartlearning-monthly-logins', methods=['GET'])
@login_required
def smartlearning_monthly_logins():
    conn = None
    try:
        today = date.today()
        current_year = today.year
        current_month = today.month

        monthly_data = []
        conn = get_ruzivo_conn()

        # conn = ruzivo_pool.connection()
        with conn.cursor() as cursor:
            for month in range(1, current_month + 1):
                # Calculate start and end dates for the month
                start_date = date(current_year, month, 1)
                end_day = monthrange(current_year, month)[1]
                end_date = date(current_year, month, end_day)

                # Run the query
                query = """
                    SELECT COUNT(DISTINCT student_id) AS asl_active30
                    FROM tblstudents_login
                    WHERE DATE(login_date) BETWEEN DATE(%s) AND DATE(%s)
                """
                cursor.execute(query, (start_date, end_date))
                result = cursor.fetchone()

                monthly_data.append({
                    "month": start_date.strftime("%B"),
                    "year": current_year,
                    "student_count": result['asl_active30'] if result['asl_active30'] is not None else 0
                })

        return jsonify({
            "year": current_year,
            "monthly_logins": monthly_data
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if conn:
    #         conn.close()






@app.route('/api/library-province-institution-logins', methods=['GET'])
def institution_logins():
    try:
        # Get province name dynamically from query param
        province = request.args.get("province")
        start_date = request.args.get("start_date")
        end_date = request.args.get("end_date")

        if not province or not start_date or not end_date:
            return jsonify({"error": "Missing required parameters"}), 400

        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                i.province,
                i.name AS institution_name,
                COUNT(DISTINCT l.user_id) AS user_count
            FROM users u
            INNER JOIN logins l ON l.user_id = u.id
            INNER JOIN institution_user iu ON iu.user_id = u.id
            INNER JOIN institutions i ON i.id = iu.institution_id
            WHERE i.province = %s
              AND l.created_at BETWEEN %s AND %s
            GROUP BY i.id, i.name, i.province
            ORDER BY user_count DESC
        """
        cursor.execute(query, (province, start_date, end_date))
        results = cursor.fetchall()

        return jsonify({
            "province": province,
            "start_date": start_date,
            "end_date": end_date,
            "institutions": results
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals() and conn:
    #         conn.close()





@app.route('/api/library-province-institution-daily-logins', methods=['GET'])
def library_province_institution_daily_logins():
    try:
        province = request.args.get("province")
        day = request.args.get("day")  # Format: YYYY-MM-DD

        if not province or not day:
            return jsonify({"error": "Missing required parameters (province, day)"}), 400

        # Construct day range
        start_date = f"{day} 00:00:00"
        end_date = f"{day} 23:59:59"

        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT 
                i.province,
                i.name AS institution_name,
                COUNT(DISTINCT l.user_id) AS user_count
            FROM users u
            INNER JOIN logins l ON l.user_id = u.id
            INNER JOIN institution_user iu ON iu.user_id = u.id
            INNER JOIN institutions i ON i.id = iu.institution_id
            WHERE i.province = %s
              AND l.created_at BETWEEN %s AND %s
            GROUP BY i.id, i.name, i.province
            ORDER BY user_count DESC
        """
        cursor.execute(query, (province, start_date, end_date))
        results = cursor.fetchall()

        return jsonify({
            "province": province,
            "day": day,
            "institutions": results
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals() and conn:
    #         conn.close()










@app.route('/akello-library-metrics', methods=['GET','POST'])
@login_required
def akello_library_metrics():
    if not request.is_json:
        return jsonify({"error": "Expected JSON body"}), 400

    data = request.get_json()
    today = date.today()
    start_date = data.get('start_date') or today.replace(day=1).strftime('%Y-%m-%d')
    end_date = data.get('end_date') or today.strftime('%Y-%m-%d')

    try:
        active_30_date = (
            datetime.strptime(end_date, '%Y-%m-%d') - timedelta(days=30)
        ).strftime('%Y-%m-%d')

        queries = {
            "Total Registrations": f"""
                SELECT COUNT(DISTINCT username) AS value
                FROM users
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND deleted_at IS NULL
            """,
            "Total Revenue (by currency)": f"""
                SELECT currency, SUM(total_amount) AS value
                FROM orders
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND status='Completed'
                GROUP BY currency
            """,
            "Unique Subscribers": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM orders
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND status='Completed'
            """,
            "Unique Readers": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM read_trackers
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND duration_minutes != 0
            """,
            "Active Users": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM logins
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
            """,
            "Active in Last 30 Days": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM last_activities
                WHERE created_at BETWEEN DATE('{active_30_date}') AND DATE('{end_date}')
            """
        }

        results = {}
        conn = get_direct_library_conn()

        with conn.cursor() as cursor:
            for label, query in queries.items():
                cursor.execute(query)
                rows = cursor.fetchall()

                if "Revenue" in label:
                    results[label] = ', '.join(
                        f"{row['currency']}: {row['value']}" for row in rows
                    ) if rows else "N/A"
                else:
                    results[label] = rows[0]['value'] if rows else 0


        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()





# old asl metrics

@app.route('/api/analytics-summary')
@login_required
def get_analytics_summary():
    try:
        # Date calculations
        today = date.today()
        start_date = today.replace(day=1)
        active_30_date = today - timedelta(days=30)

        # SQL queries
        queries = {
            "asl_unique_users": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents_login
                WHERE login_date BETWEEN DATE('{start_date}') AND DATE('{today}')
            """,
            "asl_registrations": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents
                WHERE date_added BETWEEN DATE('{start_date}') AND DATE('{today}')
            """,
            "asl_total_primary_content": f"""
                SELECT COUNT(DISTINCT ca.student_id) AS value
                FROM tblcontent_access ca
                JOIN tblstudents ts ON ts.student_id = ca.student_id
                WHERE ca.start_time BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_sec_content": f"""
                SELECT COUNT(DISTINCT ca.student_id) AS value
                FROM tblcontent_access_hs ca
                JOIN tblstudents ts ON ts.student_id = ca.student_id
                WHERE ca.start_time BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_primary_exercise": f"""
                SELECT COUNT(DISTINCT tr.student_id) AS value
                FROM tblresults tr
                JOIN tblstudents ts ON ts.student_id = tr.student_id
                WHERE tr.date_added BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_sec_exercise": f"""
                SELECT COUNT(DISTINCT tr.student_id) AS value
                FROM tblresults_hs tr
                JOIN tblstudents ts ON ts.student_id = tr.student_id
                WHERE tr.date_added BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_total_zimsec_access": f"""
                SELECT COUNT(DISTINCT zi.student_id) AS value
                FROM tblcontent_access_zimsec zi
                JOIN tblstudents ts ON ts.student_id = zi.student_id
                WHERE zi.start_time BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_teacher_set_activities": f"""
                SELECT COUNT(DISTINCT ta.student_id) AS value
                FROM tblclass_activity_results ta
                JOIN tblstudents ts ON ts.student_id = ta.student_id
                WHERE ta.date_added BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND ts.student_type != 'STAFF'
            """,
            "asl_teacher_access": f"""
                SELECT COUNT(teacher_id) AS value
                FROM tbl_teacher
                WHERE active_at BETWEEN DATE('{start_date}') AND DATE('{today}')
            """,
            "asl_revenue": f"""
                SELECT currency, SUM(amount) AS value
                FROM tblecocash_payment_order
                WHERE date_created BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND transactionOperationStatus = 'COMPLETED'
                GROUP BY currency
            """,
            "asl_unique_subscribers": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblecocash_payment_order
                WHERE date_created BETWEEN DATE('{start_date}') AND DATE('{today}')
                AND transactionOperationStatus = 'COMPLETED'
            """,
            "asl_active_30": f"""
                SELECT COUNT(DISTINCT student_id) AS value
                FROM tblstudents_login
                WHERE login_date BETWEEN DATE('{active_30_date}') AND DATE('{today}')
            """
        }

        # Use pooled DB connection
        # connection = get_ruzivo_conn()
        summary = {}

        with ruzivo_conn.cursor() as cursor:
            for key, query in queries.items():
                cursor.execute(query)
                result = cursor.fetchall()

                if key == "asl_revenue":
                    summary[key] = result  # Return list with currency + value
                else:
                    summary[key] = result[0]["value"] if result else 0

        return jsonify(summary)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

    # finally:
    #     if 'connection' in locals():
    #         connection.close()  # Return to pool





@app.route('/api/library-analytics', methods=['GET'])
@login_required
def api_libraryanalytics():
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')

    if not start_date or not end_date:
        return jsonify({"error": "Missing start_date or end_date"}), 400

    try:
        active_30_date = (
            datetime.strptime(end_date, '%Y-%m-%d') - timedelta(days=30)
        ).strftime('%Y-%m-%d')

        queries = {
            "Total Registrations": f"""
                SELECT COUNT(DISTINCT username) AS value
                FROM users
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND deleted_at IS NULL
            """,
            "Total Revenue (by currency)": f"""
                SELECT currency, SUM(total_amount) AS value
                FROM orders
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND status='Completed'
                GROUP BY currency
            """,
            "Unique Subscribers": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM orders
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND status='Completed'
            """,
            "Unique Readers": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM read_trackers
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
                AND duration_minutes != 0
            """,
            "Active Users": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM last_activities
                WHERE created_at BETWEEN DATE('{start_date}') AND DATE('{end_date}')
            """,
            "Active in Last 30 Days": f"""
                SELECT COUNT(DISTINCT user_id) AS value
                FROM last_activities
                WHERE created_at BETWEEN DATE('{active_30_date}') AND DATE('{end_date}')
            """
        }


        results = {}
        # conn = get_direct_library_conn()  # From pool with persistent tunnel

        with library_conn.cursor() as cursor:
            for label, query in queries.items():
                cursor.execute(query)
                rows = cursor.fetchall()

                if "Revenue" in label:
                    results[label] = ', '.join(
                        f"{row['currency']}: {row['value']}" for row in rows
                    ) if rows else "N/A"
                else:
                    results[label] = rows[0]['value'] if rows else 0


        return jsonify(results)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # finally:
    #     if 'conn' in locals():
    #         conn.close()  # Return to pool






from flask import request, jsonify
from datetime import date, timedelta, datetime
import pandas as pd
import traceback

@app.route("/api/library-active90", methods=["GET"])
def library_active90():
    conn = None
    try:
        # Default range: last 90 days
        today = datetime.now().date()
        start_date = today - timedelta(days=90)
        end_date = today

        # Allow custom range via query params
        if request.args.get("start_date") and request.args.get("end_date"):
            start_date = datetime.strptime(request.args["start_date"], "%Y-%m-%d").date()
            end_date = datetime.strptime(request.args["end_date"], "%Y-%m-%d").date()

        # Full datetime range
        start_dt = datetime.combine(start_date, datetime.min.time())
        end_dt = datetime.combine(end_date, datetime.max.time())

        conn = get_direct_library_conn()

        # Detect correct DictCursor based on driver
        try:
            import pymysql.cursors
            cursor = conn.cursor(pymysql.cursors.DictCursor)
        except ImportError:
            import MySQLdb.cursors
            cursor = conn.cursor(MySQLdb.cursors.DictCursor)

        query = """
            SELECT COUNT(DISTINCT user_id) AS active_users
            FROM last_activities
            WHERE created_at BETWEEN %s AND %s
        """
        cursor.execute(query, (start_dt, end_dt))
        result = cursor.fetchone()
        # cursor.close()
        # conn.close()

        active_users = result["active_users"] if result and result["active_users"] else 0

        return jsonify({
            "start_date": str(start_date),
            "end_date": str(end_date),
            "library_active90": int(active_users)
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500











import traceback


# library tables 
@app.route('/check-columns')
def check_columns():
    try:
        # Get connection from SSH-tunneled pool
        connection = get_direct_library_conn()
        cursor = connection.cursor()

        # Fetch all table names
        cursor.execute("SHOW TABLES;")
        tables = cursor.fetchall()

        html_output = "<h2>Database: akello_library_2</h2>"

        for table in tables:
            table_name = list(table.values())[0]  # pymysql.DictCursor gives a dict per row
            try:
                html_output += f"<h3>Table: <code>{table_name}</code></h3><ul>"

                cursor.execute(f"DESCRIBE `{table_name}`;")
                columns = cursor.fetchall()

                for col in columns:
                    html_output += f"<li>{col['Field']} ({col['Type']})</li>"

                html_output += "</ul><hr>"
            except Exception as e:
                html_output += f"<p style='color:red;'>Error fetching columns for table <code>{table_name}</code>: {str(e)}</p><hr>"

        return html_output

    except Exception as e:
        return f"<h3>Error:</h3><p>{str(e)}</p>"

    # finally:
    #     if 'cursor' in locals():
    #         cursor.close()
    #     if 'connection' in locals():
    #         connection.close()









@app.route('/schooltracker', methods=['GET', 'POST'])
@login_required
def schooltracker():
    # Get country filter from request
    selected_country = request.args.get('country', '') or request.form.get('country', '')
    
    countries = []
    schools = []
    champion_map = {}
    
    conn = None
    cursor = None
    
    try:
        # Get database connection
        pool = get_ruzivo_pool()
        if pool is None:
            flash('Database connection not available', 'error')
            return render_template('schooltracker.html', 
                                title='School tracker',
                                countries=[],
                                schools=[],
                                selected_country='')
        
        conn = pool.connection()
        import pymysql.cursors
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        
        # Fetch distinct countries
        cursor.execute("""
            SELECT DISTINCT school_country 
            FROM tblschools 
            WHERE school_country IS NOT NULL AND school_country != ''
            ORDER BY school_country
        """)
        countries = [row['school_country'] for row in cursor.fetchall()]
        
        # If a country is selected, fetch schools for that country
        if selected_country:
            cursor.execute("""
                SELECT school_id, school_name, school_city, school_province, school_country, contact
                FROM tblschools
                WHERE school_country = %s
                ORDER BY school_name
            """, (selected_country,))
            schools = cursor.fetchall()
            
            # Load all ChampionSchool records and create mapping
            champions = ChampionSchool.query.all()
            asl_id_to_school_id_map = {}  # Map asl_school_id to school_id for reverse lookup
            for champ in champions:
                champ_schools = champ.get_schools() or []
                for school in champ_schools:
                    asl_id = school.get('asl_school_id')
                    if asl_id:
                        # Convert to string for comparison, handle both int and str
                        asl_id_str = str(asl_id).strip()
                        if asl_id_str:
                            champion_map[asl_id_str] = f"{champ.firstname} {champ.lastname}"
                            # Also create reverse mapping: if school_name matches, map asl_id to school_id
                            school_name = school.get('school_name', '').strip()
                            if school_name:
                                asl_id_to_school_id_map[school_name.lower()] = asl_id_str
            
            # Add champion name and asl_school_id to each school
            for school in schools:
                school_id_str = str(school['school_id']).strip()
                school['champion_name'] = champion_map.get(school_id_str, '')
                # Try to find asl_school_id by matching school_name
                school_name_lower = (school.get('school_name') or '').strip().lower()
                if school_name_lower in asl_id_to_school_id_map:
                    school['asl_school_id'] = asl_id_to_school_id_map[school_name_lower]
                elif school_id_str in champion_map:
                    # If school_id matches an asl_school_id in champion_map, use it
                    school['asl_school_id'] = school_id_str
                else:
                    # If no match found, assume school_id is the same as asl_school_id
                    school['asl_school_id'] = school_id_str
        
    except Exception as e:
        app.logger.error(f"Error in schooltracker route: {e}")
        flash(f'Error loading school data: {str(e)}', 'error')
    finally:
        if cursor:
            try:
                cursor.close()
            except Exception:
                pass
        if conn:
            try:
                conn.close()
            except Exception:
                pass
    
    return render_template('schooltracker.html', 
                        title='School tracker',
                        countries=countries,
                        schools=schools,
                        selected_country=selected_country)


@app.route('/remithope', methods=['GET', 'POST'])
@login_required
def remithope():
    """
    Scrape organizations from remithope.org/campaigns/ and display them
    """
    organizations = []
    error_message = None
    refresh = request.args.get('refresh', 'false').lower() == 'true'
    
    # Check if data is cached in session (unless refresh is requested)
    if not refresh and 'remithope_organizations' in session:
        organizations = session.get('remithope_organizations', [])
        if organizations:
            return render_template('remithope.html', 
                                title='Remithope Organizations',
                                organizations=organizations,
                                error_message=None)
    
    try:
        base_url = "https://remithope.org/campaigns/"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Helper functions to identify and filter text
        def is_status_text(text):
            """Check if text is likely a status indicator"""
            if not text:
                return False
            text_lower = text.lower().strip()
            status_words = ['ended', 'active', 'closed', 'ongoing', 'completed', 'cancelled', 'pending']
            return any(word in text_lower for word in status_words) and len(text_lower) < 50
        
        def is_funding_text(text):
            """Check if text is likely funding information"""
            if not text:
                return False
            import re
            # Check for currency symbols or funding patterns
            return bool(re.search(r'[\$£€¥]|(?:raised|goal|target|funded|donated)', text, re.IGNORECASE))
        
        def is_date_text(text):
            """Check if text is likely a date"""
            if not text:
                return False
            import re
            # Check for date patterns like "2 weeks ago", "3 days ago", etc.
            return bool(re.search(r'\d+\s*(?:day|week|month|year|hour|minute)s?\s*ago', text, re.IGNORECASE))
        
        def is_likely_name(text):
            """Check if text is likely an organization name"""
            if not text:
                return False
            text = text.strip()
            # Names are usually 3-200 chars
            if len(text) < 3 or len(text) > 200:
                return False
            # Only filter out if it's clearly just status/funding (short text)
            if (is_status_text(text) and len(text) < 15) or (is_funding_text(text) and '$' in text and len(text) < 25) or is_date_text(text):
                return False
            # Be more permissive - if it's not clearly status/funding/date, it could be a name
            # Check if it starts with a capital letter (common for names) or is multi-word
            if text and (text[0].isupper() or (' ' in text and len(text.split()) <= 6)):
                return True
            return False
        
        page_num = 1
        max_pages = 20  # Limit to prevent infinite loops
        all_organizations = []
        
        while page_num <= max_pages:
            # Construct URL for current page
            if page_num == 1:
                url = base_url
            else:
                url = f"{base_url}page/{page_num}/"
            
            try:
                # Make request with timeout
                response = requests.get(url, headers=headers, timeout=15, proxies={'http': None, 'https': None})
                response.raise_for_status()
                
                # Parse HTML
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Strategy 1: WordPress-specific patterns (post, entry, wp-block)
                campaign_cards = soup.find_all(['article', 'div'], class_=lambda x: x and any(keyword in x.lower() 
                                                                                              for keyword in ['post', 'entry', 'wp-block', 'type-post', 'post-type']))
                
                # Strategy 2: Find all campaign cards by class patterns
                if not campaign_cards:
                    campaign_cards = soup.find_all(['article', 'div'], class_=lambda x: x and ('campaign' in x.lower() or 'card' in x.lower()))
                
                # Strategy 3: Try finding by other common patterns
                if not campaign_cards:
                    campaign_cards = soup.find_all('div', class_=lambda x: x and any(keyword in x.lower() for keyword in ['campaign', 'organization', 'charity', 'fund', 'cause', 'project']))
                
                # Strategy 4: Find elements with links to campaign pages
                if not campaign_cards:
                    campaign_links = soup.find_all('a', href=lambda x: x and '/campaigns/' in str(x) and x != '/campaigns/' and not x.endswith('/campaigns/'))
                    if campaign_links:
                        # Get parent containers of these links
                        campaign_cards = []
                        for link in campaign_links:
                            # Try to find a parent container (article, div, li, etc.)
                            parent = link.find_parent(['article', 'div', 'li', 'section'])
                            if parent and parent not in campaign_cards:
                                campaign_cards.append(parent)
                
                # Strategy 5: Find list items that might contain campaigns (WordPress often uses lists)
                if not campaign_cards:
                    campaign_cards = soup.find_all('li', class_=lambda x: x and any(keyword in x.lower() for keyword in ['campaign', 'item', 'card', 'post', 'entry', 'type-post']))
                
                # Strategy 6: Find all articles or divs with specific structure (has heading and link)
                if not campaign_cards:
                    all_elements = soup.find_all(['article', 'div', 'section'])
                    campaign_cards = []
                    for elem in all_elements:
                        # Check if it has a heading and a link (likely a campaign card)
                        has_heading = elem.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
                        has_link = elem.find('a', href=True)
                        has_text = len(elem.get_text(strip=True)) > 30
                        if has_heading and has_link and has_text:
                            # Avoid duplicates and very large containers
                            if len(elem.get_text(strip=True)) < 5000:  # Reasonable size limit
                                campaign_cards.append(elem)
                
                # Strategy 7: Last resort - find any div/article with heading and substantial text
                if not campaign_cards:
                    all_elements = soup.find_all(['article', 'div'])
                    campaign_cards = [card for card in all_elements 
                                    if card.find(['h2', 'h3', 'h4']) 
                                    and len(card.get_text(strip=True)) > 50 
                                    and len(card.get_text(strip=True)) < 2000]
                
                # Log for debugging
                if page_num == 1:
                    app.logger.info(f"Found {len(campaign_cards)} potential campaign cards on page {page_num}")
                    if not campaign_cards:
                        # Log a sample of the HTML structure for debugging
                        app.logger.debug(f"Page HTML sample: {str(soup)[:500]}")
                
                if not campaign_cards:
                    # If no cards found on this page, break
                    if page_num == 1:
                        error_message = "No organizations found. The website structure may have changed. Please check the website or contact support."
                    break
                
                # Extract data from each card
                page_organizations = []
                for card in campaign_cards:
                    org_data = {}
                    
                    # Debug: Log card HTML structure (first card only)
                    if page_num == 1 and len(page_organizations) == 0:
                        app.logger.info(f"Card HTML sample: {str(card)[:500]}")
                        app.logger.info(f"Card text sample: {card.get_text(separator=' | ', strip=True)[:300]}")
                    
                    # Extract organization name/title - try multiple strategies
                    # IMPORTANT: Extract name BEFORE status/funding to avoid confusion
                    # WordPress-specific patterns first
                    name = None
                    
                    # Strategy 1: WordPress-specific classes (entry-title, post-title, etc.)
                    wordpress_title_classes = ['entry-title', 'post-title', 'wp-block-post-title', 'entry-header', 
                                              'post-header', 'title', 'campaign-title', 'organization-title']
                    for class_name in wordpress_title_classes:
                        title_elem = card.find(class_=lambda x: x and class_name in x.lower())
                        if title_elem:
                            candidate = title_elem.get_text(strip=True)
                            if candidate and len(candidate) > 0:
                                if not (is_status_text(candidate) and len(candidate) < 20) and not (is_funding_text(candidate) and '$' in candidate):
                                    name = candidate
                                    break
                    
                    # Strategy 2: Find heading tags (h1-h6) - prioritize h2, h3, h4
                    if not name:
                        for tag in ['h2', 'h3', 'h4', 'h1', 'h5', 'h6']:
                            title_elem = card.find(tag)
                            if title_elem:
                                candidate = title_elem.get_text(strip=True)
                                # Filter out status and funding text, but be less strict
                                if candidate and len(candidate) > 0:
                                    # Only filter if it's clearly status/funding (not if it just contains those words)
                                    if not (is_status_text(candidate) and len(candidate) < 20) and not (is_funding_text(candidate) and '$' in candidate):
                                        name = candidate
                                        break
                    
                    # Strategy 3: WordPress entry-content or post-content title (often first heading inside)
                    if not name:
                        entry_content = card.find(class_=lambda x: x and any(keyword in x.lower() 
                                                                          for keyword in ['entry-content', 'post-content', 'wp-block-post-content', 'content']))
                        if entry_content:
                            # Find first heading in entry content
                            for tag in ['h1', 'h2', 'h3', 'h4']:
                                title_elem = entry_content.find(tag)
                                if title_elem:
                                    candidate = title_elem.get_text(strip=True)
                                    if candidate and len(candidate) > 0:
                                        if not (is_status_text(candidate) and len(candidate) < 20) and not (is_funding_text(candidate) and '$' in candidate):
                                            name = candidate
                                            break
                            if name:
                                break
                    
                    # Strategy 4: Find by class containing 'title' or 'name'
                    if not name:
                        title_elem = card.find(class_=lambda x: x and any(keyword in x.lower() 
                                                                          for keyword in ['title', 'name', 'heading', 'headline', 'campaign-title']))
                        if title_elem:
                            candidate = title_elem.get_text(strip=True)
                            if candidate and not (is_status_text(candidate) and len(candidate) < 20) and not (is_funding_text(candidate) and '$' in candidate):
                                name = candidate
                    
                    # Strategy 5: Find title in a link (common pattern for campaign cards, especially WordPress)
                    if not name:
                        links = card.find_all('a', href=True)
                        for link_elem in links:
                            link_text = link_elem.get_text(strip=True)
                            # Filter out common link texts and status/funding
                            if link_text and len(link_text) > 3 and len(link_text) < 200:
                                skip_texts = ['read more', 'view', 'click', 'learn more', 'donate', 'share', 'visit', 
                                            'see more', 'view campaign', 'donate now', 'share this', 'home', 'about']
                                if (not any(skip in link_text.lower() for skip in skip_texts) and
                                    not (is_status_text(link_text) and len(link_text) < 20) and 
                                    not (is_funding_text(link_text) and '$' in link_text)):
                                    # Check if link goes to a campaign page - more likely to be the title
                                    href = link_elem.get('href', '')
                                    if '/campaigns/' in str(href) or href.startswith('/'):
                                        name = link_text
                                        break
                    
                    # Strategy 6: Get first substantial text line (excluding navigation/common text, status, funding)
                    if not name:
                        all_text = card.get_text(separator='\n', strip=True)
                        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
                        skip_patterns = ['http', 'www', 'read more', 'view', 'click', 'learn more', 'donate', 'share', 
                                       'campaigns', 'home', 'about', 'contact', 'login', 'sign up', 'menu', 'navigation',
                                       'remithope']
                        for line in lines:
                            # Be less strict - just check it's not clearly status/funding/date
                            if (len(line) > 3 and len(line) < 200 and
                                not any(pattern in line.lower() for pattern in skip_patterns) and
                                not line.lower().startswith(('http', 'www', 'mailto:')) and
                                not (is_status_text(line) and len(line) < 20) and  # Only filter short status text
                                not (is_funding_text(line) and '$' in line and len(line) < 30) and  # Only filter short funding text
                                not is_date_text(line)):  # Filter date text
                                # Check if it looks like a name (starts with capital or is multi-word)
                                if line[0].isupper() or (' ' in line and len(line.split()) <= 5):
                                    name = line
                                    break
                    
                    # Strategy 7: Get text from first strong/bold element
                    if not name:
                        strong_elem = card.find(['strong', 'b'])
                        if strong_elem:
                            strong_text = strong_elem.get_text(strip=True)
                            if strong_text and is_likely_name(strong_text):
                                name = strong_text
                    
                    # Strategy 8: Get all text elements in order, find first that looks like a name
                    if not name:
                        # Get all text nodes in order
                        for elem in card.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'a', 'strong', 'b']):
                            text = elem.get_text(strip=True)
                            if text and len(text) > 3 and len(text) < 200:
                                # Less strict filtering
                                if not (is_status_text(text) and len(text) < 20) and not (is_funding_text(text) and '$' in text):
                                    if text[0].isupper() or (' ' in text and len(text.split()) <= 5):
                                        name = text
                                        break
                    
                    # Strategy 9: Last resort - get first non-empty line that's not status/funding
                    if not name:
                        all_text = card.get_text(separator='\n', strip=True)
                        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
                        for line in lines:
                            if (len(line) > 3 and len(line) < 200 and
                                not line.lower().startswith(('http', 'www', 'mailto:')) and
                                not (is_status_text(line) and len(line) < 15) and
                                not (is_funding_text(line) and '$' in line and len(line) < 25) and
                                not is_date_text(line)):
                                name = line
                                break
                    
                    # Strategy 10: Ultra fallback - get first word/phrase from card text
                    if not name:
                        all_text = card.get_text(separator=' ', strip=True)
                        # Split into words and try to find a name-like phrase
                        words = all_text.split()
                        # Try first 2-5 words as potential name
                        for word_count in [5, 4, 3, 2]:
                            if len(words) >= word_count:
                                candidate = ' '.join(words[:word_count])
                                if (len(candidate) > 3 and len(candidate) < 200 and
                                    not candidate.lower().startswith(('http', 'www', '$')) and
                                    not (is_status_text(candidate) and len(candidate) < 15) and
                                    not (is_funding_text(candidate) and '$' in candidate and len(candidate) < 25) and
                                    not is_date_text(candidate)):
                                    name = candidate
                                    break
                        if not name and words:
                            # Last resort: just use first word if it's capitalized
                            if words[0] and words[0][0].isupper() and len(words[0]) > 2:
                                name = words[0]
                    
                    # Only proceed if we found a name
                    if name and len(name.strip()) > 0:
                        org_data['name'] = name.strip()
                        # Debug log
                        if page_num == 1 and len(page_organizations) == 0:
                            app.logger.info(f"Extracted name: {name}")
                    else:
                        # Debug: Log why name wasn't found
                        if page_num == 1 and len(page_organizations) == 0:
                            app.logger.warning(f"No name found for card. Card text: {card.get_text(separator=' | ', strip=True)[:200]}")
                        # Don't skip - try to use a default or continue with empty name
                        # But for now, skip if no name
                        continue
                    
                    # Extract description - try multiple strategies
                    # IMPORTANT: Extract description AFTER name but BEFORE status/funding
                    # WordPress-specific patterns first
                    description = None
                    
                    # Strategy 1: WordPress entry-content or post-content (main content area)
                    wordpress_content_classes = ['entry-content', 'post-content', 'wp-block-post-content', 
                                                'entry-summary', 'post-excerpt', 'excerpt', 'content']
                    for class_name in wordpress_content_classes:
                        content_elem = card.find(class_=lambda x: x and class_name in x.lower())
                        if content_elem:
                            # Get first paragraph or substantial text from content area
                            para = content_elem.find('p')
                            if para:
                                desc_text = para.get_text(strip=True)
                            else:
                                desc_text = content_elem.get_text(separator=' ', strip=True)
                            
                            if desc_text and len(desc_text) > 10:
                                # Remove name if it appears
                                if name and name.lower() in desc_text.lower():
                                    desc_text = desc_text.replace(name, '', 1).strip()
                                # Filter but be lenient
                                if (len(desc_text) > 10 and
                                    not (is_status_text(desc_text) and len(desc_text) < 20) and 
                                    not (is_funding_text(desc_text) and '$' in desc_text and len(desc_text) < 30) and
                                    not is_date_text(desc_text)):
                                    description = desc_text
                                    break
                    
                    # Strategy 2: Find by class containing description keywords
                    if not description:
                        desc_elem = card.find(['p', 'div', 'span'], class_=lambda x: x and any(keyword in x.lower() 
                                                                                          for keyword in ['description', 'desc', 'summary', 'excerpt', 'content', 'text', 'body', 'detail']))
                        if desc_elem:
                            desc_text = desc_elem.get_text(strip=True)
                            # Filter out status, funding, and name - be less strict
                            if (desc_text and len(desc_text) > 10 and 
                                desc_text != name and 
                                not (is_status_text(desc_text) and len(desc_text) < 20) and 
                                not (is_funding_text(desc_text) and '$' in desc_text and len(desc_text) < 30) and
                                not is_date_text(desc_text)):
                                description = desc_text
                    
                    # Strategy 2: Find first paragraph with substantial text (but not the title, status, or funding)
                    if not description:
                        paragraphs = card.find_all(['p', 'div', 'span'])
                        for para in paragraphs:
                            para_text = para.get_text(strip=True)
                            # Skip if it's the same as the name, status, funding, or too short - be less strict
                            if (para_text and len(para_text) > 15 and 
                                para_text != name and 
                                not (is_status_text(para_text) and len(para_text) < 20) and
                                not (is_funding_text(para_text) and '$' in para_text and len(para_text) < 30) and
                                not is_date_text(para_text) and
                                not para_text.lower().startswith(('read more', 'view', 'click', 'donate', 'share', 'learn more')) and
                                not para_text.lower() in ['home', 'about', 'contact', 'campaigns']):
                                # Check if it's not a navigation or button text
                                if len(para_text) < 1000:  # Reasonable description length
                                    # Make sure it's not just the name repeated (but allow if description is longer)
                                    if name.lower() not in para_text.lower() or len(para_text) > len(name) * 1.3:
                                        description = para_text
                                        break
                    
                    # Strategy 3: Extract text after the name - get all text and find content after name
                    if not description:
                        # Get all text, find the name position, get text after it
                        all_text = card.get_text(separator='\n', strip=True)
                        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
                        found_name = False
                        for i, line in enumerate(lines):
                            # Check if this line contains the name
                            if name.lower() in line.lower() or line.lower() in name.lower():
                                found_name = True
                                # Look at next few lines for description (before status/funding)
                                for j in range(i + 1, min(i + 10, len(lines))):
                                    next_line = lines[j]
                                    # Skip status, funding, dates, and action buttons
                                    if (len(next_line) > 20 and 
                                        not is_status_text(next_line) and
                                        not is_funding_text(next_line) and
                                        not is_date_text(next_line) and
                                        not any(skip in next_line.lower() for skip in ['read more', 'view', 'click', 'donate', 'share', 'learn more', 'home', 'about']) and
                                        not next_line.lower().startswith(('http', 'www', '$'))):
                                        # Make sure it's not the name again
                                        if name.lower() not in next_line.lower() or len(next_line) > len(name) * 1.5:
                                            description = next_line
                                            break
                                if description:
                                    break
                    
                    # Strategy 4: Get all text content, remove name, status, funding, and action buttons
                    if not description:
                        all_text = card.get_text(separator=' ', strip=True)
                        # Remove the name from the text
                        text_without_name = all_text.replace(name, '', 1).strip()
                        
                        # Remove status text (but be less aggressive)
                        import re
                        text_without_name = re.sub(r'\b(?:ended|active|closed|ongoing|completed|cancelled|pending)\b', '', text_without_name, flags=re.IGNORECASE)
                        
                        # Remove funding patterns (but keep text that mentions funding in context)
                        text_without_name = re.sub(r'\$[\d,]+\.?\d*\s*(?:of|raised|goal)\s*\$?[\d,]+\.?\d*', '', text_without_name, flags=re.IGNORECASE)
                        text_without_name = re.sub(r'(?:raised|goal|target|funded|donated):?\s*\$?[\d,]+\.?\d*', '', text_without_name, flags=re.IGNORECASE)
                        
                        # Remove date patterns
                        text_without_name = re.sub(r'\d+\s*(?:day|week|month|year|hour|minute)s?\s*ago', '', text_without_name, flags=re.IGNORECASE)
                        
                        # Remove common action phrases
                        for phrase in ['read more', 'view campaign', 'donate now', 'share this', 'learn more']:
                            text_without_name = text_without_name.replace(phrase, '')
                        
                        text_without_name = ' '.join(text_without_name.split())
                        
                        # Take first reasonable chunk that looks like a description
                        if len(text_without_name) > 15:
                            # Get first sentence or first 200 chars
                            sentences = text_without_name.split('.')
                            for sentence in sentences:
                                sentence = sentence.strip()
                                if (len(sentence) > 15 and 
                                    not (is_status_text(sentence) and len(sentence) < 20) and
                                    not (is_funding_text(sentence) and '$' in sentence and len(sentence) < 30) and
                                    not is_date_text(sentence)):
                                    description = sentence
                                    break
                            if not description:
                                # Fallback to first 200 chars (even if it contains some funding info)
                                candidate = text_without_name[:200].strip()
                                if len(candidate) > 15:
                                    description = candidate
                    
                    # Strategy 5: Ultra fallback - get text after name in the card
                    if not description:
                        all_text = card.get_text(separator='\n', strip=True)
                        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
                        # Find line with name, get next substantial line
                        for i, line in enumerate(lines):
                            if name.lower() in line.lower():
                                # Look at next few lines
                                for j in range(i + 1, min(i + 5, len(lines))):
                                    next_line = lines[j]
                                    if (len(next_line) > 15 and
                                        not (is_status_text(next_line) and len(next_line) < 20) and
                                        not (is_funding_text(next_line) and '$' in next_line and len(next_line) < 30) and
                                        not is_date_text(next_line) and
                                        name.lower() not in next_line.lower()):
                                        description = next_line
                                        break
                                if description:
                                    break
                    
                    if description:
                        # Clean up description - limit length and remove extra whitespace
                        description = ' '.join(description.split())
                        # Only filter if description is clearly just status/funding (short text)
                        if (is_status_text(description) and len(description) < 20) or (is_funding_text(description) and '$' in description and len(description) < 30):
                            description = None
                        elif len(description) > 500:
                            description = description[:500] + '...'
                        
                    if description:
                        org_data['description'] = description
                        # Debug log
                        if page_num == 1 and len(page_organizations) == 0:
                            app.logger.info(f"Extracted description: {description[:100]}")
                    else:
                        org_data['description'] = None
                        # Debug log
                        if page_num == 1 and len(page_organizations) == 0:
                            app.logger.warning(f"No description found. Card text: {card.get_text(separator=' | ', strip=True)[:300]}")
                    
                    # Extract categories/tags
                    tags = []
                    tag_elems = card.find_all(['span', 'a', 'div'], class_=lambda x: x and ('tag' in x.lower() or 'category' in x.lower() or 'label' in x.lower()))
                    for tag_elem in tag_elems:
                        tag_text = tag_elem.get_text(strip=True)
                        if tag_text and len(tag_text) < 50:  # Reasonable tag length
                            tags.append(tag_text)
                    org_data['categories'] = tags if tags else []
                    
                    # Extract funding information
                    funding_text = card.get_text()
                    # Look for patterns like "$X of $Y" or "raised $X" or "goal $Y"
                    import re
                    funding_match = re.search(r'\$[\d,]+\.?\d*\s*(?:of|raised|goal)\s*\$[\d,]+\.?\d*', funding_text, re.IGNORECASE)
                    if funding_match:
                        org_data['funding_info'] = funding_match.group(0)
                    else:
                        # Try separate patterns
                        raised_match = re.search(r'(?:raised|raised:?)\s*\$?[\d,]+\.?\d*', funding_text, re.IGNORECASE)
                        goal_match = re.search(r'(?:goal|target):?\s*\$?[\d,]+\.?\d*', funding_text, re.IGNORECASE)
                        if raised_match or goal_match:
                            org_data['funding_info'] = f"{raised_match.group(0) if raised_match else ''} {goal_match.group(0) if goal_match else ''}".strip()
                        else:
                            org_data['funding_info'] = None
                    
                    # Extract status
                    status_elem = card.find(class_=lambda x: x and ('status' in x.lower() or 'ended' in x.lower() or 'active' in x.lower()))
                    if status_elem:
                        org_data['status'] = status_elem.get_text(strip=True)
                    else:
                        # Look for status in text
                        status_match = re.search(r'(?:ended|active|closed|ongoing).*?(?:\d+\s*(?:day|week|month|year)s?\s*ago)?', funding_text, re.IGNORECASE)
                        if status_match:
                            org_data['status'] = status_match.group(0)
                        else:
                            org_data['status'] = None
                    
                    # Extract URL/link - prioritize links to campaign pages
                    link_elem = None
                    # First try to find a link that goes to a campaign page
                    campaign_links = card.find_all('a', href=lambda x: x and '/campaigns/' in str(x))
                    if campaign_links:
                        link_elem = campaign_links[0]
                    else:
                        # Fallback to any link in the card
                        link_elem = card.find('a', href=True)
                    
                    if link_elem:
                        href = link_elem.get('href', '')
                        if href.startswith('http'):
                            org_data['url'] = href
                        elif href.startswith('/'):
                            org_data['url'] = f"https://remithope.org{href}"
                        else:
                            org_data['url'] = f"https://remithope.org/campaigns/{href}"
                    else:
                        org_data['url'] = None
                    
                    # Only add if we have at least a name
                    if org_data.get('name'):
                        page_organizations.append(org_data)
                
                if not page_organizations:
                    # No organizations found on this page, stop pagination
                    break
                
                all_organizations.extend(page_organizations)
                
                # Check for next page
                next_link = soup.find('a', class_=lambda x: x and ('next' in x.lower() if x else False))
                if not next_link:
                    # Try finding pagination with "Next" text
                    next_link = soup.find('a', string=lambda x: x and 'next' in x.lower() if x else False)
                
                if not next_link:
                    # Check if there's a page number indicator showing more pages
                    pagination = soup.find(class_=lambda x: x and ('pagination' in x.lower() if x else False))
                    if pagination:
                        page_text = pagination.get_text()
                        # Look for "Page X of Y" pattern
                        page_match = re.search(r'page\s+(\d+)\s+of\s+(\d+)', page_text, re.IGNORECASE)
                        if page_match:
                            current_page = int(page_match.group(1))
                            total_pages = int(page_match.group(2))
                            if current_page < total_pages:
                                page_num += 1
                                continue
                
                # If no next link found, break
                if not next_link:
                    break
                
                page_num += 1
                time.sleep(1)  # Be respectful, add delay between requests
                
            except requests.exceptions.RequestException as e:
                app.logger.error(f"Error fetching page {page_num}: {e}")
                if page_num == 1:
                    error_message = f"Error fetching data: {str(e)}"
                break
            except Exception as e:
                app.logger.error(f"Error parsing page {page_num}: {e}")
                if page_num == 1:
                    error_message = f"Error parsing data: {str(e)}"
                break
        
        organizations = all_organizations
        
        # Cache in session
        if organizations:
            session['remithope_organizations'] = organizations
        
    except Exception as e:
        app.logger.error(f"Error in remithope route: {e}")
        error_message = f"An error occurred while scraping: {str(e)}"
        organizations = []
    
    return render_template('remithope.html', 
                        title='Remithope Organizations',
                        organizations=organizations,
                        error_message=error_message)


@app.route('/bookallocations', methods=['GET', 'POST'])
@login_required
def bookallocations():
    allocate_form = BookAllocationForm()
    allocations = BookAllocations.query.order_by(BookAllocations.timestamp.desc()).all()


    province_summary = Counter([a.school_province for a in allocations])
    # Province Summary
    province_counter = Counter([a.school_province for a in allocations])
    province_labels = list(province_counter.keys())
    province_counts = list(province_counter.values())

    # Book Allocation Chart
    book_counter = defaultdict(int)
    for a in allocations:
        for b in a.books_allocated.split(','):
            book_counter[b.strip()] += 1

    book_labels = list(book_counter.keys())
    book_counts = list(book_counter.values())
    return render_template('book_allocations.html',allocations=allocations, province_summary=province_summary,book_labels=book_labels,
                           book_counts=book_counts,allocate_form=allocate_form,province_labels=province_labels,province_counts=province_counts, title='Book Allocations')


@app.route('/allocate_books', methods=['GET', 'POST'])
@login_required
def allocate_books():
    allocate_form = BookAllocationForm()
    
    if allocate_form.validate_on_submit():
        try:
            book_allocation = BookAllocations(
                school_name=allocate_form.school_name.data,
                school_province=allocate_form.school_province.data,
                books_allocated=','.join(allocate_form.books_allocated.data),
                allocated_by=current_user.username  # Ensure user is logged in
            )
            db.session.add(book_allocation)
            db.session.commit()
            flash('Book allocation recorded successfully!', 'success')
            return redirect(url_for('bookallocations'))  # Ensure 'metrics' route exists
        except Exception as e:
            flash(f"An error occurred: {e}", "danger")

    # Debugging: Print form errors if validation fails
    if request.method == 'POST':
        print("Form Errors:", allocate_form.errors)


# ===== Institution Analytics API Endpoints =====

@app.route('/api/institution-countries', methods=['GET'])
@login_required
def api_institution_countries():
    """Get distinct countries from institutions table"""
    try:
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT DISTINCT country 
                FROM institutions 
                WHERE country IS NOT NULL AND country != ''
                ORDER BY country ASC
            """
            cursor.execute(query)
            results = cursor.fetchall()
            
        countries = [row['country'] for row in results if row.get('country')]
        return jsonify(countries)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching countries: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/institution-provinces', methods=['GET'])
@login_required
def api_institution_provinces():
    """Get distinct provinces for a selected country"""
    try:
        country = request.args.get('country')
        if not country:
            return jsonify({"error": "Country parameter is required"}), 400
            
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT DISTINCT province 
                FROM institutions 
                WHERE country = %s AND province IS NOT NULL AND province != ''
                ORDER BY province ASC
            """
            cursor.execute(query, (country,))
            results = cursor.fetchall()
            
        provinces = [row['province'] for row in results if row.get('province')]
        return jsonify(provinces)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching provinces: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/institution-analytics', methods=['GET'])
@login_required
def api_institution_analytics():
    """Get institutions with user counts and subject groups filtered by country and province"""
    try:
        country = request.args.get('country')
        province = request.args.get('province')
        search = request.args.get('search', '').strip()
        page = int(request.args.get('page', 1))
        per_page = min(int(request.args.get('per_page', 50)), 100)  # Max 100 per page
        
        if not country or not province:
            return jsonify({"error": "Both country and province parameters are required"}), 400
        
        if page < 1:
            page = 1
        offset = (page - 1) * per_page
            
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # Build WHERE clause with optional search
            where_clause = "WHERE i.country = %s AND i.province = %s"
            params = [country, province]
            
            if search:
                where_clause += " AND i.name LIKE %s"
                params.append(f"%{search}%")
            
            # Get total count for pagination
            count_query = f"""
                SELECT COUNT(DISTINCT i.id) AS total
                FROM institutions i
                JOIN institution_user iu ON i.id = iu.institution_id
                JOIN users u ON iu.user_id = u.id
                {where_clause}
            """
            cursor.execute(count_query, params)
            total_count = cursor.fetchone()['total']
            total_pages = (total_count + per_page - 1) // per_page  # Ceiling division
            
            # Main query to get institutions with user counts (with pagination)
            query = f"""
                SELECT DISTINCT 
                    i.id AS institution_id,
                    i.name AS institution_name,
                    i.level AS institution_level,
                    COUNT(DISTINCT u.id) AS user_count
                FROM institutions i
                JOIN institution_user iu ON i.id = iu.institution_id
                JOIN users u ON iu.user_id = u.id
                {where_clause}
                GROUP BY i.id, i.name, i.level
                ORDER BY i.name ASC
                LIMIT %s OFFSET %s
            """
            params_with_limit = params + [per_page, offset]
            cursor.execute(query, params_with_limit)
            institutions = cursor.fetchall()
            
            if not institutions:
                return jsonify({
                    'results': [],
                    'pagination': {
                        'total_count': total_count,
                        'total_pages': total_pages,
                        'current_page': page,
                        'per_page': per_page
                    }
                })
            
            # Get all institution IDs
            institution_ids = [inst['institution_id'] for inst in institutions]
            placeholders = ','.join(['%s'] * len(institution_ids))
            
            # OPTIMIZED: Get all subject groups for all institutions in one query
            # This replaces the N+1 query problem
            all_sg_query = f"""
                SELECT 
                    iu.institution_id,
                    sg.id AS sg_id,
                    sg.name AS sg_name,
                    COUNT(DISTINCT u.id) AS user_count
                FROM institution_user iu
                JOIN users u ON iu.user_id = u.id
                JOIN subject_group_user sgu ON u.id = sgu.user_id
                JOIN subject_groups sg ON sgu.subject_group_id = sg.id
                WHERE iu.institution_id IN ({placeholders})
                GROUP BY iu.institution_id, sg.id, sg.name
                ORDER BY iu.institution_id, sg.name ASC
            """
            cursor.execute(all_sg_query, institution_ids)
            all_sg_results = cursor.fetchall()
            
            # OPTIMIZED: Get subject groups with books for all institutions in one query
            sg_books_query = f"""
                SELECT 
                    iu.institution_id,
                    sg.id AS sg_id,
                    sg.name AS sg_name,
                    COUNT(DISTINCT u.id) AS user_count
                FROM institution_user iu
                JOIN users u ON iu.user_id = u.id
                JOIN subject_group_user sgu ON u.id = sgu.user_id
                JOIN subject_groups sg ON sgu.subject_group_id = sg.id
                JOIN book_subject_group bsg ON sg.id = bsg.subject_group_id
                WHERE iu.institution_id IN ({placeholders})
                GROUP BY iu.institution_id, sg.id, sg.name
                ORDER BY iu.institution_id, sg.name ASC
            """
            cursor.execute(sg_books_query, institution_ids)
            sg_books_results = cursor.fetchall()
            
            # Group subject groups by institution_id in Python
            all_sg_by_inst = {}
            for row in all_sg_results:
                inst_id = row['institution_id']
                if inst_id not in all_sg_by_inst:
                    all_sg_by_inst[inst_id] = []
                all_sg_by_inst[inst_id].append({
                    'name': row['sg_name'],
                    'user_count': row['user_count']
                })
            
            sg_books_by_inst = {}
            for row in sg_books_results:
                inst_id = row['institution_id']
                if inst_id not in sg_books_by_inst:
                    sg_books_by_inst[inst_id] = []
                sg_books_by_inst[inst_id].append({
                    'name': row['sg_name'],
                    'user_count': row['user_count']
                })
            
            # Build results
            results = []
            for inst in institutions:
                inst_id = inst['institution_id']
                results.append({
                    'institution_id': inst_id,
                    'institution_name': inst['institution_name'],
                    'institution_level': inst['institution_level'] or 'N/A',
                    'user_count': inst['user_count'],
                    'all_subject_groups': all_sg_by_inst.get(inst_id, []),
                    'subject_groups_with_books': sg_books_by_inst.get(inst_id, [])
                })
            
        return jsonify({
            'results': results,
            'pagination': {
                'total_count': total_count,
                'total_pages': total_pages,
                'current_page': page,
                'per_page': per_page
            }
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching institution analytics: {e}")
        return jsonify({"error": str(e)}), 500


# ===== Quotation: resolve publisher_id to publisher name (from publishers table) =====

@app.route('/api/quotation/publisher-names', methods=['GET'])
@login_required
def quotation_publisher_names():
    """Resolve publisher IDs to names from the publishers table (for quotation uploads)."""
    ids_param = request.args.get('ids', '')
    if not ids_param:
        return jsonify({}), 200
    try:
        ids = [x.strip() for x in ids_param.split(',') if x.strip()]
        if not ids:
            return jsonify({}), 200
        conn = get_direct_library_conn()
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            placeholders = ','.join(['%s'] * len(ids))
            cursor.execute(
                "SELECT id, name FROM publishers WHERE id IN ({})".format(placeholders),
                ids
            )
            rows = cursor.fetchall()
        result = {}
        for r in rows:
            key = r['id']
            result[str(key)] = (r['name'] or '').strip()
        return jsonify(result), 200
    except Exception as e:
        app.logger.error(f"Error resolving publisher names for quotation: {e}")
        return jsonify({"error": str(e)}), 500


def _quotation_price_from_db(val):
    """Return price as-is from DB (varchar or number). No conversion; string preserved for exact display."""
    if val is None or val == '':
        return '0', 0.0
    if isinstance(val, str):
        s = val.strip().replace(',', '.')
        try:
            n = float(s)
            return s, n
        except ValueError:
            return '0', 0.0
    try:
        n = float(val)
        return str(n), n
    except (TypeError, ValueError):
        return '0', 0.0


@app.route('/api/quotation/lookup-books', methods=['POST'])
@login_required
def quotation_lookup_books():
    """Look up books and prices from the library DB for quotation. Expects JSON: { items: [ { title?, book_id?, publisher_id? }, ... ] }.
    Price column is VARCHAR; returned as string so 0.90 stays 0.90 (no conversion)."""
    try:
        data = request.get_json() or {}
        items = data.get('items') or []
        if not items:
            return jsonify({"error": "No items provided", "rows": []}), 400
        conn = get_direct_library_conn()
        rows = []
        not_found_count = 0
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            for item in items:
                book_id = item.get('book_id') or item.get('id')
                title = (item.get('title') or '').strip() if isinstance(item.get('title'), str) else ''
                publisher_id = item.get('publisher_id')
                if publisher_id is not None and publisher_id != '':
                    try:
                        publisher_id = int(publisher_id)
                    except (TypeError, ValueError):
                        publisher_id = None
                row_from_file = {
                    'name': (item.get('name') or '').strip() if isinstance(item.get('name'), str) else '',
                    'title': title or (item.get('title') or ''),
                }
                if book_id is not None and book_id != '':
                    try:
                        bid = int(book_id)
                    except (TypeError, ValueError):
                        bid = None
                    if bid is not None:
                        cursor.execute(
                            "SELECT b.id, b.title, b.price, p.name AS publisher_name FROM books b "
                            "LEFT JOIN publishers p ON b.publisher_id = p.id WHERE b.id = %s AND (COALESCE(b.is_active, 1) = 1) ORDER BY b.id LIMIT 1",
                            (bid,)
                        )
                        r = cursor.fetchone()
                        if r:
                            price_str, price_num = _quotation_price_from_db(r.get('price'))
                            rows.append({
                                'name': (r.get('publisher_name') or '').strip(),
                                'title': (r.get('title') or '').strip(),
                                'price': price_str,
                                'currency_code': 'USD',
                                'zwl_price': 0,
                                'amount': price_str,
                                'zwl_amount': 0,
                                'not_found': False,
                            })
                            continue
                if title:
                    if publisher_id is not None:
                        cursor.execute(
                            "SELECT b.id, b.title, b.price, p.name AS publisher_name FROM books b "
                            "LEFT JOIN publishers p ON b.publisher_id = p.id "
                            "WHERE TRIM(b.title) = %s AND b.publisher_id = %s AND (COALESCE(b.is_active, 1) = 1) ORDER BY b.id LIMIT 1",
                            (title, publisher_id)
                        )
                    else:
                        cursor.execute(
                            "SELECT b.id, b.title, b.price, p.name AS publisher_name FROM books b "
                            "LEFT JOIN publishers p ON b.publisher_id = p.id "
                            "WHERE TRIM(b.title) = %s AND (COALESCE(b.is_active, 1) = 1) ORDER BY b.id LIMIT 1",
                            (title,)
                        )
                    r = cursor.fetchone()
                    if r:
                        price_str, price_num = _quotation_price_from_db(r.get('price'))
                        rows.append({
                            'name': (r.get('publisher_name') or '').strip(),
                            'title': (r.get('title') or '').strip(),
                            'price': price_str,
                            'currency_code': 'USD',
                            'zwl_price': 0,
                            'amount': price_str,
                            'zwl_amount': 0,
                            'not_found': False,
                        })
                        continue
                not_found_count += 1
                rows.append({
                    'name': row_from_file.get('name') or '—',
                    'title': row_from_file.get('title') or '—',
                    'price': '0',
                    'currency_code': 'USD',
                    'zwl_price': 0,
                    'amount': '0',
                    'zwl_amount': 0,
                    'not_found': True,
                })
        # Total: parse each row price (string) to number for summing only
        total_price = 0.0
        for r in rows:
            _, n = _quotation_price_from_db(r.get('price', '0'))
            total_price += n
        total_zwl = sum(r.get('zwl_price', 0) for r in rows)
        return jsonify({
            'rows': rows,
            'totalPrice': total_price,
            'totalZwl': total_zwl,
            'hasZwl': total_zwl != 0,
            'notFoundCount': not_found_count,
        }), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error in quotation lookup-books: {e}")
        return jsonify({"error": str(e)}), 500


# ===== Book Allocation Request API Endpoints =====

@app.route('/api/book-genres', methods=['GET'])
@login_required
def get_book_genres():
    """Get books and their associated genres from the library database"""
    try:
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            query = """
                SELECT b.title, g.name AS genre_name, COALESCE(b.is_active, 1) AS is_active,
                       p.name AS publisher_name
                FROM books b
                LEFT JOIN publishers p ON b.publisher_id = p.id
                JOIN book_genre bg ON b.id = bg.book_id
                JOIN genres g ON bg.genre_id = g.id
                WHERE g.name IN ('HBC PlusOne Primary Teacher`s Guides', 'HBC PlusOne Primary')
            """
            cursor.execute(query)
            results = cursor.fetchall()
            
        return jsonify({'data': results}), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching book genres: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/book-subscriptions', methods=['GET'])
@login_required
def book_subscriptions():
    """Display most subscribed books from the library database with advanced filtering"""
    try:
        from datetime import date, timedelta
        
        # --- Filters ---
        genre_id = request.args.get('genre')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # 1. Fetch available genres
            cursor.execute("SELECT DISTINCT g.id, g.name FROM genres g JOIN book_genre bg ON g.id = bg.genre_id ORDER BY g.name")
            all_genres = cursor.fetchall()

            # 2. Main Query for Top Books
            query = """
                SELECT 
                    b.id,
                    b.title, 
                    b.author,
                    b.isbn,
                    COALESCE(b.is_active, 1) AS is_active,
                    p.name AS publisher_name,
                    COUNT(bu.book_id) as subscription_count
                FROM books b
                LEFT JOIN publishers p ON b.publisher_id = p.id
                JOIN book_user bu ON b.id = bu.book_id
                LEFT JOIN book_genre bg ON b.id = bg.book_id
                WHERE 1=1
            """
            
            params = []
            if genre_id:
                query += " AND bg.genre_id = %s"
                params.append(genre_id)
            
            # Use 'bu.created_at' if it exists, otherwise omit date filter for now
            # We'll try to use it and catch errors if the column is missing
            if start_date:
                query += " AND bu.created_at >= %s"
                params.append(start_date)
            if end_date:
                query += " AND bu.created_at <= %s"
                params.append(f"{end_date} 23:59:59")
                
            query += """
                GROUP BY b.id, b.title, b.author, b.isbn, b.is_active
                ORDER BY subscription_count DESC
            """
            
            cursor.execute(query, params)
            results = cursor.fetchall()
            
            # 3. Summary Stats
            total_subs = sum(r['subscription_count'] for r in results)
            top_book = results[0]['title'] if results else 'N/A'
            
        return render_template('book_subscriptions.html', 
                             title='Book Subscriptions', 
                             subscriptions=results,
                             genres=all_genres,
                             selected_genre=genre_id,
                             start_date=start_date,
                             end_date=end_date,
                             stats={
                                 'total': total_subs,
                                 'top_book': top_book,
                                 'count': len(results)
                             })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching book subscriptions: {e}")
        # Fallback to empty list or flash message if template expects it
        return render_template('book_subscriptions.html', 
                             title='Book Subscriptions', 
                             subscriptions=[],
                             error=str(e))


@app.route('/book-purchases', methods=['GET'])
@login_required
def book_purchases():
    """Display most purchased books from the library database (non-voucher) with advanced analytics"""
    try:
        from datetime import date, timedelta
        
        # --- Filters ---
        genre_id = request.args.get('genre')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # 1. Fetch available genres
            cursor.execute("SELECT DISTINCT g.id, g.name FROM genres g JOIN book_genre bg ON g.id = bg.genre_id ORDER BY g.name")
            all_genres = cursor.fetchall()

            # 2. Query to find most purchased books and their revenue
            query = """
                SELECT 
                    b.id,
                    b.title, 
                    b.author,
                    b.isbn,
                    COALESCE(b.is_active, 1) AS is_active,
                    SUM(bo.quantity) as purchase_count,
                    SUM(bo.quantity * b.price) as total_revenue
                FROM books b
                JOIN book_order bo ON b.id = bo.book_id
                JOIN orders o ON bo.order_id = o.id
                LEFT JOIN book_genre bg ON b.id = bg.book_id
                WHERE o.payment_method != 'Voucher'
                  AND o.status = 'Completed'
            """
            
            params = []
            if genre_id:
                query += " AND bg.genre_id = %s"
                params.append(genre_id)
            if start_date:
                query += " AND o.created_at >= %s"
                params.append(start_date)
            if end_date:
                query += " AND o.created_at <= %s"
                params.append(f"{end_date} 23:59:59")
                
            query += """
                GROUP BY b.id, b.title, b.author, b.isbn, b.is_active
                ORDER BY purchase_count DESC
            """
            cursor.execute(query, params)
            results = cursor.fetchall()
            
            # 3. Global Stats for header
            total_revenue = sum(r['total_revenue'] or 0 for r in results)
            total_qty = sum(r['purchase_count'] or 0 for r in results)
            
            # 4. Fetch Top Province Overall (for summary card)
            prov_query = """
                SELECT i.province, SUM(bo.quantity) as qty
                FROM orders o
                JOIN book_order bo ON o.id = bo.order_id
                JOIN users u ON o.user_id = u.id
                JOIN institution_user iu ON u.id = iu.user_id
                JOIN institutions i ON iu.institution_id = i.id
                WHERE o.payment_method != 'Voucher' AND o.status = 'Completed'
            """
            prov_params = []
            if start_date:
                prov_query += " AND o.created_at >= %s"
                prov_params.append(start_date)
            if end_date:
                prov_query += " AND o.created_at <= %s"
                prov_params.append(f"{end_date} 23:59:59")
                
            prov_query += " GROUP BY i.province ORDER BY qty DESC LIMIT 1"
            cursor.execute(prov_query, prov_params)
            top_prov_res = cursor.fetchone()
            top_province = top_prov_res['province'] if top_prov_res else 'N/A'
            
        return render_template('book_purchases.html', 
                             title='Book Purchases', 
                             purchases=results,
                             genres=all_genres,
                             selected_genre=genre_id,
                             start_date=start_date,
                             end_date=end_date,
                             stats={
                                 'revenue': total_revenue,
                                 'quantity': total_qty,
                                 'top_province': top_province,
                                 'count': len(results)
                             })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching book purchases: {e}")
        return render_template('book_purchases.html', 
                             title='Book Purchases', 
                             purchases=[],
                             error=str(e))


@app.route('/library-comparative-analytics', methods=['GET'])
@login_required
def library_comparative_analytics():
    """Comparative analysis between Subscriptions and Purchases (Conversion)"""
    try:
        from datetime import date
        
        genre_id = request.args.get('genre')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        conn = get_direct_library_conn()
        
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # 1. Fetch available genres
            cursor.execute("SELECT DISTINCT g.id, g.name FROM genres g JOIN book_genre bg ON g.id = bg.genre_id ORDER BY g.name")
            all_genres = cursor.fetchall()

            # 2. Main query for conversion metrics
            query = """
                SELECT 
                    b.id, b.title, b.author,
                    COALESCE(b.is_active, 1) AS is_active,
                    sub_q.sub_count,
                    pur_q.pur_count,
                    COALESCE(pur_q.pur_count, 0) / NULLIF(COALESCE(sub_q.sub_count, 0), 0) as conversion_ratio
                FROM books b
                LEFT JOIN (
                    SELECT book_id, COUNT(*) as sub_count 
                    FROM book_user 
                    WHERE 1=1
                    {sub_date_filter}
                    GROUP BY book_id
                ) sub_q ON b.id = sub_q.book_id
                LEFT JOIN (
                    SELECT bo.book_id, SUM(bo.quantity) as pur_count
                    FROM book_order bo
                    JOIN orders o ON bo.order_id = o.id
                    WHERE o.status = 'Completed' AND o.payment_method != 'Voucher'
                    {pur_date_filter}
                    GROUP BY bo.book_id
                ) pur_q ON b.id = pur_q.book_id
                LEFT JOIN book_genre bg ON b.id = bg.book_id
                WHERE (sub_q.sub_count > 0 OR pur_q.pur_count > 0)
            """
            
            sub_date_filter = ""
            pur_date_filter = ""
            params = []
            
            if start_date:
                sub_date_filter += " AND created_at >= %s"
                pur_date_filter += " AND o.created_at >= %s"
                params.append(start_date)
            if end_date:
                sub_date_filter += " AND created_at <= %s"
                pur_date_filter += " AND o.created_at <= %s"
                params.append(f"{end_date} 23:59:59")
            
            # Since we have parameters in subqueries, we must handle them carefully
            # But here we can inject the filter strings and double the params if needed
            # Actually, the params list needs to contain copies for both subqueries
            final_params = []
            if start_date: final_params.append(start_date)
            if end_date: final_params.append(f"{end_date} 23:59:59")
            if start_date: final_params.append(start_date)
            if end_date: final_params.append(f"{end_date} 23:59:59")
            
            if genre_id:
                query += " AND bg.genre_id = %s"
                final_params.append(genre_id)
                
            query = query.format(sub_date_filter=sub_date_filter, pur_date_filter=pur_date_filter)
            query += " GROUP BY b.id, b.title, b.author ORDER BY conversion_ratio DESC"
            
            cursor.execute(query, final_params)
            results = cursor.fetchall()
            
        return render_template('library_comparative_analytics.html', 
                             title='Library Conversion Analytics', 
                             data=results,
                             genres=all_genres,
                             selected_genre=genre_id,
                             start_date=start_date,
                             end_date=end_date)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        app.logger.error(f"Error fetching comparative analytics: {e}")
        return render_template('library_comparative_analytics.html', 
                             title='Library Conversion Analytics', 
                             data=[],
                             error=str(e))


@app.route('/api/schools/search', methods=['GET'])
@login_required
def api_schools_search():
    """Search schools (ASL tblschools + optional library match) for champions. Returns asl_school_id, library_school_id, school_name, province."""
    q = (request.args.get('q') or '').strip()
    province = (request.args.get('province') or '').strip()
    limit = min(int(request.args.get('limit', 50) or 50), 100)
    results = []
    conn_asl = None
    conn_lib = None
    try:
        conn_asl = get_ruzivo_conn()
        if not conn_asl:
            return jsonify({'error': 'Database connection not available'}), 500
        cursor_asl = conn_asl.cursor(pymysql.cursors.DictCursor)
        if q:
            sql = """
                SELECT school_id AS asl_school_id, school_name, school_province AS province
                FROM tblschools
                WHERE school_name LIKE %s
            """
            params = [f'%{q}%']
            if province:
                sql += " AND school_province = %s"
                params.append(province)
            sql += " ORDER BY school_name LIMIT %s"
            params.append(limit)
            cursor_asl.execute(sql, params)
        else:
            sql = """
                SELECT school_id AS asl_school_id, school_name, school_province AS province
                FROM tblschools
            """
            params = []
            if province:
                sql += " WHERE school_province = %s"
                params.append(province)
            sql += " ORDER BY school_name LIMIT %s"
            params.append(limit)
            cursor_asl.execute(sql, params)
        asl_rows = cursor_asl.fetchall()
        asl_by_name_province = {}
        for row in asl_rows:
            key = ((row.get('school_name') or '').strip().lower(), (row.get('province') or '').strip().lower())
            asl_by_name_province[key] = row
        try:
            conn_lib = get_direct_library_conn()
            if conn_lib:
                cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)
                lib_sql = "SELECT id, name, province FROM institutions"
                lib_params = []
                if q:
                    lib_sql += " WHERE name LIKE %s"
                    lib_params.append(f'%{q}%')
                if province:
                    lib_sql += " AND province = %s" if q else " WHERE province = %s"
                    lib_params.append(province)
                lib_sql += " LIMIT 200"
                cursor_lib.execute(lib_sql, lib_params)
                lib_rows = cursor_lib.fetchall()
                lib_by_name_province = {}
                for r in lib_rows:
                    key = ((r.get('name') or '').strip().lower(), (r.get('province') or '').strip().lower())
                    lib_by_name_province[key] = r.get('id')
        except Exception:
            lib_by_name_province = {}
        for row in asl_rows:
            name = (row.get('school_name') or '').strip()
            prov = (row.get('province') or '').strip()
            key = (name.lower(), prov.lower())
            library_school_id = lib_by_name_province.get(key) if lib_by_name_province else None
            if library_school_id is not None:
                library_school_id = str(library_school_id)
            results.append({
                'asl_school_id': row.get('asl_school_id') if row.get('asl_school_id') is not None else None,
                'library_school_id': library_school_id or '',
                'school_name': name,
                'province': prov
            })
        return jsonify(results)
    except Exception as e:
        logging.exception(e)
        return jsonify({'error': str(e)}), 500
    finally:
        if conn_asl:
            try:
                conn_asl.close()
            except Exception:
                pass
        if conn_lib:
            try:
                conn_lib.close()
            except Exception:
                pass


def _resolve_champion_for_current_user():
    """Resolve ChampionSchool record for current_user (Brand Ambassador). Returns (champion, error_response) or (champion, None)."""
    if getattr(current_user, 'userRole', None) != 'Brand Ambassador':
        return None, (jsonify({'error': 'Only Brand Ambassadors can request school additions'}), 403)
    champion = ChampionSchool.query.filter(
        ChampionSchool.firstname.ilike(f'%{(current_user.firstname or "").strip()}%'),
        ChampionSchool.lastname.ilike(f'%{(current_user.lastname or "").strip()}%'),
        ChampionSchool.province.ilike(f'%{(current_user.province or "").strip()}%')
    ).first()
    if not champion:
        return None, (jsonify({'error': 'No champion profile found for your account. Please contact admin.'}), 404)
    return champion, None


def _normalize_school_id(value):
    """Return stripped string or empty string if missing."""
    if value is None or value == '':
        return ''
    return str(value).strip()


@app.route('/api/champion/school-requests', methods=['POST'])
@login_required
def api_champion_school_requests_create():
    """Champion submits school(s) to add; creates Pending ChampionSchoolRequest(s) and notifies admins. Skips duplicates (already on profile, already pending, or duplicate in request)."""
    champion, err = _resolve_champion_for_current_user()
    if err:
        return err
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    schools = data.get('schools') if isinstance(data.get('schools'), list) else []
    if not schools:
        return jsonify({'error': 'Provide a non-empty list of schools (asl_school_id, library_school_id, school_name, optional province)'}), 400

    # Build "existing on profile" sets (non-empty asl_id, non-empty lib_id, and names when both IDs empty)
    profile_asl_ids = set()
    profile_lib_ids = set()
    profile_names_no_id = set()
    for entry in (champion.get_schools() or []):
        ea = _normalize_school_id(entry.get('asl_school_id'))
        el = _normalize_school_id(entry.get('library_school_id'))
        en = (entry.get('school_name') or '').strip().lower()
        if ea:
            profile_asl_ids.add(ea)
        if el:
            profile_lib_ids.add(el)
        if not ea and not el and en:
            profile_names_no_id.add(en)

    # Build "already pending" sets
    pending_asl_ids = set()
    pending_lib_ids = set()
    pending_names_no_id = set()
    for req in ChampionSchoolRequest.query.filter_by(champion_school_id=champion.id, status='Pending').all():
        pa = _normalize_school_id(req.asl_school_id)
        pl = _normalize_school_id(req.library_school_id)
        pn = (req.school_name or '').strip().lower()
        if pa:
            pending_asl_ids.add(pa)
        if pl:
            pending_lib_ids.add(pl)
        if not pa and not pl and pn:
            pending_names_no_id.add(pn)

    # Track seen in this request (first occurrence wins)
    seen_asl_ids = set()
    seen_lib_ids = set()
    seen_names_no_id = set()

    created = []
    duplicates_skipped = []

    for s in schools:
        school_name = (s.get('school_name') or '').strip()
        if not school_name:
            continue
        asl_id = _normalize_school_id(s.get('asl_school_id'))
        lib_id = _normalize_school_id(s.get('library_school_id'))
        province = (s.get('province') or '').strip() or None
        name_lower = school_name.lower()

        # Determine duplicate reason (order: in-request, profile, pending)
        reason = None
        if asl_id and asl_id in seen_asl_ids:
            reason = 'duplicate_in_request'
        elif lib_id and lib_id in seen_lib_ids:
            reason = 'duplicate_in_request'
        elif not asl_id and not lib_id and name_lower in seen_names_no_id:
            reason = 'duplicate_in_request'
        elif asl_id and asl_id in profile_asl_ids:
            reason = 'already_on_profile'
        elif lib_id and lib_id in profile_lib_ids:
            reason = 'already_on_profile'
        elif not asl_id and not lib_id and name_lower in profile_names_no_id:
            reason = 'already_on_profile'
        elif asl_id and asl_id in pending_asl_ids:
            reason = 'already_pending'
        elif lib_id and lib_id in pending_lib_ids:
            reason = 'already_pending'
        elif not asl_id and not lib_id and name_lower in pending_names_no_id:
            reason = 'already_pending'

        if reason:
            duplicates_skipped.append({
                'school_name': school_name,
                'asl_school_id': asl_id or None,
                'library_school_id': lib_id or None,
                'reason': reason
            })
            continue

        req = ChampionSchoolRequest(
            champion_school_id=champion.id,
            requested_by_user_id=current_user.id,
            asl_school_id=asl_id or None,
            library_school_id=lib_id or None,
            school_name=school_name,
            province=province,
            status='Pending'
        )
        db.session.add(req)
        created.append(req)
        if asl_id:
            seen_asl_ids.add(asl_id)
            pending_asl_ids.add(asl_id)
        if lib_id:
            seen_lib_ids.add(lib_id)
            pending_lib_ids.add(lib_id)
        if not asl_id and not lib_id and name_lower:
            seen_names_no_id.add(name_lower)
            pending_names_no_id.add(name_lower)

    if not created:
        if duplicates_skipped:
            return jsonify({
                'error': 'All selected schools are duplicates.',
                'duplicates_skipped': duplicates_skipped
            }), 400
        return jsonify({'error': 'No valid school entries to request'}), 400

    db.session.commit()
    approver_ids = get_champion_school_approvers()
    champ_name = f'{champion.firstname} {champion.lastname}'
    n = len(created)
    msg = f'Champion {champ_name} requested to add {n} school(s).'
    for uid in approver_ids:
        notif = Notification(
            user_id=uid,
            query_id=None,
            message=msg,
            notification_type='champion_school_request'
        )
        db.session.add(notif)
    db.session.commit()
    payload = {'message': f'{len(created)} school(s) requested', 'requests': [r.to_dict() for r in created]}
    if duplicates_skipped:
        payload['duplicates_skipped'] = duplicates_skipped
    return jsonify(payload), 201


@app.route('/api/champion/school-requests', methods=['GET'])
@login_required
def api_champion_school_requests_list():
    """List school requests for the current champion."""
    champion, err = _resolve_champion_for_current_user()
    if err:
        return err
    status = request.args.get('status')
    q = ChampionSchoolRequest.query.filter_by(champion_school_id=champion.id)
    if status:
        q = q.filter_by(status=status)
    q = q.order_by(ChampionSchoolRequest.created_at.desc())
    requests = q.all()
    return jsonify({'requests': [r.to_dict() for r in requests]})


@app.route('/api/champion/my-schools', methods=['GET'])
@login_required
def api_champion_my_schools():
    """Return current champion's full school list (for My schools UI)."""
    champion, err = _resolve_champion_for_current_user()
    if err:
        return err
    schools = champion.get_schools() or []
    return jsonify({'schools': schools})


@app.route('/api/champion-schools-by-user/<username>', methods=['GET'])
@login_required
def get_champion_schools_by_user(username):
    """Get schools for a specific Brand Ambassador user"""
    try:
        # Find the user
        user = User.query.filter_by(username=username).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Check if user is Brand Ambassador
        if user.userRole != 'Brand Ambassador':
            return jsonify({'error': 'User is not a Brand Ambassador'}), 403
        
        # Find ChampionSchool record matching the user
        # Match by firstname, lastname, and province
        champion = ChampionSchool.query.filter(
            ChampionSchool.firstname.ilike(f'%{user.firstname}%'),
            ChampionSchool.lastname.ilike(f'%{user.lastname}%'),
            ChampionSchool.province.ilike(f'%{user.province}%')
        ).first()
        
        if not champion:
            return jsonify({'schools': []}), 200
        
        # Get schools from the champion
        schools_data = champion.get_schools() or []
        school_names = [s.get('school_name', '') for s in schools_data if s.get('school_name')]
        
        return jsonify({'schools': school_names}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/all-champion-schools', methods=['GET'])
@login_required
def get_all_champion_schools():
    """Get all unique school names from all ChampionSchool records"""
    try:
        champions = ChampionSchool.query.all()
        school_names_set = set()
        
        for champion in champions:
            schools_data = champion.get_schools() or []
            for school in schools_data:
                school_name = school.get('school_name', '')
                if school_name:
                    school_names_set.add(school_name)
        
        return jsonify({'schools': sorted(list(school_names_set))}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/book-allocation-requests', methods=['GET'])
@login_required
def get_book_allocation_requests():
    """List book allocation requests with optional status filter"""
    try:
        status_filter = request.args.get('status', None)
        
        # Admins see all requests, users see only their own
        if current_user.userRole == 'Admin':
            query = BookAllocationRequest.query
        else:
            query = BookAllocationRequest.query.filter_by(requester_username=current_user.username)
        
        # Apply status filter if provided
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        requests = query.order_by(BookAllocationRequest.created_at.desc()).all()
        
        return jsonify({'requests': [req.to_dict() for req in requests]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/book-allocation-requests', methods=['POST'])
@login_required
def create_book_allocation_request():
    """Create a new book allocation request"""
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['school_name', 'school_province', 'school_grade']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({'error': f'{field} is required'}), 400
        
        # If user is Brand Ambassador, validate school matches their username
        if current_user.userRole == 'Brand Ambassador':
            # Get schools for this user
            champion = ChampionSchool.query.filter(
                ChampionSchool.firstname.ilike(f'%{current_user.firstname}%'),
                ChampionSchool.lastname.ilike(f'%{current_user.lastname}%'),
                ChampionSchool.province.ilike(f'%{current_user.province}%')
            ).first()
            
            if champion:
                schools_data = champion.get_schools() or []
                user_schools = [s.get('school_name', '') for s in schools_data if s.get('school_name')]
                
                if data['school_name'] not in user_schools:
                    return jsonify({'error': 'School does not match your assigned schools'}), 403
        
        # Create request
        request_obj = BookAllocationRequest(
            requester_username=current_user.username,
            school_name=data['school_name'],
            school_province=data['school_province'],
            school_grade=data['school_grade'],
            quantity=data.get('quantity', 1),
            notes=data.get('notes', None),
            requested_date=datetime.strptime(data['requested_date'], '%Y-%m-%d').date() if data.get('requested_date') else None,
            status='Not allocated'
        )
        
        db.session.add(request_obj)
        db.session.commit()
        
        return jsonify(request_obj.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/book-allocation-requests/<int:request_id>/approve', methods=['POST'])
@login_required
def approve_book_allocation_request(request_id):
    """Approve a request and convert it to an allocation"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        request_obj = BookAllocationRequest.query.get_or_404(request_id)
        
        if request_obj.status == 'Allocated':
            return jsonify({'error': 'Request already allocated'}), 400
        
        # Create BookAllocations record
        # Convert single grade to comma-separated format (for compatibility)
        books_allocated = request_obj.school_grade
        
        allocation = BookAllocations(
            school_name=request_obj.school_name,
            school_province=request_obj.school_province,
            books_allocated=books_allocated,
            allocated_by=current_user.username
        )
        
        db.session.add(allocation)
        db.session.flush()  # Get the allocation ID
        
        # Update request
        request_obj.status = 'Allocated'
        request_obj.approved_by = current_user.username
        request_obj.converted_to_allocation_id = allocation.id
        
        db.session.commit()
        
        return jsonify({
            'request': request_obj.to_dict(),
            'allocation': {
                'id': allocation.id,
                'school_name': allocation.school_name,
                'school_province': allocation.school_province,
                'books_allocated': allocation.books_allocated
            }
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    all_users = User.query.all()
    form = RegistrationForm()
    csvform = CSVUploadForm()
    number_of_users = User.query.count()
    championcsvform = ChampionCSVUploadForm()
    manual_form = ChampionSchoolForm()
    all_champions = ChampionSchool.query.all()

    all_schools = []
    # try:
    #     conn = get_ruzivo_conn()
    #     cursor = conn.cursor()

    #     # ✅ Use a dictionary cursor to get results directly as dicts
    #     cursor.execute("SELECT school_id, school_name FROM tblschools ORDER BY school_name ASC")
    #     rows = cursor.fetchall()

    #     # ✅ If rows aren't dicts, convert manually
    #     for row in rows:
    #         if isinstance(row, dict):
    #             all_schools.append({
    #                 "id": str(row['school_id']),
    #                 "name": row['school_name']
    #             })
    #         else:
    #             # fallback for tuple-style rows
    #             all_schools.append({
    #                 "id": str(row[0]),
    #                 "name": row[1]
    #             })

    # except Exception as e:
    #     flash(f"Error fetching schools: {e}", "danger")
    # finally:
    #     if 'cursor' in locals():
    #         cursor.close()
    #     if 'conn' in locals():
    #         conn.close()

    return render_template(
        'settings.html',
        form=form,
        csvform=csvform,
        all_users=all_users,
        number_of_users=number_of_users,
        manual_form=manual_form,
        championcsvform=championcsvform,
        all_champions=all_champions,
        all_schools=all_schools,
        title='Settings'
    )




@app.route('/admin', methods=['GET', 'POST'])
@login_required
def admin():
    all_users = User.query.all()
    form = RegistrationForm()
    csvform = CSVUploadForm()
    number_of_users = User.query.count()
    championcsvform = ChampionCSVUploadForm()
    manual_form = ChampionSchoolForm()
    all_champions = ChampionSchool.query.all()

    all_schools = []

    if not current_user.userRole == "Admin":
        return "Unauthorized", 403

    users = User.query.all()


    return render_template(
        'admin.html',
        form=form,
        csvform=csvform,
        all_users=all_users,
        number_of_users=number_of_users,
        manual_form=manual_form,
        championcsvform=championcsvform,
        all_champions=all_champions,
        all_schools=all_schools,
        users=users,
        title='Settings'
    )





# Get all champions
@app.route("/api/champions", methods=["GET"])
def get_champions():
    champions = ChampionSchool.query.all()
    return jsonify([c.to_dict() for c in champions])


# Add a new champion
@app.route("/api/champions", methods=["POST"])
def add_champion():
    data = request.json
    champ = ChampionSchool(
        firstname=data["firstname"],
        lastname=data["lastname"],
        province=data["province"]
    )
    if "schools" in data:
        champ.set_schools(data["schools"])
    db.session.add(champ)
    db.session.commit()
    return jsonify(champ.to_dict()), 201


# Update champion (add a school) – admin/privilege only; champions use POST /api/champion/school-requests
@app.route("/api/champions/<int:champ_id>/add_school", methods=["POST"])
@login_required
def add_school(champ_id):
    if not can_approve_champion_schools():
        return jsonify({"error": "Unauthorized"}), 403
    champ = ChampionSchool.query.get_or_404(champ_id)
    data = request.json or {}
    asl_id = data.get("asl_school_id")
    library_id = data.get("library_school_id")
    school_name = (data.get("school_name") or "").strip()
    if not school_name:
        return jsonify({"error": "school_name is required"}), 400
    champ.add_school(
        asl_id=asl_id or "",
        library_id=library_id or "",
        school_name=school_name
    )
    db.session.commit()
    return jsonify(champ.to_dict())


# Delete champion
@app.route("/api/champions/<int:champ_id>", methods=["DELETE"])
def delete_champion(champ_id):
    champ = ChampionSchool.query.get_or_404(champ_id)
    db.session.delete(champ)
    db.session.commit()
    return jsonify({"message": "Champion deleted"})



# Update champion's school details
@app.route("/api/champions/<int:champ_id>/schools/<int:school_id>", methods=["PUT"])
def update_school(champ_id, school_id):
    champ = ChampionSchool.query.get_or_404(champ_id)
    school = next((s for s in champ.schools if s.id == school_id), None)

    if not school:
        return jsonify({"error": "School not found for this champion"}), 404

    data = request.json
    if "asl_school_id" in data:
        school.asl_school_id = data["asl_school_id"]
    if "library_school_id" in data:
        school.library_school_id = data["library_school_id"]
    if "school_name" in data:
        school.school_name = data["school_name"]

    db.session.commit()
    return jsonify(champ.to_dict())




UPLOAD_FOLDER = "uploads"
BULK_CHAMP_ALLOWED_EXTENSIONS = {"csv", "xlsx"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def bulk_champ_allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in BULK_CHAMP_ALLOWED_EXTENSIONS


# Bulk upload champions
@app.route("/api/champions/bulk_upload", methods=["POST"])
def bulk_upload_champions():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "" or not bulk_champ_allowed_file(file.filename):
        return jsonify({"error": "Invalid file type (only csv/xlsx allowed)"}), 400

    filepath = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(file.filename))
    file.save(filepath)

    # Load CSV or Excel into pandas
    if file.filename.endswith(".csv"):
        df = pd.read_csv(filepath)
    else:
        df = pd.read_excel(filepath)

    required_columns = {"firstname", "lastname", "province", "school_name", "asl_school_id", "library_school_id"}
    if not required_columns.issubset(df.columns):
        return jsonify({"error": f"Missing required columns. Required: {required_columns}"}), 400

    updated_champions = {}

    for _, row in df.iterrows():
        firstname = str(row["firstname"]).strip()
        lastname = str(row["lastname"]).strip()
        province = str(row["province"]).strip()
        asl_id = str(row["asl_school_id"]).strip() if pd.notna(row["asl_school_id"]) else None
        library_id = str(row["library_school_id"]).strip() if pd.notna(row["library_school_id"]) else None
        school_name = str(row["school_name"]).strip() if pd.notna(row["school_name"]) else None

        champ = ChampionSchool.query.filter_by(
            firstname=firstname,
            lastname=lastname,
            province=province
        ).first()

        # If champion does not exist, create new
        if not champ:
            champ = ChampionSchool(
                firstname=firstname,
                lastname=lastname,
                province=province
            )
            db.session.add(champ)
            db.session.flush()

        # Fetch existing schools (stored in JSON/list in your model)
        existing_schools = champ.get_schools() or []  

        matched_school = None
        for school in existing_schools:
            if (asl_id and str(school.get("asl_school_id")) == asl_id) or \
               (library_id and str(school.get("library_school_id")) == library_id):
                matched_school = school
                break

        if matched_school:
            # Update only missing/different values
            if asl_id and str(matched_school.get("asl_school_id")) != asl_id:
                matched_school["asl_school_id"] = asl_id
            if library_id and str(matched_school.get("library_school_id")) != library_id:
                matched_school["library_school_id"] = library_id
            if school_name:
                matched_school["school_name"] = school_name
        else:
            # Add new school
            new_school = {
                "asl_school_id": asl_id,
                "library_school_id": library_id,
                "school_name": school_name
            }
            existing_schools.append(new_school)

        # Save back the modified list
        champ.set_schools(existing_schools)

        updated_champions[f"{firstname}_{lastname}_{province}"] = champ.to_dict()

    db.session.commit()

    return jsonify({
        "message": "Bulk upload successful",
        "champions": list(updated_champions.values())
    }), 201









@app.route('/add_champion_form', methods=['POST'])
@login_required
def add_champion_form():
    form = ChampionSchoolForm()
    


    if form.validate_on_submit():
        firstname = form.firstname.data
        lastname = form.lastname.data
        province = form.province.data

        # ✅ Get selected schools from hidden field
        selected_data_json = request.form.get('selected_school_data')
        try:
            selected_schools = json.loads(selected_data_json) if selected_data_json else []
        except json.JSONDecodeError:
            flash("Invalid selected school data format.", "danger")
            return redirect(url_for('settings'))

        # if not selected_schools:
        #     flash("No schools selected.", "warning")
        #     return redirect(url_for('settings'))
        
        # Before creating new champion:
        for champ in ChampionSchool.query.all():
            if champ.schools:
                used_ids = [str(s['id']) for s in json.loads(champ.schools)]
                for s in selected_schools:
                    if str(s['id']) in used_ids:
                        flash(f"School '{s['name']}' is already assigned to another champion.", "danger")
                        return redirect(url_for('settings'))

        # ✅ Store as JSON string in DB
        try:
            champ = ChampionSchool(
                firstname=firstname,
                lastname=lastname,
                province=province,
                schools=json.dumps(selected_schools)  # list of dicts: [{id, name}, ...]
            )
            db.session.add(champ)
            db.session.commit()
            flash("Champion school added successfully!", "success")
            return redirect(url_for('settings'))
        except Exception as e:
            db.session.rollback()
            flash(f"Error saving champion: {str(e)}", "danger")

    else:
        flash("Form validation failed.", "danger")

    return redirect(url_for('Administration'))




@app.route('/analyze_champion_csv', methods=['POST'])
@login_required
def analyze_champion_csv():
    """Analyze CSV for duplicates before uploading"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if not file or not file.filename:
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.lower().endswith('.csv'):
        return jsonify({'error': 'Please upload a .csv file'}), 400

    try:
        # Save the uploaded file to a temporary location so the client can
        # reference it later when confirming upload. This avoids relying on
        # client-side File objects that may be cleared when modals close.
        token = uuid.uuid4().hex
        upload_dir = os.path.join(app.instance_path, 'tmp_champion_uploads')
        os.makedirs(upload_dir, exist_ok=True)
        saved_filename = f"{token}_{secure_filename(file.filename)}"
        saved_path = os.path.join(upload_dir, saved_filename)
        file.stream.seek(0)
        file.save(saved_path)

        # Parse CSV from the saved copy - keep file open during all operations
        with open(saved_path, 'r', encoding='utf-8') as fh:
            reader = csv.DictReader(fh)

            # Validate required columns
            required_columns = ['firstname', 'lastname', 'province', 'school_name', 'asl_school_id', 'library_school_id']
            csv_columns = reader.fieldnames or []
            missing_columns = [col for col in required_columns if col not in csv_columns]
            if missing_columns:
                return jsonify({
                    'error': 'Missing required columns',
                    'missing_columns': missing_columns,
                    'required_columns': required_columns,
                    'found_columns': csv_columns
                }), 400

            # Group rows by champion identity and track schools
            grouped = {}
            csv_duplicates_within = []  # Duplicates within the CSV itself
            seen_ids = {}  # Track IDs within CSV

            for row_idx, row in enumerate(reader, start=2):  # Start at 2 (header is row 1)
                fn = (row.get('firstname') or '').strip()
                ln = (row.get('lastname') or '').strip()
                province = (row.get('province') or '').strip()
                
                if not fn or not ln or not province:
                    continue

                key = (fn, ln, province)
                entry = grouped.setdefault(key, {
                    'firstname': fn,
                    'lastname': ln,
                    'province': province,
                    'schools': [],
                    'rows': []
                })
                entry['rows'].append(row_idx)

                # Parse schools
                schools_list = []
                schools_json = row.get('schools')
                if schools_json:
                    try:
                        parsed = json.loads(schools_json)
                        if isinstance(parsed, list):
                            schools_list = parsed
                    except Exception:
                        pass

                if not schools_list:
                    school_name = (row.get('school_name') or '').strip()
                    asl_id = (row.get('asl_school_id') or '').strip()
                    library_id = (row.get('library_school_id') or '').strip()
                    if school_name or asl_id or library_id:
                        schools_list = [{
                            'school_name': school_name,
                            'asl_school_id': asl_id,
                            'library_school_id': library_id,
                            'row': row_idx
                        }]

                # Check for duplicate IDs within CSV
                for s in schools_list:
                    asl_id = s.get('asl_school_id', '').strip()
                    lib_id = s.get('library_school_id', '').strip()
                    
                    if asl_id and asl_id != '0':
                        if asl_id in seen_ids:
                            csv_duplicates_within.append({
                                'type': 'ASL ID',
                                'value': asl_id,
                                'rows': [seen_ids[asl_id], row_idx],
                                'school_name': s.get('school_name', '')
                            })
                        else:
                            seen_ids[asl_id] = row_idx
                    
                    if lib_id and lib_id != '0':
                        lib_key = f"lib_{lib_id}"
                        if lib_key in seen_ids:
                            csv_duplicates_within.append({
                                'type': 'Library ID',
                                'value': lib_id,
                                'rows': [seen_ids[lib_key], row_idx],
                                'school_name': s.get('school_name', '')
                            })
                        else:
                            seen_ids[lib_key] = row_idx

                    school = {
                        'school_name': s.get('school_name', ''),
                        'asl_school_id': s.get('asl_school_id', ''),
                        'library_school_id': s.get('library_school_id', '')
                    }
                    if school not in entry['schools']:
                        entry['schools'].append(school)

        # Check for duplicates with existing database records
        existing_champions = []
        duplicate_schools = []

        for (fn, ln, province), data in grouped.items():
            champ = ChampionSchool.query.filter_by(
                firstname=fn, lastname=ln, province=province
            ).first()
            
            if champ:
                existing_champions.append({
                    'firstname': fn,
                    'lastname': ln,
                    'province': province,
                    'existing_schools': champ.get_schools() or [],
                    'new_schools': data['schools'],
                    'csv_rows': data['rows']
                })

        # Check for duplicate school IDs in existing database
        all_existing_champs = ChampionSchool.query.all()
        for data in grouped.values():
            for new_school in data['schools']:
                asl_id = new_school.get('asl_school_id', '').strip()
                lib_id = new_school.get('library_school_id', '').strip()
                
                for existing_champ in all_existing_champs:
                    existing_schools = existing_champ.get_schools() or []
                    for ex_school in existing_schools:
                        ex_asl = str(ex_school.get('asl_school_id', '')).strip()
                        ex_lib = str(ex_school.get('library_school_id', '')).strip()
                        
                        if asl_id and asl_id != '0' and asl_id == ex_asl:
                            duplicate_schools.append({
                                'type': 'ASL ID',
                                'id_value': asl_id,
                                'csv_school': new_school.get('school_name', ''),
                                'existing_school': ex_school.get('school_name', ''),
                                'existing_champion': f"{existing_champ.firstname} {existing_champ.lastname}"
                            })
                        
                        if lib_id and lib_id != '0' and lib_id == ex_lib:
                            duplicate_schools.append({
                                'type': 'Library ID',
                                'id_value': lib_id,
                                'csv_school': new_school.get('school_name', ''),
                                'existing_school': ex_school.get('school_name', ''),
                                'existing_champion': f"{existing_champ.firstname} {existing_champ.lastname}"
                            })

        # Prepare response
        has_duplicates = len(existing_champions) > 0 or len(duplicate_schools) > 0 or len(csv_duplicates_within) > 0
        
        return jsonify({
            'success': True,
            'has_duplicates': has_duplicates,
            'csv_internal_duplicates': csv_duplicates_within,
            'existing_champion_duplicates': existing_champions,
            'duplicate_school_ids': duplicate_schools,
            'total_new_champions': len([k for k, v in grouped.items() if not ChampionSchool.query.filter_by(firstname=k[0], lastname=k[1], province=k[2]).first()]),
            'total_updates': len(existing_champions),
            'grouped_data': list(grouped.values()),  # For reference
            'file_token': token,
            'file_name': saved_filename
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/upload_csv', methods=['POST'])
@login_required
def upload_csv():
    # Admin-only action
    if current_user.userRole != 'Admin':
        flash('Unauthorized', 'danger')
        return redirect(url_for('administration'))

    form = ChampionCSVUploadForm()

    if not form.validate_on_submit():
        flash('Invalid upload. Please upload a valid CSV file.', 'warning')
        return redirect(url_for('administration'))

    try:
        file_storage = form.file.data
        # Require a file to be provided
        if not file_storage or not getattr(file_storage, 'filename', '').strip():
            flash('Please choose a CSV file to upload.', 'warning')
            return redirect(url_for('administration'))
        # Basic extension check
        if not file_storage.filename.lower().endswith('.csv'):
            flash('Please upload a .csv file.', 'warning')
            return redirect(url_for('administration'))

        # Parse CSV
        stream = io.TextIOWrapper(file_storage.stream, encoding='utf-8')
        reader = csv.DictReader(stream)

        # Group rows by champion identity
        grouped = {}
        for row in reader:
            fn = (row.get('firstname') or '').strip()
            ln = (row.get('lastname') or '').strip()
            province = (row.get('province') or '').strip()
            if not fn or not ln or not province:
                # Skip incomplete champion rows
                continue
            key = (fn, ln, province)
            entry = grouped.setdefault(key, {
                'firstname': fn,
                'lastname': ln,
                'province': province,
                'schools': []
            })

            schools_list = []
            # Support both: a JSON `schools` column OR separate columns
            schools_json = row.get('schools')
            if schools_json:
                try:
                    parsed = json.loads(schools_json)
                    if isinstance(parsed, list):
                        schools_list = parsed
                except Exception:
                    pass

            if not schools_list:
                school_name = (row.get('school_name') or '').strip()
                asl_id = (row.get('asl_school_id') or '').strip()
                library_id = (row.get('library_school_id') or '').strip()
                if school_name or asl_id or library_id:
                    schools_list = [{
                        'school_name': school_name,
                        'asl_school_id': asl_id,
                        'library_school_id': library_id
                    }]

            # Deduplicate schools before adding
            for s in schools_list:
                school = {
                    'school_name': s.get('school_name', ''),
                    'asl_school_id': s.get('asl_school_id', ''),
                    'library_school_id': s.get('library_school_id', '')
                }
                if school not in entry['schools']:
                    entry['schools'].append(school)

        # Persist grouped data, merging with existing champions
        for (fn, ln, province), data in grouped.items():
            champ = ChampionSchool.query.filter_by(
                firstname=fn, lastname=ln, province=province
            ).first()
            incoming_schools = data['schools']
            if champ:
                existing = champ.get_schools() or []
                for s in incoming_schools:
                    if s not in existing:
                        existing.append(s)
                champ.set_schools(existing)
            else:
                champ = ChampionSchool(
                    firstname=fn,
                    lastname=ln,
                    province=province
                )
                champ.set_schools(incoming_schools)
                db.session.add(champ)

        db.session.commit()
        flash('CSV uploaded and champions updated successfully!', 'success')
        return redirect(url_for('administration'))

    except Exception as e:
        db.session.rollback()
        flash(f'CSV upload failed: {str(e)}', 'danger')
        return redirect(url_for('administration'))


@app.route('/upload_csv_confirmed', methods=['POST'])
@login_required
def upload_csv_confirmed():
    """Upload CSV with user's duplicate resolution choices"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403

    # Support either a freshly uploaded file OR a previously analyzed file
    # referenced by `file_token` (returned by /analyze_champion_csv).
    file = None
    saved_path = None
    file_token = request.form.get('file_token')
    if file_token:
        upload_dir = os.path.join(app.instance_path, 'tmp_champion_uploads')
        if os.path.isdir(upload_dir):
            for fname in os.listdir(upload_dir):
                if fname.startswith(f"{file_token}_"):
                    saved_path = os.path.join(upload_dir, fname)
                    break
        if not saved_path or not os.path.exists(saved_path):
            return jsonify({'error': 'File token not found'}), 400
    else:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        file = request.files['file']
    skip_champions = request.form.get('skip_champions', '[]')
    skip_schools = request.form.get('skip_schools', '[]')
    
    try:
        skip_champions = json.loads(skip_champions)
        skip_schools = json.loads(skip_schools)
    except:
        skip_champions = []
        skip_schools = []

    try:
        # Parse CSV from either the uploaded FileStorage or the saved temp file
        if saved_path:
            with open(saved_path, 'r', encoding='utf-8') as fh:
                reader = csv.DictReader(fh)
                
                # Group rows by champion identity
                grouped = {}
                for row in reader:
                    fn = (row.get('firstname') or '').strip()
                    ln = (row.get('lastname') or '').strip()
                    province = (row.get('province') or '').strip()
                    if not fn or not ln or not province:
                        continue
                    key = (fn, ln, province)
                    entry = grouped.setdefault(key, {
                        'firstname': fn,
                        'lastname': ln,
                        'province': province,
                        'schools': []
                    })

                    schools_list = []
                    schools_json = row.get('schools')
                    if schools_json:
                        try:
                            parsed = json.loads(schools_json)
                            if isinstance(parsed, list):
                                schools_list = parsed
                        except Exception:
                            pass

                    if not schools_list:
                        school_name = (row.get('school_name') or '').strip()
                        asl_id = (row.get('asl_school_id') or '').strip()
                        library_id = (row.get('library_school_id') or '').strip()
                        if school_name or asl_id or library_id:
                            schools_list = [{
                                'school_name': school_name,
                                'asl_school_id': asl_id,
                                'library_school_id': library_id
                            }]

                    for s in schools_list:
                        school = {
                            'school_name': s.get('school_name', ''),
                            'asl_school_id': s.get('asl_school_id', ''),
                            'library_school_id': s.get('library_school_id', '')
                        }
                        # Check if school should be skipped
                        asl_id = school['asl_school_id'].strip()
                        lib_id = school['library_school_id'].strip()
                        should_skip = False
                        for skip_id in skip_schools:
                            if (asl_id and asl_id == skip_id) or (lib_id and lib_id == skip_id):
                                should_skip = True
                                break
                        
                        if not should_skip and school not in entry['schools']:
                            entry['schools'].append(school)
        else:
            stream = io.TextIOWrapper(file.stream, encoding='utf-8')
            reader = csv.DictReader(stream)
            
            # Group rows by champion identity
            grouped = {}
            for row in reader:
                fn = (row.get('firstname') or '').strip()
                ln = (row.get('lastname') or '').strip()
                province = (row.get('province') or '').strip()
                if not fn or not ln or not province:
                    continue
                key = (fn, ln, province)
                entry = grouped.setdefault(key, {
                    'firstname': fn,
                    'lastname': ln,
                    'province': province,
                    'schools': []
                })

                schools_list = []
                schools_json = row.get('schools')
                if schools_json:
                    try:
                        parsed = json.loads(schools_json)
                        if isinstance(parsed, list):
                            schools_list = parsed
                    except Exception:
                        pass

                if not schools_list:
                    school_name = (row.get('school_name') or '').strip()
                    asl_id = (row.get('asl_school_id') or '').strip()
                    library_id = (row.get('library_school_id') or '').strip()
                    if school_name or asl_id or library_id:
                        schools_list = [{
                            'school_name': school_name,
                            'asl_school_id': asl_id,
                            'library_school_id': library_id
                        }]

                for s in schools_list:
                    school = {
                        'school_name': s.get('school_name', ''),
                        'asl_school_id': s.get('asl_school_id', ''),
                        'library_school_id': s.get('library_school_id', '')
                    }
                    # Check if school should be skipped
                    asl_id = school['asl_school_id'].strip()
                    lib_id = school['library_school_id'].strip()
                    should_skip = False
                    for skip_id in skip_schools:
                        if (asl_id and asl_id == skip_id) or (lib_id and lib_id == skip_id):
                            should_skip = True
                            break
                    
                    if not should_skip and school not in entry['schools']:
                        entry['schools'].append(school)

        # Persist grouped data
        created = 0
        updated = 0
        skipped = 0
        
        for (fn, ln, province), data in grouped.items():
            # Check if champion should be skipped
            champ_key = f"{fn}|{ln}|{province}"
            if champ_key in skip_champions:
                skipped += 1
                continue
            
            champ = ChampionSchool.query.filter_by(
                firstname=fn, lastname=ln, province=province
            ).first()
            incoming_schools = data['schools']
            
            if champ:
                existing = champ.get_schools() or []
                for s in incoming_schools:
                    if s not in existing:
                        existing.append(s)
                champ.set_schools(existing)
                updated += 1
            else:
                champ = ChampionSchool(
                    firstname=fn,
                    lastname=ln,
                    province=province
                )
                champ.set_schools(incoming_schools)
                db.session.add(champ)
                created += 1

        db.session.commit()
        
        # Clean up temp file if used
        if saved_path:
            try:
                os.remove(saved_path)
            except Exception:
                pass
                
        return jsonify({
            'success': True,
            'message': f'Upload complete! Created: {created}, Updated: {updated}, Skipped: {skipped}',
            'created': created,
            'updated': updated,
            'skipped': skipped
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


def _get_profile_school_tables_data(username):
    """Compute school tables data for profile API. Returns dict with truly_active_schools, expired_schools, schools_without_scholarship, counts (JSON-serializable)."""
    user = db.first_or_404(sa.select(User).where(User.username == username))
    champ = ChampionSchool.query.filter_by(firstname=user.firstname, lastname=user.lastname, province=user.province).first()
    existing_school_dicts = []
    if champ and champ.schools:
        try:
            existing_school_dicts = json.loads(champ.schools)
        except Exception:
            existing_school_dicts = []
    preselected_asl_ids = [s["asl_school_id"] for s in existing_school_dicts if "asl_school_id" in s]
    today = datetime.today().date()
    first_of_month, end_of_month = _mtd_month_bounds(today)

    def normalize_active_flag(value):
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        if isinstance(value, str):
            normalized = value.strip().lower()
            return normalized in {"1", "true", "t", "yes", "y"}
        return False

    duration_map = {}
    if preselected_asl_ids:
        try:
            conn = get_ruzivo_conn()
            cursor = conn.cursor()
            placeholders = ','.join(['%s'] * len(preselected_asl_ids))
            duration_query = f"""
                SELECT SUM(x.duration) AS total_months, x.school_id, MIN(x.awarded_at) AS first_awarded_at
                FROM tblscholarships_schools x
                WHERE x.school_id IN ({placeholders})
                AND x.awarded_at >= %s
                GROUP BY x.school_id
            """
            cursor.execute(duration_query, (*preselected_asl_ids, '2025-01-01'))
            duration_rows = cursor.fetchall()
            cursor.close()
            conn.close()
            for duration_row in duration_rows:
                if isinstance(duration_row, dict):
                    school_key = duration_row.get("school_id")
                    total_months = duration_row.get("total_months") or 0
                    first_awarded = duration_row.get("first_awarded_at")
                else:
                    school_key = duration_row[1]
                    total_months = duration_row[0] or 0
                    first_awarded = duration_row[2]
                if school_key:
                    is_over_12 = False
                    if first_awarded:
                        fa_date = first_awarded.date() if hasattr(first_awarded, 'date') else first_awarded
                        if isinstance(fa_date, date) and fa_date < today - timedelta(days=365):
                            is_over_12 = True
                    duration_map[school_key] = {"total_months": total_months, "first_awarded_at": first_awarded, "is_over_12_months": is_over_12}
        except Exception as e:
            print(f"Error calculating duration data: {e}")
            duration_map = {}

    from app.models import AppSetting
    exclude_12_months = AppSetting.get_value('asl_mtd_exclude_12_months', 'true') == 'true'
    exclude_1_year_awarded = AppSetting.get_value('asl_mtd_exclude_1_year_awarded', 'true') == 'true'
    schools_to_exclude_from_mtd = set()
    for school_id, duration_info in duration_map.items():
        total_months = duration_info.get("total_months", 0)
        is_over_12 = duration_info.get("is_over_12_months", False)
        should_exclude = False
        if exclude_12_months and total_months > 12:
            should_exclude = True
        if exclude_1_year_awarded and is_over_12:
            should_exclude = True
        if should_exclude:
            schools_to_exclude_from_mtd.add(str(school_id))
    filtered_asl_ids = [sid for sid in preselected_asl_ids if str(sid) not in schools_to_exclude_from_mtd]

    asl_mtd_by_school = {}
    smartlearning_champ_mtd = 0
    if filtered_asl_ids:
        period_end_by_school = {}
        for sid in filtered_asl_ids:
            dinfo = duration_map.get(sid, {})
            fa = dinfo.get("first_awarded_at")
            period_end_by_school[sid] = _mtd_period_end(fa, end_of_month)
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        placeholders = ','.join(['%s'] * len(filtered_asl_ids))
        query = f"""
            SELECT school_id, student_id, last_login
            FROM vwstudent
            WHERE school_id IN ({placeholders})
            AND last_login BETWEEN %s AND %s
        """
        cursor.execute(query, (*filtered_asl_ids, first_of_month, end_of_month))
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        counted = defaultdict(set)
        for row in rows:
            sid = row[0] if isinstance(row, (tuple, list)) else row.get("school_id")
            student_id = row[1] if isinstance(row, (tuple, list)) else row.get("student_id")
            last_login = row[2] if isinstance(row, (tuple, list)) else row.get("last_login")
            login_date = last_login.date() if hasattr(last_login, 'date') else last_login
            pe = period_end_by_school.get(sid, end_of_month)
            if login_date <= pe:
                counted[sid].add(student_id)
        smartlearning_champ_mtd = sum(len(s) for s in counted.values())
        asl_mtd_by_school = {sid: len(counted[sid]) for sid in counted}

    al_mtd_by_asl_id = {}
    library_ids = [s.get("library_school_id") for s in existing_school_dicts if s.get("library_school_id")]
    if library_ids and end_of_month and first_of_month:
        try:
            lib_id_by_asl = {}
            for s in existing_school_dicts:
                a, lib = s.get("asl_school_id"), s.get("library_school_id")
                if a is None or lib is None:
                    continue
                try:
                    lib_id_by_asl[int(a)] = int(lib)
                except (TypeError, ValueError):
                    pass
            period_end_by_lib_id = {}
            for asl_id, lib_id in lib_id_by_asl.items():
                fa = duration_map.get(asl_id, {}).get("first_awarded_at")
                period_end_by_lib_id[lib_id] = _mtd_period_end(fa, end_of_month)
            conn_lib = get_direct_library_conn()
            if conn_lib:
                import pymysql.cursors
                cursor_lib = conn_lib.cursor(pymysql.cursors.DictCursor)
                library_ids_int = [int(sid) for sid in library_ids if sid]
                if library_ids_int:
                    in_ph = ', '.join(['%s'] * len(library_ids_int))
                    start_dt = datetime.combine(first_of_month, datetime.min.time())
                    end_dt = datetime.combine(end_of_month, datetime.max.time())
                    q = f"""
                        SELECT iu.institution_id, la.user_id, la.created_at
                        FROM logins la
                        JOIN institution_user iu ON la.user_id = iu.user_id
                        WHERE iu.institution_id IN ({in_ph})
                          AND la.created_at BETWEEN %s AND %s
                    """
                    cursor_lib.execute(q, library_ids_int + [start_dt, end_dt])
                    al_rows = cursor_lib.fetchall()
                    cursor_lib.close()
                    conn_lib.close()
                    al_counted = defaultdict(set)
                    for row in al_rows:
                        iid = row.get("institution_id")
                        uid = row.get("user_id")
                        created = row.get("created_at")
                        created_date = created.date() if hasattr(created, 'date') else created
                        pe = period_end_by_lib_id.get(iid, end_of_month)
                        if created_date <= pe:
                            al_counted[iid].add(uid)
                    al_mtd_by_lib_id = {iid: len(s) for iid, s in al_counted.items()}
                    for asl_id, lib_id in lib_id_by_asl.items():
                        al_mtd_by_asl_id[asl_id] = al_mtd_by_lib_id.get(lib_id, 0)
        except Exception as e:
            print(f"Profile AL MTD per school: {e}")

    active_scholarship_schools = []
    inactive_scholarship_schools = []
    schools_without_scholarship = []
    school_duration_map = duration_map
    schools_with_scholarship_ids = set()

    if preselected_asl_ids:
        try:
            conn = get_ruzivo_conn()
            cursor = conn.cursor()
            placeholders = ','.join(['%s'] * len(preselected_asl_ids))
            scholarship_query = f"""
                SELECT ts.school_id, sc.school_name, sc.school_province, sc.school_city,
                       ts.awarded_at, ts.awarded_by, ts.school_type, ts.scholarship_type,
                       ts.scholarship_grades, ts.expiry_date, ts.duration, ts.active,
                       ts.date_modified
                FROM tblscholarships_schools ts
                LEFT JOIN tblschools sc ON sc.school_id = ts.school_id
                WHERE ts.school_id IN ({placeholders})
            """
            cursor.execute(scholarship_query, preselected_asl_ids)
            scholarship_results = cursor.fetchall()
            cursor.close()
            conn.close()

            for row in scholarship_results:
                if isinstance(row, dict):
                    school_data = {
                        "id": row.get("school_id"), "name": row.get("school_name"), "province": row.get("school_province"), "city": row.get("school_city"),
                        "awarded_at": row.get("awarded_at"), "awarded_by": row.get("awarded_by"), "school_type": row.get("school_type"),
                        "scholarship_type": row.get("scholarship_type"), "scholarship_grades": row.get("scholarship_grades"),
                        "expiry_date": row.get("expiry_date"), "duration": row.get("duration"), "active": normalize_active_flag(row.get("active")),
                        "date_modified": row.get("date_modified"),
                        "total_months": duration_map.get(row.get("school_id"), {}).get("total_months", 0),
                        "first_awarded_at": duration_map.get(row.get("school_id"), {}).get("first_awarded_at", "N/A")
                    }
                    if school_data["expiry_date"]:
                        if isinstance(school_data["expiry_date"], str):
                            try:
                                school_data["expiry_date"] = datetime.strptime(school_data["expiry_date"], "%Y-%m-%d").date()
                            except Exception:
                                try:
                                    school_data["expiry_date"] = datetime.strptime(school_data["expiry_date"], "%Y-%m-%d %H:%M:%S").date()
                                except Exception:
                                    pass
                        elif hasattr(school_data["expiry_date"], 'date'):
                            school_data["expiry_date"] = school_data["expiry_date"].date()
                else:
                    school_active_value = row[11]
                    school_data = {
                        "id": row[0], "name": row[1], "province": row[2], "city": row[3],
                        "awarded_at": row[4], "awarded_by": row[5], "school_type": row[6], "scholarship_type": row[7],
                        "scholarship_grades": row[8], "expiry_date": row[9], "duration": row[10],
                        "active": normalize_active_flag(school_active_value), "date_modified": row[12],
                        "total_months": duration_map.get(row[0], {}).get("total_months", 0),
                        "first_awarded_at": duration_map.get(row[0], {}).get("first_awarded_at", "N/A")
                    }
                duration_info = duration_map.get(school_data["id"], {})
                school_data["is_over_12_months"] = duration_info.get("is_over_12_months", False)
                school_data["months_left_until_12"] = None
                fa_raw = duration_info.get("first_awarded_at")
                if fa_raw and fa_raw != "N/A":
                    fa_d = fa_raw.date() if hasattr(fa_raw, 'date') else fa_raw
                    if isinstance(fa_d, date) and fa_d <= today:
                        months_elapsed = (today.year - fa_d.year) * 12 + (today.month - fa_d.month)
                        months_elapsed = max(0, months_elapsed)
                        if months_elapsed < 12:
                            school_data["months_left_until_12"] = 12 - months_elapsed
                if school_data["expiry_date"]:
                    if isinstance(school_data["expiry_date"], str):
                        try:
                            school_data["expiry_date"] = datetime.strptime(school_data["expiry_date"], "%Y-%m-%d").date()
                        except Exception:
                            try:
                                school_data["expiry_date"] = datetime.strptime(school_data["expiry_date"], "%Y-%m-%d %H:%M:%S").date()
                            except Exception:
                                pass
                    elif hasattr(school_data["expiry_date"], 'date'):
                        school_data["expiry_date"] = school_data["expiry_date"].date()
                _dur = duration_map.get(school_data["id"], {})
                _fa = school_data.get("first_awarded_at") or _dur.get("first_awarded_at")
                school_data["first_awarded_2025"] = "—"
                school_data["days_left_until_12"] = None
                if _fa and _fa != "N/A":
                    _fa_d = _fa.date() if hasattr(_fa, 'date') else _fa
                    if isinstance(_fa_d, date):
                        if _fa_d.year >= 2025:
                            school_data["first_awarded_2025"] = _fa_d.strftime("%Y-%m-%d")
                        if _fa_d <= today:
                            _days_elapsed = (today - _fa_d).days
                            if _days_elapsed < 365:
                                school_data["days_left_until_12"] = 365 - _days_elapsed
                _sid = school_data.get("id")
                school_data["asl_mtd"] = asl_mtd_by_school.get(_sid, 0)
                school_data["al_mtd"] = al_mtd_by_asl_id.get(_sid, 0)
                school_data["total_mtd"] = school_data["asl_mtd"] + school_data["al_mtd"]
                if school_data["active"]:
                    is_expired = school_data["expiry_date"] and school_data["expiry_date"] < today
                    if is_expired and school_data["id"]:
                        try:
                            first_day_of_month = today.replace(day=1)
                            first_day_next_month = (today.replace(year=today.year + 1, month=1, day=1) if today.month == 12 else today.replace(month=today.month + 1, day=1))
                            conn_sub = get_ruzivo_conn()
                            cursor_sub = conn_sub.cursor()
                            cursor_sub.execute(
                                "SELECT COUNT(DISTINCT p.student_id) FROM tblpoints_purchase p JOIN tblstudents s ON p.student_id = s.student_id WHERE s.school_id = %s AND p.expiry_date >= %s AND p.expiry_date < %s",
                                (school_data["id"], first_day_of_month, first_day_next_month)
                            )
                            result = cursor_sub.fetchone()
                            cursor_sub.close()
                            conn_sub.close()
                            school_data["active_subscription_count"] = (result[0] if isinstance(result, (tuple, list)) else result.get("active_count", 0)) or 0 if result else 0
                        except Exception:
                            school_data["active_subscription_count"] = 0
                    else:
                        school_data["active_subscription_count"] = 0
                    active_scholarship_schools.append(school_data)
                else:
                    school_data["active_subscription_count"] = 0
                    if school_data["id"]:
                        try:
                            first_day_of_month = today.replace(day=1)
                            first_day_next_month = (today.replace(year=today.year + 1, month=1, day=1) if today.month == 12 else today.replace(month=today.month + 1, day=1))
                            conn_sub = get_ruzivo_conn()
                            cursor_sub = conn_sub.cursor()
                            cursor_sub.execute(
                                "SELECT COUNT(DISTINCT p.student_id) FROM tblpoints_purchase p JOIN tblstudents s ON p.student_id = s.student_id WHERE s.school_id = %s AND p.expiry_date >= %s AND p.expiry_date < %s",
                                (school_data["id"], first_day_of_month, first_day_next_month)
                            )
                            result = cursor_sub.fetchone()
                            cursor_sub.close()
                            conn_sub.close()
                            school_data["active_subscription_count"] = (result[0] if isinstance(result, (tuple, list)) else result.get("active_count", 0)) or 0 if result else 0
                        except Exception:
                            school_data["active_subscription_count"] = 0
                    inactive_scholarship_schools.append(school_data)
                schools_with_scholarship_ids.add(school_data["id"])

            for s_dict in existing_school_dicts:
                asl_id = s_dict.get("asl_school_id")
                if asl_id and int(asl_id) not in schools_with_scholarship_ids:
                    try:
                        _aid = int(asl_id)
                    except (TypeError, ValueError):
                        _aid = asl_id
                    _asl_mtd = asl_mtd_by_school.get(_aid, 0)
                    _al_mtd = al_mtd_by_asl_id.get(_aid, 0)
                    schools_without_scholarship.append({
                        "id": asl_id, "name": s_dict.get("school_name", "Unknown School"), "active": False, "scholarship_type": "No Scholarship",
                        "active_subscription_count": 0, "first_awarded_2025": "—", "days_left_until_12": None,
                        "asl_mtd": _asl_mtd, "al_mtd": _al_mtd, "total_mtd": _asl_mtd + _al_mtd,
                    })
        except Exception as e:
            print(f"Error fetching scholarship data: {e}")

    total_active_sub_students = 0
    schools_with_active_subs_count = 0
    for s in active_scholarship_schools + inactive_scholarship_schools:
        subs = s.get("active_subscription_count", 0)
        if subs > 0:
            total_active_sub_students += subs
            schools_with_active_subs_count += 1

    count_valid_scholarship = 0
    count_elapsed_scholarship = 0
    all_processed = active_scholarship_schools + inactive_scholarship_schools
    for s in all_processed:
        sid = s.get("id")
        is_expired = s.get("expiry_date") and s.get("expiry_date") < today
        is_active = s.get("active", False)
        duration_info = school_duration_map.get(sid, {})
        is_over_12 = duration_info.get("is_over_12_months", False)
        total_m = duration_info.get("total_months", 0)
        if not is_active or is_expired or is_over_12 or (total_m > 12):
            count_elapsed_scholarship += 1
        elif is_active:
            count_valid_scholarship += 1
    count_no_scholarship = len(schools_without_scholarship)

    try:
        active_schools_asl_mtd_total = sum(
            s.get("asl_mtd", 0) for s in active_scholarship_schools
            if s.get("expiry_date") is None or s.get("expiry_date") >= today
        )
    except (TypeError, ValueError):
        active_schools_asl_mtd_total = 0
    if active_schools_asl_mtd_total is None:
        active_schools_asl_mtd_total = 0

    truly_active_schools = [s for s in active_scholarship_schools if s.get("expiry_date") is None or s.get("expiry_date") >= today]
    expired_schools = [s for s in active_scholarship_schools if s.get("expiry_date") and s.get("expiry_date") < today]

    library_id_by_asl_id = {}
    for _school in existing_school_dicts:
        asl_value = _school.get("asl_school_id")
        lib_value = _school.get("library_school_id")
        if asl_value is None:
            continue
        library_id_by_asl_id[str(asl_value)] = lib_value

    def _serialize_school(s):
        school_id = s.get("id")
        library_id = library_id_by_asl_id.get(str(school_id))
        return {
            "id": school_id,
            "asl_school_id": school_id,
            "library_id": library_id,
            "name": s.get("name") or "—",
            "first_awarded_2025": s.get("first_awarded_2025") or "—",
            "days_left_until_12": s.get("days_left_until_12"),
            "asl_mtd": s.get("asl_mtd", 0),
            "al_mtd": s.get("al_mtd", 0),
            "total_mtd": s.get("total_mtd", 0),
        }

    return {
        "truly_active_schools": [_serialize_school(s) for s in truly_active_schools],
        "expired_schools": [_serialize_school(s) for s in expired_schools],
        "schools_without_scholarship": [_serialize_school(s) for s in schools_without_scholarship],
        "count_valid_scholarship": count_valid_scholarship,
        "count_elapsed_scholarship": count_elapsed_scholarship,
        "count_no_scholarship_schools": count_no_scholarship,
        "active_schools_asl_mtd_total": active_schools_asl_mtd_total,
        "total_active_sub_students": total_active_sub_students,
        "schools_with_active_subs_count": schools_with_active_subs_count,
        "smartlearning_champ_mtd": smartlearning_champ_mtd,
    }


@app.route('/api/profile/<username>/school-tables', methods=['GET'])
@login_required
def api_profile_school_tables(username):
    """Return school tables data for profile page (loaded async with loading indicators)."""
    try:
        data = _get_profile_school_tables_data(username)
        return jsonify(data)
    except Exception as e:
        logging.exception("profile school-tables API: %s", e)
        return jsonify({"error": str(e)}), 500


def _visit_api_error(error_code, message, status_code=400, data=None):
    payload = {
        "success": False,
        "error": error_code,
        "message": message,
    }
    if data is not None:
        payload["data"] = data
    return jsonify(payload), status_code


def _visit_api_iso_utc(value):
    if not value:
        return None
    return value.replace(microsecond=0).isoformat() + 'Z'


def _require_json_request():
    if not request.is_json:
        return _visit_api_error(
            'ValidationError',
            'Content-Type must be application/json',
            400
        )
    return None


def _is_visit_admin_viewer():
    role = (getattr(current_user, 'userRole', '') or '').strip().lower()
    return role in {
        'admin',
        'super admin',
        'superadmin',
        'super-admin',
        'manager',
    }


def _visit_admin_metadata(log):
    return {
        'checkin_device_install_id': log.checkin_device_install_id,
        'checkout_device_install_id': log.checkout_device_install_id,
        'checkin_latitude': log.checkin_latitude,
        'checkin_longitude': log.checkin_longitude,
        'checkout_latitude': log.checkout_latitude,
        'checkout_longitude': log.checkout_longitude,
        'location_source': log.location_source,
        'location_text': log.location_text,
        'checkout_location_source': log.checkout_location_source,
        'checkout_location_text': log.checkout_location_text,
        'checkin_device_platform': log.checkin_device_platform,
    }


def _apply_visit_context_to_log(log, vc):
    """Apply optional visit_context fields to a SchoolVisitLog instance. Returns an error string or None."""
    if vc is None:
        return None
    if not isinstance(vc, dict):
        return 'visit_context must be an object'
    if 'grades' in vc:
        raw = vc.get('grades')
        if raw is None or raw == []:
            log.visit_grades_json = None
        else:
            if not isinstance(raw, list):
                return 'grades must be a list of integers or null'
            ints = []
            for g in raw:
                try:
                    gi = int(g)
                except (TypeError, ValueError):
                    return 'grades must be a list of integers'
                if gi < 4 or gi > 13:
                    return 'each grade must be between 4 and 13'
                ints.append(gi)
            uniq = sorted(set(ints))
            log.visit_grades_json = json.dumps(uniq)
    if 'teacher_name' in vc:
        raw = vc.get('teacher_name')
        if raw is None:
            log.teacher_name = None
        else:
            s = str(raw).strip()
            if not s:
                log.teacher_name = None
            elif len(s) > 255:
                return 'teacher_name must be 255 characters or fewer'
            else:
                log.teacher_name = s
    if 'teacher_contact' in vc:
        raw = vc.get('teacher_contact')
        if raw is None:
            log.teacher_contact = None
        else:
            s = str(raw).strip()
            if not s:
                log.teacher_contact = None
            elif len(s) > 255:
                return 'teacher_contact must be 255 characters or fewer'
            else:
                log.teacher_contact = s
    return None


def _visit_dict_for_viewer(log, *, include_admin_fields=False):
    payload = {
        'visit_id': log.id,
        'username': log.username,
        'asl_school_id': log.asl_school_id,
        'library_id': log.library_id,
        'school_name': log.school_name,
        'status': log.status,
        'checkin_at': _visit_api_iso_utc(log.checkin_at),
        'checkout_at': _visit_api_iso_utc(log.checkout_at),
        'duration_minutes': log.duration_minutes,
        'notes': log.notes,
        'visit_context': log.visit_context_dict(),
    }
    if include_admin_fields:
        payload['admin_metadata'] = _visit_admin_metadata(log)
    return payload


def _school_visit_log_report_row(
    log,
    *,
    champion_name=None,
    champion_province=None,
    champion_email=None,
    include_admin_fields=False,
):
    row = {
        "visit_id": log.id,
        "champion": champion_name or log.username,
        "username": log.username,
        "champion_email": champion_email,
        "province": champion_province,
        "school_name": log.school_name,
        "asl_school_id": log.asl_school_id,
        "library_id": log.library_id,
        "status": log.status,
        "checkin_at": _visit_api_iso_utc(log.checkin_at),
        "checkout_at": _visit_api_iso_utc(log.checkout_at),
        "duration_minutes": log.duration_minutes,
        "visit_grades": log.visit_grades_json or "",
        "teacher_name": log.teacher_name or "",
        "teacher_contact": log.teacher_contact or "",
    }
    if include_admin_fields:
        row["admin_metadata"] = _visit_admin_metadata(log)
        row.update({
            "location_source": log.location_source,
            "location_text": log.location_text,
            "checkin_latitude": log.checkin_latitude,
            "checkin_longitude": log.checkin_longitude,
            "checkout_location_source": log.checkout_location_source,
            "checkout_location_text": log.checkout_location_text,
            "checkout_latitude": log.checkout_latitude,
            "checkout_longitude": log.checkout_longitude,
            "checkin_device_install_id": log.checkin_device_install_id,
            "checkout_device_install_id": log.checkout_device_install_id,
        })
    return row


def _validate_device_payload(device, *, required=True):
    if not isinstance(device, dict):
        if required:
            return None, None, "device object is required"
        return None, None, None

    install_id = str((device.get('install_id') or '')).strip()
    platform = str((device.get('platform') or 'web')).strip().lower() or 'web'

    if required and not install_id:
        return None, None, "device.install_id is required"
    if install_id and len(install_id) > 64:
        return None, None, "device.install_id must be 64 characters or fewer"
    if platform and len(platform) > 32:
        return None, None, "device.platform must be 32 characters or fewer"

    return install_id or None, platform, None


def _validate_location_payload(location):
    if not isinstance(location, dict):
        return None, None, None, None, "location object is required"

    latitude_raw = location.get("latitude")
    longitude_raw = location.get("longitude")
    location_text = (location.get("location_text") or "").strip() or None
    location_source = (location.get("source") or "unknown").strip().lower() or "unknown"

    lat = None
    lng = None

    if latitude_raw is not None or longitude_raw is not None:
        try:
            lat = float(latitude_raw)
            lng = float(longitude_raw)
        except (TypeError, ValueError):
            return None, None, None, None, "location.latitude and location.longitude must be numeric values"
        if lat < -90 or lat > 90:
            return None, None, None, None, "location.latitude must be between -90 and 90"
        if lng < -180 or lng > 180:
            return None, None, None, None, "location.longitude must be between -180 and 180"

    if lat is None and lng is None and not location_text:
        return None, None, None, None, "location must include either latitude/longitude or location_text"

    return lat, lng, location_text, location_source, None


@app.route('/api/profile/school-visit/checkin', methods=['POST'])
@login_required
def api_profile_school_visit_checkin():
    json_error = _require_json_request()
    if json_error:
        return json_error

    data = request.get_json() or {}
    asl_school_id = str((data.get('asl_school_id') or '')).strip()
    library_id = str((data.get('library_id') or '')).strip()
    school_name = str((data.get('school_name') or '')).strip()

    if not asl_school_id or not library_id or not school_name:
        return _visit_api_error('ValidationError', 'asl_school_id, library_id and school_name are required', 400)

    lat, lng, location_text, location_source, location_error = _validate_location_payload(data.get('location'))
    if location_error:
        return _visit_api_error('ValidationError', location_error, 400)

    install_id, platform, device_error = _validate_device_payload(data.get('device'), required=True)
    if device_error:
        return _visit_api_error('ValidationError', device_error, 400)

    active_visit = SchoolVisitLog.query.filter_by(
        user_id=current_user.id,
        status='checked_in'
    ).order_by(SchoolVisitLog.checkin_at.desc()).first()
    if active_visit:
        return _visit_api_error(
            'ActiveSessionExists',
            'User already checked in at another school.',
            409,
            {
                'active_visit_id': active_visit.id,
                'school_name': active_visit.school_name,
                'checkin_at': _visit_api_iso_utc(active_visit.checkin_at),
            }
        )

    user_agent = (request.headers.get('User-Agent') or '').strip() or None
    if user_agent and len(user_agent) > 512:
        user_agent = user_agent[:512]

    visit = SchoolVisitLog(
        user_id=current_user.id,
        username=current_user.username,
        asl_school_id=asl_school_id,
        library_id=library_id,
        school_name=school_name,
        status='checked_in',
        checkin_at=datetime.utcnow(),
        checkin_latitude=lat,
        checkin_longitude=lng,
        location_text=location_text,
        location_source=location_source,
        checkin_device_install_id=install_id,
        checkin_device_platform=platform,
        checkin_user_agent=user_agent,
    )
    if 'visit_context' in data:
        ctx_err = _apply_visit_context_to_log(visit, data.get('visit_context'))
        if ctx_err:
            return _visit_api_error('ValidationError', ctx_err, 400)
    db.session.add(visit)
    db.session.commit()

    include_admin_fields = _is_visit_admin_viewer()
    return jsonify({
        'success': True,
        'message': 'Checked in successfully.',
        'data': _visit_dict_for_viewer(visit, include_admin_fields=include_admin_fields)
    }), 201


@app.route('/api/profile/school-visit/active', methods=['PATCH'])
@login_required
def api_profile_school_visit_active_patch():
    json_error = _require_json_request()
    if json_error:
        return json_error

    data = request.get_json() or {}
    active_visit = SchoolVisitLog.query.filter_by(
        user_id=current_user.id,
        status='checked_in'
    ).order_by(SchoolVisitLog.checkin_at.desc()).first()
    if not active_visit:
        return _visit_api_error('NoActiveSession', 'No active check-in found for this user.', 404)

    ctx_err = _apply_visit_context_to_log(active_visit, data.get('visit_context'))
    if ctx_err:
        return _visit_api_error('ValidationError', ctx_err, 400)
    db.session.commit()

    include_admin_fields = _is_visit_admin_viewer()
    return jsonify({
        'success': True,
        'message': 'Visit updated.',
        'data': _visit_dict_for_viewer(active_visit, include_admin_fields=include_admin_fields)
    }), 200


@app.route('/api/profile/school-visit/checkout', methods=['POST'])
@login_required
def api_profile_school_visit_checkout():
    json_error = _require_json_request()
    if json_error:
        return json_error

    data = request.get_json() or {}
    notes = (data.get('notes') or '').strip() or None

    lat, lng, location_text, location_source, location_error = _validate_location_payload(data.get('location'))
    if location_error:
        return _visit_api_error('ValidationError', location_error, 400)

    install_id, platform, device_error = _validate_device_payload(data.get('device'), required=True)
    if device_error:
        return _visit_api_error('ValidationError', device_error, 400)

    active_visit = SchoolVisitLog.query.filter_by(
        user_id=current_user.id,
        status='checked_in'
    ).order_by(SchoolVisitLog.checkin_at.desc()).first()
    if not active_visit:
        return _visit_api_error('NoActiveSession', 'No active check-in found for this user.', 404)

    if active_visit.checkin_device_install_id and install_id and active_visit.checkin_device_install_id != install_id:
        return _visit_api_error(
            'DeviceMismatch',
            'Checkout device does not match the device used for check-in.',
            409,
            {
                'checkin_device_install_id': active_visit.checkin_device_install_id,
                'checkout_device_install_id': install_id,
            }
        )

    now = datetime.utcnow()
    duration_minutes = None
    if active_visit.checkin_at:
        duration_minutes = int(max(0, (now - active_visit.checkin_at).total_seconds() // 60))

    active_visit.checkout_at = now
    active_visit.duration_minutes = duration_minutes
    active_visit.status = 'checked_out'
    active_visit.notes = notes
    active_visit.checkout_latitude = lat
    active_visit.checkout_longitude = lng
    active_visit.checkout_location_text = location_text
    active_visit.checkout_location_source = location_source
    active_visit.checkout_device_install_id = install_id
    if platform:
        active_visit.checkin_device_platform = active_visit.checkin_device_platform or platform
    db.session.commit()

    return jsonify({
        'success': True,
        'message': 'Checked out successfully.',
        'data': active_visit.to_dict()
    }), 200


@app.route('/api/profile/<username>/school-visit/status', methods=['GET'])
@login_required
def api_profile_school_visit_status(username):
    target_user = User.query.filter_by(username=username).first()
    if not target_user:
        return _visit_api_error('NotFound', 'User not found', 404)

    if current_user.username != username and current_user.userRole != 'Admin':
        return _visit_api_error('Unauthorized', 'Unauthorized', 403)

    active_visit = SchoolVisitLog.query.filter_by(
        user_id=target_user.id,
        status='checked_in'
    ).order_by(SchoolVisitLog.checkin_at.desc()).first()

    last_visit = SchoolVisitLog.query.filter_by(
        user_id=target_user.id,
        status='checked_out'
    ).order_by(SchoolVisitLog.checkout_at.desc(), SchoolVisitLog.id.desc()).first()

    return jsonify({
        'success': True,
        'data': {
            'active_visit': active_visit.to_dict() if active_visit else None,
            'last_visit': last_visit.to_dict() if last_visit else None,
        }
    }), 200


@app.route('/profile/<username>', methods=['GET', 'POST'])
@login_required
def profile(username):
    user = db.first_or_404(sa.select(User).where(User.username == username))
    form = ChampionSchoolForm()
    user_province = user.province

    # Fetch existing ChampionSchool record
    champ = ChampionSchool.query.filter_by(
        firstname=user.firstname,
        lastname=user.lastname,
        province=user.province
    ).first()

    # ✅ Load previously selected schools (stored as list of dicts)
    existing_school_dicts = []
    if champ and champ.schools:
        try:
            existing_school_dicts = json.loads(champ.schools)
        except Exception as e:
            print("Error decoding champ.schools JSON:", e)
            existing_school_dicts = []

    # ✅ Extract ASL school IDs for preselection
    existing_asl_ids = set()
    try:
        existing_asl_ids = {
            int(s["asl_school_id"]) for s in existing_school_dicts if "asl_school_id" in s
        }
    except Exception as e:
        print("Error parsing existing_asl_ids:", e)

    # ✅ Handle form submission
    if request.method == 'POST' and form.validate_on_submit():
        try:
            # Get selected schools from hidden JSON input
            raw_school_data = request.form.get('selected_school_data', '[]')
            selected_school_data = json.loads(raw_school_data)

            # Save to DB
            if champ:
                champ.schools = json.dumps(selected_school_data)
            else:
                champ = ChampionSchool(
                    firstname=form.firstname.data,
                    lastname=form.lastname.data,
                    province=form.province.data,
                    schools=json.dumps(selected_school_data)
                )
                db.session.add(champ)

            db.session.commit()
            flash('Champion school updated successfully!', 'success')
            return redirect(url_for('user', username=username))

        except Exception as e:
            db.session.rollback()
            flash(f'Error saving champion schools: {str(e)}', 'danger')
            return redirect(url_for('user', username=username))

    # ✅ Pre-fill form fields
    form.firstname.data = user.firstname
    form.lastname.data = user.lastname
    form.province.data = user.province

    # For JS and template rendering
    preselected_asl_ids = [s["asl_school_id"] for s in existing_school_dicts if "asl_school_id" in s]
    count_user_schools = len(preselected_asl_ids)
    smartlearning_champ_mtd = 0

    today = datetime.today().date()

    # School tables and stats loaded async via API; pass empty/zeros for initial render
    schools_for_template = []
    for s in existing_school_dicts:
        if "asl_school_id" in s and "school_name" in s:
            schools_for_template.append({"id": s["asl_school_id"], "name": s["school_name"]})

    active_scholarship_schools = []
    inactive_scholarship_schools = []
    schools_without_scholarship = []
    school_duration_map = {}
    count_valid_scholarship = 0
    count_elapsed_scholarship = 0
    count_no_scholarship_schools = 0
    active_schools_asl_mtd_total = 0
    total_active_sub_students = 0
    schools_with_active_subs_count = 0

    # Heavy computation moved to _get_profile_school_tables_data; data loaded via GET /api/profile/<username>/school-tables

    # (removed STEP 1 duration_map, STEP 2 filter, STEP 3 ASL MTD, AL MTD, scholarship query, counts)

    return render_template('userProfile.html',
                           count_valid_scholarship=count_valid_scholarship,
                           count_elapsed_scholarship=count_elapsed_scholarship,
                           count_no_scholarship_schools=count_no_scholarship_schools,
                           user=user,
                            form=form,
                            champ=champ,
                            schools_for_template=schools_for_template,   # ✅ new var for frontend
                            preselected_school_ids=preselected_asl_ids,
                            count_user_schools=count_user_schools,
                            smartlearning_champ_mtd=smartlearning_champ_mtd,
                            active_schools_asl_mtd_total=active_schools_asl_mtd_total,
                           active_scholarship_schools=active_scholarship_schools,
                           inactive_scholarship_schools=inactive_scholarship_schools,
                           schools_without_scholarship=schools_without_scholarship,
                           school_duration_map=school_duration_map,
                           current_date=today,  # ✅ Pass current date for expiry comparison
                           total_active_sub_students=total_active_sub_students,
                           schools_with_active_subs_count=schools_with_active_subs_count,
                           title="Profile")



@app.route('/user/<username>', methods=['GET', 'POST'])
@login_required
def user(username):
    user = db.first_or_404(sa.select(User).where(User.username == username))
    form = ChampionSchoolForm()
    user_province = user.province

    # Fetch existing ChampionSchool record
    champ = ChampionSchool.query.filter_by(
        firstname=user.firstname,
        lastname=user.lastname,
        province=user.province
    ).first()

    # ✅ Load previously selected schools (stored as list of dicts)
    existing_school_dicts = []
    if champ and champ.schools:
        try:
            existing_school_dicts = json.loads(champ.schools)
        except Exception as e:
            print("Error decoding champ.schools JSON:", e)
            existing_school_dicts = []

    # ✅ Extract ASL school IDs for preselection
    existing_asl_ids = set()
    try:
        existing_asl_ids = {
            int(s["asl_school_id"]) for s in existing_school_dicts if "asl_school_id" in s
        }
    except Exception as e:
        print("Error parsing existing_asl_ids:", e)

    # ✅ Handle form submission
    if request.method == 'POST' and form.validate_on_submit():
        try:
            # Get selected schools from hidden JSON input
            raw_school_data = request.form.get('selected_school_data', '[]')
            selected_school_data = json.loads(raw_school_data)

            # Save to DB
            if champ:
                champ.schools = json.dumps(selected_school_data)
            else:
                champ = ChampionSchool(
                    firstname=form.firstname.data,
                    lastname=form.lastname.data,
                    province=form.province.data,
                    schools=json.dumps(selected_school_data)
                )
                db.session.add(champ)

            db.session.commit()
            flash('Champion school updated successfully!', 'success')
            return redirect(url_for('user', username=username))

        except Exception as e:
            db.session.rollback()
            flash(f'Error saving champion schools: {str(e)}', 'danger')
            return redirect(url_for('user', username=username))

    # ✅ Pre-fill form fields
    form.firstname.data = user.firstname
    form.lastname.data = user.lastname
    form.province.data = user.province

    # For JS and template rendering
    preselected_asl_ids = [s["asl_school_id"] for s in existing_school_dicts if "asl_school_id" in s]
    count_user_schools = len(preselected_asl_ids)
    smartlearning_champ_mtd = 0

    # ✅ Date range (default: this month until today)
    today = datetime.today().date()
    start_date = today.replace(day=1)

    # ✅ Batch query instead of loop - much faster
    if preselected_asl_ids:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()

        # Create placeholders for IN clause
        placeholders = ','.join(['%s'] * len(preselected_asl_ids))
        query = f"""
            SELECT SUM(student_count) AS total_count
            FROM (
                SELECT COUNT(*) AS student_count
                FROM vwstudent
                WHERE school_id IN ({placeholders})
                AND last_login BETWEEN %s AND %s
            ) AS subquery
        """
        
        # Execute with all IDs at once
        cursor.execute(query, (*preselected_asl_ids, start_date, today))
        row = cursor.fetchone()
        cursor.close()
        conn.close()

        if row:
            # ✅ Handle both tuple and dict
            if isinstance(row, dict):
                smartlearning_champ_mtd = row.get("total_count", 0) or 0
            else:
                smartlearning_champ_mtd = row[0] or 0

    # ✅ Prepare clean school list for template rendering
    schools_for_template = []
    for s in existing_school_dicts:
        if "asl_school_id" in s and "school_name" in s:
            schools_for_template.append({
                "id": s["asl_school_id"],
                "name": s["school_name"]
            })

    # Final total after loop
    # print(f"\nTotal SmartLearning Champs MTD: {smartlearning_champ_mtd}")

    return render_template(
        "user.html",
        user=user,
        form=form,
        champ=champ,
        schools_for_template=schools_for_template,   # ✅ new var for frontend
        preselected_school_ids=preselected_asl_ids,
        count_user_schools=count_user_schools,
        smartlearning_champ_mtd=smartlearning_champ_mtd,
        title="Profile"
    )




def _mtd_month_bounds(today=None):
    """Return (first_of_month, end_of_month) for current month (date objects)."""
    if today is None:
        today = datetime.today().date()
    first_of_month = today.replace(day=1)
    _, last_day = calendar.monthrange(first_of_month.year, first_of_month.month)
    end_of_month = first_of_month.replace(day=last_day)
    return first_of_month, end_of_month


def _mtd_period_end(first_awarded_date, end_of_month):
    """Return period end date for MTD: min(end_of_month, first_awarded + 365 days). If no first_awarded, return end_of_month."""
    if not first_awarded_date:
        return end_of_month
    fa = first_awarded_date.date() if hasattr(first_awarded_date, 'date') else first_awarded_date
    if not isinstance(fa, date):
        return end_of_month
    scholarship_end = fa + timedelta(days=365)
    return min(end_of_month, scholarship_end)


@app.route('/api/champion/<username>/asl_active', methods=['GET'])
@login_required
def get_champion_asl_active(username):
    """Return total ASL active learners for all schools under a champion. Uses current month with per-school period (1st to min(end of month, first_awarded+365)). Ignores date params for MTD count."""
    user = db.first_or_404(sa.select(User).where(User.username == username))

    today = datetime.today().date()
    first_of_month, end_of_month = _mtd_month_bounds(today)

    champ = ChampionSchool.query.filter_by(
        firstname=user.firstname,
        lastname=user.lastname,
        province=user.province
    ).first()

    if not champ or not champ.schools:
        return jsonify({
            "username": username,
            "asl_student_count": 0,
            "school_count": 0
        })

    try:
        schools_data = json.loads(champ.schools)
    except Exception as e:
        print("Error decoding champ.schools JSON:", e)
        return jsonify({"error": "Invalid champ.schools JSON"}), 400

    asl_school_ids = [s.get("asl_school_id") for s in schools_data if "asl_school_id" in s]
    if not asl_school_ids:
        return jsonify({
            "username": username,
            "asl_student_count": 0,
            "school_count": 0
        })

    # Get first_awarded_at per school from Ruzivo
    period_end_by_school = {}
    conn = None
    cursor = None
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        placeholders = ','.join(['%s'] * len(asl_school_ids))
        duration_query = f"""
            SELECT x.school_id, MIN(x.awarded_at) AS first_awarded_at
            FROM tblscholarships_schools x
            WHERE x.school_id IN ({placeholders}) AND x.awarded_at >= %s
            GROUP BY x.school_id
        """
        cursor.execute(duration_query, (*asl_school_ids, '2025-01-01'))
        for row in cursor.fetchall():
            sid = row[0] if isinstance(row, (tuple, list)) else row.get("school_id")
            fa = row[1] if isinstance(row, (tuple, list)) else row.get("first_awarded_at")
            period_end_by_school[sid] = _mtd_period_end(fa, end_of_month)
    except Exception as e:
        logging.error("ASL duration query for champion %s: %s", username, e)
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

    for sid in asl_school_ids:
        if sid not in period_end_by_school:
            period_end_by_school[sid] = end_of_month

    # One query: all logins in month for these schools; then aggregate in Python by per-school period_end
    smartlearning_count = 0
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        placeholders = ','.join(['%s'] * len(asl_school_ids))
        query = f"""
            SELECT school_id, student_id, last_login
            FROM vwstudent
            WHERE school_id IN ({placeholders})
            AND last_login BETWEEN %s AND %s
        """
        cursor.execute(query, (*asl_school_ids, first_of_month, end_of_month))
        rows = cursor.fetchall()
        # Per-school distinct students within that school's period
        counted = defaultdict(set)
        for row in rows:
            sid = row[0] if isinstance(row, (tuple, list)) else row.get("school_id")
            student_id = row[1] if isinstance(row, (tuple, list)) else row.get("student_id")
            last_login = row[2] if isinstance(row, (tuple, list)) else row.get("last_login")
            login_date = last_login.date() if hasattr(last_login, 'date') else last_login
            pe = period_end_by_school.get(sid, end_of_month)
            if login_date <= pe:
                counted[sid].add(student_id)
        smartlearning_count = sum(len(s) for s in counted.values())
    except Exception as e:
        logging.error("ASL query error for champion %s: %s", username, e)
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

    return jsonify({
        "username": username,
        "asl_student_count": smartlearning_count,
        "school_count": len(asl_school_ids)
    })


@app.route('/api/champion/<username>/al_active', methods=['GET'])
@login_required
def get_champion_al_active(username):
    """Return total AL active learners for all schools under a champion this month."""
    user = db.first_or_404(sa.select(User).where(User.username == username))

    # Fetch Champion record
    champ = ChampionSchool.query.filter_by(
        firstname=user.firstname,
        lastname=user.lastname,
        province=user.province
    ).first()

    if not champ or not champ.schools:
        return jsonify({
            "username": username,
            "al_active_total": 0,
            "school_breakdown": []
        })


    try:
        schools_data = json.loads(champ.schools)
        # print(schools_data)
    except Exception as e:
        print("Error decoding champ.schools JSON:", e)
        return jsonify({"error": "Invalid champ.schools JSON"}), 400

    # Extract AL school IDs
    al_school_ids = [s.get("library_school_id") for s in schools_data if "library_school_id" in s]
    # print(al_school_ids)

    if not al_school_ids:
        return jsonify({
            "username": username,
            "al_active_total": 0,
            "school_breakdown": []
        })

    today = datetime.today().date()
    start_date = today.replace(day=1)

    conn = get_direct_library_conn()
    cursor = conn.cursor()

    total_active = 0
    school_breakdown = []

    # ✅ Batch query with GROUP BY instead of loop - much faster
    placeholders = ','.join(['%s'] * len(al_school_ids))
    query = f"""
        SELECT iu.institution_id, COUNT(DISTINCT la.user_id) AS al_active
        FROM logins la
        JOIN institution_user iu ON la.user_id = iu.user_id
        WHERE iu.institution_id IN ({placeholders})
        AND la.created_at BETWEEN %s AND %s
        GROUP BY iu.institution_id
    """

    cursor.execute(query, (*al_school_ids, start_date, today))
    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    # Build a mapping of school_id -> count
    school_counts = {}
    for row in rows:
        if isinstance(row, dict):
            school_id = row.get("institution_id")
            count = row.get("al_active", 0)
        else:
            school_id = row[0]
            count = row[1]
        school_counts[school_id] = count

    # Create school breakdown with names
    for al_id in al_school_ids:
        count = school_counts.get(al_id, 0)
        total_active += count

        school_name = next(
            (s.get("school_name") for s in schools_data if s.get("library_school_id") == al_id),
            "Unknown"
        )
        school_breakdown.append({
            "al_school_id": al_id,
            "school_name": school_name,
            "al_active": count
        })

    return jsonify({
        "username": username,
        "al_active_total": total_active,
        "school_breakdown": school_breakdown
    })




@app.route('/api/champion/<username>/remove_school', methods=['DELETE'])
@login_required
def remove_champion_school(username):
    """Remove a school (by asl_school_id, library_school_id, or both) from champion's assigned schools."""
    user = db.first_or_404(sa.select(User).where(User.username == username))

    champ = ChampionSchool.query.filter_by(
        firstname=user.firstname,
        lastname=user.lastname,
        province=user.province
    ).first()

    if not champ or not champ.schools:
        return jsonify({"error": "No schools found for this champion"}), 404

    try:
        schools_data = json.loads(champ.schools)
    except Exception as e:
        return jsonify({"error": f"Invalid JSON in champ.schools: {e}"}), 400

    # Extract query params
    asl_school_id = request.args.get("asl_school_id", type=int)
    library_school_id = request.args.get("library_school_id", type=int)

    if not asl_school_id and not library_school_id:
        return jsonify({"error": "Provide at least asl_school_id or library_school_id"}), 400

    # Filter out matching schools
    updated_schools = []
    removed = []

    for s in schools_data:
        match_asl = asl_school_id and s.get("asl_school_id") == asl_school_id
        match_lib = library_school_id and s.get("library_school_id") == library_school_id

        # Keep only schools that don't match the removal criteria
        if (asl_school_id and library_school_id and match_asl and match_lib) \
           or (asl_school_id and not library_school_id and match_asl) \
           or (library_school_id and not asl_school_id and match_lib):
            removed.append(s)
        else:
            updated_schools.append(s)

    if not removed:
        return jsonify({"message": "No matching school found for removal"}), 404

    # Save changes
    champ.schools = json.dumps(updated_schools)
    db.session.commit()

    return jsonify({
        "message": "School(s) removed successfully",
        "removed": removed,
        "remaining_schools": updated_schools
    })




    


@app.route('/api/schools/<province>')
@login_required
def get_schools_by_province(province):
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)  # ensures dict results

        cursor.execute("""
            SELECT school_id, school_name
            FROM tblschools
            WHERE school_province = %s
              AND (flag IS NULL OR flag <> 'd')
            ORDER BY school_name ASC
        """, (province,))

        rows = cursor.fetchall()

        schools = [{'id': int(row['school_id']), 'name': row['school_name']} for row in rows]

        return jsonify(schools)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'{type(e).__name__}: {str(e)}'}), 500


    # finally:
    #     if 'cursor' in locals() and cursor: cursor.close()
    #     if 'conn' in locals() and conn: conn.close()





UPLOAD_FOLDER = 'static/uploads/avatars'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Help desk uploads
HELP_DESK_UPLOAD_FOLDER = 'static/uploads/helpdesk'
app.config['HELP_DESK_UPLOAD_FOLDER'] = HELP_DESK_UPLOAD_FOLDER
os.makedirs(HELP_DESK_UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/upload_avatar', methods=['POST'])
@login_required
def upload_avatar():
    file = request.files.get('avatar')
    if not file or file.filename == '':
        flash('No file selected.')
        return redirect(request.referrer)

    if not allowed_file(file.filename):
        flash('File type not allowed.')
        return redirect(request.referrer)

    filename = secure_filename(f"user_{current_user.id}_" + file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    file.save(file_path)

    current_user.avatar_filename = f'uploads/avatars/{filename}'
    db.session.commit()

    flash('Avatar updated!')
    return redirect(url_for('index'))  # Change to your profile route



@app.route('/reports', methods=['GET', 'POST'])
def reports():
    # Get all users
    users = User.query.all()
    user_department = current_user.department

    # Get selected user (default to first user if none chosen)
    selected_user_id = request.args.get("user_id")
    if not selected_user_id and users:
        selected_user_id = users[0].id
    else:
        selected_user_id = int(selected_user_id) if selected_user_id else None

    reports = []
    if selected_user_id:
        reports = WeeklyReport.query.filter_by(user_id=selected_user_id)\
                                    .order_by(WeeklyReport.created_at.desc())\
                                    .all()

    selected_report_user = User.query.get_or_404(selected_user_id)

    return render_template(
        'reports.html',
        reports=reports,
        user_department =user_department ,
        users=users,
        selected_user_id=selected_user_id,
        selected_report_user=selected_report_user,
        title='Reports'
    )


@app.route('/akello_weekly_reports', methods=['GET'])
@login_required
def akello_weekly_reports():
    # Reuse the same data preparation as /reports
    users = User.query.all()
    user_department = current_user.department

    selected_user_id = request.args.get('user_id')
    if not selected_user_id and users:
        selected_user_id = users[0].id
    else:
        selected_user_id = int(selected_user_id) if selected_user_id else None

    reports = []
    if selected_user_id:
        reports = WeeklyReport.query.filter_by(user_id=selected_user_id) \
                                    .order_by(WeeklyReport.created_at.desc()) \
                                    .all()

    selected_report_user = User.query.get_or_404(selected_user_id)

    return render_template(
        'akello_weekly_reports.html',
        reports=reports,
        user_department=user_department,
        users=users,
        selected_user_id=selected_user_id,
        selected_report_user=selected_report_user,
        title='Reports'
    )


@app.route("/submit_weekly_report", methods=["GET", "POST"])
@login_required
def submit_report():
    if request.method == "POST":
        week_start = request.form["week_start"]
        week_start_date = datetime.strptime(week_start, "%Y-%m-%d").date()
        if week_start_date.weekday() != 0:  # Must be Monday
            flash("Week start date must be a Monday!", "error")
            return redirect(url_for("submit_report"))

        report = WeeklyReport(
            user_id=current_user.id,
            department=current_user.department or 'General',
            week_start=week_start,
            work_done=request.form["work_done"],
            work_next=request.form["work_next"],
            challenges=request.form["challenges"]
        )
        db.session.add(report)
        db.session.commit()
        flash("Report submitted successfully!", "success")
        return redirect(url_for("akello_weekly_reports"))

    return redirect(url_for('akello_weekly_reports'))



@app.route("/edit_report/<int:report_id>", methods=["GET", "POST"])
@login_required
def edit_report(report_id):

    report = WeeklyReport.query.get_or_404(report_id)

    # Ensure user owns this report
    if report.user_id != current_user.id:
        flash("You are not authorized to edit this report.", "error")
        return redirect(url_for("akello_weekly_reports"))

    # Check if report is older than 7 days
    if datetime.utcnow() > report.created_at + timedelta(days=7):
        flash("You can only edit reports within 7 days of submission.", "error")
        return redirect(url_for("akello_weekly_reports"))

    if request.method == "POST":
        report.department = current_user.department or report.department or 'General'
        report.week_start = request.form["week_start"]
        report.work_done = request.form["work_done"]
        report.work_next = request.form["work_next"]
        report.challenges = request.form["challenges"]

        db.session.commit()
        flash("Report updated successfully!", "success")
        return redirect(url_for("akello_weekly_reports"))

@app.route('/workspaces', methods=['GET', 'POST'])
@login_required
def workspaces():
    if current_user.userRole == 'Brand Ambassador':
        return redirect(url_for('index'))
    else:
        newspaceform = WorkspaceForm()

        my_workspaces = current_user.memberships # This fetches workspaces the current user is a member of
        all_spaces = Workspace.query.all() # This fetches ALL workspaces in the database

    return render_template(
        'workspaces.html',
        newspaceform=newspaceform,
        all_spaces=all_spaces, # This list contains ALL workspaces, and you can access their .members
        workspaces=my_workspaces, # This list contains only workspaces the current user is a member of
        title='Workspaces'
    )

@app.route('/workspace/create', methods=['GET', 'POST'])
@login_required
def create_workspace():
    if current_user.userRole == 'Brand Ambassador':
        return redirect(url_for('index'))
    else:
        form = WorkspaceForm()
        if form.validate_on_submit():
            workspace = Workspace(name=form.name.data, description=form.description.data, created_by=current_user.id)
            workspace.members.append(current_user)
            db.session.add(workspace)
            db.session.commit()
            flash('Workspace created successfully!', 'success')
    return redirect(url_for('workspaces'))



@app.route('/workspace/<int:workspace_id>/invite', methods=['POST'])
@login_required
def invite_user_to_workspace(workspace_id):
    if current_user.userRole == 'Brand Ambassador':
        return redirect(url_for('index'))
    else:
        workspace = Workspace.query.get_or_404(workspace_id)
        if current_user not in workspace.members:
            flash('You are not a member of this workspace.', 'danger')
            return redirect(url_for('workspaces'))

        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()

        if not user:
            flash('User not found.', 'danger')
        elif user in workspace.members:
            flash('User is already a member of the workspace.', 'info')
        else:
            workspace.members.append(user)
            db.session.commit()
            flash(f'{user.username} added to the workspace!', 'success')

    return redirect(url_for('view_workspace', workspace_id=workspace_id))





@app.route('/workspace/<int:workspace_id>')
@login_required
def view_workspace(workspace_id):
    taskform=TaskForm()
    projectform=ProjectForm()
    workspace = Workspace.query.get_or_404(workspace_id)
    available_users = User.query.filter(~User.id.in_([member.id for member in workspace.members])).all()
    # available_users = User.query.filter(User.id.in_([member.id for member in workspace.members])).all()
    if current_user not in workspace.members:
        return redirect(url_for('workspaces'))
    else:
        show_gantt = request.args.get('gantt') == '1'

        gantt_projects_data = []
        gantt_data = []
        if show_gantt:
            for project in workspace.projects:
                if project.start_date and project.end_date:
                    gantt_data.append({
                        'Task': project.title,
                        'Start': project.start_date.strftime('%Y-%m-%d'),
                        'Finish': project.end_date.strftime('%Y-%m-%d'),
                        'StartLabel': project.start_date.strftime('%Y-%m-%d'),
                        'FinishLabel': project.end_date.strftime('%Y-%m-%d'),
                        'Resource': 'Project',
                        'Color': '#1f77b4',
                        'TaskID': None
                    })

                for task in project.tasks:
                    start_date = task.start_date.strftime('%Y-%m-%d') if task.start_date else (
                        task.due_date.strftime('%Y-%m-%d') if task.due_date else None)
                    finish_date = task.due_date.strftime('%Y-%m-%d') if task.due_date else start_date
                    color = '#28a745' if task.status == 'Done' else '#ff7f0e'

                    gantt_data.append({
                        'Task': f"↳ {project.title}: {task.title}",
                        'Start': start_date,
                        'Finish': finish_date,
                        'StartLabel': start_date,
                        'FinishLabel': finish_date,
                        'Resource': 'Task',
                        'Color': color,
                        'TaskID': task.id
                    })


        page = request.args.get('page', 1, type=int)
        per_page = 5
        for project in workspace.projects:
            project.paginated_tasks = Task.query.filter_by(project_id=project.id).paginate(page=page, per_page=per_page, error_out=False)

    return render_template('workspace_detail.html',taskform=taskform,projectform=projectform,available_users=available_users, workspace=workspace,show_gantt=show_gantt,gantt_projects_data=gantt_projects_data, gantt_data=gantt_data if show_gantt else [], title='Workspaces')

@app.route('/workspace/<int:workspace_id>/project/create', methods=['GET', 'POST'])
@login_required
def create_project(workspace_id):
    return redirect(url_for('projectmanagement'))


@app.route('/project/<int:project_id>/edit', methods=['POST'])
@login_required
def edit_project(project_id):
    return redirect(url_for('projectmanagement', project=project_id))


@app.route('/project/<int:project_id>/delete', methods=['POST'])
@login_required
def delete_project(project_id):
    return redirect(url_for('projectmanagement'))


@app.route('/project/<int:project_id>/task/create', methods=['GET', 'POST'])
@login_required
def create_task(project_id):
    return redirect(url_for('projectmanagement', project=project_id))


@app.route('/task/<int:task_id>/update_dates', methods=['POST'])
@login_required
def update_task_dates(task_id):
    task = Task.query.get_or_404(task_id)
    data = request.json
    try:
        task.start_date = datetime.strptime(data['start'], '%Y-%m-%d')
        task.due_date = datetime.strptime(data['end'], '%Y-%m-%d')
        db.session.commit()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400


@app.route('/task/<int:task_id>/delete', methods=['POST'])
@login_required
def delete_task(task_id):
    task = Task.query.get_or_404(task_id)
    workspace_id = task.project.workspace_id
    db.session.delete(task)
    db.session.commit()
    flash('Task deleted successfully!', 'success')
    return redirect(url_for('view_workspace', workspace_id=workspace_id))


@app.route('/task/<int:task_id>/edit', methods=['POST'])
@login_required
def edit_task(task_id):
    task = Task.query.get_or_404(task_id)
    form = TaskForm()
    if form.validate_on_submit():
        task.title = form.title.data
        task.description = form.description.data
        task.start_date = form.start_date.data
        task.due_date = form.due_date.data
        task.status = form.status.data
        task.progress=form.progress.data
        db.session.commit()
        flash('Task updated successfully!', 'success')
    return redirect(url_for('view_workspace', workspace_id=task.project.workspace_id))



@app.route('/contracting', methods=['GET', 'POST'])
@login_required
def contracting():
    if current_user.userRole == 'Brand Ambassador':
        return redirect(url_for('index'))
    else:
        username = (current_user.username or '').strip()
        role = current_user.userRole
        scorecards = Scorecard.query.filter(Scorecard.employee_name.ilike(username)).all()
    return render_template('contracting.html', username=username, role=role, title='Contracting', scorecards=scorecards)


@app.route('/save_scorecard_row', methods=['POST'])
@login_required
def save_scorecard_row():
    """Create or update a scorecard row"""
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
        
    try:
        row_id = data.get('id')
        current_username = (current_user.username or '').strip()
        required_fields = {
            'keyFocusArea': 'Key focus area is required',
            'strategicObjective': 'Strategic objective is required',
            'performanceMeasure': 'Performance measure is required',
            'unitOfMeasure': 'Unit of measure is required',
            'target': 'Target is required',
            'weight': 'Weight is required'
        }
        for key, message in required_fields.items():
            value = data.get(key)
            if value is None or str(value).strip() == '':
                return jsonify({'error': message}), 400

        try:
            weight = float(data.get('weight'))
        except (TypeError, ValueError):
            return jsonify({'error': 'Weight must be a numeric value'}), 400
        if weight < 0:
            return jsonify({'error': 'Weight must be greater than or equal to 0'}), 400
        
        if row_id:
            # Update existing row
            row = Scorecard.query.get(row_id)
            if not row:
                return jsonify({'error': 'Row not found'}), 404
            
            # Authorization check: only the owner or an admin can update
            row_owner = (row.employee_name or '').strip().lower()
            if row_owner != current_username.lower() and current_user.userRole != 'Admin':
                return jsonify({'error': 'Unauthorized'}), 403
        else:
            # Create new row
            row = Scorecard(employee_name=current_username)
            db.session.add(row)
            
        # Update fields
        row.key_focus_area = data.get('keyFocusArea')
        row.strategic_objective = data.get('strategicObjective')
        row.performance_measure = data.get('performanceMeasure')
        row.unit_of_measure = data.get('unitOfMeasure')
        row.target = str(data.get('target')).strip()
        row.weight = weight
        
        db.session.commit()
        return jsonify({
            'message': 'Row saved successfully',
            'id': row.id,
            'new_id': row.id,
            'row': {
                'keyFocusArea': row.key_focus_area,
                'strategicObjective': row.strategic_objective,
                'performanceMeasure': row.performance_measure,
                'unitOfMeasure': row.unit_of_measure,
                'target': row.target,
                'weight': row.weight
            }
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/delete_scorecard_row/<int:id>', methods=['POST'])
@login_required
def delete_scorecard_row(id):
    """Delete a scorecard row"""
    try:
        row = Scorecard.query.get(id)
        if not row:
            return jsonify({'error': 'Row not found'}), 404
            
        # Authorization check
        row_owner = (row.employee_name or '').strip().lower()
        current_username = (current_user.username or '').strip().lower()
        if row_owner != current_username and current_user.userRole != 'Admin':
            return jsonify({'error': 'Unauthorized'}), 403
            
        db.session.delete(row)
        db.session.commit()
        return jsonify({'message': 'Row deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/scorecards/me', methods=['GET'])
@login_required
def my_scorecards():
    username = (current_user.username or '').strip()
    try:
        rows = (
            Scorecard.query
            .filter(Scorecard.employee_name.ilike(username))
            .order_by(Scorecard.id.asc())
            .all()
        )
        return jsonify({
            'rows': [{
                'id': row.id,
                'keyFocusArea': row.key_focus_area,
                'strategicObjective': row.strategic_objective,
                'performanceMeasure': row.performance_measure,
                'unitOfMeasure': row.unit_of_measure,
                'target': row.target,
                'weight': row.weight
            } for row in rows]
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


from psycopg2.extras import DictCursor

@app.route('/school_profile/<int:school_id>', methods=['GET', 'POST'])
@login_required
def school_profile(school_id):
    username = current_user.username
    role = current_user.userRole

    conn = get_ruzivo_conn()
    cursor = conn.cursor()

    query = "SELECT school_name, school_province FROM tblschools WHERE school_id = %s"
    cursor.execute(query, (school_id,))
    row = cursor.fetchone()

    # cursor.close()
    # conn.close()

    if not row:
        return "School not found", 404

    try:
        school_name = row['school_name']
        school_province = row['school_province']
    except (TypeError, KeyError):
        # Fallback if row is a tuple instead of a dict
        school_name = row[0]
        school_province = row[1]

    return render_template(
        'school_profile.html',
        username=username,
        role=role,
        school_name=school_name,
        school_province=school_province,
        title='School tracker'
    )





@app.route('/school_profile_usage/<int:school_id>', methods=['GET', 'POST'])
@login_required
def school_profile_usage(school_id):
    username = current_user.username
    role = current_user.userRole
    asl_school_id = school_id

    conn = get_ruzivo_conn()
    cursor = conn.cursor()

    query = "SELECT school_name, school_province FROM tblschools WHERE school_id = %s"
    cursor.execute(query, (school_id,))
    row = cursor.fetchone()

    # cursor.close()
    # conn.close()

    if not row:
        return "School not found", 404

    try:
        school_name = row['school_name']
        school_province = row['school_province']
    except (TypeError, KeyError):
        # Fallback if row is a tuple instead of a dict
        school_name = row[0]
        school_province = row[1]

    return render_template(
        'simone_school_usage.html',
        username=username,
        role=role,
        school_name=school_name,
        school_province=school_province,
        asl_school_id=asl_school_id,
        title='School tracker'
    )


@app.route('/api/learner-profile/<username>', methods=['GET'])
@login_required
def get_learner_profile(username):
    """Get learner profile data by username"""
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        query = """
            SELECT 
                s.username,
                s.name,
                s.grade,
                s.last_login,
                s.school_id,
                s.access_sdate,
                s.access_edate,
                t.school_name,
                t.school_province
            FROM vwstudent s
            JOIN tblschools t ON s.school_id = t.school_id
            WHERE s.username = %s
            LIMIT 1
        """
        cursor.execute(query, (username,))
        row = cursor.fetchone()
        
        if not row:
            return jsonify({'error': 'Learner not found'}), 404
        
        # Format datetime fields - rows are dicts from DictCursor
        def format_datetime(dt):
            if dt:
                if hasattr(dt, 'strftime'):
                    return dt.strftime('%Y-%m-%d %H:%M:%S')
            return None
        
        profile_data = {
            'username': row.get('username'),
            'name': row.get('name'),
            'grade': row.get('grade'),
            'last_login': format_datetime(row.get('last_login')),
            'school_id': row.get('school_id'),
            'access_sdate': format_datetime(row.get('access_sdate')),
            'access_edate': format_datetime(row.get('access_edate')),
            'school_name': row.get('school_name'),
            'school_province': row.get('school_province')
        }
        
        return jsonify(profile_data), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/learner-profile/<username>', methods=['GET'])
@login_required
def learner_profile(username):
    """Display learner profile page"""
    current_username = current_user.username
    role = current_user.userRole
    
    return render_template(
        'learner_profile.html',
        username=current_username,
        role=role,
        learner_username=username,
        title='Learner Profile'
    )


@app.route('/api/search-learners', methods=['GET'])
@login_required
def search_learners():
    """Search learners by name, surname, or username"""
    try:
        search_term = request.args.get('q', '').strip()
        if not search_term:
            return jsonify({'error': 'Search term is required'}), 400
        
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        # Split search term into words to handle name and surname searches
        search_words = search_term.split()
        search_pattern = f'%{search_term}%'
        
        # Build query to search by username or name
        # Parse name field to extract first name and surname
        query = """
            SELECT 
                s.username,
                s.name,
                s.grade,
                s.last_login,
                s.school_id,
                t.school_name,
                t.school_province
            FROM vwstudent s
            JOIN tblschools t ON s.school_id = t.school_id
            WHERE s.username LIKE %s
               OR s.name LIKE %s
        """
        
        # If search term has multiple words, also try matching first word and last word separately
        if len(search_words) >= 2:
            # Multiple words: try matching first word and last word in the name
            first_word_pattern = f'%{search_words[0]}%'
            last_word_pattern = f'%{search_words[-1]}%'
            query += """
               OR (s.name LIKE %s AND s.name LIKE %s)
            """
            cursor.execute(query, (
                search_pattern, search_pattern,
                first_word_pattern, last_word_pattern
            ))
        else:
            # Single word: search in username or name
            cursor.execute(query, (search_pattern, search_pattern))
        
        rows = cursor.fetchall()
        
        learners = []
        for row in rows:
            full_name = row.get('name') or ''
            
            # Parse name to extract first name and surname
            name_parts = full_name.strip().split() if full_name else []
            if len(name_parts) >= 2:
                firstname = name_parts[0]
                surname = ' '.join(name_parts[1:])  # In case surname has multiple words
            elif len(name_parts) == 1:
                firstname = name_parts[0]
                surname = ''
            else:
                firstname = ''
                surname = ''
            
            learners.append({
                'username': row.get('username'),
                'name': firstname,
                'surname': surname,
                'full_name': full_name,
                'grade': row.get('grade'),
                'last_login': row.get('last_login').strftime('%Y-%m-%d %H:%M:%S') if row.get('last_login') else None,
                'school_id': row.get('school_id'),
                'school_name': row.get('school_name'),
                'school_province': row.get('school_province')
            })
        
        return jsonify({'learners': learners, 'count': len(learners)}), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/learner-platform-stats/<username>', methods=['GET'])
@login_required
def get_learner_platform_stats(username):
    """Get platform statistics for a learner across Library and SmartLearning platforms"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not start_date or not end_date:
            return jsonify({'error': 'start_date and end_date are required'}), 400
        
        # Get student_id and name from username for SmartLearning
        conn_ruzivo = get_ruzivo_conn()
        cursor_ruzivo = conn_ruzivo.cursor()
        
        student_query = """
            SELECT student_id, name FROM vwstudent WHERE username = %s LIMIT 1
        """
        cursor_ruzivo.execute(student_query, (username,))
        student_row = cursor_ruzivo.fetchone()
        
        if not student_row:
            return jsonify({'error': 'Learner not found'}), 404
        
        student_id = student_row.get('student_id') if isinstance(student_row, dict) else student_row[0]
        student_name = student_row.get('name') if isinstance(student_row, dict) else student_row[1]
        
        if not student_id:
            return jsonify({'error': 'Student ID not found for this learner'}), 404
        
        # Parse name into first_name and last_name for library search
        # Handle cases where name might have multiple parts
        name_parts = student_name.strip().split() if student_name else []
        if len(name_parts) >= 2:
            first_name = name_parts[0].strip()
            last_name = ' '.join(name_parts[1:]).strip()  # All remaining parts as last_name
        elif len(name_parts) == 1:
            first_name = name_parts[0].strip()
            last_name = ''  # No last name available
        else:
            first_name = ''
            last_name = ''
        
        # Excluded student IDs
        excluded_ids = [356,905,30843,265227,265228,265230,3223972,406978,516518,577527,1032617,1621500,1632975,1660731,1660834,1661007,1661053,1661353,1662420,1662839,1664013,1664021,1664698,1684433,1685102]
        
        # SmartLearning Platform Queries
        smartlearning_stats = {}
        
        # Build NOT IN clause with placeholders
        excluded_placeholders = ','.join(['%s'] * len(excluded_ids))
        
        # 1. Primary content access - check if student accessed primary content
        primary_content_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_primary_content 
            FROM tblcontent_access 
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(primary_content_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_primary_content'] = result.get('total_primary_content') if result else 0
        
        # 2. Secondary content access
        sec_content_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_sec_content 
            FROM tblcontent_access_hs 
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(sec_content_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_sec_content'] = result.get('total_sec_content') if result else 0
        
        # 3. Primary exercises by category
        primary_exercise_query = f"""
            SELECT category, COUNT(DISTINCT student_id) AS total_primary_exercise 
            FROM tblresults 
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
            GROUP BY category
        """
        cursor_ruzivo.execute(primary_exercise_query, (start_date, end_date, student_id, *excluded_ids))
        primary_exercises = cursor_ruzivo.fetchall()
        smartlearning_stats['primary_exercises'] = {row.get('category'): row.get('total_primary_exercise') for row in primary_exercises}
        
        # 4. Secondary exercises by category
        sec_exercise_query = f"""
            SELECT category, COUNT(DISTINCT student_id) AS total_sec_exercise 
            FROM tblresults_hs 
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
            GROUP BY category
        """
        cursor_ruzivo.execute(sec_exercise_query, (start_date, end_date, student_id, *excluded_ids))
        sec_exercises = cursor_ruzivo.fetchall()
        smartlearning_stats['secondary_exercises'] = {row.get('category'): row.get('total_sec_exercise') for row in sec_exercises}
        
        # 5. Zimsec access
        zimsec_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_zimsec_access 
            FROM tblcontent_access_zimsec
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(zimsec_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_zimsec_access'] = result.get('total_zimsec_access') if result else 0
        
        # 6. Teacher access
        teacher_access_query = f"""
            SELECT COUNT(DISTINCT student_id) AS teacher_access
            FROM tblclass_activity_results
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(teacher_access_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['teacher_access'] = result.get('teacher_access') if result else 0
        
        # Library Platform Query
        library_stats = {}
        try:
            conn_library = get_direct_library_conn()
            cursor_library = conn_library.cursor()
            
            library_user_id = None
            search_attempts = []
            
            # Strategy 1: Search by username (exact match)
            library_user_query = """
                SELECT id FROM users 
                WHERE TRIM(LOWER(username)) = TRIM(LOWER(%s))
                LIMIT 1
            """
            cursor_library.execute(library_user_query, (username,))
            library_user_row = cursor_library.fetchone()
            search_attempts.append(f"Username match: '{username}'")
            
            if library_user_row:
                library_user_id = library_user_row.get('id') if isinstance(library_user_row, dict) else library_user_row[0]
            
            # Strategy 2: If no match by username, try by first_name and last_name (fallback)
            if not library_user_id and first_name and last_name:
                library_user_query = """
                    SELECT id FROM users 
                    WHERE TRIM(LOWER(first_name)) = TRIM(LOWER(%s)) 
                    AND TRIM(LOWER(last_name)) = TRIM(LOWER(%s))
                    LIMIT 1
                """
                cursor_library.execute(library_user_query, (first_name, last_name))
                library_user_row = cursor_library.fetchone()
                search_attempts.append(f"Name fallback: '{first_name}' '{last_name}'")
                
                if library_user_row:
                    library_user_id = library_user_row.get('id') if isinstance(library_user_row, dict) else library_user_row[0]
            
            if library_user_id:
                # Query total duration and average duration for this user
                library_duration_query = """
                    SELECT 
                        SUM(duration_minutes) AS total_duration_minutes,
                        AVG(duration_minutes) AS avg_duration_minutes,
                        COUNT(*) AS reading_sessions
                    FROM read_trackers 
                    WHERE duration_minutes != 0
                    AND user_id = %s
                    AND DATE(created_at) BETWEEN DATE(%s) AND DATE(%s)
                """
                cursor_library.execute(library_duration_query, (library_user_id, start_date, end_date))
                result = cursor_library.fetchone()
                
                total_minutes = float(result.get('total_duration_minutes')) if result and result.get('total_duration_minutes') else 0
                avg_minutes = float(result.get('avg_duration_minutes')) if result and result.get('avg_duration_minutes') else 0
                sessions = int(result.get('reading_sessions')) if result and result.get('reading_sessions') else 0
                
                # Convert to hours and minutes for display
                total_hours = int(total_minutes // 60)
                total_mins = int(total_minutes % 60)
                
                library_stats['total_duration_minutes'] = total_minutes
                library_stats['avg_duration_minutes'] = avg_minutes
                library_stats['reading_sessions'] = sessions
                library_stats['formatted_time'] = f"{total_hours}h {total_mins}m" if total_hours > 0 else f"{total_mins}m"
                library_stats['found'] = True
            else:
                library_stats['avg_duration_minutes'] = 0
                library_stats['note'] = f'User not found in library platform. Searched with: {"; ".join(search_attempts)}. Full name from SmartLearning: "{student_name}"'
            
            if cursor_library:
                cursor_library.close()
            if conn_library:
                conn_library.close()
        except Exception as e:
            import traceback
            traceback.print_exc()
            library_stats['error'] = str(e)
            library_stats['avg_duration_minutes'] = 0
        
        return jsonify({
            'smartlearning': smartlearning_stats,
            'library': library_stats
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/learner-time-spent/<username>', methods=['GET'])
@login_required
def get_learner_time_spent(username):
    """Get time spent on platform for a learner in a specific time period"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not start_date or not end_date:
            return jsonify({'error': 'start_date and end_date are required'}), 400
        
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        # Get student_id from vwstudent (it has student_id based on other queries)
        student_query = """
            SELECT student_id FROM vwstudent WHERE username = %s LIMIT 1
        """
        cursor.execute(student_query, (username,))
        student_row = cursor.fetchone()
        
        if not student_row:
            return jsonify({'error': 'Learner not found'}), 404
        
        student_id = student_row.get('student_id') if isinstance(student_row, dict) else student_row[0]
        
        if not student_id:
            return jsonify({'error': 'Student ID not found for this learner'}), 404
        
        # Calculate total time spent from content access tables
        # Sum up durations from primary, secondary, and zimsec content access
        time_query = """
            SELECT 
                COALESCE(SUM(
                    TIMESTAMPDIFF(SECOND, 
                        COALESCE(ca1.start_time, ca2.start_time, ca3.start_time),
                        COALESCE(ca1.end_time, ca2.end_time, ca3.end_time)
                    )
                ), 0) AS total_seconds
            FROM (
                SELECT student_id, start_time, end_time, NULL as dummy1, NULL as dummy2
                FROM tblcontent_access
                WHERE student_id = %s 
                  AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT student_id, start_time, end_time, NULL, NULL
                FROM tblcontent_access_hs
                WHERE student_id = %s 
                  AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT student_id, start_time, end_time, NULL, NULL
                FROM tblcontent_access_zimsec
                WHERE student_id = %s 
                  AND start_time BETWEEN %s AND %s
            ) AS combined
            LEFT JOIN tblcontent_access ca1 ON combined.student_id = ca1.student_id 
                AND combined.start_time = ca1.start_time
            LEFT JOIN tblcontent_access_hs ca2 ON combined.student_id = ca2.student_id 
                AND combined.start_time = ca2.start_time
            LEFT JOIN tblcontent_access_zimsec ca3 ON combined.student_id = ca3.student_id 
                AND combined.start_time = ca3.start_time
        """
        
        # Simplified approach - calculate from each table separately
        total_seconds = 0
        
        # Primary content access
        primary_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(primary_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        # Secondary content access
        secondary_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access_hs
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(secondary_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        # Zimsec content access
        zimsec_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access_zimsec
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(zimsec_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        # Convert to hours, minutes, seconds
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        seconds = total_seconds % 60
        
        # Calculate longest streak (consecutive days with activity)
        # Get all unique dates with activity (from login or content access)
        active_dates_query = """
            SELECT DISTINCT DATE(activity_date) AS active_date
            FROM (
                SELECT DATE(login_date) AS activity_date
                FROM tblstudents_login
                WHERE student_id = %s AND login_date BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_activities
            ORDER BY active_date ASC
        """
        cursor.execute(active_dates_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        active_dates = cursor.fetchall()
        
        # Calculate longest streak
        longest_streak = 0
        current_streak = 0
        prev_date = None
        
        for row in active_dates:
            active_date = row.get('active_date') if isinstance(row, dict) else row[0]
            if prev_date is None:
                current_streak = 1
            else:
                # Check if consecutive
                from datetime import timedelta
                if active_date == prev_date + timedelta(days=1):
                    current_streak += 1
                else:
                    longest_streak = max(longest_streak, current_streak)
                    current_streak = 1
            prev_date = active_date
        
        longest_streak = max(longest_streak, current_streak)
        
        # Find most active day (day with most time spent)
        most_active_day_query = """
            SELECT 
                DATE(start_time) AS activity_date,
                SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS total_seconds
            FROM (
                SELECT start_time, end_time
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_content
            GROUP BY DATE(start_time)
            ORDER BY total_seconds DESC
            LIMIT 1
        """
        cursor.execute(most_active_day_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        most_active_day_result = cursor.fetchone()
        
        most_active_day = None
        most_active_day_seconds = 0
        if most_active_day_result:
            most_active_day = most_active_day_result.get('activity_date') if isinstance(most_active_day_result, dict) else most_active_day_result[0]
            most_active_day_seconds = most_active_day_result.get('total_seconds') if isinstance(most_active_day_result, dict) else most_active_day_result[1]
            if most_active_day:
                most_active_day = most_active_day.strftime('%Y-%m-%d') if hasattr(most_active_day, 'strftime') else str(most_active_day)
        
        # Find most active month (month with most time spent in the date range)
        most_active_month_query = """
            SELECT 
                DATE_FORMAT(start_time, '%%Y-%%m') AS activity_month,
                SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS total_seconds
            FROM (
                SELECT start_time, end_time
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_content
            GROUP BY DATE_FORMAT(start_time, '%%Y-%%m')
            ORDER BY total_seconds DESC
            LIMIT 1
        """
        cursor.execute(most_active_month_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        most_active_month_result = cursor.fetchone()
        
        most_active_month = None
        most_active_month_seconds = 0
        if most_active_month_result:
            most_active_month = most_active_month_result.get('activity_month') if isinstance(most_active_month_result, dict) else most_active_month_result[0]
            most_active_month_seconds = most_active_month_result.get('total_seconds') if isinstance(most_active_month_result, dict) else most_active_month_result[1]
        
        # Format most active day time
        most_active_day_hours = most_active_day_seconds // 3600
        most_active_day_minutes = (most_active_day_seconds % 3600) // 60
        most_active_day_formatted = f"{most_active_day_hours}h {most_active_day_minutes}m" if most_active_day_hours > 0 else f"{most_active_day_minutes}m"
        
        # Format most active month time
        most_active_month_hours = most_active_month_seconds // 3600
        most_active_month_minutes = (most_active_month_seconds % 3600) // 60
        most_active_month_formatted = f"{most_active_month_hours}h {most_active_month_minutes}m" if most_active_month_hours > 0 else f"{most_active_month_minutes}m"
        
        # Format month name
        if most_active_month:
            try:
                from datetime import datetime
                month_date = datetime.strptime(most_active_month, '%Y-%m')
                most_active_month_display = month_date.strftime('%B %Y')
            except:
                most_active_month_display = most_active_month
        else:
            most_active_month_display = None
        
        return jsonify({
            'total_seconds': total_seconds,
            'hours': hours,
            'minutes': minutes,
            'seconds': seconds,
            'formatted': f"{hours}h {minutes}m {seconds}s" if hours > 0 else f"{minutes}m {seconds}s",
            'longest_streak': longest_streak,
            'most_active_day': most_active_day,
            'most_active_day_seconds': most_active_day_seconds,
            'most_active_day_formatted': most_active_day_formatted if most_active_day else None,
            'most_active_month': most_active_month_display,
            'most_active_month_seconds': most_active_month_seconds,
            'most_active_month_formatted': most_active_month_formatted if most_active_month else None
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/search-learners', methods=['GET'])
@login_required
def search_learners_page():
    """Display learner search page"""
    username = current_user.username
    role = current_user.userRole
    
    return render_template(
        'search_learners.html',
        username=username,
        role=role,
        title='Search Learners'
    )


def _fetch_learner_profile_data(username):
    """Helper function to fetch basic learner profile data"""
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        query = """
            SELECT 
                s.username,
                s.name,
                s.grade,
                s.last_login,
                s.school_id,
                s.access_sdate,
                s.access_edate,
                t.school_name,
                t.school_province
            FROM vwstudent s
            JOIN tblschools t ON s.school_id = t.school_id
            WHERE s.username = %s
            LIMIT 1
        """
        cursor.execute(query, (username,))
        row = cursor.fetchone()
        
        if not row:
            return None
        
        def format_datetime(dt):
            if dt:
                if hasattr(dt, 'strftime'):
                    return dt.strftime('%Y-%m-%d %H:%M:%S')
            return None
        
        return {
            'username': row.get('username'),
            'name': row.get('name'),
            'grade': row.get('grade'),
            'last_login': format_datetime(row.get('last_login')),
            'school_id': row.get('school_id'),
            'access_sdate': format_datetime(row.get('access_sdate')),
            'access_edate': format_datetime(row.get('access_edate')),
            'school_name': row.get('school_name'),
            'school_province': row.get('school_province')
        }
    except Exception as e:
        return None


def _fetch_learner_time_spent_data(username, start_date, end_date):
    """Helper function to fetch time spent data for a learner"""
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        student_query = """
            SELECT student_id FROM vwstudent WHERE username = %s LIMIT 1
        """
        cursor.execute(student_query, (username,))
        student_row = cursor.fetchone()
        
        if not student_row:
            return None
        
        student_id = student_row.get('student_id') if isinstance(student_row, dict) else student_row[0]
        
        if not student_id:
            return None
        
        total_seconds = 0
        
        # Primary content access
        primary_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(primary_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        # Secondary content access
        secondary_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access_hs
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(secondary_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        # Zimsec content access
        zimsec_query = """
            SELECT SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS seconds
            FROM tblcontent_access_zimsec
            WHERE student_id = %s AND start_time BETWEEN %s AND %s
        """
        cursor.execute(zimsec_query, (student_id, start_date, end_date))
        result = cursor.fetchone()
        if result and result.get('seconds'):
            total_seconds += result.get('seconds') or 0
        
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        seconds = total_seconds % 60
        
        # Calculate longest streak
        active_dates_query = """
            SELECT DISTINCT DATE(activity_date) AS active_date
            FROM (
                SELECT DATE(login_date) AS activity_date
                FROM tblstudents_login
                WHERE student_id = %s AND login_date BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION
                
                SELECT DATE(start_time) AS activity_date
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_activities
            ORDER BY active_date ASC
        """
        cursor.execute(active_dates_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        active_dates = cursor.fetchall()
        
        longest_streak = 0
        current_streak = 0
        prev_date = None
        
        for row in active_dates:
            active_date = row.get('active_date') if isinstance(row, dict) else row[0]
            if prev_date is None:
                current_streak = 1
            else:
                if active_date == prev_date + timedelta(days=1):
                    current_streak += 1
                else:
                    longest_streak = max(longest_streak, current_streak)
                    current_streak = 1
            prev_date = active_date
        
        longest_streak = max(longest_streak, current_streak)
        
        # Most active day
        most_active_day_query = """
            SELECT 
                DATE(start_time) AS activity_date,
                SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS total_seconds
            FROM (
                SELECT start_time, end_time
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_content
            GROUP BY DATE(start_time)
            ORDER BY total_seconds DESC
            LIMIT 1
        """
        cursor.execute(most_active_day_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        most_active_day_result = cursor.fetchone()
        
        most_active_day = None
        most_active_day_seconds = 0
        if most_active_day_result:
            most_active_day = most_active_day_result.get('activity_date') if isinstance(most_active_day_result, dict) else most_active_day_result[0]
            most_active_day_seconds = most_active_day_result.get('total_seconds') if isinstance(most_active_day_result, dict) else most_active_day_result[1]
            if most_active_day:
                most_active_day = most_active_day.strftime('%Y-%m-%d') if hasattr(most_active_day, 'strftime') else str(most_active_day)
        
        # Most active month
        most_active_month_query = """
            SELECT 
                DATE_FORMAT(start_time, '%%Y-%%m') AS activity_month,
                SUM(TIMESTAMPDIFF(SECOND, start_time, COALESCE(end_time, start_time))) AS total_seconds
            FROM (
                SELECT start_time, end_time
                FROM tblcontent_access
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_hs
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
                
                UNION ALL
                
                SELECT start_time, end_time
                FROM tblcontent_access_zimsec
                WHERE student_id = %s AND start_time BETWEEN %s AND %s
            ) AS all_content
            GROUP BY DATE_FORMAT(start_time, '%%Y-%%m')
            ORDER BY total_seconds DESC
            LIMIT 1
        """
        cursor.execute(most_active_month_query, (
            student_id, start_date, end_date,
            student_id, start_date, end_date,
            student_id, start_date, end_date
        ))
        most_active_month_result = cursor.fetchone()
        
        most_active_month = None
        most_active_month_seconds = 0
        if most_active_month_result:
            most_active_month = most_active_month_result.get('activity_month') if isinstance(most_active_month_result, dict) else most_active_month_result[0]
            most_active_month_seconds = most_active_month_result.get('total_seconds') if isinstance(most_active_month_result, dict) else most_active_month_result[1]
        
        most_active_day_hours = most_active_day_seconds // 3600
        most_active_day_minutes = (most_active_day_seconds % 3600) // 60
        most_active_day_formatted = f"{most_active_day_hours}h {most_active_day_minutes}m" if most_active_day_hours > 0 else f"{most_active_day_minutes}m"
        
        most_active_month_hours = most_active_month_seconds // 3600
        most_active_month_minutes = (most_active_month_seconds % 3600) // 60
        most_active_month_formatted = f"{most_active_month_hours}h {most_active_month_minutes}m" if most_active_month_hours > 0 else f"{most_active_month_minutes}m"
        
        if most_active_month:
            try:
                month_date = datetime.strptime(most_active_month, '%Y-%m')
                most_active_month_display = month_date.strftime('%B %Y')
            except:
                most_active_month_display = most_active_month
        else:
            most_active_month_display = None
        
        return {
            'total_seconds': total_seconds,
            'hours': hours,
            'minutes': minutes,
            'seconds': seconds,
            'formatted': f"{hours}h {minutes}m {seconds}s" if hours > 0 else f"{minutes}m {seconds}s",
            'longest_streak': longest_streak,
            'most_active_day': most_active_day,
            'most_active_day_seconds': most_active_day_seconds,
            'most_active_day_formatted': most_active_day_formatted if most_active_day else None,
            'most_active_month': most_active_month_display,
            'most_active_month_seconds': most_active_month_seconds,
            'most_active_month_formatted': most_active_month_formatted if most_active_month else None
        }
    except Exception as e:
        return None


def _fetch_learner_platform_stats_data(username, start_date, end_date):
    """Helper function to fetch platform statistics for a learner"""
    try:
        conn_ruzivo = get_ruzivo_conn()
        cursor_ruzivo = conn_ruzivo.cursor()
        
        student_query = """
            SELECT student_id, name FROM vwstudent WHERE username = %s LIMIT 1
        """
        cursor_ruzivo.execute(student_query, (username,))
        student_row = cursor_ruzivo.fetchone()
        
        if not student_row:
            return None
        
        student_id = student_row.get('student_id') if isinstance(student_row, dict) else student_row[0]
        student_name = student_row.get('name') if isinstance(student_row, dict) else student_row[1]
        
        if not student_id:
            return None
        
        name_parts = student_name.strip().split() if student_name else []
        if len(name_parts) >= 2:
            first_name = name_parts[0].strip()
            last_name = ' '.join(name_parts[1:]).strip()
        elif len(name_parts) == 1:
            first_name = name_parts[0].strip()
            last_name = ''
        else:
            first_name = ''
            last_name = ''
        
        excluded_ids = [356,905,30843,265227,265228,265230,3223972,406978,516518,577527,1032617,1621500,1632975,1660731,1660834,1661007,1661053,1661353,1662420,1662839,1664013,1664021,1664698,1684433,1685102]
        smartlearning_stats = {}
        excluded_placeholders = ','.join(['%s'] * len(excluded_ids))
        
        # Primary content access
        primary_content_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_primary_content 
            FROM tblcontent_access 
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(primary_content_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_primary_content'] = result.get('total_primary_content') if result else 0
        
        # Secondary content access
        sec_content_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_sec_content 
            FROM tblcontent_access_hs 
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(sec_content_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_sec_content'] = result.get('total_sec_content') if result else 0
        
        # Primary exercises by category
        primary_exercise_query = f"""
            SELECT category, COUNT(DISTINCT student_id) AS total_primary_exercise 
            FROM tblresults 
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
            GROUP BY category
        """
        cursor_ruzivo.execute(primary_exercise_query, (start_date, end_date, student_id, *excluded_ids))
        primary_exercises = cursor_ruzivo.fetchall()
        smartlearning_stats['primary_exercises'] = {row.get('category'): row.get('total_primary_exercise') for row in primary_exercises}
        
        # Secondary exercises by category
        sec_exercise_query = f"""
            SELECT category, COUNT(DISTINCT student_id) AS total_sec_exercise 
            FROM tblresults_hs 
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
            GROUP BY category
        """
        cursor_ruzivo.execute(sec_exercise_query, (start_date, end_date, student_id, *excluded_ids))
        sec_exercises = cursor_ruzivo.fetchall()
        smartlearning_stats['secondary_exercises'] = {row.get('category'): row.get('total_sec_exercise') for row in sec_exercises}
        
        # Zimsec access
        zimsec_query = f"""
            SELECT COUNT(DISTINCT student_id) AS total_zimsec_access 
            FROM tblcontent_access_zimsec
            WHERE start_time BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(zimsec_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['total_zimsec_access'] = result.get('total_zimsec_access') if result else 0
        
        # Teacher access
        teacher_access_query = f"""
            SELECT COUNT(DISTINCT student_id) AS teacher_access
            FROM tblclass_activity_results
            WHERE date_added BETWEEN %s AND %s
            AND student_id = %s
            AND student_id NOT IN ({excluded_placeholders})
        """
        cursor_ruzivo.execute(teacher_access_query, (start_date, end_date, student_id, *excluded_ids))
        result = cursor_ruzivo.fetchone()
        smartlearning_stats['teacher_access'] = result.get('teacher_access') if result else 0
        
        # Library Platform Query
        library_stats = {}
        try:
            conn_library = get_direct_library_conn()
            cursor_library = conn_library.cursor()
            
            library_user_id = None
            
            # Strategy 1: Search by username
            library_user_query = """
                SELECT id FROM users 
                WHERE TRIM(LOWER(username)) = TRIM(LOWER(%s))
                LIMIT 1
            """
            cursor_library.execute(library_user_query, (username,))
            library_user_row = cursor_library.fetchone()
            
            if library_user_row:
                library_user_id = library_user_row.get('id') if isinstance(library_user_row, dict) else library_user_row[0]
            
            # Strategy 2: Try by first_name and last_name
            if not library_user_id and first_name and last_name:
                library_user_query = """
                    SELECT id FROM users 
                    WHERE TRIM(LOWER(first_name)) = TRIM(LOWER(%s)) 
                    AND TRIM(LOWER(last_name)) = TRIM(LOWER(%s))
                    LIMIT 1
                """
                cursor_library.execute(library_user_query, (first_name, last_name))
                library_user_row = cursor_library.fetchone()
                
                if library_user_row:
                    library_user_id = library_user_row.get('id') if isinstance(library_user_row, dict) else library_user_row[0]
            
            if library_user_id:
                library_duration_query = """
                    SELECT 
                        SUM(duration_minutes) AS total_duration_minutes,
                        AVG(duration_minutes) AS avg_duration_minutes,
                        COUNT(*) AS reading_sessions
                    FROM read_trackers 
                    WHERE duration_minutes != 0
                    AND user_id = %s
                    AND DATE(created_at) BETWEEN DATE(%s) AND DATE(%s)
                """
                cursor_library.execute(library_duration_query, (library_user_id, start_date, end_date))
                result = cursor_library.fetchone()
                
                total_minutes = float(result.get('total_duration_minutes')) if result and result.get('total_duration_minutes') else 0
                avg_minutes = float(result.get('avg_duration_minutes')) if result and result.get('avg_duration_minutes') else 0
                sessions = int(result.get('reading_sessions')) if result and result.get('reading_sessions') else 0
                
                total_hours = int(total_minutes // 60)
                total_mins = int(total_minutes % 60)
                
                library_stats['total_duration_minutes'] = total_minutes
                library_stats['avg_duration_minutes'] = avg_minutes
                library_stats['reading_sessions'] = sessions
                library_stats['formatted_time'] = f"{total_hours}h {total_mins}m" if total_hours > 0 else f"{total_mins}m"
                library_stats['found'] = True
            else:
                library_stats['avg_duration_minutes'] = 0
                library_stats['note'] = f'User not found in library platform'
            
            if cursor_library:
                cursor_library.close()
            if conn_library:
                conn_library.close()
        except Exception as e:
            library_stats['error'] = str(e)
            library_stats['avg_duration_minutes'] = 0
        
        return {
            'smartlearning': smartlearning_stats,
            'library': library_stats
        }
    except Exception as e:
        return None


def _normalize_column_name(col_name):
    """Normalize column name for flexible matching"""
    if pd.isna(col_name):
        return ""
    # Convert to string, lowercase, remove all spaces and special chars
    normalized = str(col_name).lower().strip()
    # Remove spaces, underscores, hyphens for comparison
    normalized = normalized.replace(' ', '').replace('_', '').replace('-', '')
    return normalized


def _find_column(df, possible_names):
    """Find column in dataframe that matches any of the possible names"""
    # First try exact match after basic normalization (lowercase, strip)
    df_cols_basic_normalized = {str(col).lower().strip(): col for col in df.columns}
    
    for possible in possible_names:
        normalized_possible = str(possible).lower().strip()
        if normalized_possible in df_cols_basic_normalized:
            return df_cols_basic_normalized[normalized_possible]
    
    # Try aggressive normalization (remove spaces, underscores, hyphens)
    df_cols_normalized = {_normalize_column_name(col): col for col in df.columns}
    
    for possible in possible_names:
        normalized_possible = _normalize_column_name(possible)
        if normalized_possible in df_cols_normalized:
            return df_cols_normalized[normalized_possible]
    
    # Try contains matching (check if normalized column contains or is contained by normalized possible)
    for col in df.columns:
        col_normalized = _normalize_column_name(col)
        for possible in possible_names:
            possible_normalized = _normalize_column_name(possible)
            if possible_normalized in col_normalized or col_normalized in possible_normalized:
                return col
    
    return None


@app.route('/api/bulk-search-learners', methods=['POST'])
@login_required
def bulk_search_learners():
    """Bulk search learners from Excel file upload"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file or not file.filename:
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        if not (file.filename.lower().endswith('.xlsx') or file.filename.lower().endswith('.xls')):
            return jsonify({'error': 'Please upload an Excel file (.xlsx or .xls)'}), 400
        
        # Get date range (optional)
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')
        include_stats = request.form.get('include_stats', 'false').lower() == 'true'
        
        # Validate date range if stats are requested
        if include_stats:
            if not start_date or not end_date:
                return jsonify({'error': 'start_date and end_date are required when including statistics'}), 400
            if start_date > end_date:
                return jsonify({'error': 'start_date must be before or equal to end_date'}), 400
        
        # Read Excel file
        try:
            df = pd.read_excel(file)
        except Exception as e:
            return jsonify({'error': f'Error reading Excel file: {str(e)}'}), 400
        
        # Map possible column names
        column_mapping = {
            'firstname': ['firstname', 'first_name', 'first name'],
            'lastname': ['lastname', 'last_name', 'last name', 'surname'],
            'username': ['username', 'user_name', 'user name']
        }
        
        # Find actual column names using flexible matching
        firstname_col = _find_column(df, column_mapping['firstname'])
        lastname_col = _find_column(df, column_mapping['lastname'])
        username_col = _find_column(df, column_mapping['username'])
        
        # Validate required columns
        if not username_col and (not firstname_col or not lastname_col):
            # Build helpful error message
            found_cols = list(df.columns)
            error_msg = 'Missing required columns. '
            if username_col:
                error_msg += 'Found "username" column. '
            elif firstname_col and not lastname_col:
                error_msg += f'Found "firstname" column ("{firstname_col}") but missing "lastname". '
            elif lastname_col and not firstname_col:
                error_msg += f'Found "lastname" column ("{lastname_col}") but missing "firstname". '
            else:
                error_msg += 'Need either "username" OR both "firstname" and "lastname" columns. '
            
            error_msg += f'Found columns: {", ".join(found_cols)}. '
            error_msg += 'Acceptable column names: firstname/first_name/first name, lastname/last_name/last name/surname, username/user_name/user name'
            
            return jsonify({
                'error': error_msg,
                'required': 'Either "username" OR both "firstname" and "lastname" columns are required',
                'found_columns': found_cols
            }), 400
        
        # Process each row
        results = []
        conn = get_ruzivo_conn()
        cursor = conn.cursor()
        
        for idx, row in df.iterrows():
            row_result = {
                'row_number': idx + 2,  # +2 because Excel rows start at 1 and we skip header
                'input_data': {},
                'error': None,
                'profile': None,
                'time_spent': None,
                'platform_stats': None
            }
            
            # Extract input data
            if username_col:
                username_input = str(row[username_col]).strip() if pd.notna(row[username_col]) else ''
                row_result['input_data']['username'] = username_input
            else:
                firstname_input = str(row[firstname_col]).strip() if pd.notna(row[firstname_col]) else ''
                lastname_input = str(row[lastname_col]).strip() if pd.notna(row[lastname_col]) else ''
                row_result['input_data']['firstname'] = firstname_input
                row_result['input_data']['lastname'] = lastname_input
            
            # Find learner
            matched_username = None
            
            if username_col and username_input:
                # Match by username
                query = """
                    SELECT username FROM vwstudent WHERE username = %s LIMIT 1
                """
                cursor.execute(query, (username_input,))
                match = cursor.fetchone()
                if match:
                    matched_username = match.get('username') if isinstance(match, dict) else match[0]
            elif firstname_col and lastname_col and firstname_input and lastname_input:
                # Match by firstname + lastname
                full_name = f"{firstname_input} {lastname_input}".strip()
                query = """
                    SELECT username FROM vwstudent 
                    WHERE name LIKE %s
                    LIMIT 1
                """
                cursor.execute(query, (f"%{full_name}%",))
                match = cursor.fetchone()
                if match:
                    matched_username = match.get('username') if isinstance(match, dict) else match[0]
            
            if not matched_username:
                row_result['error'] = 'Learner not found'
                results.append(row_result)
                continue
            
            # Fetch profile data
            profile_data = _fetch_learner_profile_data(matched_username)
            if not profile_data:
                row_result['error'] = 'Error fetching profile data'
                results.append(row_result)
                continue
            
            row_result['profile'] = profile_data
            
            # Fetch time spent and platform stats if requested
            if include_stats and start_date and end_date:
                time_spent_data = _fetch_learner_time_spent_data(matched_username, start_date, end_date)
                row_result['time_spent'] = time_spent_data
                
                platform_stats_data = _fetch_learner_platform_stats_data(matched_username, start_date, end_date)
                row_result['platform_stats'] = platform_stats_data
            
            results.append(row_result)
        
        return jsonify({
            'results': results,
            'total_rows': len(results),
            'successful': len([r for r in results if r.get('profile')]),
            'failed': len([r for r in results if r.get('error')])
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-bulk-learner-results', methods=['POST'])
@login_required
def download_bulk_learner_results():
    """Download bulk learner results as Excel or CSV"""
    try:
        data = request.get_json()
        if not data or 'results' not in data:
            return jsonify({'error': 'No results data provided'}), 400
        
        format_type = request.args.get('format', 'excel').lower()
        if format_type not in ['excel', 'csv']:
            format_type = 'excel'
        
        results = data['results']
        
        # Prepare data for export
        export_rows = []
        for result in results:
            if result.get('error'):
                # Include error rows
                export_rows.append({
                    'Row Number': result.get('row_number', ''),
                    'Error': result.get('error', ''),
                    'Input Username': result.get('input_data', {}).get('username', ''),
                    'Input First Name': result.get('input_data', {}).get('firstname', ''),
                    'Input Last Name': result.get('input_data', {}).get('lastname', ''),
                    'Username': '',
                    'Full Name': '',
                    'Grade': '',
                    'School Name': '',
                    'School Province': '',
                    'Last Login': '',
                    'Access Start Date': '',
                    'Access End Date': '',
                    'Total Time Spent': '',
                    'Hours': '',
                    'Minutes': '',
                    'Seconds': '',
                    'Longest Streak': '',
                    'Most Active Day': '',
                    'Most Active Month': '',
                    'Primary Content Access': '',
                    'Secondary Content Access': '',
                    'Zimsec Access': '',
                    'Teacher Activities': '',
                    'Library Reading Time': '',
                    'Library Sessions': '',
                    'Library Avg Duration (min)': ''
                })
            else:
                profile = result.get('profile', {})
                time_spent = result.get('time_spent', {})
                platform_stats = result.get('platform_stats', {})
                smartlearning = platform_stats.get('smartlearning', {}) if platform_stats else {}
                library = platform_stats.get('library', {}) if platform_stats else {}
                
                export_rows.append({
                    'Row Number': result.get('row_number', ''),
                    'Error': '',
                    'Input Username': result.get('input_data', {}).get('username', ''),
                    'Input First Name': result.get('input_data', {}).get('firstname', ''),
                    'Input Last Name': result.get('input_data', {}).get('lastname', ''),
                    'Username': profile.get('username', ''),
                    'Full Name': profile.get('name', ''),
                    'Grade': profile.get('grade', ''),
                    'School Name': profile.get('school_name', ''),
                    'School Province': profile.get('school_province', ''),
                    'Last Login': profile.get('last_login', ''),
                    'Access Start Date': profile.get('access_sdate', ''),
                    'Access End Date': profile.get('access_edate', ''),
                    'Total Time Spent': time_spent.get('formatted', '') if time_spent else '',
                    'Hours': time_spent.get('hours', '') if time_spent else '',
                    'Minutes': time_spent.get('minutes', '') if time_spent else '',
                    'Seconds': time_spent.get('seconds', '') if time_spent else '',
                    'Longest Streak': time_spent.get('longest_streak', '') if time_spent else '',
                    'Most Active Day': time_spent.get('most_active_day', '') if time_spent else '',
                    'Most Active Month': time_spent.get('most_active_month', '') if time_spent else '',
                    'Primary Content Access': smartlearning.get('total_primary_content', ''),
                    'Secondary Content Access': smartlearning.get('total_sec_content', ''),
                    'Zimsec Access': smartlearning.get('total_zimsec_access', ''),
                    'Teacher Activities': smartlearning.get('teacher_access', ''),
                    'Library Reading Time': library.get('formatted_time', '') if library else '',
                    'Library Sessions': library.get('reading_sessions', '') if library else '',
                    'Library Avg Duration (min)': round(library.get('avg_duration_minutes', 0), 2) if library else ''
                })
        
        df = pd.DataFrame(export_rows)
        
        if format_type == 'excel':
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Learner Results')
            output.seek(0)
            filename = f'bulk_learner_results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                           as_attachment=True, download_name=filename)
        else:
            output = io.StringIO()
            df.to_csv(output, index=False)
            output.seek(0)
            filename = f'bulk_learner_results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            return Response(
                output.getvalue(),
                mimetype='text/csv',
                headers={'Content-Disposition': f'attachment; filename={filename}'}
            )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/word-to-html-converter', methods=['GET'])
@login_required
def word_to_html_converter():
    """Display Word to HTML converter page"""
    username = current_user.username
    role = current_user.userRole
    
    return render_template(
        'word_to_html_converter.html',
        username=username,
        role=role,
        title='Word to HTML Converter'
    )


@app.route('/api/convert-word-to-html', methods=['POST'])
@login_required
def api_convert_word_to_html():
    """Convert Word document or text to HTML"""
    try:
        # Check if file was uploaded
        if 'file' in request.files:
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': 'No file selected'}), 400
            
            if not file.filename.endswith('.docx'):
                return jsonify({'error': 'Only .docx files are supported'}), 400
            
            # Save uploaded file temporarily
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                file.save(tmp_file.name)
                tmp_path = tmp_file.name
            
            try:
                # Parse Word document
                doc = Document(tmp_path)
                content_lines = []
                images = []
                
                # Extract images
                image_dir = os.path.join(app.root_path, 'static', 'uploads', 'word_images')
                os.makedirs(image_dir, exist_ok=True)
                
                # Extract images from document
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_data = rel.target_part.blob
                        image_ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                        image_filename = f"{secrets.token_hex(8)}.{image_ext}"
                        image_path = os.path.join(image_dir, image_filename)
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        image_url = f"/static/uploads/word_images/{image_filename}"
                        images.append({'url': image_url, 'path': image_path})
                
                # Extract text content
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_lines.append(paragraph.text)
                
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            content_lines.append(row_text)
                
                content = '\n'.join(content_lines)
                
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        # Check if text was pasted
        elif 'text_content' in request.form:
            content = request.form.get('text_content', '').strip()
            images = []
        
        else:
            return jsonify({'error': 'Either file or text_content must be provided'}), 400
        
        if not content:
            return jsonify({'error': 'No content to convert'}), 400
        
        # Parse and convert to HTML
        html_output = _generate_html_from_content(content, images)
        
        # Post-process HTML to match bash script behavior
        html_output = _post_process_html(html_output)
        
        return jsonify({
            'html': html_output,
            'images': [img['url'] for img in images],
            'success': True
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


def _generate_html_from_content(content, images):
    """Generate HTML from parsed content - matches expected structure"""
    lines = [l.strip() for l in content.split('\n') if l.strip()]
    html_parts = []
    image_index = 0
    
    # Add required scripts at the top
    html_parts.append('<script src="https://ajax.googleapis.com/ajax/libs/jquery/3.6.0/jquery.min.js"></script>')
    html_parts.append('<script src="https://maxcdn.bootstrapcdn.com/bootstrap/3.4.1/js/bootstrap.min.js"></script>')
    html_parts.append('<!-- CONTENT STARTS HERE -->')
    html_parts.append('<div class="containerx">')
    html_parts.append('    <div class="row">')
    
    i = 0
    main_header = None
    in_objectives = False
    objectives_content = []
    lesson_topic = None
    in_main_content = False
    main_content_start_idx = None
    
    # Find main header (first line, or first ALL CAPS line)
    if len(lines) > 0:
        first_line = lines[0]
        if first_line.isupper() and len(first_line) > 2:
            main_header = first_line
        else:
            # Use first line as header even if not all caps
            main_header = first_line
        html_parts.append('        <div class="col-md-12">')
        html_parts.append(f'            <h2 class="bannerG">{main_header}</h2>')
        html_parts.append('        </div>')
        i = 1
    
    # Find lesson topic (line before "Objectives")
    objectives_idx = None
    for j in range(i, len(lines)):
        if 'objectives' in lines[j].lower():
            objectives_idx = j
            # Look backwards for lesson topic (1-2 lines before objectives)
            if j > 0:
                for k in range(max(0, j-2), j):
                    potential_topic = lines[k]
                    if (potential_topic and len(potential_topic) > 5 and 
                        not potential_topic.lower().startswith('objectives') and
                        not potential_topic.isupper() and
                        not potential_topic.startswith('-') and
                        not potential_topic.startswith('•')):
                        lesson_topic = potential_topic
                        break
            break
    
    # Find and process Objectives
    while i < len(lines):
        line = lines[i]
        
        if 'objectives' in line.lower():
            in_objectives = True
            objectives_content = []
            i += 1
            continue
        
        if in_objectives:
            if 'by the end' in line.lower():
                # Skip the "By the end" line, it's in the template
                i += 1
                continue
            elif any(line.strip().startswith(verb) for verb in ['List', 'Identify', 'Explain', 'Demonstrate', 'Describe', 'Analyze', 'Compare', 'Define']):
                # Objective item - remove leading dashes/bullets/numbers if any
                clean_line = re.sub(r'^[-•]\s*|\d+[\.\)]\s*', '', line).strip()
                if clean_line:
                    objectives_content.append(clean_line)
                i += 1
                continue
            elif (line.lower().startswith('the ') or line.lower().startswith('these ') or 
                  line.lower().startswith('as the ') or 
                  line.lower().startswith('natural ') or
                  line.lower().startswith('types of') or
                  len(line) > 30):
                # End of objectives, format it
                img_url = images[image_index]['url'] if image_index < len(images) else None
                html_parts.append(_format_objectives_section(objectives_content, img_url))
                if image_index < len(images):
                    image_index += 1
                in_objectives = False
                
                # Add lesson topic if we found it
                if lesson_topic:
                    html_parts.append('        <div class="col-md-12">')
                    html_parts.append('            <div class="col-md-12">')
                    html_parts.append(f'                <h2 class="bannerG">{lesson_topic}</h2>')
                    html_parts.append('            </div>')
                    html_parts.append('            <div class="col-md-12">')
                
                in_main_content = True
                main_content_start_idx = i
                break
            else:
                # Still collecting objectives or might be continuation
                if objectives_content:  # If we have objectives, this might be end
                    # Check if this looks like content (not an objective)
                    if (len(line) > 30 or 
                        line.lower().startswith('the ') or 
                        line.lower().startswith('they ') or
                        line.lower().startswith('types')):
                        # End objectives
                        img_url = images[image_index]['url'] if image_index < len(images) else None
                        html_parts.append(_format_objectives_section(objectives_content, img_url))
                        if image_index < len(images):
                            image_index += 1
                        in_objectives = False
                        if lesson_topic:
                            html_parts.append('        <div class="col-md-12">')
                            html_parts.append('            <div class="col-md-12">')
                            html_parts.append(f'                <h2 class="bannerG">{lesson_topic}</h2>')
                            html_parts.append('            </div>')
                            html_parts.append('            <div class="col-md-12">')
                        in_main_content = True
                        main_content_start_idx = i
                        break
                i += 1
                continue
    
    # Process main content
    if in_main_content and main_content_start_idx is not None:
        main_content_lines = lines[main_content_start_idx:]
        html_parts.extend(_process_main_content_advanced(main_content_lines, images, image_index))
        html_parts.append('            </div>')
        html_parts.append('        </div>')
    elif not in_objectives:
        # Process remaining lines as main content
        remaining_lines = lines[i:]
        if remaining_lines:
            if not lesson_topic:
                # Try to find lesson topic in remaining lines
                for line in remaining_lines[:5]:
                    if line and len(line) > 15 and not line.isupper() and not line.lower().startswith('objectives'):
                        lesson_topic = line
                        html_parts.append('        <div class="col-md-12">')
                        html_parts.append('            <div class="col-md-12">')
                        html_parts.append(f'                <h2 class="bannerG">{lesson_topic}</h2>')
                        html_parts.append('            </div>')
                        html_parts.append('            <div class="col-md-12">')
                        remaining_lines = remaining_lines[1:]
                        break
            
            if lesson_topic or any('the ' in l.lower()[:10] for l in remaining_lines[:3]):
                html_parts.extend(_process_main_content_advanced(remaining_lines, images, image_index))
                html_parts.append('            </div>')
                html_parts.append('        </div>')
    
    # Close container divs
    html_parts.append('        <!--end of st1-->')
    html_parts.append('    </div>')
    html_parts.append('</div>')
    
    return '\n'.join(html_parts)


def _post_process_html(html):
    """Post-process HTML to match bash script transformations"""
    import re
    
    # Remove blockquotes
    html = re.sub(r'<blockquote>', '', html)
    html = re.sub(r'</blockquote>', '', html)
    
    # Convert <strong> to <b> and <em> to <i>
    html = re.sub(r'<strong>', '<b>', html)
    html = re.sub(r'</strong>', '</b>', html)
    html = re.sub(r'<em>', '<i>', html)
    html = re.sub(r'</em>', '</i>', html)
    
    # Remove <p> tags inside <li> tags (but keep the content)
    # Pattern: <li><p>content</p></li> -> <li>content</li>
    html = re.sub(r'<li>\s*<p>(.*?)</p>\s*</li>', r'<li>\1</li>', html, flags=re.DOTALL)
    
    # Handle paragraphs containing only images - convert to divs
    # Pattern: <p><img ...></p> -> <div><img ...></div>
    html = re.sub(r'<p>\s*(<img[^>]*>)\s*</p>', r'<div>\1</div>', html, flags=re.IGNORECASE)
    
    # Handle "Fig X" captions - add line break after <b>Fig X</b>
    # Pattern: <p><b>Fig X</b> caption</p> -> <p><b>Fig X</b><br>caption</p>
    html = re.sub(r'(<p><b>Fig\s+\d+[^<]*</b>)([^<])', r'\1<br>\2', html, flags=re.IGNORECASE)
    
    # Clean up any double line breaks
    html = re.sub(r'<br>\s*<br>', '<br>', html)
    
    return html


def _find_lesson_topic_before_objectives(lines, obj_idx):
    """Find the lesson topic before objectives"""
    # Look backwards from objectives (check 1-5 lines before)
    for i in range(max(0, obj_idx - 5), obj_idx):
        if i < len(lines):
            line = lines[i]
            # Lesson topic is usually a substantial line that's not ALL CAPS, not objectives, not a list item
            if (line and len(line) > 15 and not line.isupper() and 
                not line.lower().startswith('objectives') and
                not line.lower().startswith('by the end') and
                not line.startswith('-') and not line.startswith('•') and
                not re.match(r'^\d+[\.\)]', line) and
                not any(line.startswith(verb) for verb in ['List', 'Identify', 'Explain', 'Demonstrate'])):
                return line
    return None


def _process_main_content_advanced(lines, images, image_index):
    """Process main content - simplified and robust version"""
    html_parts = []
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Check for "Fig X" pattern
        fig_match = re.match(r'^Fig\s+(\d+)\s*(.*)$', line, re.IGNORECASE)
        if fig_match:
            fig_num = fig_match.group(1)
            caption = fig_match.group(2).strip()
            img_url = images[image_index]['url'] if image_index < len(images) else "https://smartlearning.akello.co/public/uploads/content/HBC%20Social%20Science%20Grade%207/Lesson%2001%20Respiratory%20System_files/media/image1.png"
            html_parts.append(f'                <div><img src="{img_url}" style="width: 4.95689in; height: 4.18535in; display: block; margin-left: auto; margin-right: auto;" alt="image1.png (240 KB)" width="60%" height="40%" caption="false" /></div>')
            html_parts.append(f'                <div style="text-align: center;">Fig {fig_num} {caption}</div>')
            if image_index < len(images):
                image_index += 1
            i += 1
            continue
        
        # Check for Activity
        activity_match = re.match(r'^Activity\s+(\d+)\s*(.*)$', line, re.IGNORECASE)
        if activity_match:
            activity_num = activity_match.group(1)
            activity_content = activity_match.group(2).strip()
            html_parts.append('                <div class="col-md-12">')
            html_parts.append('                    <div class="row banner-area1">')
            html_parts.append('                        <div class="col-md-12">')
            html_parts.append(f'                            <h2 class="primary-activity-banner">Activity {activity_num}</h2>')
            html_parts.append('                        </div>')
            html_parts.append('                        <div style="background-color: #f2f2f2; justify-content: center; width: 96%; margin: auto;">')
            html_parts.append(f'                            {activity_content}')
            html_parts.append('                        </div>')
            html_parts.append('                    </div>')
            html_parts.append('                </div>')
            html_parts.append('                <br /><br />')
            i += 1
            continue
        
        # Check if line ends with colon and introduces a list
        if line.endswith(':'):
            html_parts.append(f'                <p>{line}</p>')
            i += 1
            # Collect list items
            list_items = []
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    continue
                # Check if it's a list item
                if (next_line.startswith('-') or next_line.startswith('•') or
                    re.match(r'^[A-Z][^:]*:', next_line) or  # "Earthquakes:", "Floods:", etc.
                    (len(next_line.split()) <= 6 and next_line[0].isupper() and 
                     not next_line.endswith('.') and not 'Fig' in next_line and
                     not 'Activity' in next_line)):
                    list_items.append(next_line)
                    i += 1
                else:
                    break
            
            if list_items:
                html_parts.append(_format_simple_list(list_items))
                continue
        
        # Regular paragraph
        html_parts.append(f'                <p>{line}</p>')
        i += 1
    
    return html_parts


def _format_simple_list(list_items):
    """Format a simple list"""
    html_parts = []
    html_parts.append('                <ul>')
    for item in list_items:
        clean = re.sub(r'^[-•]\s*', '', item).strip()
        # If item has colon, format as bold title
        if ':' in clean:
            parts = clean.split(':', 1)
            html_parts.append(f'                    <li><b>{parts[0]}:</b>{parts[1] if len(parts) > 1 else ""}</li>')
        else:
            html_parts.append(f'                    <li>{clean}</li>')
    html_parts.append('                </ul>')
    return '\n'.join(html_parts)


def _format_roman_list(list_items, all_lines, next_idx):
    """Format a simple lower-roman list"""
    html_parts = []
    html_parts.append('                <ol style="list-style-type: lower-roman;">')
    
    for item in list_items:
        clean = re.sub(r'^[-•]\s*|[a-z]\)\s*', '', item, flags=re.IGNORECASE).strip()
        html_parts.append(f'                    <li>{clean}</li>')
    
    html_parts.append('                </ol>')
    return '\n'.join(html_parts)


def _format_nested_roman_list(main_items, all_lines):
    """Format nested list with lower-roman outer and nested ul/ol inner"""
    html_parts = []
    html_parts.append('                <ol style="list-style-type: lower-roman;">')
    
    for idx, (line_idx, title) in enumerate(main_items):
        html_parts.append('                    <li>')
        html_parts.append(f'                        <p>{title}</p>')
        
        # Collect content for this item
        start_idx = line_idx + 1
        end_idx = main_items[idx + 1][0] if idx + 1 < len(main_items) else len(all_lines)
        
        item_content = []
        for j in range(start_idx, end_idx):
            if j < len(all_lines):
                item_content.append(all_lines[j])
        
        if item_content:
            html_parts.append('                        <ul>')
            for content_line in item_content:
                # Skip empty lines and figure captions (they'll be handled separately)
                if not content_line.strip() or content_line.strip().startswith('Fig '):
                    continue
                clean = re.sub(r'^[-•]\s*', '', content_line).strip()
                if clean:
                    html_parts.append(f'                            <li>{clean}</li>')
            html_parts.append('                        </ul>')
        
        html_parts.append('                    </li>')
    
    html_parts.append('                </ol>')
    return '\n'.join(html_parts)


def _format_objectives_section(content, image_url=None):
    """Format objectives section with banner"""
    html = []
    html.append('        <div class="col-md-12">')
    html.append('            <div class="row banner-area">')
    
    if image_url:
        html.append(f'                <div class="col-md-6  banner-image"><img class="img-fluid" src="{image_url}" alt="Objective Icons" style="width: 100%;" caption="false" /></div>')
    else:
        html.append('                <div class="col-md-6  banner-image"><img class="img-fluid" src="https://asladmin.akello.co/public/uploads/content/Maths%20Grade%206/Objectives%20Icon%20/__54.png" alt="Objective Icons form 4-43.png (148 KB)" style="width: 100%;" caption="false" /></div>')
    
    html.append('                <div class="col-md-6">')
    html.append('                    <h2>Objectives</h2>')
    html.append('                    <h4 class="padding-5">By the end of this lesson, you should be able to:</h4>')
    html.append('                    <ol type="1">')
    for item in content:
        if item.strip():
            html.append(f'                        <li>{item.strip()}</li>')
    html.append('                    </ol>')
    html.append('                </div>')
    html.append('            </div>')
    html.append('        </div>')
    return '\n'.join(html)


@app.route('/api/top-learners', methods=['GET'])
@login_required
def get_top_learners():
    """Get top 20 learners across platforms with various filters"""
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        platform = request.args.get('platform', 'overall')  # overall, smartlearning, library
        province = request.args.get('province')  # Optional province filter
        gender = request.args.get('gender')  # Optional gender filter: male, female, unknown
        
        if not start_date or not end_date:
            return jsonify({'error': 'start_date and end_date are required'}), 400
        
        excluded_ids = [356,905,30843,265227,265228,265230,3223972,406978,516518,577527,1032617,1621500,1632975,1660731,1660834,1661007,1661053,1661353,1662420,1662839,1664013,1664021,1664698,1684433,1685102]
        excluded_placeholders = ','.join(['%s'] * len(excluded_ids))
        
        results = []
        
        # SmartLearning Platform Top Users
        if platform in ['overall', 'smartlearning']:
            conn_ruzivo = get_ruzivo_conn()
            cursor_ruzivo = conn_ruzivo.cursor()
            
            # Build WHERE clause for province and gender
            where_clauses = [
                f"student_id NOT IN ({excluded_placeholders})",
                "start_time BETWEEN %s AND %s"
            ]
            params = list(excluded_ids) + [start_date, end_date]
            
            if province:
                where_clauses.append("s.school_province = %s")
                params.append(province)
            
            if gender:
                gender_condition = {
                    'male': "UPPER(TRIM(vs.gender)) IN ('MALE','M')",
                    'female': "UPPER(TRIM(vs.gender)) IN ('FEMALE','F')",
                    'unknown': "UPPER(TRIM(vs.gender)) NOT IN ('MALE','M','FEMALE','F') OR vs.gender IS NULL"
                }
                if gender in gender_condition:
                    where_clauses.append(gender_condition[gender])
            
            where_sql = " AND ".join(where_clauses)
            
            # Calculate total time spent per student
            smartlearning_query = f"""
                SELECT 
                    vs.student_id,
                    vs.username,
                    vs.name,
                    s.school_province,
                    CASE
                        WHEN UPPER(TRIM(vs.gender)) IN ('MALE','M') THEN 'Male'
                        WHEN UPPER(TRIM(vs.gender)) IN ('FEMALE','F') THEN 'Female'
                        ELSE 'Unknown'
                    END AS gender,
                    COALESCE(SUM(
                        TIMESTAMPDIFF(SECOND, 
                            COALESCE(ca1.start_time, ca2.start_time, ca3.start_time),
                            COALESCE(ca1.end_time, ca2.end_time, ca3.end_time)
                        )
                    ), 0) AS total_seconds
                FROM (
                    SELECT student_id, start_time, end_time
                    FROM tblcontent_access
                    WHERE {where_sql.replace('s.school_province', '1=1').replace('vs.gender', '1=1')}
                    
                    UNION ALL
                    
                    SELECT student_id, start_time, end_time
                    FROM tblcontent_access_hs
                    WHERE {where_sql.replace('s.school_province', '1=1').replace('vs.gender', '1=1')}
                    
                    UNION ALL
                    
                    SELECT student_id, start_time, end_time
                    FROM tblcontent_access_zimsec
                    WHERE {where_sql.replace('s.school_province', '1=1').replace('vs.gender', '1=1')}
                ) AS combined
                LEFT JOIN tblcontent_access ca1 ON combined.student_id = ca1.student_id 
                    AND combined.start_time = ca1.start_time
                LEFT JOIN tblcontent_access_hs ca2 ON combined.student_id = ca2.student_id 
                    AND combined.start_time = ca2.start_time
                LEFT JOIN tblcontent_access_zimsec ca3 ON combined.student_id = ca3.student_id 
                    AND combined.start_time = ca3.start_time
                JOIN vwstudent vs ON vs.student_id = combined.student_id
                JOIN tblschools s ON s.school_id = vs.school_id
                WHERE {where_sql}
                GROUP BY vs.student_id, vs.username, vs.name, s.school_province, vs.gender
                ORDER BY total_seconds DESC
                LIMIT 20
            """
            
            # Simplified approach - calculate from each table separately and join
            smartlearning_simple_query = f"""
                SELECT 
                    vs.student_id,
                    vs.username,
                    vs.name,
                    s.school_province,
                    CASE
                        WHEN UPPER(TRIM(vs.gender)) IN ('MALE','M') THEN 'Male'
                        WHEN UPPER(TRIM(vs.gender)) IN ('FEMALE','F') THEN 'Female'
                        ELSE 'Unknown'
                    END AS gender,
                    COALESCE(SUM(
                        TIMESTAMPDIFF(SECOND, ca.start_time, COALESCE(ca.end_time, ca.start_time))
                    ), 0) AS total_seconds
                FROM tblcontent_access ca
                JOIN vwstudent vs ON vs.student_id = ca.student_id
                JOIN tblschools s ON s.school_id = vs.school_id
                WHERE ca.start_time BETWEEN %s AND %s
                    AND ca.student_id NOT IN ({excluded_placeholders})
            """
            query_params = [start_date, end_date] + excluded_ids
            
            if province:
                smartlearning_simple_query += " AND s.school_province = %s"
                query_params.append(province)
            
            if gender:
                if gender == 'male':
                    smartlearning_simple_query += " AND UPPER(TRIM(vs.gender)) IN ('MALE','M')"
                elif gender == 'female':
                    smartlearning_simple_query += " AND UPPER(TRIM(vs.gender)) IN ('FEMALE','F')"
                elif gender == 'unknown':
                    smartlearning_simple_query += " AND (UPPER(TRIM(vs.gender)) NOT IN ('MALE','M','FEMALE','F') OR vs.gender IS NULL)"
            
            smartlearning_simple_query += """
                GROUP BY vs.student_id, vs.username, vs.name, s.school_province, vs.gender
            """
            
            # Also include secondary and zimsec
            smartlearning_query_final = f"""
                SELECT 
                    student_id,
                    username,
                    name,
                    school_province,
                    gender,
                    SUM(total_seconds) AS total_seconds
                FROM (
                    {smartlearning_simple_query.replace('tblcontent_access', 'tblcontent_access').replace('ca.', 'ca1.')}
                    UNION ALL
                    {smartlearning_simple_query.replace('tblcontent_access', 'tblcontent_access_hs').replace('ca.', 'ca2.')}
                    UNION ALL
                    {smartlearning_simple_query.replace('tblcontent_access', 'tblcontent_access_zimsec').replace('ca.', 'ca3.')}
                ) AS combined
                GROUP BY student_id, username, name, school_province, gender
                ORDER BY total_seconds DESC
                LIMIT 20
            """
            
            # Use simpler approach - query each table and combine
            all_sl_users = {}
            
            for table in ['tblcontent_access', 'tblcontent_access_hs', 'tblcontent_access_zimsec']:
                table_query = f"""
                    SELECT 
                        vs.student_id,
                        vs.username,
                        vs.name,
                        s.school_province,
                        CASE
                            WHEN UPPER(TRIM(vs.gender)) IN ('MALE','M') THEN 'Male'
                            WHEN UPPER(TRIM(vs.gender)) IN ('FEMALE','F') THEN 'Female'
                            ELSE 'Unknown'
                        END AS gender,
                        SUM(TIMESTAMPDIFF(SECOND, ca.start_time, COALESCE(ca.end_time, ca.start_time))) AS total_seconds
                    FROM {table} ca
                    JOIN vwstudent vs ON vs.student_id = ca.student_id
                    JOIN tblschools s ON s.school_id = vs.school_id
                    WHERE ca.start_time BETWEEN %s AND %s
                        AND ca.student_id NOT IN ({excluded_placeholders})
                """
                table_params = [start_date, end_date] + excluded_ids
                
                if province:
                    table_query += " AND s.school_province = %s"
                    table_params.append(province)
                
                if gender:
                    if gender == 'male':
                        table_query += " AND UPPER(TRIM(vs.gender)) IN ('MALE','M')"
                    elif gender == 'female':
                        table_query += " AND UPPER(TRIM(vs.gender)) IN ('FEMALE','F')"
                    elif gender == 'unknown':
                        table_query += " AND (UPPER(TRIM(vs.gender)) NOT IN ('MALE','M','FEMALE','F') OR vs.gender IS NULL)"
                
                table_query += " GROUP BY vs.student_id, vs.username, vs.name, s.school_province, vs.gender"
                
                cursor_ruzivo.execute(table_query, table_params)
                rows = cursor_ruzivo.fetchall()
                
                for row in rows:
                    student_id = row.get('student_id') if isinstance(row, dict) else row[0]
                    if student_id not in all_sl_users:
                        all_sl_users[student_id] = {
                            'student_id': student_id,
                            'username': row.get('username') if isinstance(row, dict) else row[1],
                            'name': row.get('name') if isinstance(row, dict) else row[2],
                            'school_province': row.get('school_province') if isinstance(row, dict) else row[3],
                            'gender': row.get('gender') if isinstance(row, dict) else row[4],
                            'total_seconds': 0,
                            'platform': 'SmartLearning'
                        }
                    all_sl_users[student_id]['total_seconds'] += row.get('total_seconds') if isinstance(row, dict) else (row[5] or 0)
            
            for user_data in all_sl_users.values():
                hours = user_data['total_seconds'] // 3600
                minutes = (user_data['total_seconds'] % 3600) // 60
                user_data['formatted_time'] = f"{hours}h {minutes}m" if hours > 0 else f"{minutes}m"
                user_data['hours'] = hours
                user_data['minutes'] = minutes
            
            if platform == 'smartlearning':
                results = sorted(list(all_sl_users.values()), key=lambda x: x['total_seconds'], reverse=True)[:20]
            elif platform == 'overall':
                # Add SmartLearning users to overall results
                results.extend(sorted(list(all_sl_users.values()), key=lambda x: x['total_seconds'], reverse=True)[:20])
        
        # Library Platform Top Users
        if platform in ['overall', 'library']:
            try:
                conn_library = get_direct_library_conn()
                cursor_library = conn_library.cursor()
                
                library_query = """
                    SELECT 
                        u.id AS user_id,
                        u.first_name,
                        u.last_name,
                        u.username,
                        AVG(rt.duration_minutes) AS avg_duration_minutes,
                        SUM(rt.duration_minutes) AS total_duration_minutes,
                        COUNT(*) AS reading_sessions
                    FROM read_trackers rt
                    JOIN users u ON u.id = rt.user_id
                    WHERE rt.duration_minutes != 0
                        AND DATE(rt.created_at) BETWEEN DATE(%s) AND DATE(%s)
                """
                library_params = [start_date, end_date]
                
                # Note: Library database might not have province/gender fields easily accessible
                # We'll add them if the schema supports it
                
                library_query += """
                    GROUP BY u.id, u.first_name, u.last_name, u.username
                    ORDER BY total_duration_minutes DESC
                    LIMIT 20
                """
                
                cursor_library.execute(library_query, library_params)
                library_rows = cursor_library.fetchall()
                
                for row in library_rows:
                    user_id = row.get('user_id') if isinstance(row, dict) else row[0]
                    total_minutes = row.get('total_duration_minutes') if isinstance(row, dict) else (row[5] or 0)
                    total_seconds = int(total_minutes * 60)
                    hours = total_seconds // 3600
                    minutes = (total_seconds % 3600) // 60
                    
                    library_user = {
                        'user_id': user_id,
                        'username': row.get('username') if isinstance(row, dict) else row[3],
                        'name': f"{row.get('first_name', '') if isinstance(row, dict) else row[1]} {row.get('last_name', '') if isinstance(row, dict) else row[2]}".strip(),
                        'school_province': 'N/A',  # Library might not have province
                        'gender': 'Unknown',  # Library might not have gender
                        'total_seconds': total_seconds,
                        'formatted_time': f"{hours}h {minutes}m" if hours > 0 else f"{minutes}m",
                        'hours': hours,
                        'minutes': minutes,
                        'platform': 'Library',
                        'avg_duration_minutes': float(row.get('avg_duration_minutes')) if isinstance(row, dict) else (row[4] or 0),
                        'reading_sessions': row.get('reading_sessions') if isinstance(row, dict) else (row[6] or 0)
                    }
                    
                    if platform == 'library':
                        results.append(library_user)
                    elif platform == 'overall':
                        # For overall, we need to match by name and combine
                        # This is complex, so we'll just add library users separately
                        results.append(library_user)
                
                if cursor_library:
                    cursor_library.close()
                if conn_library:
                    conn_library.close()
            except Exception as e:
                import traceback
                traceback.print_exc()
                # Continue even if library query fails
        
        # Sort overall results
        if platform == 'overall':
            results = sorted(results, key=lambda x: x.get('total_seconds', 0), reverse=True)[:20]
        elif platform == 'smartlearning':
            results = sorted(results, key=lambda x: x.get('total_seconds', 0), reverse=True)[:20]
        elif platform == 'library':
            results = sorted(results, key=lambda x: x.get('total_seconds', 0), reverse=True)[:20]
        
        return jsonify({
            'learners': results,
            'count': len(results),
            'platform': platform,
            'province': province,
            'gender': gender,
            'start_date': start_date,
            'end_date': end_date
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500





@app.route('/download_user_scorecard', methods=['GET'])
def download_user_scorecard():
    """
    Downloads the scorecard table as a CSV file.

    Returns:
        Response: A Flask Response object containing the downloaded CSV data.
    """
    scorecards = Scorecard.query.all()

    # Create a CSV string in memory
    csv_string = io.StringIO()
    writer = csv.writer(csv_string)

    # Write header row
    writer.writerow([
        'Key Focus Area', 'Strategic Objective', 'Performance Measure',
        'Unit Of Measure', 'Target', 'Weight'
    ])

    # Write data rows
    for scorecard in scorecards:
        writer.writerow([
            scorecard.key_focus_area, scorecard.strategic_objective,
            scorecard.performance_measure, scorecard.unit_of_measure,
            scorecard.target, scorecard.weight
        ])

    csv_string.seek(0)
    return Response(
        csv_string.read(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=scorecard.csv'}
    )



@app.route('/save_scorecard', methods=['POST'])
def save_scorecard():
    key_focus_areas = request.form.getlist('keyFocusArea[]')
    strategic_objectives = request.form.getlist('strategicObjective[]')
    performance_measures = request.form.getlist('performanceMeasure[]')
    unit_of_measures = request.form.getlist('unitOfMeasure[]')
    targets = request.form.getlist('target[]')
    weights = request.form.getlist('weight[]')
    employee = current_user.username

    for i in range(len(key_focus_areas)):
        scorecard = Scorecard(
            key_focus_area=key_focus_areas[i],
            strategic_objective=strategic_objectives[i],
            performance_measure=performance_measures[i],
            unit_of_measure=unit_of_measures[i],
            target=targets[i],
            weight=float(weights[i]),
            employee_name = employee
        )
        db.session.add(scorecard)

    db.session.commit()
    flash('Scorecards saved successfully!', 'success')
    return redirect(url_for('contracting'))





# ---------- Password reset helpers ----------

def _get_serializer():
    secret_key = app.config.get('SECRET_KEY')
    return URLSafeTimedSerializer(secret_key)


def generate_reset_token(email: str) -> str:
    serializer = _get_serializer()
    salt = app.config.get('SECURITY_PASSWORD_SALT', 'change-this-salt')
    return serializer.dumps(email, salt=salt)


def verify_reset_token(token: str, max_age: int = 3600):
    serializer = _get_serializer()
    salt = app.config.get('SECURITY_PASSWORD_SALT', 'change-this-salt')
    try:
        email = serializer.loads(token, salt=salt, max_age=max_age)
        return email
    except (SignatureExpired, BadSignature):
        return None


def send_email(subject: str, recipient: str, html_body: str, text_body: str = None):
    if app.config.get('MAIL_SUPPRESS_SEND'):
        # Suppressed sending: log instead
        app.logger.info(f"[MAIL_SUPPRESS_SEND] To: {recipient}\nSubject: {subject}\n{text_body or ''}\n{html_body}")
        return True

    sender = app.config.get('MAIL_DEFAULT_SENDER', app.config.get('MAIL_USERNAME'))
    if not sender:
        app.logger.warning('Email not sent: MAIL_USERNAME/PASSWORD/SENDER not fully configured')
        return False

    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = sender
    msg['To'] = recipient
    if text_body:
        msg.set_content(text_body)
    msg.add_alternative(html_body, subtype='html')
    return _send_email_message(msg)


@app.route('/forgot-password', methods=['GET'])
def forgot_password():
    # Simple page with an email field that POSTs to /request-password-reset (already exists as template)
    return render_template('forgot-password.html')


@app.route('/request-password-reset', methods=['POST'])
def request_password_reset():
    data = request.get_json(silent=True) or {}
    email = (data.get('email') or '').strip()
    if not email:
        return jsonify({"message": "Email is required"}), 400
    user = db.session.scalar(sa.select(User).where(User.email == email))
    # Always respond with success message to avoid user enumeration
    if user:
        token = generate_reset_token(email)
        reset_url = url_for('reset_password', token=token, _external=True)
        subject = 'Your password reset link'
        html_body = f"""
            <p>Hello {user.firstname or user.username},</p>
            <p>You requested a password reset. Click the link below to set a new password:</p>
            <p><a href="{reset_url}">{reset_url}</a></p>
            <p>If you did not request this, please ignore this email.</p>
        """
        text_body = f"Hello {user.firstname or user.username},\n\nVisit this link to reset your password: {reset_url}\nIf you did not request this, ignore this email."
        send_email(subject, email, html_body, text_body)
    return jsonify({"message": "If that email exists, a reset link has been sent."}), 200


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    from app.forms import ResetPasswordForm
    email = verify_reset_token(token)
    if not email:
        flash('The reset link is invalid or has expired.')
        return redirect(url_for('forgot_password'))

    form = ResetPasswordForm()
    if form.validate_on_submit():
        user = db.session.scalar(sa.select(User).where(User.email == email))
        if not user:
            flash('User not found.')
            return redirect(url_for('forgot_password'))
        user.set_password(form.password.data)
        db.session.commit()
        flash('Your password has been reset. Please log in.')
        return redirect(url_for('login'))

    return render_template('reset-password.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        if current_user.userRole == 'Brand Ambassador':
                return redirect(url_for('profile',username=current_user.username))
        else:
            return redirect(url_for('overview'))
    form = LoginForm()
    if form.validate_on_submit():
        user = db.session.scalar(
            sa.select(User).where(User.username == form.username.data))
        if user is None or not user.check_password(form.password.data):
            flash('Invalid username or password')
            return redirect(url_for('login'))
        login_user(user, remember=form.remember_me.data)
        next_page = request.args.get('next')
        # Ensure non-Brand Ambassador users land on /overview (not /welcome) after login.
        # Keep Brand Ambassador users redirected to their /profile.
        if not next_page or urlsplit(next_page).netloc != '' or next_page == url_for('overview'):
            if (user.userRole == 'Brand Ambassador') or (getattr(current_user, 'userRole', None) == 'Brand Ambassador'):
                next_page = url_for('profile', username=user.username)
            else:
                next_page = url_for('overview')
        return redirect(next_page)
    return render_template('login.html', title='Sign In', form=form)



@app.route('/helpdesk')
def helpdesk():

    return render_template('help_desk.html', title='Help desk')


# ---------- Email helpdesk tickets (IMAP) ----------
@app.route('/dashboard')
@login_required
def email_tickets_dashboard():
    """Show all email-derived tickets in a table."""
    tickets = Ticket.query.order_by(Ticket.created_at.desc()).all()
    return render_template('dashboard.html', title='Email tickets', tickets=tickets)


@app.route('/ticket/<int:id>')
@login_required
def ticket_detail(id):
    """View a single ticket by id."""
    ticket = Ticket.query.get_or_404(id)
    return render_template('ticket.html', title=f'Ticket #{ticket.id}', ticket=ticket)


@app.route('/check-email', methods=['GET', 'POST'])
@login_required
def check_email():
    """Manually trigger email fetch and create tickets."""
    from app.email_fetcher import fetch_emails_and_create_tickets
    count, err = fetch_emails_and_create_tickets(app)
    if err:
        flash(f'Email check failed: {err}', 'warning')
    else:
        flash(f'Fetched {count} new ticket(s) from email.', 'success')
    return redirect(url_for('email_tickets_dashboard'))


@app.route('/api/tickets')
@login_required
def api_tickets():
    """Return all tickets as JSON."""
    tickets = Ticket.query.order_by(Ticket.created_at.desc()).all()
    return jsonify([{
        'id': t.id,
        'sender_email': t.sender_email,
        'subject': t.subject,
        'message': t.message,
        'status': t.status,
        'created_at': t.created_at.isoformat() if t.created_at else None,
    } for t in tickets])


@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('welcome'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    """
    Handles user registration.
    If the user is already logged in, they are redirected to the index page.
    On form submission, it validates the data and creates a new user
    with all the required fields from the updated User model.
    """
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = RegistrationForm()
    
    if form.validate_on_submit():
        # Create a new User instance with the data from the registration form
        user = User(
            username=form.username.data,
            email=form.email.data,
            firstname=form.firstname.data,
            lastname=form.lastname.data,
            userRole=form.userRole.data,
            department=form.department.data,
            province=form.province.data,
            # Privileges are usually not set at registration.
            # This can be handled by an admin later.
            # We can set a default value if needed, e.g., privileges={}
        )
        
        # Set the user's password
        user.set_password(form.password.data)
        
        # Add the new user to the database
        db.session.add(user)
        db.session.commit()
        
        flash('Congratulations, you are now a registered user!')
        return redirect(url_for('login'))
        
    return render_template('register.html', title='Register', form=form)





# VTL DILY REPORT 
import pymysql
import pandas as pd
from flask import Flask, request, render_template, jsonify, send_file
import pandas as pd
from openpyxl import load_workbook
import os
from datetime import datetime


# ---------- DB Fetch ----------

import re

today = datetime.today()

# ---------- Excel Update ----------
# def update_excel(data, template_path="report_template.xlsx", reports_dir="reports"):
#     wb = load_workbook(template_path)
#     ws = wb.worksheets[0]  # Access the first worksheet

#     # Debugging info
#     print("DataFrame head:\n", data.head())
#     print("Dtypes:", data.dtypes)

#     # Extract numeric safely
#     value = pd.to_numeric(data['total_count'], errors="coerce").fillna(0).astype(int).iloc[0]
#     ws["C7"].value = value
#     ws["M7"].value = '=(H7*31)/21'

#     # Get the current formula or value in H7
#     h7_value = ws["H7"].value

#     # Extract the last integer from H7 value if it contains a formula
#     if isinstance(h7_value, str) and '+' in h7_value:
#         last_integer = re.findall(r'\d+', h7_value.strip())[-1]  # Find all integers and get the last one
#         ws["F7"].value = int(last_integer)  # Enter the last integer into F7
        

#     # Append +C7 to whatever is in H7
#     if h7_value:
#         ws["H7"].value = f"{h7_value}"
#     else:
#         ws["H7"].value = f"{ws['C7'].value}"

    

#     os.makedirs(reports_dir, exist_ok=True)

#     output_filename = f"daily_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
#     output_path = os.path.join(reports_dir, output_filename)

#     wb.save(output_path)
#     return output_filename, output_path


def update_excel(data, template_path="report_template.xlsx", reports_dir="reports", report_day=None):
    wb = load_workbook(template_path)
    ws = wb.worksheets[0]   # or ws["SheetName"]

    # Debugging info
    print("DataFrame head:\n", data.head())
    print("Dtypes:", data.dtypes)

    # Extract numeric safely
    value = pd.to_numeric(data['total_count'], errors="coerce").fillna(0).astype(int).iloc[0]
    day_of_month = int(report_day or (datetime.today() - timedelta(days=1)).day)
    ws["C7"].value = value
    ws["M7"].value = f'=(H7*31)/{day_of_month}'
    ws["I7"].value = f'=D7*{day_of_month}'

    # Handle H7 formula
    if ws["H7"].value:
        h7_formula = str(ws["H7"].value)

        # Extract all integers from the formula
        numbers = re.findall(r"\d+", h7_formula)
        if numbers:
            last_int = int(numbers[-1])
            ws["F7"].value = last_int  # put last number into F7

        # Append +C7 reference to formula
        ws["H7"].value = f"{h7_formula}+{value}"
    else:
        ws["H7"].value = f"{value}"

    os.makedirs(reports_dir, exist_ok=True)

    output_filename = f"daily_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    output_path = os.path.join(reports_dir, output_filename)

    wb.save(output_path)
    return output_filename, output_path



# ---------- Routes ----------
@app.route('/vtl_template', methods=['GET', 'POST'])
def vtl_template():
    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        if not file.filename.lower().endswith(".xlsx"):
            return jsonify({"error": "Only .xlsx files are supported"}), 400

        template_path = os.path.join("instance", "uploaded_template.xlsx")
        os.makedirs("instance", exist_ok=True)
        file.save(template_path)

        return jsonify({
            "ok": True,
            "message": "Template uploaded successfully",
            "stored_as": "uploaded_template.xlsx"
        })

    # GET → render upload page
    return render_template('vtl_template.html', template_filename="uploaded_template.xlsx")


import pandas as pd
import os
from datetime import datetime
from flask import send_file, jsonify

# ---------- DB Fetch ----------
def fetch_data(as_of=None):
    if as_of is None:
        as_of = datetime.today() - timedelta(days=1)

    target_day_start = as_of.replace(hour=0, minute=0, second=0, microsecond=0)
    target_day_end = as_of.replace(hour=23, minute=59, second=59, microsecond=0)
    month_start = as_of.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    conn = get_ruzivo_conn()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT COUNT(DISTINCT sl.student_id) AS total_count
        FROM tblstudents_login sl
        WHERE sl.student_id NOT IN (
            SELECT DISTINCT tp.student_id
            FROM tblecocash_payment_order tp
            WHERE transactionOperationStatus='COMPLETED'
        )
        AND sl.student_id NOT IN (
            SELECT DISTINCT sl2.student_id
            FROM tblstudents_login sl2
            WHERE sl2.login_date BETWEEN %s AND %s
        )
        AND sl.login_date BETWEEN %s AND %s;
    """, (month_start, target_day_end, target_day_start, target_day_end))
    rows = cursor.fetchall()
    df = pd.DataFrame(rows, columns=[desc[0] for desc in cursor.description])
    cursor.close()
    conn.close()
    return df




def generate_and_send_report(template_path, as_of=None):
    """
    Run fetch_data(), update the Excel template at C7,
    and save the result into /app/reports.
    """
    df = fetch_data(as_of=as_of)

    # ✅ Always use absolute /app/reports directory
    reports_dir = os.path.join(os.getcwd(), "app", "reports")
    os.makedirs(reports_dir, exist_ok=True)

    # ✅ update_excel already handles saving
    report_day = int((as_of or (datetime.today() - timedelta(days=1))).day)
    filename, path = update_excel(df, template_path, reports_dir=reports_dir, report_day=report_day)

    return filename, path, df


def _send_revenue_report_email(*, report_path, filename, total_count, as_of, triggered_by):
    """Send revenue report attachment plus KPI summary, if configured."""
    from app.models import AppSetting

    auto_email_enabled = AppSetting.get_value('revenue_reports_auto_email_enabled', 'false') == 'true'
    if not auto_email_enabled:
        return {"status": "skipped", "reason": "auto_email_disabled"}

    recipient_mode = AppSetting.get_value('revenue_reports_email_recipient_mode', 'custom_group_later')
    if recipient_mode == 'custom_group_later':
        return {"status": "skipped", "reason": "recipients_pending_configuration"}

    # Future recipient strategies can be added here.
    return {"status": "skipped", "reason": f"unsupported_recipient_mode:{recipient_mode}"}


def _send_revenue_report_email_with_attachment(*, recipients, report_path, filename, total_count, as_of, triggered_by):
    """SMTP send helper for revenue report with attachment and KPI summary body."""
    if app.config.get('MAIL_SUPPRESS_SEND'):
        app.logger.info("[MAIL_SUPPRESS_SEND] Revenue report email suppressed for %s", recipients)
        return True, ""

    smtp_host = app.config.get('MAIL_SERVER', 'smtp.gmail.com')
    smtp_port = app.config.get('MAIL_PORT', 587)
    use_tls = app.config.get('MAIL_USE_TLS', True)
    username = app.config.get('MAIL_USERNAME')
    password = app.config.get('MAIL_PASSWORD')
    sender = app.config.get('MAIL_DEFAULT_SENDER', username)
    if not (username and password and sender):
        return False, 'MAIL_USERNAME/PASSWORD/SENDER not fully configured'

    run_day = (as_of or (datetime.today() - timedelta(days=1))).strftime('%Y-%m-%d')
    subject = f"Revenue Report - {run_day}"
    text_body = (
        f"Revenue report generated successfully.\n\n"
        f"Run date (as_of): {run_day}\n"
        f"C7 total_count: {total_count}\n"
        f"File: {filename}\n"
        f"Triggered by: {triggered_by}\n"
    )
    html_body = f"""
    <p>Revenue report generated successfully.</p>
    <ul>
      <li><strong>Run date (as_of):</strong> {run_day}</li>
      <li><strong>C7 total_count:</strong> {total_count}</li>
      <li><strong>File:</strong> {filename}</li>
      <li><strong>Triggered by:</strong> {triggered_by}</li>
    </ul>
    """

    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = sender
    msg['To'] = ', '.join(recipients)
    msg.set_content(text_body)
    msg.add_alternative(html_body, subtype='html')
    with open(report_path, 'rb') as f:
        payload = f.read()
        msg.add_attachment(
            payload,
            maintype='application',
            subtype='vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            filename=filename
        )

    try:
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            if use_tls:
                server.starttls()
            server.login(username, password)
            server.send_message(msg)
        return True, ""
    except Exception as exc:
        return False, str(exc)


def run_revenue_report_job(triggered_by="manual", as_of=None):
    """Run section-1 revenue report and persist latest status metadata."""
    from app.models import AppSetting

    source_mode = AppSetting.get_value('revenue_reports_source_mode', 'db_template')
    active_mode = source_mode if source_mode in ('db_template', 'hybrid') else 'db_template'

    template_path = os.path.join("instance", "uploaded_template.xlsx")
    if not os.path.exists(template_path):
        error_msg = "No template uploaded yet (instance/uploaded_template.xlsx)"
        AppSetting.set_value('revenue_reports_last_status', 'failed')
        AppSetting.set_value('revenue_reports_last_error', error_msg)
        raise FileNotFoundError(error_msg)

    filename, path, df = generate_and_send_report(template_path, as_of=as_of)
    total_count = int(pd.to_numeric(df['total_count'], errors="coerce").fillna(0).astype(int).iloc[0])

    AppSetting.set_value('revenue_reports_last_status', 'success')
    AppSetting.set_value('revenue_reports_last_error', '')
    AppSetting.set_value('revenue_reports_last_filename', filename)
    AppSetting.set_value('revenue_reports_last_total_count', str(total_count))
    AppSetting.set_value('revenue_reports_last_run_at', datetime.now(timezone.utc).isoformat())
    AppSetting.set_value('revenue_reports_last_triggered_by', triggered_by)
    AppSetting.set_value('revenue_reports_last_source_mode', source_mode)
    AppSetting.set_value('revenue_reports_last_active_mode', active_mode)

    # Email delivery is non-blocking for report success.
    email_result = {"status": "skipped", "reason": "not_attempted"}
    try:
        app.logger.info("Revenue report email step starting for as_of=%s", (as_of or (datetime.today() - timedelta(days=1))).strftime('%Y-%m-%d'))
        email_result = _send_revenue_report_email(
            report_path=path,
            filename=filename,
            total_count=total_count,
            as_of=as_of,
            triggered_by=triggered_by,
        )
    except Exception as email_exc:
        email_result = {"status": "failed", "reason": str(email_exc)}

    AppSetting.set_value('revenue_reports_last_email_status', email_result.get("status", "skipped"))
    AppSetting.set_value('revenue_reports_last_email_reason', email_result.get("reason", ""))
    AppSetting.set_value('revenue_reports_last_email_at', datetime.now(timezone.utc).isoformat())

    return {
        "filename": filename,
        "path": path,
        "total_count": total_count,
        "source_mode": source_mode,
        "active_mode": active_mode,
        "triggered_by": triggered_by,
        "email_status": email_result.get("status", "skipped"),
        "email_reason": email_result.get("reason", ""),
    }


def _normalize_table_value(value):
    if value is None:
        return ""
    return str(value)


def _parse_workbook_used_ranges(workbook_path):
    wb_formula = load_workbook(workbook_path, data_only=False)
    wb_cached = load_workbook(workbook_path, data_only=True)
    sheets_payload = []
    for ws_idx, ws in enumerate(wb_formula.worksheets):
        ws_cached = wb_cached.worksheets[ws_idx] if ws_idx < len(wb_cached.worksheets) else None
        min_row, min_col, max_row, max_col = 1, 1, 1, 1
        if ws.max_row and ws.max_column:
            min_row, min_col, max_row, max_col = (
                ws.min_row or 1,
                ws.min_column or 1,
                ws.max_row or 1,
                ws.max_column or 1,
            )

        matrix = []
        for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            parsed_row = []
            for cell in row:
                formula_or_value = cell.value
                if isinstance(formula_or_value, str) and formula_or_value.startswith('='):
                    cached_value = None
                    if ws_cached is not None:
                        cached_value = ws_cached[cell.coordinate].value
                    # Prefer cached/computed workbook value. If unavailable, leave empty.
                    parsed_row.append(_normalize_table_value(cached_value))
                else:
                    parsed_row.append(_normalize_table_value(formula_or_value))
            matrix.append(parsed_row)

        if not matrix:
            matrix = [[]]

        col_count = len(matrix[0]) if matrix and matrix[0] else max_col - min_col + 1

        # Pick the best header row from the first few rows.
        # Prefer rows containing known header words (e.g., Daily Actual) over dense numeric data rows.
        header_scan_limit = min(len(matrix), 12)
        header_row_idx = 0
        best_score = -1
        header_keywords = (
            "daily actual", "daily forecast", "budget variance", "prior day",
            "mtd", "runrate", "run rate", "ytd", "forecast"
        )
        for idx in range(header_scan_limit):
            row_values = [str(v).strip() for v in matrix[idx]]
            non_empty_score = sum(1 for v in row_values if v != "")
            keyword_score = 0
            for v in row_values:
                lowered = v.lower()
                if any(k in lowered for k in header_keywords):
                    keyword_score += 5
            score = non_empty_score + keyword_score
            if score > best_score:
                best_score = score
                header_row_idx = idx

        candidate_headers = matrix[header_row_idx] if matrix else []
        headers = []
        for i in range(col_count):
            raw = candidate_headers[i] if i < len(candidate_headers) else ""
            label = str(raw).strip()
            headers.append(label if label else f"Col {i + 1}")

        blank_columns = []
        for col_idx in range(col_count):
            col_values = [str(row[col_idx]).strip() if col_idx < len(row) else "" for row in matrix]
            if all(v == "" for v in col_values):
                blank_columns.append({
                    "index": col_idx,
                    "name": headers[col_idx] if col_idx < len(headers) else f"Col {col_idx + 1}",
                })

        sheets_payload.append({
            "sheet_name": ws.title,
            "headers": headers,
            "rows": matrix,
            "row_count": len(matrix),
            "column_count": col_count,
            "blank_columns": blank_columns,
        })
    wb_formula.close()
    wb_cached.close()
    return sheets_payload


def _as_float(value):
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(",", "")
    if text == "":
        return None
    try:
        return float(text)
    except Exception:
        return None


def _format_num(value):
    if value is None:
        return "—"
    try:
        num = float(value)
    except Exception:
        return str(value)
    if abs(num - round(num)) < 1e-9:
        return f"{int(round(num)):,}"
    return f"{num:,.2f}"


def _format_pct(value):
    if value is None:
        return "—"
    try:
        num = float(value) * 100.0
        return f"{num:.0f}%"
    except Exception:
        return "—"


def _eval_add_formula(formula, ws):
    """Evaluate simple '=a+b+C7+...' formulas used in revenue drivers sheet."""
    if not isinstance(formula, str) or not formula.startswith('='):
        return _as_float(formula)
    expr = formula[1:].strip()
    parts = [p.strip() for p in expr.split('+')]
    total = 0.0
    for token in parts:
        if re.match(r"^[A-Z]+[0-9]+$", token):
            v = _as_float(ws[token].value)
            total += (v or 0.0)
        else:
            total += (_as_float(token) or 0.0)
    return total


def _extract_multiplier_from_formula(formula, default_value):
    if isinstance(formula, str):
        m = re.search(r"\*([0-9]+(?:\.[0-9]+)?)", formula.replace(" ", ""))
        if m:
            return _as_float(m.group(1)) or default_value
    return default_value


def _extract_divisor_from_formula(formula, default_value):
    if isinstance(formula, str):
        m = re.search(r"/([0-9]+(?:\.[0-9]+)?)", formula.replace(" ", ""))
        if m:
            return _as_float(m.group(1)) or default_value
    return default_value


def _build_revenue_drivers_sheet(workbook_path, prev_day_values=None):
    wb = load_workbook(workbook_path, data_only=False)
    ws = wb.worksheets[0]

    # Fixed layout (B:R) that matches the Revenue Drivers report block.
    headers = [
        "Category",
        "Daily Actual",
        "Daily Forecast",
        "% Budget Variance",
        "Prior day Actual",
        "Prior day Variance%",
        "Aug FY26 Actual",
        "MTD FY26 Forecast (Budget)",
        "% Budget Variance",
        "Prior Month MTD FY26",
        "% Prior Month Variance",
        "MTD Runrate",
        "Aug Forecast FY26 (Budget)",
        "Run rate % Variance",
        "YTD FY26 ACTUAL",
        "YTD Forecast FY26",
        "FY26 %Budget Variance",
    ]

    month_days = _extract_divisor_from_formula(ws["D7"].value, 31.0)
    report_day = _extract_multiplier_from_formula(ws["I7"].value, float((datetime.today() - timedelta(days=1)).day))
    if report_day == 0:
        report_day = 1.0

    prev_smartlearning = _as_float((prev_day_values or {}).get("smartlearning_value")) or 0.0
    prev_library = _as_float((prev_day_values or {}).get("library_value")) or 0.0

    def row_values(row_num):
        if row_num == 7:
            c = prev_smartlearning
        elif row_num == 8:
            c = prev_library
        else:
            c = _as_float(ws[f"C{row_num}"].value) or 0.0
        f = _as_float(ws[f"F{row_num}"].value) or 0.0
        k = _as_float(ws[f"K{row_num}"].value) or 0.0
        n = _as_float(ws[f"N{row_num}"].value) or 0.0

        d = n / month_days if month_days else 0.0
        e = ((c - d) / d) if d else None
        g = ((c - f) / f) if f else None

        h = _eval_add_formula(ws[f"H{row_num}"].value, ws)
        if h is None:
            h = _as_float(ws[f"H{row_num}"].value) or 0.0

        i = d * report_day
        j = ((h - i) / i) if i else None
        l = ((h - k) / k) if k else None
        m = (h * month_days) / report_day if report_day else 0.0
        o = ((m - n) / n) if n else None

        p_base = 0.0
        q_base = 0.0
        p_formula = ws[f"P{row_num}"].value
        q_formula = ws[f"Q{row_num}"].value
        if isinstance(p_formula, str) and "+" in p_formula:
            p_base = _as_float(p_formula[1:].split("+")[0]) or 0.0
        if isinstance(q_formula, str) and "+" in q_formula:
            q_base = _as_float(q_formula[1:].split("+")[0]) or 0.0
        p = p_base + h
        q = q_base + i
        r = ((p - q) / q) if q else None

        return {
            "C": c, "D": d, "E": e, "F": f, "G": g, "H": h, "I": i, "J": j,
            "K": k, "L": l, "M": m, "N": n, "O": o, "P": p, "Q": q, "R": r
        }

    r7 = row_values(7)
    r8 = row_values(8)
    r9 = {k: (r7[k] + r8[k]) if k not in ("E", "G", "J", "L", "O", "R") else None for k in r7.keys()}
    r9["E"] = ((r9["C"] - r9["D"]) / r9["D"]) if r9["D"] else None
    r9["G"] = ((r9["C"] - r9["F"]) / r9["F"]) if r9["F"] else None
    r9["J"] = ((r9["H"] - r9["I"]) / r9["I"]) if r9["I"] else None
    r9["L"] = ((r9["H"] - r9["K"]) / r9["K"]) if r9["K"] else None
    r9["O"] = ((r9["M"] - r9["N"]) / r9["N"]) if r9["N"] else None
    r9["R"] = ((r9["P"] - r9["Q"]) / r9["Q"]) if r9["Q"] else None

    rows = [
        ["Smart learning", _format_num(r7["C"]), _format_num(r7["D"]), _format_pct(r7["E"]), _format_num(r7["F"]), _format_pct(r7["G"]),
         _format_num(r7["H"]), _format_num(r7["I"]), _format_pct(r7["J"]), _format_num(r7["K"]), _format_pct(r7["L"]),
         _format_num(r7["M"]), _format_num(r7["N"]), _format_pct(r7["O"]), _format_num(r7["P"]), _format_num(r7["Q"]), _format_pct(r7["R"])],
        ["Library", _format_num(r8["C"]), _format_num(r8["D"]), _format_pct(r8["E"]), _format_num(r8["F"]), _format_pct(r8["G"]),
         _format_num(r8["H"]), _format_num(r8["I"]), _format_pct(r8["J"]), _format_num(r8["K"]), _format_pct(r8["L"]),
         _format_num(r8["M"]), _format_num(r8["N"]), _format_pct(r8["O"]), _format_num(r8["P"]), _format_num(r8["Q"]), _format_pct(r8["R"])],
        ["Total", _format_num(r9["C"]), _format_num(r9["D"]), _format_pct(r9["E"]), _format_num(r9["F"]), _format_pct(r9["G"]),
         _format_num(r9["H"]), _format_num(r9["I"]), _format_pct(r9["J"]), _format_num(r9["K"]), _format_pct(r9["L"]),
         _format_num(r9["M"]), _format_num(r9["N"]), _format_pct(r9["O"]), _format_num(r9["P"]), _format_num(r9["Q"]), _format_pct(r9["R"])],
    ]

    # Compare a couple of key values against workbook values (if cached) for diagnostics.
    diagnostics = {"mismatch_count": 0, "mismatches": []}
    checks = [
        ("H7", r7["H"]), ("I7", r7["I"]), ("R7", r7["R"]),
        ("H8", r8["H"]), ("I8", r8["I"]), ("R8", r8["R"]),
    ]
    for addr, computed in checks:
        workbook_val = _as_float(ws[addr].value)
        if workbook_val is None:
            continue
        if abs(workbook_val - (computed or 0.0)) > 0.51:
            diagnostics["mismatch_count"] += 1
            diagnostics["mismatches"].append({
                "cell": addr,
                "workbook": workbook_val,
                "computed": round(computed or 0.0, 4),
            })

    wb.close()
    sheet_payload = {
        "sheet_name": ws.title,
        "headers": headers,
        "rows": rows,
        "row_count": len(rows),
        "column_count": len(headers),
        "blank_columns": [],
    }
    return sheet_payload, diagnostics


def _latest_report_file_by_tokens(tokens):
    reports_dir = os.path.join(os.getcwd(), "app", "reports")
    if not os.path.exists(reports_dir):
        return None
    candidates = []
    for name in os.listdir(reports_dir):
        lowered = name.lower()
        if name.lower().endswith(".xlsx") and any(token in lowered for token in tokens):
            candidates.append(os.path.join(reports_dir, name))
    if not candidates:
        return None
    return max(candidates, key=os.path.getctime)


def _resolve_revenue_tab_file(tab_id, table_source, drivers_template_path, flash_template_path):
    if tab_id == "revenue_drivers":
        template_path = drivers_template_path
        generated_tokens = ["daily_report", "revenue_drivers"]
    else:
        template_path = flash_template_path
        generated_tokens = ["akello_flash", "flash_revenue", "flash_report"]

    if table_source == "latest_generated":
        generated_path = _latest_report_file_by_tokens(generated_tokens)
        if generated_path:
            return "latest_generated", generated_path
        return "template_fallback", template_path
    return "template", template_path


def _parse_revenue_series_datetime(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, "%d-%m-%Y %H-%M-%S")
    except Exception as exc:
        raise ValueError("Invalid datetime format. Use DD-MM-YYYY HH-MM-SS.") from exc


def _current_week_bounds(now_dt=None):
    current = now_dt or datetime.now()
    week_start = datetime(current.year, current.month, current.day) - timedelta(days=current.weekday())
    week_end = week_start + timedelta(days=6, hours=23, minutes=59, seconds=59)
    return week_start, week_end


def _fetch_daily_actuals_by_day(start_dt, end_dt):
    start_day = datetime(start_dt.year, start_dt.month, start_dt.day)
    end_day = datetime(end_dt.year, end_dt.month, end_dt.day, 23, 59, 59)
    series_by_date = {}

    ruzivo_conn = get_ruzivo_conn()
    with ruzivo_conn.cursor() as cursor:
        try:
            cursor.execute(
                """
                SELECT DATE(last_login) AS login_date, COUNT(*) AS login_count
                FROM tblstudents_info
                WHERE last_login BETWEEN %s AND %s
                GROUP BY DATE(last_login)
                ORDER BY login_date ASC
                """,
                (start_day, end_day),
            )
            for row in cursor.fetchall():
                day = row.get("login_date")
                if day is None:
                    continue
                key = day.isoformat() if hasattr(day, "isoformat") else str(day)
                bucket = series_by_date.setdefault(key, {"smartlearning": 0.0, "library": 0.0})
                bucket["smartlearning"] = float(row.get("login_count") or 0.0)
        finally:
            try:
                ruzivo_conn.close()
            except Exception:
                pass

    library_conn_local = get_direct_library_conn()
    with library_conn_local.cursor(pymysql.cursors.DictCursor) as cursor:
        try:
            cursor.execute(
                """
                SELECT DATE(o.created_at) AS login_date, COUNT(DISTINCT o.user_id) AS user_count
                FROM orders o
                WHERE o.created_at BETWEEN %s AND %s
                AND o.status = 'Completed'
                GROUP BY DATE(o.created_at)
                ORDER BY login_date ASC
                """,
                (start_day, end_day),
            )
            for row in cursor.fetchall():
                day = row.get("login_date")
                if day is None:
                    continue
                key = day.isoformat() if hasattr(day, "isoformat") else str(day)
                bucket = series_by_date.setdefault(key, {"smartlearning": 0.0, "library": 0.0})
                bucket["library"] = float(row.get("user_count") or 0.0)
        finally:
            try:
                library_conn_local.close()
            except Exception:
                pass

    return series_by_date


def _get_previous_day_daily_actual_values(reference_dt=None):
    base = reference_dt or datetime.now()
    target_day = (base - timedelta(days=1)).date()
    target_start = datetime(target_day.year, target_day.month, target_day.day, 0, 0, 0)
    target_end = datetime(target_day.year, target_day.month, target_day.day, 23, 59, 59)
    series_by_date = _fetch_daily_actuals_by_day(target_start, target_end)
    bucket = series_by_date.get(target_day.isoformat(), {"smartlearning": 0.0, "library": 0.0})
    smartlearning_value = float(bucket.get("smartlearning", 0.0))
    library_value = float(bucket.get("library", 0.0))
    return {
        "date": target_day.isoformat(),
        "smartlearning_value": smartlearning_value,
        "library_value": library_value,
        "total_value": smartlearning_value + library_value,
    }


def _build_drivers_daily_actual_series(start_at, end_at):
    series_by_date = _fetch_daily_actuals_by_day(start_at, end_at)

    points = []
    for key in sorted(series_by_date.keys()):
        day_start_dt = datetime.strptime(key, "%Y-%m-%d")
        day_end_dt = day_start_dt + timedelta(hours=23, minutes=59, seconds=59)
        if day_end_dt < start_at or day_start_dt > end_at:
            continue
        smartlearning_value = float(series_by_date[key]["smartlearning"])
        library_value = float(series_by_date[key]["library"])
        total_value = smartlearning_value + library_value
        points.append({
            "timestamp_iso": day_start_dt.isoformat(),
            "timestamp_label": day_start_dt.strftime("%d-%m-%Y %H-%M-%S"),
            "smartlearning_value": smartlearning_value,
            "library_value": library_value,
            "total_value": total_value,
        })

    return points


def _normalize_match_text(value):
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _apply_flash_daily_actual_overrides(sheets, prev_day_values, zig_exchange):
    diagnostics = {
        "usd_sheet_found": False,
        "zig_sheet_found": False,
        "usd_daily_actual_col_found": False,
        "zig_daily_actual_col_found": False,
        "usd_smartlearning_row_found": False,
        "usd_library_row_found": False,
        "zig_smartlearning_row_found": False,
        "warnings": [],
    }
    if not sheets:
        diagnostics["warnings"].append("No flash sheets available to map.")
        return sheets, diagnostics

    smartlearning_base = float((prev_day_values or {}).get("smartlearning_value") or 0.0)
    library_base = float((prev_day_values or {}).get("library_value") or 0.0)
    smartlearning_zig = smartlearning_base * float(zig_exchange or 37.0)

    def is_daily_actual_header(text):
        normalized = _normalize_match_text(text)
        return normalized == "dailyactual" or "dailyactual" in normalized

    def find_daily_idx(headers, rows):
        for idx, header in enumerate(headers or []):
            if is_daily_actual_header(header):
                return idx
        scan_rows = rows[:20] if rows else []
        for row in scan_rows:
            for idx, cell in enumerate(row):
                if is_daily_actual_header(cell):
                    return idx
        return -1

    def row_kind(row):
        if not row:
            return ""
        probe = " ".join(str(cell or "") for cell in row[:3]).lower()
        normalized = _normalize_match_text(probe)
        if "smartlearning" in normalized:
            return "smartlearning"
        if "library" in normalized:
            return "library"
        return ""

    def sheet_kind(normalized_sheet, headers, rows):
        if "usd" in normalized_sheet:
            return "usd"
        if "zig" in normalized_sheet:
            return "zig"
        header_blob = " ".join(str(h or "") for h in (headers or [])).lower()
        if "usd" in header_blob:
            return "usd"
        if "zig" in header_blob:
            return "zig"
        sample_blob = " ".join(
            " ".join(str(c or "") for c in row[:8]) for row in (rows[:20] if rows else [])
        ).lower()
        if " usd " in f" {sample_blob} ":
            return "usd"
        if " zig " in f" {sample_blob} ":
            return "zig"
        return ""

    inferred_usd_assigned = False
    inferred_zig_assigned = False
    for sheet in sheets:
        sheet_name = str(sheet.get("sheet_name", ""))
        normalized_sheet = _normalize_match_text(sheet_name)
        headers = sheet.get("headers") or []
        rows = sheet.get("rows") or []

        daily_idx = find_daily_idx(headers, rows)
        if daily_idx < 0:
            continue

        kind = sheet_kind(normalized_sheet, headers, rows)
        if not kind and not inferred_usd_assigned:
            kind = "usd"
            inferred_usd_assigned = True
        elif not kind and not inferred_zig_assigned:
            kind = "zig"
            inferred_zig_assigned = True

        if kind == "usd":
            diagnostics["usd_sheet_found"] = True
            diagnostics["usd_daily_actual_col_found"] = True
            for row in rows:
                kind = row_kind(row)
                if not kind:
                    continue
                while len(row) <= daily_idx:
                    row.append("")
                if kind == "smartlearning":
                    row[daily_idx] = _format_num(smartlearning_base)
                    diagnostics["usd_smartlearning_row_found"] = True
                elif kind == "library":
                    row[daily_idx] = _format_num(library_base)
                    diagnostics["usd_library_row_found"] = True

        if kind == "zig":
            diagnostics["zig_sheet_found"] = True
            diagnostics["zig_daily_actual_col_found"] = True
            for row in rows:
                kind = row_kind(row)
                if kind != "smartlearning":
                    continue
                while len(row) <= daily_idx:
                    row.append("")
                row[daily_idx] = _format_num(smartlearning_zig)
                diagnostics["zig_smartlearning_row_found"] = True

    if diagnostics["usd_sheet_found"] and not diagnostics["usd_smartlearning_row_found"]:
        diagnostics["warnings"].append("USD Smartlearning row not found for Daily Actual mapping.")
    if diagnostics["usd_sheet_found"] and not diagnostics["usd_library_row_found"]:
        diagnostics["warnings"].append("USD Library row not found for Daily Actual mapping.")
    if diagnostics["zig_sheet_found"] and not diagnostics["zig_smartlearning_row_found"]:
        diagnostics["warnings"].append("ZiG Smartlearning row not found for Daily Actual mapping.")
    if not diagnostics["usd_sheet_found"]:
        diagnostics["warnings"].append("USD sheet not found for Flash mapping.")
    if not diagnostics["zig_sheet_found"]:
        diagnostics["warnings"].append("ZiG sheet not found for Flash mapping.")

    return sheets, diagnostics


@app.route('/api/revenue-reports/tables', methods=['GET'])
@login_required
def api_revenue_reports_tables():
    if not can_view_revenue_reports():
        return jsonify({"error": "Unauthorized"}), 403

    from app.models import AppSetting

    table_source = AppSetting.get_value('revenue_reports_table_source', 'latest_generated')
    if table_source not in ('latest_generated', 'template'):
        table_source = 'latest_generated'
    zig_exchange = _as_float(AppSetting.get_value('revenue_reports_zig_exchange', '37')) or 37.0
    prev_day_values = _get_previous_day_daily_actual_values()

    drivers_template_path = AppSetting.get_value(
        'revenue_reports_template_path_1',
        os.path.join("instance", "uploaded_template.xlsx"),
    )
    flash_template_path = AppSetting.get_value(
        'revenue_reports_template_path_2',
        r"C:\Users\quincy.mashava\Downloads\Akello Flash Revenue Report as at  18 August 2025.xlsx",
    )

    tabs = [
        {"id": "revenue_drivers", "title": "Revenue Drivers Report"},
        {"id": "akello_flash", "title": "Akello Flash Revenue Report"},
    ]

    payload_tabs = []
    for tab in tabs:
        resolved_source, resolved_path = _resolve_revenue_tab_file(
            tab["id"], table_source, drivers_template_path, flash_template_path
        )

        tab_payload = {
            "id": tab["id"],
            "title": tab["title"],
            "source_type": resolved_source,
            "selected_path": resolved_path,
            "file_exists": bool(resolved_path and os.path.exists(resolved_path)),
            "error": "",
            "sheets": [],
            "diagnostics": {"blank_columns": []},
        }

        if not tab_payload["file_exists"]:
            tab_payload["error"] = f"Workbook not found: {resolved_path}"
        else:
            try:
                if tab["id"] == "revenue_drivers":
                    drivers_sheet, drivers_diag = _build_revenue_drivers_sheet(resolved_path, prev_day_values=prev_day_values)
                    tab_payload["sheets"] = [drivers_sheet]
                    tab_payload["diagnostics"]["mismatch_count"] = drivers_diag.get("mismatch_count", 0)
                    tab_payload["diagnostics"]["mismatches"] = drivers_diag.get("mismatches", [])
                    tab_payload["diagnostics"]["previous_day_source"] = prev_day_values
                    try:
                        raw_sheets = _parse_workbook_used_ranges(resolved_path)
                        tab_payload["diagnostics"]["raw_sheet"] = raw_sheets[0] if raw_sheets else None
                    except Exception:
                        tab_payload["diagnostics"]["raw_sheet"] = None
                else:
                    tab_payload["sheets"] = _parse_workbook_used_ranges(resolved_path)
                    tab_payload["sheets"], flash_mapping_diag = _apply_flash_daily_actual_overrides(
                        tab_payload["sheets"],
                        prev_day_values=prev_day_values,
                        zig_exchange=zig_exchange
                    )
                    tab_payload["diagnostics"]["daily_actual_mapping"] = flash_mapping_diag
                    tab_payload["diagnostics"]["previous_day_source"] = prev_day_values
                    tab_payload["diagnostics"]["zig_exchange"] = zig_exchange
                all_blank_columns = []
                for sheet in tab_payload["sheets"]:
                    for col in sheet.get("blank_columns", []):
                        all_blank_columns.append({
                            "sheet_name": sheet.get("sheet_name", ""),
                            "index": col.get("index", 0),
                            "name": col.get("name", ""),
                        })
                tab_payload["diagnostics"]["blank_columns"] = all_blank_columns
            except Exception as exc:
                tab_payload["error"] = f"Failed to parse workbook: {exc}"

        payload_tabs.append(tab_payload)

    return jsonify({
        "success": True,
        "table_source": table_source,
        "tabs": payload_tabs,
    })


@app.route('/api/revenue-reports/drivers/daily-actual-series', methods=['GET'])
@login_required
def api_revenue_drivers_daily_actual_series():
    if not can_view_revenue_reports():
        return jsonify({"error": "Unauthorized"}), 403

    start_raw = request.args.get("start_at", "").strip()
    end_raw = request.args.get("end_at", "").strip()

    try:
        start_at = _parse_revenue_series_datetime(start_raw) if start_raw else None
        end_at = _parse_revenue_series_datetime(end_raw) if end_raw else None
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400

    if start_at is None and end_at is None:
        start_at, end_at = _current_week_bounds()
    elif start_at is None:
        start_at, _ = _current_week_bounds(end_at)
    elif end_at is None:
        _, end_at = _current_week_bounds(start_at)

    if start_at > end_at:
        return jsonify({"error": "Invalid range: start_at must be earlier than or equal to end_at."}), 400

    points = _build_drivers_daily_actual_series(start_at, end_at)
    return jsonify({
        "success": True,
        "start_at": start_at.strftime("%d-%m-%Y %H-%M-%S"),
        "end_at": end_at.strftime("%d-%m-%Y %H-%M-%S"),
        "point_count": len(points),
        "points": points,
    })


@app.route('/api/revenue-reports/flash/upload', methods=['POST'])
@login_required
def api_revenue_flash_upload():
    if not can_view_revenue_reports():
        return jsonify({"error": "Unauthorized"}), 403

    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if not file or file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if not file.filename.lower().endswith('.xlsx'):
        return jsonify({"error": "Only .xlsx files are supported"}), 400

    from app.models import AppSetting
    os.makedirs("instance", exist_ok=True)
    stored_path = os.path.join("instance", "uploaded_flash_template.xlsx")
    file.save(stored_path)

    AppSetting.set_value(
        'revenue_reports_template_path_2',
        stored_path,
        getattr(current_user, 'id', None),
        'Active uploaded template path for Akello Flash report tab'
    )

    # Ensure future renders use the uploaded workbook directly.
    AppSetting.set_value(
        'revenue_reports_table_source',
        'template',
        getattr(current_user, 'id', None),
        'Data source for revenue report tables: latest_generated or template'
    )

    return jsonify({
        "success": True,
        "message": "Flash template uploaded successfully",
        "stored_as": "uploaded_flash_template.xlsx",
        "stored_path": stored_path
    })


@app.route('/revenue_reports', methods=['GET'])
@login_required
def revenue_reports():
    if not can_view_revenue_reports():
        return "Unauthorized", 403
    return render_template('revenue_reports.html', title='Revenue Reports', is_admin=(current_user.userRole == 'Admin'))




# ---------- Run Report ----------
@app.route('/api/run-report', methods=['POST'])
def api_run_report():
    template_path = os.path.join("instance", "uploaded_template.xlsx")

    if not os.path.exists(template_path):
        return jsonify({"error": "No template uploaded yet"}), 400

    filename, path, _ = generate_and_send_report(template_path)

    if not os.path.exists(path):
        return jsonify({"error": "Report not generated"}), 500

    return send_file(path, as_attachment=True, download_name=filename)


@app.route('/api/revenue-reports/run', methods=['POST'])
@login_required
def api_revenue_reports_run():
    if not can_view_revenue_reports():
        return jsonify({"error": "Unauthorized"}), 403

    try:
        result = run_revenue_report_job(triggered_by=f"user:{current_user.username}")
        return jsonify({"success": True, **result})
    except Exception as e:
        from app.models import AppSetting
        AppSetting.set_value('revenue_reports_last_status', 'failed')
        AppSetting.set_value('revenue_reports_last_error', str(e))
        AppSetting.set_value('revenue_reports_last_run_at', datetime.now(timezone.utc).isoformat())
        return jsonify({"error": str(e)}), 500


@app.route('/api/revenue-reports/section1/latest', methods=['GET'])
@login_required
def api_revenue_reports_latest():
    if not can_view_revenue_reports():
        return jsonify({"error": "Unauthorized"}), 403

    from app.models import AppSetting

    reports_dir = os.path.join(os.getcwd(), "app", "reports")
    latest_download_url = None
    if os.path.exists(reports_dir) and os.listdir(reports_dir):
        latest_download_url = url_for('api_download_latest')

    payload = {
        "last_status": AppSetting.get_value('revenue_reports_last_status', 'not_run'),
        "last_error": AppSetting.get_value('revenue_reports_last_error', ''),
        "last_filename": AppSetting.get_value('revenue_reports_last_filename', ''),
        "last_total_count": AppSetting.get_value('revenue_reports_last_total_count', '0'),
        "last_run_at": AppSetting.get_value('revenue_reports_last_run_at', ''),
        "last_triggered_by": AppSetting.get_value('revenue_reports_last_triggered_by', ''),
        "source_mode": AppSetting.get_value('revenue_reports_source_mode', 'db_template'),
        "schedule_time": AppSetting.get_value('revenue_reports_schedule_time', '06:00'),
        "last_email_status": AppSetting.get_value('revenue_reports_last_email_status', 'not_sent'),
        "last_email_reason": AppSetting.get_value('revenue_reports_last_email_reason', ''),
        "last_email_at": AppSetting.get_value('revenue_reports_last_email_at', ''),
        "latest_download_url": latest_download_url,
    }
    return jsonify({"success": True, "data": payload})







@app.route('/api/download/<filename>')
def api_download(filename):
    path = os.path.join("reports", filename)
    if not os.path.exists(path):
        return jsonify({"error": "File not found"}), 404
    return send_file(path, as_attachment=True)


@app.route('/api/download-latest')
def api_download_latest():
    reports_dir = os.path.join(os.getcwd(), "app", "reports")
    if not os.path.exists(reports_dir) or not os.listdir(reports_dir):
        return jsonify({"error": "No reports found"}), 404

    latest = max(
        [os.path.join(reports_dir, f) for f in os.listdir(reports_dir)],
        key=os.path.getctime
    )
    return send_file(latest, as_attachment=True)






@app.route('/conversiontemplatescript', methods=['GET', 'POST'])
@login_required
def conversiontemplatescript():


        return render_template('conversiontemplatescript.html',
                           title='Analytics')




import docx
import os

def convert_docx_to_html(docx_path, output_path):
    doc = docx.Document(docx_path)

    # Placeholders for extracted content
    h1 = ""
    h2 = ""
    h3 = ""
    objectives = []
    content = []
    content2 = []
    images_html = []
    examples_html = []

    # Loop through document elements
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""

        # Extract headings
        if style.startswith("Heading 1"):
            h1 = para.text.strip()

        elif style.startswith("Heading 2"):
            h2 = para.text.strip()

        elif style.startswith("Heading 3"):
            h3 = para.text.strip()

        # Detect example text
        elif "example" in para.text.lower():
            examples_html.append(f"""
            <div class="new-example-container">
                <img class="new-example-image" src="https://smartlearning.akello.co/public/uploads/content/maths-icons/Group%20572.png" alt="" /> 
                <span class="new-example-text"><b>{para.text.strip()}</b></span>
            </div>
            """)

        # Otherwise treat as normal content
        else:
            # Put first items under objectives, rest under content
            if "objective" in para.text.lower():
                objectives.append(f"<li>{para.text.strip()}</li>")
            else:
                if not h2:
                    content.append(f"<li>{para.text.strip()}</li>")
                else:
                    content2.append(f"<li>{para.text.strip()}</li>")

    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            # Word images are embedded, here we simulate with placeholder HTML
            images_html.append("""
            <div class="opacity-70" style="margin: 30px; background: #CAE3F6; border-radius: 20px; width: fit-content; display: flex; flex-direction: row; align-items: center;">
                <img src="https://smartlearning.akello.co/public/uploads/content/maths-icons/icon.png" 
                     style="border-radius: 50%; margin-left: -30px;" 
                     alt="icon.png (1 KB)" width="44" height="44" caption="false" />
                <p>{Insert image description or label here}</p>
            </div>
            """)

    # Build final HTML
    html_template = f"""
<div class="row">
  <div class="col-md-12">
    <h2 class="primary-top-banner">{h1}</h2>
  </div>

  <div class="col-md-12">
    <div class="row banner-area">
      <div class="col-md-5 banner-image">
        <img class="img-fluid" src="https://smartlearning.akello.co/public/uploads/content/Grade%207%20English%20New/Lesson%201%20Structure%20of%20a%20Paragraph/_-01_result_result.webp" style="width:60%;" />
      </div>
      <div class="col-md-7">
        <h2>Objectives</h2>
        <h4 class="padding-5">By the end of the lesson, you should be able to:</h4>
        <ol>
          {''.join(objectives)}
        </ol>
      </div>
    </div>
  </div>

  <div class="col-md-12">
    <h2 class="primary-top-banner">{h2}</h2>
    <div class="col-md-12">
      <ul>
        {''.join(content)}
      </ul>

      <h2>{h3}</h2>
      <ul>
        {''.join(content2)}
      </ul>

      {''.join(images_html)}
      {''.join(examples_html)}
    </div>
  </div>
</div>
"""

    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_template)

    print(f"✅ Conversion complete! HTML saved to {output_path}")



UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return "No file uploaded", 400

    file = request.files["file"]
    if file.filename == "":
        return "No selected file", 400

    # Save uploaded file
    filename = secure_filename(file.filename)
    docx_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(docx_path)

    # Output path
    output_filename = filename.replace(".docx", ".html")
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    try:
        # Run conversion
        convert_docx_to_html(docx_path, output_path)
    except Exception as e:
        # Log the error for debugging purposes and return a user-friendly message
        print(f"Error during conversion: {e}")
        return "Error converting the document. Please try again.", 500

    # Check if the file was created successfully before attempting to send it
    if os.path.exists(output_path):
        # Return file for download
        return send_file(output_path, as_attachment=True)
    else:
        # If the file doesn't exist after the conversion attempt, something went wrong.
        return "Conversion succeeded, but the output file could not be found.", 500








@app.route('/api/countries', methods=['GET'])
def get_countries():
    try:
        conn = get_ruzivo_conn()
        with conn.cursor() as cursor:
            query = """
                SELECT 
                    id,
                    iso,
                    name,
                    nicename,
                    iso3,
                    numcode,
                    phonecode,
                    flag
                FROM tblcountries;
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            print("DB Rows:", rows)  # Debugging output

            if not rows:
                return jsonify({"message": "No countries found"}), 404

            col_names = [desc[0] for desc in cursor.description]
            data = [dict(zip(col_names, row)) for row in rows]

        return jsonify(data)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/lesotho-asl-registrations', methods=['GET'])
def lesotho_asl_registrations():
    conn = None
    try:
        conn = get_ruzivo_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        query = """
            SELECT COUNT(DISTINCT student_id) AS asl_registrations
            FROM tblstudents
            WHERE country_id = 119
        """
        cursor.execute(query)
        result = cursor.fetchone()

        return jsonify({
            "country_id": 119,
            "asl_registrations": result["asl_registrations"] if result else 0
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if conn:
            conn.close()






import imaplib
import email
import smtplib
import ssl
from email.header import decode_header
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import json
import msal



# ===== Microsoft Graph API Configuration =====
def get_microsoft_config():
    """Get Microsoft Graph API configuration from environment"""
    return {
        'client_id': os.getenv('MICROSOFT_CLIENT_ID'),
        'client_secret': os.getenv('MICROSOFT_CLIENT_SECRET'),
        'tenant_id': os.getenv('MICROSOFT_TENANT_ID'),
        'redirect_uri': os.getenv('MICROSOFT_REDIRECT_URI', 'http://localhost:5000/microsoft-callback'),
        'authority': f"https://login.microsoftonline.com/{os.getenv('MICROSOFT_TENANT_ID', '')}",
        'scopes': ['User.Read', 'Mail.Read']
    }

def build_msal_app(cache=None):
    """Build MSAL ConfidentialClientApplication"""
    config = get_microsoft_config()
    if not config['client_id'] or not config['client_secret'] or not config['tenant_id']:
        return None
    return msal.ConfidentialClientApplication(
        config['client_id'],
        authority=config['authority'],
        client_credential=config['client_secret'],
        token_cache=cache
    )

def get_token_from_cache():
    """Get token from Flask session"""
    return session.get('microsoft_token')

def get_graph_api_token():
    """Get valid access token, refreshing if necessary"""
    token = get_token_from_cache()
    if not token:
        return None
    
    # Check if token is expired (with 5 minute buffer)
    expires_on = token.get('expires_on', 0)
    current_time = datetime.now().timestamp()
    
    # Handle both Unix timestamp and datetime string formats
    if isinstance(expires_on, str):
        try:
            expires_on = datetime.fromisoformat(expires_on.replace('Z', '+00:00')).timestamp()
        except:
            expires_on = 0
    
    if expires_on and expires_on < (current_time + 300):
        # Token expired or about to expire, try to refresh
        msal_app = build_msal_app()
        if msal_app:
            accounts = msal_app.get_accounts()
            if accounts:
                result = msal_app.acquire_token_silent(
                    get_microsoft_config()['scopes'],
                    account=accounts[0]
                )
                if result and 'access_token' in result:
                    session['microsoft_token'] = result
                    return result.get('access_token')
                elif result and 'error' in result:
                    # Silent refresh failed, token needs re-authentication
                    session.pop('microsoft_token', None)
                    return None
    
    return token.get('access_token') if token else None


# ===== Email Query Storage =====
# Store email statuses in a simple JSON file or database
# For simplicity, using a JSON file approach
EMAIL_STATUS_FILE = 'email_query_statuses.json'

def load_email_statuses():
    """Load email statuses from file"""
    try:
        with open(EMAIL_STATUS_FILE, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def save_email_status(email_id, status):
    """Save email status to file"""
    statuses = load_email_statuses()
    statuses[email_id] = {
        'status': status,
        'updated_at': datetime.now().isoformat()
    }
    with open(EMAIL_STATUS_FILE, 'w') as f:
        json.dump(statuses, f, indent=2)


def ensure_json_response(data, status_code=200):
    """
    Helper function to ensure all API responses are valid JSON with proper headers.
    Validates the data can be serialized to JSON and creates a proper Flask response.
    
    Args:
        data: Dictionary or data structure to return as JSON
        status_code: HTTP status code (default: 200)
    
    Returns:
        Flask Response object with JSON content and proper headers
    """
    # #region agent log
    serialize_start = time.time()
    log_entry = {
        'sessionId': 'debug-session',
        'runId': 'run1',
        'hypothesisId': 'B',
        'location': 'routes.py:ensure_json_response:entry',
        'message': 'ensure_json_response called',
        'data': {'status_code': status_code, 'data_type': type(data).__name__, 'has_emails_key': 'emails' in data if isinstance(data, dict) else False, 'emails_count': len(data.get('emails', [])) if isinstance(data, dict) and 'emails' in data else 0},
        'timestamp': int(time.time() * 1000)
    }
    try:
        write_debug_log(log_entry)
        app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
    except Exception as e:
        app.logger.error(f"[DEBUG] Failed to write log: {e}")
    # #endregion
    
    try:
        # Validate that data can be serialized to JSON
        json_str = json.dumps(data)
        serialize_time = (time.time() - serialize_start) * 1000
        
        # #region agent log
        log_entry = {
            'sessionId': 'debug-session',
            'runId': 'run1',
            'hypothesisId': 'B',
            'location': 'routes.py:ensure_json_response:after_serialize',
            'message': 'JSON serialization completed',
            'data': {'serialization_time_ms': serialize_time, 'json_size_bytes': len(json_str), 'status_code': status_code},
            'timestamp': int(time.time() * 1000)
        }
        try:
            write_debug_log(log_entry)
            app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
        except Exception as e:
            app.logger.error(f"[DEBUG] Failed to write log: {e}")
        # #endregion
        
        # Create response with explicit Content-Type
        response = jsonify(data)
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        response.status_code = status_code
        
        # #region agent log
        log_entry = {
            'sessionId': 'debug-session',
            'runId': 'run1',
            'hypothesisId': 'B',
            'location': 'routes.py:ensure_json_response:before_return',
            'message': 'Response object created, about to return',
            'data': {'status_code': status_code, 'content_type': response.headers.get('Content-Type'), 'response_size_bytes': len(json_str)},
            'timestamp': int(time.time() * 1000)
        }
        try:
            write_debug_log(log_entry)
            app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
        except Exception as e:
            app.logger.error(f"[DEBUG] Failed to write log: {e}")
        # #endregion
        
        return response
    except (TypeError, ValueError) as e:
        # If data can't be serialized, return error response
        print(f"Error serializing response to JSON: {str(e)}")
        error_response = {
            'error': 'Internal server error: Failed to format response',
            'details': str(e) if app.debug else None
        }
        response = jsonify(error_response)
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        response.status_code = 500
        return response
    except Exception as e:
        # Catch any other unexpected errors
        print(f"Unexpected error creating JSON response: {str(e)}")
        error_response = {
            'error': 'Internal server error: Unexpected error formatting response',
            'details': str(e) if app.debug else None
        }
        response = jsonify(error_response)
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        response.status_code = 500
        return response


# ===== API Endpoints =====

@app.route('/api/help-desk/outlook-config', methods=['GET'])
@login_required
def get_help_desk_outlook_config():
    """Return whether Outlook uses IMAP (shared mailbox) or OAuth. When IMAP is configured, UI should not show Connect to Outlook."""
    # TODO: remove hardcoded fallback after testing; use .env only in production
    help_desk_email = os.getenv('HELP_DESK_EMAIL') or 'library@akello.co'
    help_desk_password = os.getenv('HELP_DESK_EMAIL_APP_PASSWORD') or 'wdmqpbqtpdkqsmbr'
    use_imap = bool(help_desk_email and help_desk_password)
    out = {'use_imap': use_imap}
    if use_imap:
        out['mailbox'] = help_desk_email
    return jsonify(out)


@app.route('/api/email-queries', methods=['GET'])
@login_required
def get_email_queries():
    """Fetch emails from Outlook (Graph API) or Gmail (IMAP) based on source parameter"""
    # #region agent log
    route_start_time = time.time()
    log_entry = {
        'sessionId': 'debug-session',
        'runId': 'run1',
        'hypothesisId': 'A',
        'location': 'routes.py:get_email_queries:entry',
        'message': 'get_email_queries route called',
        'data': {'source': request.args.get('source', 'outlook')},
        'timestamp': int(time.time() * 1000)
    }
    try:
        write_debug_log(log_entry)
        app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
    except Exception as e:
        app.logger.error(f"[DEBUG] Failed to write log: {e}")
    # #endregion
    
    try:
        # Get email source from query parameter (default: 'outlook')
        email_source = request.args.get('source', 'outlook').lower()
        
        try:
            # #region agent log
            before_fetch_time = time.time()
            log_entry = {
                'sessionId': 'debug-session',
                'runId': 'run1',
                'hypothesisId': 'A',
                'location': 'routes.py:get_email_queries:before_fetch',
                'message': 'About to call fetch function',
                'data': {'source': email_source, 'elapsed_ms': (before_fetch_time - route_start_time) * 1000},
                'timestamp': int(time.time() * 1000)
            }
            try:
                write_debug_log(log_entry)
                app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
            except Exception as e:
                app.logger.error(f"[DEBUG] Failed to write log: {e}")
            # #endregion
            
            if email_source == 'gmail':
                result = fetch_gmail_emails()
            elif email_source == 'outlook' and (os.getenv('HELP_DESK_EMAIL') or 'library@akello.co') and (os.getenv('HELP_DESK_EMAIL_APP_PASSWORD') or 'wdmqpbqtpdkqsmbr'):
                result = fetch_outlook_emails_imap()
            else:
                result = fetch_outlook_emails()
            
            # #region agent log
            after_fetch_time = time.time()
            log_entry = {
                'sessionId': 'debug-session',
                'runId': 'run1',
                'hypothesisId': 'A',
                'location': 'routes.py:get_email_queries:after_fetch',
                'message': 'Fetch function returned',
                'data': {'source': email_source, 'fetch_time_ms': (after_fetch_time - before_fetch_time) * 1000, 'result_type': type(result).__name__, 'is_tuple': isinstance(result, tuple), 'total_elapsed_ms': (after_fetch_time - route_start_time) * 1000},
                'timestamp': int(time.time() * 1000)
            }
            try:
                write_debug_log(log_entry)
                app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
            except Exception as e:
                app.logger.error(f"[DEBUG] Failed to write log: {e}")
            # #endregion
            
            # ensure_json_response returns a Response object with status_code already set
            # If result is a tuple, extract response and status_code
            # #region agent log
            before_return_time = time.time()
            log_entry = {
                'sessionId': 'debug-session',
                'runId': 'run1',
                'hypothesisId': 'A',
                'location': 'routes.py:get_email_queries:before_return',
                'message': 'About to return response',
                'data': {'is_tuple': isinstance(result, tuple), 'result_type': type(result).__name__, 'total_elapsed_ms': (before_return_time - route_start_time) * 1000},
                'timestamp': int(time.time() * 1000)
            }
            try:
                write_debug_log(log_entry)
                app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
            except Exception as e:
                app.logger.error(f"[DEBUG] Failed to write log: {e}")
            # #endregion
            
            if isinstance(result, tuple):
                response_obj, status_code = result
                # Ensure Content-Type is set
                if hasattr(response_obj, 'headers'):
                    response_obj.headers['Content-Type'] = 'application/json; charset=utf-8'
                response_obj.status_code = status_code
                
                # #region agent log
                log_entry = {
                    'sessionId': 'debug-session',
                    'runId': 'run1',
                    'hypothesisId': 'A',
                    'location': 'routes.py:get_email_queries:returning_tuple',
                    'message': 'Returning tuple response',
                    'data': {'status_code': status_code, 'content_type': response_obj.headers.get('Content-Type'), 'total_elapsed_ms': (time.time() - route_start_time) * 1000},
                    'timestamp': int(time.time() * 1000)
                }
                try:
                    write_debug_log(log_entry)
                    app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
                except Exception as e:
                    app.logger.error(f"[DEBUG] Failed to write log: {e}")
                # #endregion
                
                return response_obj, status_code
            else:
                # Result is already a Response object from ensure_json_response
                # Just ensure Content-Type is set (it should already be set)
                if hasattr(result, 'headers'):
                    if 'Content-Type' not in result.headers or 'application/json' not in result.headers.get('Content-Type', ''):
                        result.headers['Content-Type'] = 'application/json; charset=utf-8'
                
                # #region agent log
                log_entry = {
                    'sessionId': 'debug-session',
                    'runId': 'run1',
                    'hypothesisId': 'A',
                    'location': 'routes.py:get_email_queries:returning_response',
                    'message': 'Returning Response object',
                    'data': {'status_code': result.status_code, 'content_type': result.headers.get('Content-Type'), 'total_elapsed_ms': (time.time() - route_start_time) * 1000},
                    'timestamp': int(time.time() * 1000)
                }
                try:
                    write_debug_log(log_entry)
                    app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
                except Exception as e:
                    app.logger.error(f"[DEBUG] Failed to write log: {e}")
                # #endregion
                
                return result
            
        except Exception as fetch_error:
            # Catch any errors from fetch functions and ensure JSON response
            print(f"Error in fetch function ({email_source}): {str(fetch_error)}")
            import traceback
            traceback.print_exc()
            return ensure_json_response({
                'error': f'Error fetching emails from {email_source}: {str(fetch_error)}',
                'source': email_source
            }, 500)
            
    except Exception as e:
        print(f"Unexpected error in get_email_queries: {str(e)}")
        import traceback
        traceback.print_exc()
        return ensure_json_response({'error': f'Unexpected error: {str(e)}'}, 500)


def fetch_outlook_emails():
    """Fetch emails from Outlook using Microsoft Graph API"""
    try:
        # Check if user is authenticated with Microsoft
        access_token = get_graph_api_token()
        if not access_token:
            return jsonify({
                'error': 'Not authenticated with Microsoft. Please connect to Outlook first.',
                'requires_auth': True
            }), 401
        
        # Get sender filter from environment
        sender_filter = os.getenv('EMAIL_SENDER_FILTER', 'mashavaquincy@gmail.com')
        
        # Build Graph API endpoint with filter
        graph_endpoint = "https://graph.microsoft.com/v1.0/me/messages"
        
        # Add filter for emails from specific sender
        if sender_filter:
            # URL encode the filter
            from urllib.parse import quote
            filter_query = f"$filter=from/emailAddress/address eq '{sender_filter}'"
            graph_endpoint += f"?{filter_query}"
        
        # Add orderby and top to get most recent emails
        separator = "&" if sender_filter else "?"
        graph_endpoint += f"{separator}$orderby=receivedDateTime desc&$top=50&$select=id,subject,from,toRecipients,receivedDateTime,bodyPreview,body"
        
        # Make request to Graph API
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }
        
        print(f"Fetching emails from Graph API: {graph_endpoint}")
        response = requests.get(graph_endpoint, headers=headers)
        
        if response.status_code == 401:
            # Token expired or invalid, clear session
            session.pop('microsoft_token', None)
            return jsonify({
                'error': 'Authentication expired. Please reconnect to Outlook.',
                'requires_auth': True
            }), 401
        
        if not response.ok:
            error_data = response.json() if response.content else {}
            error_msg = error_data.get('error', {}).get('message', f'Graph API error: {response.status_code}')
            print(f"Graph API error: {error_msg}")
            return jsonify({
                'error': f'Failed to fetch emails: {error_msg}',
                'details': error_data
            }), response.status_code
        
        # Parse response
        data = response.json()
        graph_emails = data.get('value', [])
        
        # Load existing statuses
        email_statuses = load_email_statuses()
        
        # Map Graph API response to our email format
        emails = []
        for graph_email in graph_emails:
            try:
                email_id = graph_email.get('id', '')
                from_info = graph_email.get('from', {}).get('emailAddress', {})
                from_email = from_info.get('address', 'Unknown')
                from_name = from_info.get('name', '')
                
                # Get recipients
                to_recipients = graph_email.get('toRecipients', [])
                to_email = ', '.join([r.get('emailAddress', {}).get('address', '') for r in to_recipients])
                
                # Get subject
                subject = graph_email.get('subject', '(No Subject)')
                
                # Get date
                received_date = graph_email.get('receivedDateTime', '')
                date_str = received_date
                
                # Get body preview and full body
                body_preview = graph_email.get('bodyPreview', '')
                body_content = graph_email.get('body', {})
                body = body_content.get('content', body_preview) if isinstance(body_content, dict) else body_preview
                
                # Clean HTML from body if present
                if body_content.get('contentType') == 'html':
                    from bs4 import BeautifulSoup
                    try:
                        soup = BeautifulSoup(body, 'html.parser')
                        body = soup.get_text(separator=' ', strip=True)
                    except:
                        pass  # Keep original if parsing fails
                
                # Get status from storage
                email_status_data = email_statuses.get(email_id, {})
                current_status = email_status_data.get('status', 'Not started')
                
                emails.append({
                    'id': email_id,
                    'subject': subject,
                    'from': from_email,
                    'to': to_email,
                    'date': date_str,
                    'preview': body_preview[:200] if body_preview else '',
                    'body': body,
                    'status': current_status
                })
            except Exception as e:
                print(f"Error processing email {graph_email.get('id', 'unknown')}: {str(e)}")
                continue
        
        return jsonify({'emails': emails}), 200
        
    except Exception as e:
        print(f"Unexpected error fetching Outlook emails: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def fetch_outlook_emails_imap():
    """Fetch emails from Outlook using IMAP (HELP_DESK_EMAIL + HELP_DESK_EMAIL_APP_PASSWORD)"""
    mail = None
    try:
        # TODO: remove hardcoded fallback after testing
        help_desk_email = os.getenv('HELP_DESK_EMAIL') or 'library@akello.co'
        help_desk_password = os.getenv('HELP_DESK_EMAIL_APP_PASSWORD') or 'wdmqpbqtpdkqsmbr'
        imap_server = os.getenv('HELP_DESK_IMAP_SERVER', 'outlook.office365.com')
        imap_port = int(os.getenv('HELP_DESK_IMAP_PORT', '993'))

        if not help_desk_email or not help_desk_password:
            return jsonify({
                'error': 'Help desk Outlook IMAP not configured. Set HELP_DESK_EMAIL and HELP_DESK_EMAIL_APP_PASSWORD in .env',
                'requires_config': True
            }), 500

        socket.setdefaulttimeout(30)
        mail = imaplib.IMAP4_SSL(imap_server, imap_port)
        mail.sock.settimeout(30)
        mail.login(help_desk_email, help_desk_password)
        mail.select('INBOX')

        status, messages = mail.search(None, 'ALL')
        if status != 'OK':
            if mail:
                try:
                    mail.close()
                    mail.logout()
                except Exception:
                    pass
            return jsonify({'error': 'Failed to search Outlook inbox'}), 500

        email_seq_ids = messages[0].split() if messages and messages[0] else []
        max_emails = 50
        emails_to_fetch = email_seq_ids[-max_emails:] if len(email_seq_ids) > max_emails else email_seq_ids
        email_statuses = load_email_statuses()
        emails = []

        for seq_id in emails_to_fetch:
            try:
                status, msg_parts = mail.fetch(seq_id, '(UID RFC822)')
                if status != 'OK' or not msg_parts or not msg_parts[0]:
                    continue
                part = msg_parts[0]
                if isinstance(part, tuple) and len(part) >= 2:
                    meta, raw_msg = part[0], part[1]
                else:
                    continue
                uid = None
                if meta:
                    meta_str = meta.decode() if isinstance(meta, bytes) else str(meta)
                    uid_m = re.search(r'UID\s+(\d+)', meta_str, re.IGNORECASE)
                    if uid_m:
                        uid = int(uid_m.group(1))
                if uid is None:
                    continue
                msg = email.message_from_bytes(raw_msg)

                subject_header = msg.get('Subject', '')
                if subject_header:
                    subject_decoded = decode_header(subject_header)[0][0]
                    subject = subject_decoded.decode('utf-8', errors='ignore') if isinstance(subject_decoded, bytes) else (subject_decoded or '(No Subject)')
                else:
                    subject = '(No Subject)'

                from_header = msg.get('From', '')
                from_email = from_header
                if '<' in from_header and '>' in from_header:
                    from_email = from_header[from_header.index('<') + 1:from_header.index('>')].strip()

                to_header = msg.get('To', '')
                to_email = to_header
                if '<' in to_header and '>' in to_header:
                    to_email = to_header[to_header.index('<') + 1:to_header.index('>')].strip()

                date_str = msg.get('Date', '')
                received_date = date_str
                try:
                    from email.utils import parsedate_to_datetime
                    dt = parsedate_to_datetime(date_str)
                    received_date = dt.isoformat() if dt else date_str
                except Exception:
                    pass

                body = ''
                body_preview = ''
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == 'text/plain':
                            payload = part.get_payload(decode=True)
                            if payload:
                                body = payload.decode('utf-8', errors='ignore')
                                body_preview = body[:200] if body else ''
                                break
                        elif content_type == 'text/html' and not body:
                            payload = part.get_payload(decode=True)
                            if payload:
                                html_body = payload.decode('utf-8', errors='ignore')
                                try:
                                    soup = BeautifulSoup(html_body, 'html.parser')
                                    body = soup.get_text(separator=' ', strip=True)
                                    body_preview = body[:200] if body else ''
                                except Exception:
                                    body = html_body
                                    body_preview = body[:200] if body else ''
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        body_preview = body[:200] if body else ''

                email_id_str = f'outlook_imap_{uid}'
                current_status = email_statuses.get(email_id_str, {}).get('status', 'Not started')

                emails.append({
                    'id': email_id_str,
                    'subject': subject,
                    'from': from_email,
                    'to': to_email,
                    'date': received_date,
                    'preview': body_preview,
                    'body': body,
                    'status': current_status
                })
            except Exception as e:
                print(f"Error processing Outlook IMAP email {seq_id}: {str(e)}")
                continue

        if mail:
            try:
                mail.close()
                mail.logout()
            except Exception:
                pass

        return jsonify({'emails': emails, 'mailbox': help_desk_email, 'source_config': 'imap'}), 200

    except imaplib.IMAP4.error as imap_error:
        err_msg = str(imap_error)
        app.logger.warning(f"Outlook IMAP login error: {err_msg}")
        print(f"[Outlook IMAP] Login failed: {err_msg}")
        if mail:
            try:
                mail.close()
                mail.logout()
            except Exception:
                pass
        if 'LOGIN' in err_msg or 'AUTHENTICATE' in err_msg:
            return jsonify({
                'error': f'Outlook IMAP login failed. {err_msg} Check HELP_DESK_EMAIL and HELP_DESK_EMAIL_APP_PASSWORD. If library@akello.co is a shared mailbox, use a user account that has full access instead (shared mailboxes often do not support direct IMAP login).',
                'requires_auth': True
            }), 401
        return jsonify({'error': f'Outlook IMAP error: {err_msg}'}), 500
    except Exception as e:
        if mail:
            try:
                mail.close()
                mail.logout()
            except Exception:
                pass
        print(f"Unexpected error in fetch_outlook_emails_imap: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def fetch_gmail_emails():
    """Fetch emails from Gmail using IMAP"""
    # #region agent log
    request_start_time = time.time()
    log_entry = {
        'sessionId': 'debug-session',
        'runId': 'run1',
        'hypothesisId': 'A',
        'location': 'routes.py:fetch_gmail_emails:entry',
        'message': 'fetch_gmail_emails called',
        'data': {'timestamp': request_start_time},
        'timestamp': int(time.time() * 1000)
    }
    try:
        write_debug_log(log_entry)
        app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
    except Exception as e:
        app.logger.error(f"[DEBUG] Failed to write log: {e}")
    # #endregion
    
    mail = None
    try:
        # Get Gmail configuration from environment
        gmail_imap_server = os.getenv('GMAIL_IMAP_SERVER', 'imap.gmail.com')
        gmail_imap_port = int(os.getenv('GMAIL_IMAP_PORT', '993'))
        gmail_email = os.getenv('GMAIL_EMAIL_ADDRESS', 'mashavaquincy@gmail.com')
        # Try both uppercase and lowercase variable names
        gmail_password = os.getenv('GOOGLE_APP_PASSWORD') or os.getenv('google_app_password')
        gmail_sender_filter = os.getenv('GMAIL_SENDER_FILTER', 'quincy.mashava@akello.co')
        
        # Debug logging - check both variable names
        gmail_password_upper = os.getenv('GOOGLE_APP_PASSWORD')
        gmail_password_lower = os.getenv('google_app_password')
        print(f"Gmail Config Check - Email: {gmail_email}")
        print(f"  GOOGLE_APP_PASSWORD (upper): {'✓ SET' if gmail_password_upper else '✗ NOT SET'}")
        print(f"  google_app_password (lower): {'✓ SET' if gmail_password_lower else '✗ NOT SET'}")
        print(f"  Final password: {'✓ SET' if gmail_password else '✗ NOT SET'}")
        print(f"  Filter: {gmail_sender_filter}")
        
        if not gmail_email or not gmail_password:
            error_msg = f'Gmail configuration not set. Email: {"✓" if gmail_email else "✗"}, Password: {"✓" if gmail_password else "✗"}'
            if not gmail_password:
                error_msg += '\n\nTo fix this:\n1. Get a Google App Password from: https://myaccount.google.com/apppasswords\n2. Add GOOGLE_APP_PASSWORD=your-password to your .env file\n3. Restart the Flask application'
            print(f"Gmail config error: {error_msg}")
            return ensure_json_response({
                'error': error_msg,
                'requires_config': True
            }, 500)
        
        # Connect to Gmail IMAP server with error handling
        try:
            # #region agent log
            imap_connect_start = time.time()
            log_entry = {
                'sessionId': 'debug-session',
                'runId': 'run1',
                'hypothesisId': 'A',
                'location': 'routes.py:fetch_gmail_emails:before_imap_connect',
                'message': 'About to connect to IMAP',
                'data': {'elapsed_ms': (imap_connect_start - request_start_time) * 1000},
                'timestamp': int(time.time() * 1000)
            }
            try:
                write_debug_log(log_entry)
                app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
            except Exception as e:
                app.logger.error(f"[DEBUG] Failed to write log: {e}")
            # #endregion
            
            # Create IMAP connection with socket timeout to prevent hanging
            # Set socket timeout to 30 seconds to prevent Cloudflare 502 errors
            socket.setdefaulttimeout(30)
            mail = imaplib.IMAP4_SSL(gmail_imap_server, gmail_imap_port)
            # Set timeout on the IMAP connection itself
            mail.sock.settimeout(30)
            mail.login(gmail_email, gmail_password)
            mail.select('inbox')
            
            # #region agent log
            imap_connect_time = (time.time() - imap_connect_start) * 1000
            log_entry = {
                'sessionId': 'debug-session',
                'runId': 'run1',
                'hypothesisId': 'A',
                'location': 'routes.py:fetch_gmail_emails:after_imap_connect',
                'message': 'IMAP connection established',
                'data': {'imap_connect_time_ms': imap_connect_time, 'total_elapsed_ms': (time.time() - request_start_time) * 1000},
                'timestamp': int(time.time() * 1000)
            }
            try:
                write_debug_log(log_entry)
                app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
            except Exception as e:
                app.logger.error(f"[DEBUG] Failed to write log: {e}")
            # #endregion
        except imaplib.IMAP4.error as imap_error:
            error_msg = str(imap_error)
            print(f"Gmail IMAP connection error: {error_msg}")
            if mail:
                try:
                    mail.close()
                    mail.logout()
                except:
                    pass
            if 'LOGIN' in error_msg or 'AUTHENTICATE' in error_msg:
                return ensure_json_response({
                    'error': 'Gmail authentication failed. Please check your Gmail app password.',
                    'requires_config': True
                }, 401)
            return ensure_json_response({'error': f'Gmail IMAP connection error: {error_msg}'}, 500)
        except Exception as conn_error:
            error_msg = str(conn_error)
            print(f"Gmail connection error: {error_msg}")
            if mail:
                try:
                    mail.close()
                    mail.logout()
                except:
                    pass
            return ensure_json_response({'error': f'Gmail connection error: {error_msg}'}, 500)
        
        # Search for emails from specific sender
        try:
            if gmail_sender_filter:
                search_query = f'FROM "{gmail_sender_filter}"'
                print(f"Searching for emails: {search_query}")
                status, messages = mail.search(None, search_query)
            else:
                print("Searching for all emails")
                status, messages = mail.search(None, 'ALL')
            
            if status != 'OK':
                error_msg = f'Failed to search emails. Status: {status}'
                print(f"Gmail search error: {error_msg}")
                if mail:
                    try:
                        mail.close()
                        mail.logout()
                    except:
                        pass
                return ensure_json_response({'error': error_msg}, 500)
        except Exception as search_error:
            error_msg = str(search_error)
            print(f"Gmail search error: {error_msg}")
            if mail:
                try:
                    mail.close()
                    mail.logout()
                except:
                    pass
            return ensure_json_response({'error': f'Gmail search error: {error_msg}'}, 500)
        
        email_ids = messages[0].split() if messages and messages[0] else []
        print(f"Found {len(email_ids)} email(s) matching the filter")
        
        # Load existing statuses
        email_statuses = load_email_statuses()
        
        # Limit to last 20 emails to prevent timeout (reduced from 50)
        # This ensures faster response times and prevents Cloudflare 502 errors
        max_emails = 20
        emails_to_fetch = email_ids[-max_emails:] if len(email_ids) > max_emails else email_ids
        print(f"Fetching {len(emails_to_fetch)} email(s) (limited from {len(email_ids)})")
        
        emails = []
        for email_id in emails_to_fetch:
            try:
                email_id_str = email_id.decode() if isinstance(email_id, bytes) else str(email_id)
                status, msg_data = mail.fetch(email_id, '(RFC822)')
                
                if status != 'OK' or not msg_data or not msg_data[0]:
                    continue
                
                msg = email.message_from_bytes(msg_data[0][1])
                
                # Decode subject
                subject_header = msg['Subject']
                if subject_header:
                    subject_decoded = decode_header(subject_header)[0][0]
                    if isinstance(subject_decoded, bytes):
                        subject = subject_decoded.decode('utf-8', errors='ignore')
                    else:
                        subject = subject_decoded
                else:
                    subject = '(No Subject)'
                
                # Get sender
                from_header = msg.get('From', '')
                from_email = from_header
                if '<' in from_header and '>' in from_header:
                    from_email = from_header[from_header.index('<')+1:from_header.index('>')]
                
                # Get recipients
                to_header = msg.get('To', '')
                to_email = to_header
                if '<' in to_header and '>' in to_header:
                    to_email = to_header[to_header.index('<')+1:to_header.index('>')]
                
                # Get date
                date_str = msg.get('Date', '')
                received_date = date_str
                
                # Parse date to ISO format if possible
                try:
                    from email.utils import parsedate_to_datetime
                    dt = parsedate_to_datetime(date_str)
                    received_date = dt.isoformat() if dt else date_str
                except:
                    pass
                
                # Get body preview and full body
                body = ''
                body_preview = ''
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == 'text/plain':
                            payload = part.get_payload(decode=True)
                            if payload:
                                body = payload.decode('utf-8', errors='ignore')
                                body_preview = body[:200] if body else ''
                                break
                        elif content_type == 'text/html' and not body:
                            payload = part.get_payload(decode=True)
                            if payload:
                                html_body = payload.decode('utf-8', errors='ignore')
                                from bs4 import BeautifulSoup
                                try:
                                    soup = BeautifulSoup(html_body, 'html.parser')
                                    body = soup.get_text(separator=' ', strip=True)
                                    body_preview = body[:200] if body else ''
                                except:
                                    body = html_body
                                    body_preview = body[:200] if body else ''
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        body_preview = body[:200] if body else ''
                
                # Get status from storage (use email_id_str as key)
                email_status_data = email_statuses.get(email_id_str, {})
                current_status = email_status_data.get('status', 'Not started')
                
                emails.append({
                    'id': email_id_str,
                    'subject': subject,
                    'from': from_email,
                    'to': to_email,
                    'date': received_date,
                    'preview': body_preview,
                    'body': body,
                    'status': current_status
                })
            except Exception as e:
                print(f"Error processing Gmail email {email_id}: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
        # Clean up IMAP connection
        # #region agent log
        before_close_time = time.time()
        log_entry = {
            'sessionId': 'debug-session',
            'runId': 'run1',
            'hypothesisId': 'C',
            'location': 'routes.py:fetch_gmail_emails:before_close',
            'message': 'About to close IMAP connection',
            'data': {'emails_count': len(emails), 'total_elapsed_ms': (before_close_time - request_start_time) * 1000, 'mail_exists': mail is not None},
            'timestamp': int(time.time() * 1000)
        }
        try:
            write_debug_log(log_entry)
            app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
        except Exception as e:
            app.logger.error(f"[DEBUG] Failed to write log: {e}")
        # #endregion
        
        if mail:
            try:
                mail.close()
                mail.logout()
                # #region agent log
                log_entry = {
                    'sessionId': 'debug-session',
                    'runId': 'run1',
                    'hypothesisId': 'C',
                    'location': 'routes.py:fetch_gmail_emails:after_close',
                    'message': 'IMAP connection closed',
                    'data': {'close_time_ms': (time.time() - before_close_time) * 1000},
                    'timestamp': int(time.time() * 1000)
                }
                try:
                    write_debug_log(log_entry)
                    app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
                except Exception as e:
                    app.logger.error(f"[DEBUG] Failed to write log: {e}")
                # #endregion
            except Exception as close_error:
                # #region agent log
                log_entry = {
                    'sessionId': 'debug-session',
                    'runId': 'run1',
                    'hypothesisId': 'C',
                    'location': 'routes.py:fetch_gmail_emails:close_error',
                    'message': 'Error closing IMAP connection',
                    'data': {'error': str(close_error)},
                    'timestamp': int(time.time() * 1000)
                }
                try:
                    write_debug_log(log_entry)
                    app.logger.error(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
                except Exception as e:
                    app.logger.error(f"[DEBUG] Failed to write log: {e}")
                # #endregion
                pass
        
        # Sort by date (most recent first)
        emails.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        # #region agent log
        before_response_time = time.time()
        total_elapsed = (before_response_time - request_start_time) * 1000
        emails_data_size = len(json.dumps(emails)) if emails else 0
        log_entry = {
            'sessionId': 'debug-session',
            'runId': 'run1',
            'hypothesisId': 'A,D',
            'location': 'routes.py:fetch_gmail_emails:before_response',
            'message': 'About to create response',
            'data': {'emails_count': len(emails), 'total_elapsed_ms': total_elapsed, 'emails_data_size_bytes': emails_data_size},
            'timestamp': int(time.time() * 1000)
        }
        try:
            write_debug_log(log_entry)
            app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
        except Exception as e:
            app.logger.error(f"[DEBUG] Failed to write log: {e}")
        # #endregion
        
        print(f"Successfully fetched {len(emails)} Gmail email(s)")
        response = ensure_json_response({'emails': emails}, 200)
        
        # #region agent log
        after_response_time = time.time()
        log_entry = {
            'sessionId': 'debug-session',
            'runId': 'run1',
            'hypothesisId': 'A,B',
            'location': 'routes.py:fetch_gmail_emails:after_response',
            'message': 'Response created, about to return',
            'data': {'response_creation_time_ms': (after_response_time - before_response_time) * 1000, 'total_elapsed_ms': (after_response_time - request_start_time) * 1000},
            'timestamp': int(time.time() * 1000)
        }
        try:
            write_debug_log(log_entry)
            app.logger.info(f"[DEBUG] {log_entry['location']}: {log_entry['message']} - {json.dumps(log_entry['data'])}")
        except Exception as e:
            app.logger.error(f"[DEBUG] Failed to write log: {e}")
        # #endregion
        
        return response
        
    except imaplib.IMAP4.error as e:
        error_msg = str(e)
        print(f"Gmail IMAP error: {error_msg}")
        if mail:
            try:
                mail.close()
                mail.logout()
            except:
                pass
        if 'LOGIN' in error_msg or 'AUTHENTICATE' in error_msg:
            return ensure_json_response({
                'error': 'Gmail authentication failed. Please check your Gmail app password.',
                'requires_config': True
            }, 401)
        return ensure_json_response({'error': f'Gmail IMAP error: {error_msg}'}, 500)
    except Exception as e:
        print(f"Unexpected error fetching Gmail emails: {str(e)}")
        import traceback
        traceback.print_exc()
        if mail:
            try:
                mail.close()
                mail.logout()
            except:
                pass
        return ensure_json_response({'error': f'Unexpected error: {str(e)}'}, 500)


@app.route('/api/email-queries/<email_id>', methods=['GET'])
@login_required
def get_email_query_details(email_id):
    """Get details of a specific email using Microsoft Graph API or Gmail IMAP"""
    try:
        # Get email source from query parameter (default: 'outlook')
        email_source = request.args.get('source', 'outlook').lower()
        
        if email_source == 'gmail':
            return get_gmail_email_details(email_id)
        if email_source == 'outlook' and str(email_id).startswith('outlook_imap_'):
            return get_outlook_imap_email_details(email_id)
        return get_outlook_email_details(email_id)
    except Exception as e:
        print(f"Unexpected error in get_email_query_details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def get_outlook_imap_email_details(email_id):
    """Get details of a specific email from Outlook using IMAP (by UID)"""
    try:
        # TODO: remove hardcoded fallback after testing
        help_desk_email = os.getenv('HELP_DESK_EMAIL') or 'library@akello.co'
        help_desk_password = os.getenv('HELP_DESK_EMAIL_APP_PASSWORD') or 'wdmqpbqtpdkqsmbr'
        imap_server = os.getenv('HELP_DESK_IMAP_SERVER', 'outlook.office365.com')
        imap_port = int(os.getenv('HELP_DESK_IMAP_PORT', '993'))

        if not help_desk_email or not help_desk_password:
            return jsonify({
                'error': 'Help desk Outlook IMAP not configured.',
                'requires_config': True
            }), 500

        uid_str = str(email_id).replace('outlook_imap_', '', 1)
        try:
            uid = int(uid_str)
        except ValueError:
            return jsonify({'error': 'Invalid email id'}), 400

        mail = imaplib.IMAP4_SSL(imap_server, imap_port)
        mail.login(help_desk_email, help_desk_password)
        mail.select('INBOX')
        status, msg_parts = mail.fetch(str(uid), '(UID RFC822)')
        mail.close()
        mail.logout()

        if status != 'OK' or not msg_parts or not msg_parts[0]:
            return jsonify({'error': 'Email not found'}), 404

        part = msg_parts[0]
        if isinstance(part, tuple) and len(part) >= 2:
            raw_msg = part[1]
        else:
            return jsonify({'error': 'Email not found'}), 404

        msg = email.message_from_bytes(raw_msg)

        subject_header = msg.get('Subject', '')
        if subject_header:
            subject_decoded = decode_header(subject_header)[0][0]
            subject = subject_decoded.decode('utf-8', errors='ignore') if isinstance(subject_decoded, bytes) else (subject_decoded or '(No Subject)')
        else:
            subject = '(No Subject)'

        from_header = msg.get('From', '')
        from_email = from_header
        if '<' in from_header and '>' in from_header:
            from_email = from_header[from_header.index('<') + 1:from_header.index('>')].strip()

        to_header = msg.get('To', '')
        to_email = to_header
        if '<' in to_header and '>' in to_header:
            to_email = to_header[to_header.index('<') + 1:to_header.index('>')].strip()

        date_str = msg.get('Date', '')
        received_date = date_str
        try:
            from email.utils import parsedate_to_datetime
            dt = parsedate_to_datetime(date_str)
            received_date = dt.isoformat() if dt else date_str
        except Exception:
            pass

        body = ''
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        break
                elif content_type == 'text/html' and not body:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_body = payload.decode('utf-8', errors='ignore')
                        try:
                            soup = BeautifulSoup(html_body, 'html.parser')
                            body = soup.get_text(separator='\n', strip=True)
                        except Exception:
                            body = html_body
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode('utf-8', errors='ignore')

        return jsonify({
            'email': {
                'id': email_id,
                'subject': subject,
                'from': from_email,
                'to': to_email,
                'date': received_date,
                'body': body
            }
        }), 200

    except imaplib.IMAP4.error as e:
        return jsonify({'error': f'Outlook IMAP error: {str(e)}'}), 500
    except Exception as e:
        print(f"Unexpected error in get_outlook_imap_email_details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def get_outlook_email_details(email_id):
    """Get details of a specific email using Microsoft Graph API"""
    try:
        # Check if user is authenticated with Microsoft
        access_token = get_graph_api_token()
        if not access_token:
            return jsonify({
                'error': 'Not authenticated with Microsoft. Please connect to Outlook first.',
                'requires_auth': True
            }), 401
        
        # Build Graph API endpoint for specific email
        graph_endpoint = f"https://graph.microsoft.com/v1.0/me/messages/{email_id}"
        
        # Make request to Graph API
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }
        
        print(f"Fetching email details from Graph API: {graph_endpoint}")
        response = requests.get(graph_endpoint, headers=headers)
        
        if response.status_code == 401:
            # Token expired or invalid, clear session
            session.pop('microsoft_token', None)
            return jsonify({
                'error': 'Authentication expired. Please reconnect to Outlook.',
                'requires_auth': True
            }), 401
        
        if not response.ok:
            error_data = response.json() if response.content else {}
            error_msg = error_data.get('error', {}).get('message', f'Graph API error: {response.status_code}')
            return jsonify({
                'error': f'Failed to fetch email: {error_msg}',
                'details': error_data
            }), response.status_code
        
        # Parse response
        graph_email = response.json()
        
        # Extract email data
        from_info = graph_email.get('from', {}).get('emailAddress', {})
        from_email = from_info.get('address', 'Unknown')
        from_name = from_info.get('name', '')
        
        # Get recipients
        to_recipients = graph_email.get('toRecipients', [])
        to_email = ', '.join([r.get('emailAddress', {}).get('address', '') for r in to_recipients])
        
        # Get subject
        subject = graph_email.get('subject', '(No Subject)')
        
        # Get date
        received_date = graph_email.get('receivedDateTime', '')
        
        # Get body
        body_content = graph_email.get('body', {})
        body = body_content.get('content', '') if isinstance(body_content, dict) else ''
        
        # Clean HTML from body if present
        if body_content.get('contentType') == 'html':
            from bs4 import BeautifulSoup
            try:
                soup = BeautifulSoup(body, 'html.parser')
                body = soup.get_text(separator='\n', strip=True)
            except:
                pass  # Keep original if parsing fails
        
        return jsonify({
            'email': {
                'id': email_id,
                'subject': subject,
                'from': from_email,
                'to': to_email,
                'date': received_date,
                'body': body
            }
        }), 200
        
    except Exception as e:
        print(f"Unexpected error fetching Outlook email details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def get_gmail_email_details(email_id):
    """Get details of a specific email from Gmail using IMAP"""
    try:
        # Get Gmail configuration from environment
        gmail_imap_server = os.getenv('GMAIL_IMAP_SERVER', 'imap.gmail.com')
        gmail_imap_port = int(os.getenv('GMAIL_IMAP_PORT', '993'))
        gmail_email = os.getenv('GMAIL_EMAIL_ADDRESS', 'mashavaquincy@gmail.com')
        # Try both uppercase and lowercase variable names
        gmail_password = os.getenv('GOOGLE_APP_PASSWORD') or os.getenv('google_app_password')
        
        if not gmail_email or not gmail_password:
            return jsonify({
                'error': 'Gmail configuration not set. Please configure GMAIL_EMAIL_ADDRESS and GOOGLE_APP_PASSWORD.',
                'requires_config': True
            }), 500
        
        # Connect to Gmail IMAP server
        mail = imaplib.IMAP4_SSL(gmail_imap_server, gmail_imap_port)
        mail.login(gmail_email, gmail_password)
        mail.select('inbox')
        
        # Fetch the specific email
        status, msg_data = mail.fetch(email_id.encode(), '(RFC822)')
        
        if status != 'OK':
            mail.close()
            mail.logout()
            return jsonify({'error': 'Email not found'}), 404
        
        msg = email.message_from_bytes(msg_data[0][1])
        
        # Decode subject
        subject_header = msg['Subject']
        if subject_header:
            subject_decoded = decode_header(subject_header)[0][0]
            if isinstance(subject_decoded, bytes):
                subject = subject_decoded.decode('utf-8', errors='ignore')
            else:
                subject = subject_decoded
        else:
            subject = '(No Subject)'
        
        # Get sender
        from_header = msg.get('From', '')
        from_email = from_header
        if '<' in from_header and '>' in from_header:
            from_email = from_header[from_header.index('<')+1:from_header.index('>')]
        
        # Get recipients
        to_header = msg.get('To', '')
        to_email = to_header
        if '<' in to_header and '>' in to_header:
            to_email = to_header[to_header.index('<')+1:to_header.index('>')]
        
        # Get date
        date_str = msg.get('Date', '')
        received_date = date_str
        try:
            from email.utils import parsedate_to_datetime
            dt = parsedate_to_datetime(date_str)
            received_date = dt.isoformat() if dt else date_str
        except:
            pass
        
        # Get body
        body = ''
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        break
                elif content_type == 'text/html' and not body:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_body = payload.decode('utf-8', errors='ignore')
                        from bs4 import BeautifulSoup
                        try:
                            soup = BeautifulSoup(html_body, 'html.parser')
                            body = soup.get_text(separator='\n', strip=True)
                        except:
                            body = html_body
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode('utf-8', errors='ignore')
        
        mail.close()
        mail.logout()
        
        return jsonify({
            'email': {
                'id': email_id,
                'subject': subject,
                'from': from_email,
                'to': to_email,
                'date': received_date,
                'body': body
            }
        }), 200
        
    except imaplib.IMAP4.error as e:
        error_msg = str(e)
        if 'LOGIN' in error_msg or 'AUTHENTICATE' in error_msg:
            return jsonify({
                'error': 'Gmail authentication failed. Please check your Gmail app password.',
                'requires_config': True
            }), 401
        return jsonify({'error': f'Gmail IMAP error: {error_msg}'}), 500
    except Exception as e:
        print(f"Unexpected error fetching Gmail email details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


@app.route('/api/email-queries/<email_id>/status', methods=['PATCH'])
@login_required
def update_email_query_status(email_id):
    """Update the status of an email query"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.get_json()
        status = data.get('status')
        
        if not status or status not in ['Not started', 'Looking into it', 'Resolved']:
            return jsonify({'error': 'Invalid status'}), 400
        
        # Save the status
        save_email_status(email_id, status)
        
        return jsonify({'success': True, 'status': status}), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/email-queries/<email_id>/send-resolution', methods=['POST'])
@login_required
def send_resolution_email(email_id):
    """Send resolution notification email to customer"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.get_json()
        recipient = data.get('recipient')
        message_body = data.get('message')
        
        if not recipient or not message_body:
            return jsonify({'error': 'Recipient and message are required'}), 400
        
        # Get SMTP configuration
        smtp_server = os.getenv('EMAIL_SMTP_SERVER', 'smtp.office365.com')
        smtp_port = int(os.getenv('EMAIL_SMTP_PORT', '587'))
        email_address = os.getenv('EMAIL_ADDRESS')
        email_password = os.getenv('EMAIL_APP_PASSWORD')
        
        if not email_address or not email_password:
            return jsonify({'error': 'Email configuration not set'}), 500
        
        # Create email message
        msg = MIMEMultipart()
        msg['From'] = email_address
        msg['To'] = recipient
        msg['Subject'] = 'Query Resolution - Akello Support'
        
        # Email body
        body_html = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background: linear-gradient(135deg, #4f46e5 0%, #06b6d4 100%); padding: 20px; text-align: center; border-radius: 8px 8px 0 0;">
                    <h2 style="color: white; margin: 0;">Akello Support</h2>
                </div>
                <div style="background: #ffffff; padding: 30px; border: 1px solid #e5e7eb; border-radius: 0 0 8px 8px;">
                    <p style="white-space: pre-wrap;">{message_body}</p>
                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 20px 0;">
                    <p style="font-size: 12px; color: #6b7280;">
                        This is an automated message from Akello Support.<br>
                        For further assistance, please contact us at: <a href="mailto:info@akello.co">info@akello.co</a>
                    </p>
                </div>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(body_html, 'html'))
        
        # Send email via SMTP
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_address, email_password)
        server.send_message(msg)
        server.quit()
        
        return jsonify({'success': True, 'message': 'Email sent successfully'}), 200
        
    except Exception as e:
        print(f"Error sending email: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/email-queries/<email_id>/convert', methods=['POST'])
@login_required
def convert_email_to_query(email_id):
    """Convert an email to a help desk query using Microsoft Graph API or Gmail IMAP"""
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        # Get email source from request body or query parameter (default: 'outlook')
        data = request.get_json() or {}
        email_source = data.get('source') or request.args.get('source', 'outlook').lower()
        
        if email_source == 'gmail':
            return convert_gmail_email_to_query(email_id)
        if email_source == 'outlook' and str(email_id).startswith('outlook_imap_'):
            return convert_outlook_imap_email_to_query(email_id)
        return convert_outlook_email_to_query(email_id)
    except Exception as e:
        db.session.rollback()
        print(f"Unexpected error in convert_email_to_query: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def convert_outlook_imap_email_to_query(email_id):
    """Convert an Outlook IMAP email (outlook_imap_* id) to a help desk query; uses HELP_DESK_EMAIL app password."""
    try:
        # TODO: remove hardcoded fallback after testing
        help_desk_email = os.getenv('HELP_DESK_EMAIL') or 'library@akello.co'
        help_desk_password = os.getenv('HELP_DESK_EMAIL_APP_PASSWORD') or 'wdmqpbqtpdkqsmbr'
        imap_server = os.getenv('HELP_DESK_IMAP_SERVER', 'outlook.office365.com')
        imap_port = int(os.getenv('HELP_DESK_IMAP_PORT', '993'))
        if not help_desk_email or not help_desk_password:
            return jsonify({'error': 'Help desk Outlook IMAP not configured.', 'requires_config': True}), 500

        uid_str = str(email_id).replace('outlook_imap_', '', 1)
        try:
            uid = int(uid_str)
        except ValueError:
            return jsonify({'error': 'Invalid email id'}), 400

        mail = imaplib.IMAP4_SSL(imap_server, imap_port)
        mail.login(help_desk_email, help_desk_password)
        mail.select('INBOX')
        status, msg_parts = mail.fetch(str(uid), '(UID RFC822)')
        mail.close()
        mail.logout()

        if status != 'OK' or not msg_parts or not msg_parts[0]:
            return jsonify({'error': 'Email not found'}), 404
        part = msg_parts[0]
        if not isinstance(part, tuple) or len(part) < 2:
            return jsonify({'error': 'Email not found'}), 404
        msg = email.message_from_bytes(part[1])

        subject_header = msg.get('Subject', '')
        subject = '(No Subject)'
        if subject_header:
            subject_decoded = decode_header(subject_header)[0][0]
            subject = subject_decoded.decode('utf-8', errors='ignore') if isinstance(subject_decoded, bytes) else (subject_decoded or '(No Subject)')

        from_header = msg.get('From', '')
        from_email = from_header
        if '<' in from_header and '>' in from_header:
            from_email = from_header[from_header.index('<') + 1:from_header.index('>')].strip()

        body = ''
        if msg.is_multipart():
            for p in msg.walk():
                ct = p.get_content_type()
                if ct == 'text/plain':
                    payload = p.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        break
                elif ct == 'text/html' and not body:
                    payload = p.get_payload(decode=True)
                    if payload:
                        html_body = payload.decode('utf-8', errors='ignore')
                        try:
                            body = BeautifulSoup(html_body, 'html.parser').get_text(separator='\n', strip=True)
                        except Exception:
                            body = html_body
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode('utf-8', errors='ignore')

        from app.models import HelpDeskQuery, Notification, User
        from sqlalchemy import inspect
        data = request.get_json() or {}
        assigned_user_ids = data.get('assigned_user_ids', [])

        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False

        if not has_resolved_at:
            result = db.session.execute(
                text("""
                    INSERT INTO helpdesk_queries (query_title, query_description, timestamp, query_type, created_by, image_path, status)
                    VALUES (:title, :description, :timestamp, :type, :created_by, :image_path, :status)
                """),
                {
                    'title': subject,
                    'description': f"From: {from_email}\n\n{body}",
                    'timestamp': datetime.utcnow(),
                    'type': 'Email',
                    'created_by': from_email,
                    'image_path': None,
                    'status': 'Not started'
                }
            )
            db.session.flush()
            query_id = result.lastrowid
            query_row = db.session.execute(
                text("SELECT id, query_title, query_description, timestamp, query_type, created_by, image_path, status FROM helpdesk_queries WHERE id = :id"),
                {'id': query_id}
            ).fetchone()
            if query_row:
                row_dict = dict(query_row._mapping)
                if 'timestamp' in row_dict and isinstance(row_dict.get('timestamp'), str):
                    try:
                        row_dict['timestamp'] = datetime.fromisoformat(row_dict['timestamp'].replace('Z', '+00:00'))
                    except Exception:
                        try:
                            row_dict['timestamp'] = datetime.strptime(row_dict['timestamp'], '%Y-%m-%d %H:%M:%S')
                        except Exception:
                            row_dict['timestamp'] = None
                query = HelpDeskQuery(**row_dict)
            else:
                raise Exception("Failed to retrieve created query")
        else:
            query = HelpDeskQuery(
                query_title=subject,
                query_description=f"From: {from_email}\n\n{body}",
                query_type='Email',
                created_by=from_email,
                status='Not started'
            )
            db.session.add(query)
            db.session.flush()

        if assigned_user_ids:
            for user_id in assigned_user_ids:
                user = User.query.get(user_id)
                if user:
                    try:
                        db.session.execute(
                            text("INSERT INTO query_assignees (query_id, user_id) VALUES (:query_id, :user_id)"),
                            {'query_id': query.id, 'user_id': user_id}
                        )
                    except Exception:
                        try:
                            query.assignees.append(user)
                        except Exception as e:
                            app.logger.warning(f"Could not assign user: {str(e)}")
                    try:
                        notification = Notification(
                            user_id=user_id,
                            query_id=query.id,
                            message=f"You have been assigned a new query: '{subject}'",
                            notification_type='assignment'
                        )
                        db.session.add(notification)
                    except Exception as e:
                        app.logger.warning(f"Could not create notification: {str(e)}")

        db.session.commit()
        return jsonify({'success': True, 'query_id': query.id}), 201
    except Exception as e:
        db.session.rollback()
        print(f"Unexpected error in convert_outlook_imap_email_to_query: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


def convert_outlook_email_to_query(email_id):
    """Convert an Outlook email to a help desk query using Microsoft Graph API"""
    try:
        # Check if user is authenticated with Microsoft
        access_token = get_graph_api_token()
        if not access_token:
            return jsonify({
                'error': 'Not authenticated with Microsoft. Please connect to Outlook first.',
                'requires_auth': True
            }), 401
        
        # Build Graph API endpoint for specific email
        graph_endpoint = f"https://graph.microsoft.com/v1.0/me/messages/{email_id}"
        
        # Make request to Graph API
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }
        
        print(f"Fetching email for conversion from Graph API: {graph_endpoint}")
        response = requests.get(graph_endpoint, headers=headers)
        
        if response.status_code == 401:
            # Token expired or invalid, clear session
            session.pop('microsoft_token', None)
            return jsonify({
                'error': 'Authentication expired. Please reconnect to Outlook.',
                'requires_auth': True
            }), 401
        
        if not response.ok:
            error_data = response.json() if response.content else {}
            error_msg = error_data.get('error', {}).get('message', f'Graph API error: {response.status_code}')
            return jsonify({
                'error': f'Failed to fetch email: {error_msg}',
                'details': error_data
            }), response.status_code
        
        # Parse response
        graph_email = response.json()
        
        # Extract email data
        from_info = graph_email.get('from', {}).get('emailAddress', {})
        from_email = from_info.get('address', 'Unknown')
        from_name = from_info.get('name', '')
        
        # Get subject
        subject = graph_email.get('subject', 'Email Query')
        
        # Get body
        body_content = graph_email.get('body', {})
        body = body_content.get('content', '') if isinstance(body_content, dict) else ''
        
        # Clean HTML from body if present
        if body_content.get('contentType') == 'html':
            from bs4 import BeautifulSoup
            try:
                soup = BeautifulSoup(body, 'html.parser')
                body = soup.get_text(separator='\n', strip=True)
            except:
                pass  # Keep original if parsing fails
        
        # Create a help desk query from the email
        from app.models import HelpDeskQuery, Notification, User
        from sqlalchemy import inspect
        
        # Get assigned user IDs from request
        data = request.get_json() or {}
        assigned_user_ids = data.get('assigned_user_ids', [])
        
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Create query using raw SQL if column doesn't exist to avoid INSERT error
        if not has_resolved_at:
            # Use raw SQL to insert without resolved_at column
            result = db.session.execute(
                text("""
                    INSERT INTO helpdesk_queries (query_title, query_description, timestamp, query_type, created_by, image_path, status)
                    VALUES (:title, :description, :timestamp, :type, :created_by, :image_path, :status)
                """),
                {
                    'title': subject,
                    'description': f"From: {from_email}\n\n{body}",
                    'timestamp': datetime.utcnow(),
                    'type': 'Email',
                    'created_by': from_email,
                    'image_path': None,
                    'status': 'Not started'
                }
            )
            db.session.flush()  # Don't commit yet, need to handle assignments
            query_id = result.lastrowid
            # Fetch query using raw SQL to avoid resolved_at column
            query_row = db.session.execute(
                text("SELECT id, query_title, query_description, timestamp, query_type, created_by, image_path, status FROM helpdesk_queries WHERE id = :id"),
                {'id': query_id}
            ).fetchone()
            if query_row:
                row_dict = dict(query_row._mapping)
                # Convert timestamp string to datetime if needed
                if 'timestamp' in row_dict and isinstance(row_dict['timestamp'], str):
                    try:
                        row_dict['timestamp'] = datetime.fromisoformat(row_dict['timestamp'].replace('Z', '+00:00'))
                    except Exception:
                        try:
                            row_dict['timestamp'] = datetime.strptime(row_dict['timestamp'], '%Y-%m-%d %H:%M:%S')
                        except Exception:
                            row_dict['timestamp'] = None
                query = HelpDeskQuery(**row_dict)
            else:
                raise Exception("Failed to retrieve created query")
        else:
            # Use ORM when column exists
            query = HelpDeskQuery(
                query_title=subject,
                query_description=f"From: {from_email}\n\n{body}",
                query_type='Email',
                created_by=from_email,
                status='Not started'
            )
            db.session.add(query)
            db.session.flush()  # Get query.id before commit
        
        # Assign users if provided
        if assigned_user_ids:
            for user_id in assigned_user_ids:
                user = User.query.get(user_id)
                if user:
                    # Use raw SQL to insert into query_assignees if table exists
                    try:
                        db.session.execute(
                            text("INSERT INTO query_assignees (query_id, user_id) VALUES (:query_id, :user_id)"),
                            {'query_id': query.id, 'user_id': user_id}
                        )
                    except Exception:
                        # Table might not exist, try using relationship
                        try:
                            query.assignees.append(user)
                        except Exception as e:
                            app.logger.warning(f"Could not assign user: {str(e)}")
                    
                    # Create notification for assigned user
                    try:
                        notification = Notification(
                            user_id=user_id,
                            query_id=query.id,
                            message=f"You have been assigned a new query: '{subject}'",
                            notification_type='assignment'
                        )
                        db.session.add(notification)
                    except Exception as e:
                        app.logger.warning(f"Could not create notification: {str(e)}")
        
        db.session.commit()
        
        return jsonify({'success': True, 'query_id': query.id}), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"Unexpected error converting Outlook email to query: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500


def convert_gmail_email_to_query(email_id):
    """Convert a Gmail email to a help desk query using IMAP"""
    try:
        # Get Gmail configuration from environment
        gmail_imap_server = os.getenv('GMAIL_IMAP_SERVER', 'imap.gmail.com')
        gmail_imap_port = int(os.getenv('GMAIL_IMAP_PORT', '993'))
        gmail_email = os.getenv('GMAIL_EMAIL_ADDRESS', 'mashavaquincy@gmail.com')
        # Try both uppercase and lowercase variable names
        gmail_password = os.getenv('GOOGLE_APP_PASSWORD') or os.getenv('google_app_password')
        
        if not gmail_email or not gmail_password:
            return jsonify({
                'error': 'Gmail configuration not set. Please configure GMAIL_EMAIL_ADDRESS and GOOGLE_APP_PASSWORD.',
                'requires_config': True
            }), 500
        
        # Connect to Gmail IMAP server
        mail = imaplib.IMAP4_SSL(gmail_imap_server, gmail_imap_port)
        mail.login(gmail_email, gmail_password)
        mail.select('inbox')
        
        # Fetch the specific email
        status, msg_data = mail.fetch(email_id.encode(), '(RFC822)')
        
        if status != 'OK':
            mail.close()
            mail.logout()
            return jsonify({'error': 'Email not found'}), 404
        
        msg = email.message_from_bytes(msg_data[0][1])
        
        # Decode subject
        subject_header = msg['Subject']
        if subject_header:
            subject_decoded = decode_header(subject_header)[0][0]
            if isinstance(subject_decoded, bytes):
                subject = subject_decoded.decode('utf-8', errors='ignore')
            else:
                subject = subject_decoded
        else:
            subject = 'Email Query'
        
        # Get sender
        from_header = msg.get('From', '')
        from_email = from_header
        if '<' in from_header and '>' in from_header:
            from_email = from_header[from_header.index('<')+1:from_header.index('>')]
        
        # Get body
        body = ''
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                        break
                elif content_type == 'text/html' and not body:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_body = payload.decode('utf-8', errors='ignore')
                        from bs4 import BeautifulSoup
                        try:
                            soup = BeautifulSoup(html_body, 'html.parser')
                            body = soup.get_text(separator='\n', strip=True)
                        except:
                            body = html_body
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode('utf-8', errors='ignore')
        
        mail.close()
        mail.logout()
        
        # Create a help desk query from the email
        from app.models import HelpDeskQuery, Notification, User
        from sqlalchemy import inspect
        
        # Get assigned user IDs from request
        data = request.get_json() or {}
        assigned_user_ids = data.get('assigned_user_ids', [])
        
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Create query using raw SQL if column doesn't exist to avoid INSERT error
        if not has_resolved_at:
            # Use raw SQL to insert without resolved_at column
            result = db.session.execute(
                text("""
                    INSERT INTO helpdesk_queries (query_title, query_description, timestamp, query_type, created_by, image_path, status)
                    VALUES (:title, :description, :timestamp, :type, :created_by, :image_path, :status)
                """),
                {
                    'title': subject,
                    'description': f"From: {from_email}\n\n{body}",
                    'timestamp': datetime.utcnow(),
                    'type': 'Email',
                    'created_by': from_email,
                    'image_path': None,
                    'status': 'Not started'
                }
            )
            db.session.flush()  # Don't commit yet, need to handle assignments
            query_id = result.lastrowid
            # Fetch query using raw SQL to avoid resolved_at column
            query_row = db.session.execute(
                text("SELECT id, query_title, query_description, timestamp, query_type, created_by, image_path, status FROM helpdesk_queries WHERE id = :id"),
                {'id': query_id}
            ).fetchone()
            if query_row:
                row_dict = dict(query_row._mapping)
                # Convert timestamp string to datetime if needed
                if 'timestamp' in row_dict and isinstance(row_dict['timestamp'], str):
                    try:
                        row_dict['timestamp'] = datetime.fromisoformat(row_dict['timestamp'].replace('Z', '+00:00'))
                    except Exception:
                        try:
                            row_dict['timestamp'] = datetime.strptime(row_dict['timestamp'], '%Y-%m-%d %H:%M:%S')
                        except Exception:
                            row_dict['timestamp'] = None
                query = HelpDeskQuery(**row_dict)
            else:
                raise Exception("Failed to retrieve created query")
        else:
            # Use ORM when column exists
            query = HelpDeskQuery(
                query_title=subject,
                query_description=f"From: {from_email}\n\n{body}",
                query_type='Email',
                created_by=from_email,
                status='Not started'
            )
            db.session.add(query)
            db.session.flush()  # Get query.id before commit
        
        # Assign users if provided
        if assigned_user_ids:
            for user_id in assigned_user_ids:
                user = User.query.get(user_id)
                if user:
                    # Use raw SQL to insert into query_assignees if table exists
                    try:
                        db.session.execute(
                            text("INSERT INTO query_assignees (query_id, user_id) VALUES (:query_id, :user_id)"),
                            {'query_id': query.id, 'user_id': user_id}
                        )
                    except Exception:
                        # Table might not exist, try using relationship
                        try:
                            query.assignees.append(user)
                        except Exception as e:
                            app.logger.warning(f"Could not assign user: {str(e)}")
                    
                    # Create notification for assigned user
                    try:
                        notification = Notification(
                            user_id=user_id,
                            query_id=query.id,
                            message=f"You have been assigned a new query: '{subject}'",
                            notification_type='assignment'
                        )
                        db.session.add(notification)
                    except Exception as e:
                        app.logger.warning(f"Could not create notification: {str(e)}")
        
        db.session.commit()
        
        return jsonify({'success': True, 'query_id': query.id}), 201
        
    except imaplib.IMAP4.error as e:
        db.session.rollback()
        error_msg = str(e)
        if 'LOGIN' in error_msg or 'AUTHENTICATE' in error_msg:
            return jsonify({
                'error': 'Gmail authentication failed. Please check your Gmail app password.',
                'requires_config': True
            }), 401
        return jsonify({'error': f'Gmail IMAP error: {error_msg}'}), 500
    except Exception as e:
        db.session.rollback()
        print(f"Unexpected error converting Gmail email to query: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500







@app.route('/akello_sim_calendar', methods=['GET'])
@login_required
def akello_sim_calendar():
    form = AkelloSimEventForm()
    return render_template('akello_sim_calendar.html', form=form, title='Akello SIM Calendar')


@app.route('/api/sim-calendar/events', methods=['GET'])
@login_required
def sim_calendar_events():
    evs = AkelloSimEvent.query.order_by(AkelloSimEvent.date.asc()).all()
    return jsonify([
        {
            'id': e.id,
            'title': e.calendar_title,
            'start': e.date.isoformat(),
            'status': e.status,
            'request_collateral': e.request_collateral,
            'collateral_items': e.collateral_items or []
        } for e in evs
    ])


@app.route('/api/sim-calendar/events', methods=['POST'])
@login_required
def sim_calendar_create():
    data = request.get_json() or {}
    try:
        date_str = data.get('date')
        dt = datetime.fromisoformat(date_str) if date_str else None
        if not dt:
            return jsonify({'error': 'date is required (ISO 8601)'}), 400
        evt = AkelloSimEvent(
            calendar_title=data.get('calendar_title') or data.get('title') or 'Untitled',
            description=data.get('description', ''),
            date=dt,
            status=data.get('status', 'Confirmed'),
            created_by=current_user.username,
            request_collateral=bool(data.get('request_collateral', False)),
            collateral_items=data.get('collateral_items') or []
        )
        db.session.add(evt)
        db.session.commit()
        return jsonify({'id': evt.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/sim-calendar/events/<int:event_id>', methods=['PATCH'])
@login_required
def sim_calendar_update(event_id):
    evt = AkelloSimEvent.query.get_or_404(event_id)
    data = request.get_json() or {}
    if 'calendar_title' in data or 'title' in data:
        evt.calendar_title = data.get('calendar_title') or data.get('title')
    if 'description' in data:
        evt.description = data.get('description')
    if 'date' in data and data['date']:
        try:
            evt.date = datetime.fromisoformat(data['date'])
        except Exception:
            return jsonify({'error': 'Invalid date format'}), 400
    if 'status' in data:
        evt.status = data['status']
    if 'request_collateral' in data:
        evt.request_collateral = bool(data['request_collateral'])
    if 'collateral_items' in data:
        evt.collateral_items = data['collateral_items'] or []
    db.session.commit()
    return jsonify({'status': 'ok'})


@app.route('/api/sim-calendar/events/<int:event_id>', methods=['DELETE'])
@login_required
def sim_calendar_delete(event_id):
    evt = AkelloSimEvent.query.get_or_404(event_id)
    db.session.delete(evt)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/akello_monitoring', methods=['GET'])
@login_required
def akello_monitoring():
    return render_template('akello_monitor.html', title='Analytics')

@app.route('/akello_monitor', methods=['GET'])
@login_required
def akello_monitor():
    return render_template('simone_monitor.html', title='Analytics')


@app.route('/help-desk', methods=['GET', 'POST'])
@login_required
def help_desk():
    """Legacy Help Desk entry — redirect into the support hub."""
    qid = request.args.get('query', type=int)
    if qid:
        return redirect(url_for('help_desk.detail', query_id=qid))
    return redirect(url_for('help_desk.index'))


@app.route('/api/help-desk/stats', methods=['GET'])
@login_required
def help_desk_stats():
    """Get help desk statistics for dashboard"""
    try:
        from app.models import HelpDeskQuery
        from datetime import datetime, timedelta
        
        # Check if resolved_at column exists
        try:
            from sqlalchemy import inspect
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Basic counts - use raw SQL to avoid resolved_at column if it doesn't exist
        if has_resolved_at:
            total_queries = HelpDeskQuery.query.count()
            resolved_queries = HelpDeskQuery.query.filter_by(status='Resolved').count()
            not_started = HelpDeskQuery.query.filter_by(status='Not started').count()
            looking_into = HelpDeskQuery.query.filter_by(status='Looking into it').count()
            
            # Queries over time (last 30 days)
            thirty_days_ago = datetime.utcnow() - timedelta(days=30)
            seven_days_ago = datetime.utcnow() - timedelta(days=7)
            
            queries_last_30_days = HelpDeskQuery.query.filter(
                HelpDeskQuery.timestamp >= thirty_days_ago
            ).count()
            
            queries_last_7_days = HelpDeskQuery.query.filter(
                HelpDeskQuery.timestamp >= seven_days_ago
            ).count()
            
            # Daily breakdown for last 7 days
            daily_data = []
            for i in range(6, -1, -1):
                day_start = (datetime.utcnow() - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
                day_end = day_start + timedelta(days=1)
                count = HelpDeskQuery.query.filter(
                    HelpDeskQuery.timestamp >= day_start,
                    HelpDeskQuery.timestamp < day_end
                ).count()
                daily_data.append({
                    'date': day_start.strftime('%Y-%m-%d'),
                    'count': count
                })
            
            # Average resolution time (for resolved queries)
            avg_resolution_hours = 0
            try:
                resolved_with_times = HelpDeskQuery.query.filter(
                    HelpDeskQuery.status == 'Resolved',
                    HelpDeskQuery.resolved_at.isnot(None)
                ).all()
                
                if resolved_with_times:
                    total_hours = 0
                    count = 0
                    for query in resolved_with_times:
                        if query.resolved_at and query.timestamp:
                            delta = query.resolved_at - query.timestamp
                            total_hours += delta.total_seconds() / 3600
                            count += 1
                    if count > 0:
                        avg_resolution_hours = total_hours / count
            except Exception:
                avg_resolution_hours = 0
        else:
            # Use raw SQL queries to avoid resolved_at column
            total_queries = db.session.execute(text("SELECT COUNT(*) FROM helpdesk_queries")).scalar()
            resolved_queries = db.session.execute(text("SELECT COUNT(*) FROM helpdesk_queries WHERE status = 'Resolved'")).scalar()
            not_started = db.session.execute(text("SELECT COUNT(*) FROM helpdesk_queries WHERE status = 'Not started'")).scalar()
            looking_into = db.session.execute(text("SELECT COUNT(*) FROM helpdesk_queries WHERE status = 'Looking into it'")).scalar()
            
            # Queries over time (last 30 days)
            thirty_days_ago = datetime.utcnow() - timedelta(days=30)
            seven_days_ago = datetime.utcnow() - timedelta(days=7)
            
            queries_last_30_days = db.session.execute(
                text("SELECT COUNT(*) FROM helpdesk_queries WHERE timestamp >= :date"),
                {'date': thirty_days_ago}
            ).scalar()
            
            queries_last_7_days = db.session.execute(
                text("SELECT COUNT(*) FROM helpdesk_queries WHERE timestamp >= :date"),
                {'date': seven_days_ago}
            ).scalar()
            
            # Daily breakdown for last 7 days
            daily_data = []
            for i in range(6, -1, -1):
                day_start = (datetime.utcnow() - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
                day_end = day_start + timedelta(days=1)
                count = db.session.execute(
                    text("SELECT COUNT(*) FROM helpdesk_queries WHERE timestamp >= :start AND timestamp < :end"),
                    {'start': day_start, 'end': day_end}
                ).scalar()
                daily_data.append({
                    'date': day_start.strftime('%Y-%m-%d'),
                    'count': count
                })
            
            avg_resolution_hours = 0
        
        unresolved_queries = total_queries - resolved_queries
        success_rate = (resolved_queries / total_queries * 100) if total_queries > 0 else 0
        
        return jsonify({
            'success': True,
            'stats': {
                'total_queries': total_queries,
                'resolved_queries': resolved_queries,
                'unresolved_queries': unresolved_queries,
                'success_rate': round(success_rate, 2),
                'status_breakdown': {
                    'not_started': not_started,
                    'looking_into_it': looking_into,
                    'resolved': resolved_queries
                },
                'queries_last_30_days': queries_last_30_days,
                'queries_last_7_days': queries_last_7_days,
                'daily_data': daily_data,
                'avg_resolution_hours': round(avg_resolution_hours, 2)
            }
        }), 200
        
    except Exception as e:
        app.logger.error(f"Error fetching help desk stats: {str(e)}")
        import traceback
        app.logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': f'Error fetching stats: {str(e)}'
        }), 500


@app.route('/api/search-users', methods=['GET'])
@login_required
def search_users():
    """Search registered users for assignment purposes"""
    search_term = request.args.get('q', '').strip()
    
    if not search_term or len(search_term) < 2:
        return jsonify({'success': False, 'error': 'Search term must be at least 2 characters'}), 400
    
    try:
        from app.models import User
        # Search by username, email, firstname, or lastname
        search_pattern = f"%{search_term}%"
        users = User.query.filter(
            (User.username.ilike(search_pattern)) |
            (User.email.ilike(search_pattern)) |
            (User.firstname.ilike(search_pattern)) |
            (User.lastname.ilike(search_pattern))
        ).limit(20).all()
        
        users_data = []
        for user in users:
            users_data.append({
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'firstname': user.firstname,
                'lastname': user.lastname,
                'userRole': user.userRole
            })
        
        return jsonify({
            'success': True,
            'users': users_data,
            'count': len(users_data)
        }), 200
        
    except Exception as e:
        app.logger.error(f"Error searching users: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error searching users: {str(e)}'
        }), 500


@app.route('/help-desk/update-status/<int:query_id>', methods=['POST'])
@login_required
def update_help_desk_status(query_id):
    """Update status of a help desk query"""
    try:
        from app.models import HelpDeskQuery
        from datetime import datetime
        from sqlalchemy import inspect
        data = request.json
        new_status = data.get('status')
        
        if not new_status:
            return jsonify({'success': False, 'message': 'Status required'}), 400
            
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
            
        if has_resolved_at:
            query = HelpDeskQuery.query.get_or_404(query_id)
            query.status = new_status
            if new_status == 'Resolved':
                query.resolved_at = datetime.utcnow()
            else:
                query.resolved_at = None
            db.session.commit()
        else:
            # Raw SQL fallback
            if new_status == 'Resolved':
                db.session.execute(
                    text("UPDATE helpdesk_queries SET status = :status WHERE id = :id"),
                    {'status': new_status, 'id': query_id}
                )
            else:
                db.session.execute(
                    text("UPDATE helpdesk_queries SET status = :status WHERE id = :id"),
                    {'status': new_status, 'id': query_id}
                )
            db.session.commit()
            
        return jsonify({'success': True, 'message': 'Status updated'})
    except Exception as e:
        app.logger.error(f"Error updating status: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/help-desk/assign-query', methods=['POST'])
@login_required
def assign_help_desk_query():
    """Assign a query to a user"""
    try:
        data = request.json
        query_id = data.get('queryId')
        user_email = data.get('userEmail')
        
        if not query_id or not user_email:
            return jsonify({'success': False, 'message': 'Missing data'}), 400
            
        # Get user ID
        from app.models import User
        user = User.query.filter_by(email=user_email).first()
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
            
        # Check if already assigned
        exists = db.session.execute(
            text("SELECT 1 FROM query_assignees WHERE query_id = :qid AND user_id = :uid"),
            {'qid': query_id, 'uid': user.id}
        ).scalar()
        
        if exists:
            return jsonify({'success': True, 'message': 'Already assigned'})
            
        # Assign
        db.session.execute(
            text("INSERT INTO query_assignees (query_id, user_id) VALUES (:qid, :uid)"),
            {'qid': query_id, 'uid': user.id}
        )
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Assigned successfully'})
    except Exception as e:
        app.logger.error(f"Error assigning query: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/help-desk/query/<int:query_id>/assignments', methods=['PUT'])
@login_required
def update_help_desk_assignments(query_id):
    """Update assignments for a help desk query (Multi-User)"""
    try:
        data = request.json
        user_ids = data.get('user_ids', [])
        
        # Clear existing assignments
        db.session.execute(
            text("DELETE FROM query_assignees WHERE query_id = :qid"),
            {'qid': query_id}
        )
        
        # Add new assignments
        if user_ids:
            for uid in user_ids:
                db.session.execute(
                    text("INSERT INTO query_assignees (query_id, user_id) VALUES (:qid, :uid)"),
                    {'qid': query_id, 'uid': uid}
                )
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Assignments updated'})
    except Exception as e:
        app.logger.error(f"Error updating assignments: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/search-library-users', methods=['GET'])
@login_required
def search_library_users():
    """Search library users by email, first_name, and/or last_name"""
    email = request.args.get('email', '').strip()
    first_name = request.args.get('first_name', '').strip()
    last_name = request.args.get('last_name', '').strip()
    
    # At least one search parameter is required
    if not email and not first_name and not last_name:
        return jsonify({'success': False, 'error': 'At least one search parameter (email, first_name, or last_name) is required'}), 400
    
    conn = None
    cursor = None
    try:
        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        
        # Build dynamic query based on provided parameters
        conditions = []
        params = []
        
        if email:
            conditions.append("email = %s")
            params.append(email)
        
        if first_name:
            conditions.append("first_name LIKE %s")
            params.append(f"%{first_name}%")
        
        if last_name:
            conditions.append("last_name LIKE %s")
            params.append(f"%{last_name}%")
        
        # Construct the query
        query = f"SELECT * FROM users WHERE {' AND '.join(conditions)}"
        cursor.execute(query, tuple(params))
        results = cursor.fetchall()
        
        # Convert results to list of dictionaries and fetch books for each user
        users = []
        for row in results:
            # Convert any datetime objects to strings
            user_dict = {}
            for key, value in row.items():
                if hasattr(value, 'isoformat'):  # datetime objects
                    user_dict[key] = value.isoformat()
                else:
                    user_dict[key] = value
            
            # Fetch books for this user through orders
            user_id = row.get('id')
            books = []
            if user_id:
                try:
                    # Debug: Log user_id and type
                    app.logger.info(f"Fetching books for user_id: {user_id} (type: {type(user_id)})")
                    
                    # Ensure user_id is the correct type (int)
                    user_id_param = int(user_id) if user_id else None
                    if user_id_param:
                        # Query to get books through orders: users → orders → book_order → books
                        books_query = """
                            SELECT 
                                b.id, b.author, b.title, b.price,
                                COALESCE(b.is_active, 1) AS is_active,
                                bo.quantity, bo.total_cost,
                                o.id as order_id, o.status, o.payment_method, o.total_amount, 
                                o.created_at as order_created_at, o.updated_at as order_updated_at, o.phone
                            FROM orders o
                            INNER JOIN book_order bo ON o.id = bo.order_id
                            INNER JOIN books b ON bo.book_id = b.id
                            WHERE o.user_id = %s
                            ORDER BY o.created_at DESC, b.title ASC
                        """
                        cursor.execute(books_query, (user_id_param,))
                        books_results = cursor.fetchall()
                        
                        app.logger.info(f"Query executed for user_id {user_id_param}, found {len(books_results)} book records")
                        
                        # Convert books results to list of dictionaries
                        for book_row in books_results:
                            book_dict = {}
                            for key, value in book_row.items():
                                if hasattr(value, 'isoformat'):  # datetime objects
                                    book_dict[key] = value.isoformat()
                                else:
                                    book_dict[key] = value
                            books.append(book_dict)
                    else:
                        app.logger.warning(f"Invalid user_id: {user_id}")
                except Exception as e:
                    error_msg = f"Error fetching books for user {user_id}: {str(e)}"
                    app.logger.error(error_msg)
                    import traceback
                    app.logger.error(traceback.format_exc())
                    # Continue without books if query fails
            
            user_dict['books'] = books
            users.append(user_dict)
        
        return jsonify({
            'success': True,
            'users': users,
            'count': len(users)
        }), 200
        
    except Exception as e:
        app.logger.error(f"Error searching library users: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error searching library users: {str(e)}'
        }), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


# ===== Microsoft Graph API OAuth Routes =====

@app.route('/microsoft-login')
@login_required
def microsoft_login():
    """Initiate Microsoft OAuth login flow"""
    try:
        msal_app = build_msal_app()
        if not msal_app:
            flash('Microsoft Graph API is not configured. Please check your .env file.', 'error')
            return redirect(url_for('help_desk'))
        
        config = get_microsoft_config()
        auth_url = msal_app.get_authorization_request_url(
            config['scopes'],
            redirect_uri=config['redirect_uri']
        )
        return redirect(auth_url)
    except Exception as e:
        print(f"Error initiating Microsoft login: {str(e)}")
        flash(f'Error connecting to Microsoft: {str(e)}', 'error')
        return redirect(url_for('help_desk'))


@app.route('/microsoft-callback')
@login_required
def microsoft_callback():
    """Handle Microsoft OAuth callback"""
    try:
        error = request.args.get('error')
        if error:
            error_description = request.args.get('error_description', 'Unknown error')
            flash(f'Microsoft authentication failed: {error_description}', 'error')
            return redirect(url_for('help_desk'))
        
        code = request.args.get('code')
        if not code:
            flash('No authorization code received', 'error')
            return redirect(url_for('help_desk'))
        
        msal_app = build_msal_app()
        if not msal_app:
            flash('Microsoft Graph API is not configured.', 'error')
            return redirect(url_for('help_desk'))
        
        config = get_microsoft_config()
        result = msal_app.acquire_token_by_authorization_code(
            code,
            scopes=config['scopes'],
            redirect_uri=config['redirect_uri']
        )
        
        if 'access_token' in result:
            session['microsoft_token'] = result
            flash('Successfully connected to Microsoft Outlook!', 'success')
        else:
            error_msg = result.get('error_description', result.get('error', 'Unknown error'))
            flash(f'Failed to get access token: {error_msg}', 'error')
        
        return redirect(url_for('help_desk'))
    except Exception as e:
        print(f"Error in Microsoft callback: {str(e)}")
        flash(f'Error during authentication: {str(e)}', 'error')
        return redirect(url_for('help_desk'))


@app.route('/microsoft-logout')
@login_required
def microsoft_logout():
    """Logout from Microsoft (clear session token)"""
    session.pop('microsoft_token', None)
    flash('Disconnected from Microsoft Outlook', 'info')
    return redirect(url_for('help_desk'))


@app.route('/api/microsoft-auth-status')
@login_required
def microsoft_auth_status():
    """Check Microsoft authentication status"""
    token = get_graph_api_token()
    return jsonify({
        'authenticated': token is not None,
        'has_token': get_token_from_cache() is not None
    })


# Notification endpoints
@app.route('/api/notifications', methods=['GET'])
@login_required
def get_notifications():
    """Get current user's notifications"""
    try:
        from app.models import Notification
        notifications = db.session.query(Notification).filter_by(user_id=current_user.id)\
            .order_by(Notification.created_at.desc())\
            .limit(50).all()
        
        notifications_data = []
        for notif in notifications:
            pm_project_id = None
            if getattr(notif, 'task_id', None):
                from app.models import TaskA, ColumnA
                t = TaskA.query.get(notif.task_id)
                if t and t.column_id:
                    col = ColumnA.query.get(t.column_id)
                    if col:
                        pm_project_id = col.project_id
            notifications_data.append({
                'id': notif.id,
                'query_id': notif.query_id,
                'task_id': getattr(notif, 'task_id', None),
                'pm_project_id': pm_project_id,
                'meeting_note_id': getattr(notif, 'meeting_note_id', None),
                'action_item_id': getattr(notif, 'action_item_id', None),
                'stakeholder_lead_id': getattr(notif, 'stakeholder_lead_id', None),
                'message': notif.message,
                'notification_type': notif.notification_type,
                'read': notif.read,
                'created_at': notif.created_at.isoformat() if notif.created_at else None
            })
        
        return jsonify({
            'success': True,
            'notifications': notifications_data
        }), 200
    except Exception as e:
        app.logger.error(f"Error fetching notifications: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/help-desk/notifications', methods=['GET'])
@login_required
def get_helpdesk_notifications():
    """Get query-related notifications for current user (help desk specific)"""
    try:
        from app.models import Notification
        # Get only notifications that have a query_id (help desk related)
        notifications = db.session.query(Notification).filter(
            Notification.user_id == current_user.id,
            Notification.query_id.isnot(None)
        ).order_by(Notification.created_at.desc()).limit(50).all()
        
        notifications_data = []
        for notif in notifications:
            notifications_data.append({
                'id': notif.id,
                'query_id': notif.query_id,
                'message': notif.message,
                'notification_type': notif.notification_type,
                'read': notif.read,
                'created_at': notif.created_at.isoformat() if notif.created_at else None
            })
        
        return jsonify({
            'success': True,
            'notifications': notifications_data
        }), 200
    except Exception as e:
        app.logger.error(f"Error fetching help desk notifications: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/notifications/unread-count', methods=['GET'])
@login_required
def get_unread_notification_count():
    """Get count of unread notifications for current user"""
    try:
        from app.models import Notification
        count = db.session.query(Notification).filter_by(user_id=current_user.id, read=False).count()
        return jsonify({
            'success': True,
            'count': count
        }), 200
    except Exception as e:
        app.logger.error(f"Error getting unread count: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/notifications/<int:notification_id>/read', methods=['PATCH'])
@login_required
def mark_notification_read(notification_id):
    """Mark a notification as read"""
    try:
        from app.models import Notification
        notification = db.session.get(Notification, notification_id)
        if not notification:
            return jsonify({'error': 'Notification not found'}), 404
        
        # Ensure user owns this notification
        if notification.user_id != current_user.id:
            return jsonify({'error': 'Unauthorized'}), 403
        
        notification.read = True
        db.session.commit()
        
        return jsonify({'success': True}), 200
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error marking notification as read: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# Update help desk query status (Admin only)
@app.route('/help-desk/<int:query_id>/status', methods=['PATCH'])
@login_required
def update_helpdesk_status(query_id):
    # Admins and users with Admin Queries Access privilege can update status
    if not can_view_all_queries():
        return jsonify({'error': 'Unauthorized'}), 403
    from app.models import HelpDeskQuery, Notification, User
    from sqlalchemy import text, inspect
    
    try:
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        data = request.get_json() or {}
        new_status = data.get('status')
        allowed = ['Not started', 'Looking into it', 'Resolved']
        if new_status not in allowed:
            return jsonify({'error': 'Invalid status'}), 400
        
        if has_resolved_at:
            # Use ORM when column exists
            q = HelpDeskQuery.query.get_or_404(query_id)
            old_status = q.status
            q.status = new_status
            
            # If status changed to "Resolved", set resolved_at
            if new_status == 'Resolved' and old_status != 'Resolved':
                if hasattr(q, 'resolved_at'):
                    try:
                        q.resolved_at = datetime.utcnow()
                    except Exception:
                        pass
        else:
            # Use raw SQL when column doesn't exist
            # First get the old status and other info
            result = db.session.execute(
                text("SELECT status, query_title, created_by FROM helpdesk_queries WHERE id = :query_id"),
                {'query_id': query_id}
            ).fetchone()
            
            if not result:
                return jsonify({'error': 'Query not found'}), 404
            
            old_status = result.status
            query_title = result.query_title
            created_by = result.created_by
            
            # Update status using raw SQL
            db.session.execute(
                text("UPDATE helpdesk_queries SET status = :status WHERE id = :query_id"),
                {'status': new_status, 'query_id': query_id}
            )
            
            # Create a minimal query object for notification purposes
            class QueryObj:
                def __init__(self):
                    self.id = query_id
                    self.query_title = query_title
                    self.created_by = created_by
            
            q = QueryObj()
        
        # Get all assignees for this query
        assignees = []
        try:
            if has_resolved_at and hasattr(q, 'assignees'):
                # Use ORM relationship
                assignees = q.assignees
            else:
                # Use raw SQL to get assignees
                assignees_result = db.session.execute(
                    text("SELECT u.id, u.username, u.email FROM query_assignees qa JOIN user u ON qa.user_id = u.id WHERE qa.query_id = :query_id"),
                    {'query_id': query_id}
                ).fetchall()
                assignees = [User.query.get(row.id) for row in assignees_result if row.id]
        except Exception as e:
            app.logger.warning(f"Could not load assignees: {str(e)}")
            assignees = []
        
        # Create notifications for all assignees when status changes
        if old_status != new_status:
            status_messages = {
                'Not started': 'has been set to Not started',
                'Looking into it': 'is now being looked into',
                'Resolved': 'has been resolved'
            }
            message_template = status_messages.get(new_status, f'status has been changed to {new_status}')
            
            for assignee in assignees:
                if assignee:
                    try:
                        notification = Notification(
                            user_id=assignee.id,
                            query_id=q.id,
                            message=f"Query '{q.query_title}' {message_template}.",
                            notification_type='status_change' if new_status != 'Resolved' else 'resolution'
                        )
                        db.session.add(notification)
                    except Exception as e:
                        app.logger.warning(f"Could not create notification for assignee: {str(e)}")
        
        # If status changed to "Resolved", also create notification for query creator
        if new_status == 'Resolved' and old_status != 'Resolved':
            # Find the creator user (created_by might be username or email)
            creator_user = None
            if q.created_by:
                # Try to find by username first
                creator_user = User.query.filter_by(username=q.created_by).first()
                # If not found, try by email
                if not creator_user:
                    creator_user = User.query.filter_by(email=q.created_by).first()
            
            # Only notify creator if they're not already an assignee (to avoid duplicate notifications)
            if creator_user and creator_user not in assignees:
                try:
                    notification = Notification(
                        user_id=creator_user.id,
                        query_id=q.id,
                        message=f"Your query '{q.query_title}' has been resolved.",
                        notification_type='resolution'
                    )
                    db.session.add(notification)
                except Exception as e:
                    app.logger.warning(f"Could not create notification for creator: {str(e)}")
        
        db.session.commit()
        return jsonify({'message': 'Status updated', 'status': new_status})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating query status: {str(e)}")
        return jsonify({'error': 'Failed to update status'}), 500


# Delete help desk query (Admin only)
@app.route('/help-desk/<int:query_id>', methods=['DELETE'])
@login_required
def delete_helpdesk_query(query_id):
    if current_user.userRole != 'Admin':
        return jsonify({'error': 'Unauthorized'}), 403
    from app.models import HelpDeskQuery
    from sqlalchemy import text, inspect
    
    try:
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Get image_path before deleting (needed for file removal)
        image_path = None
        if has_resolved_at:
            # Use ORM when column exists
            q = HelpDeskQuery.query.get_or_404(query_id)
            image_path = q.image_path
            db.session.delete(q)
        else:
            # Use raw SQL when column doesn't exist
            # First get the image_path
            result = db.session.execute(
                text("SELECT image_path FROM helpdesk_queries WHERE id = :query_id"),
                {'query_id': query_id}
            ).fetchone()
            
            if not result:
                return jsonify({'error': 'Query not found'}), 404
            
            image_path = result.image_path if result else None
            
            # Delete using raw SQL
            db.session.execute(
                text("DELETE FROM helpdesk_queries WHERE id = :query_id"),
                {'query_id': query_id}
            )
        
        # Optionally remove file
        try:
            if image_path and image_path.startswith('/static/uploads/helpdesk/'):
                fs_path = image_path.lstrip('/')
                if os.path.exists(fs_path):
                    os.remove(fs_path)
        except Exception:
            pass
        
        db.session.commit()
        return jsonify({'message': 'Deleted'})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting query: {str(e)}")
        return jsonify({'error': 'Failed to delete query'}), 500


# Get help desk query details
@app.route('/api/help-desk/query/<int:query_id>', methods=['GET'])
@login_required
def get_helpdesk_query_details(query_id):
    """Get detailed information about a specific help desk query"""
    from app.models import HelpDeskQuery, User
    from sqlalchemy import text, inspect
    
    try:
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Use raw SQL to avoid loading non-existent columns
        if has_resolved_at:
            # Use ORM when column exists
            q = HelpDeskQuery.query.get_or_404(query_id)
            query_id_val = q.id
            created_by_val = q.created_by
            resolved_at_val = q.resolved_at if hasattr(q, 'resolved_at') else None
            
            # Check permissions: admins and privilege holders can see all, others can only see their own or anonymous queries
            if not can_view_all_queries():
                if created_by_val != current_user.username and created_by_val != 'anonymous':
                    return jsonify({'error': 'Unauthorized'}), 403
        else:
            # Use raw SQL when column doesn't exist
            result = db.session.execute(
                text("SELECT id, query_title, query_description, query_type, created_by, timestamp, image_path, status FROM helpdesk_queries WHERE id = :query_id"),
                {'query_id': query_id}
            ).fetchone()
            
            if not result:
                return jsonify({'error': 'Query not found'}), 404
            
            row_dict = dict(result._mapping)
            query_id_val = row_dict['id']
            created_by_val = row_dict['created_by']
            resolved_at_val = None
            
            # Check permissions: admins and privilege holders can see all, others can only see their own or anonymous queries
            if not can_view_all_queries():
                if created_by_val != current_user.username and created_by_val != 'anonymous':
                    return jsonify({'error': 'Unauthorized'}), 403
            
            # Build response from raw SQL result
            assignees_list = []
            try:
                assignees_result = db.session.execute(
                    text("SELECT u.id, u.username, u.email, u.firstname, u.lastname, u.userRole FROM query_assignees qa JOIN user u ON qa.user_id = u.id WHERE qa.query_id = :query_id"),
                    {'query_id': query_id_val}
                ).fetchall()
                assignees_list = [dict(row._mapping) for row in assignees_result]
            except Exception as e:
                app.logger.warning(f"Could not load assignees: {str(e)}")
                assignees_list = []
            
            # Convert timestamp string to datetime if needed
            timestamp_val = row_dict.get('timestamp')
            if isinstance(timestamp_val, str):
                try:
                    timestamp_val = datetime.fromisoformat(timestamp_val.replace('Z', '+00:00'))
                except Exception:
                    try:
                        timestamp_val = datetime.strptime(timestamp_val, '%Y-%m-%d %H:%M:%S')
                    except Exception:
                        timestamp_val = None
            
            response_data = {
                'id': query_id_val,
                'query_title': row_dict.get('query_title', ''),
                'query_description': row_dict.get('query_description', ''),
                'query_type': row_dict.get('query_type', ''),
                'created_by': created_by_val,
                'status': row_dict.get('status') or 'Not started',
                'timestamp': timestamp_val.isoformat() if timestamp_val else None,
                'image_path': row_dict.get('image_path'),
                'assignees': assignees_list,
                'resolved_at': None
            }
            
            return jsonify(response_data), 200
        
        # Check permissions: admins can see all, others can only see their own or anonymous queries
        if current_user.userRole != 'Admin':
            if created_by_val != current_user.username and created_by_val != 'anonymous':
                return jsonify({'error': 'Unauthorized'}), 403
        
        # Load assignees with their roles
        assignees_list = []
        try:
            # Try using ORM first
            if hasattr(q, 'assignees') and q.assignees:
                assignees_list = [{
                    'id': u.id,
                    'username': u.username,
                    'email': u.email,
                    'firstname': u.firstname,
                    'lastname': u.lastname,
                    'userRole': u.userRole
                } for u in q.assignees]
            else:
                # Fallback to raw SQL
                assignees_result = db.session.execute(
                    text("SELECT u.id, u.username, u.email, u.firstname, u.lastname, u.userRole FROM query_assignees qa JOIN user u ON qa.user_id = u.id WHERE qa.query_id = :query_id"),
                    {'query_id': query_id_val}
                ).fetchall()
                assignees_list = [dict(row._mapping) for row in assignees_result]
        except Exception as e:
            app.logger.warning(f"Could not load assignees: {str(e)}")
            assignees_list = []
        
        # Build response
        response_data = {
            'id': q.id,
            'query_title': q.query_title,
            'query_description': q.query_description,
            'query_type': q.query_type,
            'created_by': q.created_by,
            'status': q.status or 'Not started',
            'timestamp': q.timestamp.isoformat() if q.timestamp else None,
            'image_path': q.image_path,
            'assignees': assignees_list,
            'resolved_at': resolved_at_val.isoformat() if resolved_at_val else None
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        app.logger.error(f"Error fetching query details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to fetch query details'}), 500


# Update query assignments
@app.route('/api/help-desk/query/<int:query_id>/assignments', methods=['PUT'])
@login_required
def update_query_assignments(query_id):
    """Update assignees for a help desk query"""
    from app.models import HelpDeskQuery, User, Notification
    from sqlalchemy import text, inspect
    
    try:
        # Check permissions: admins and users with Admin Queries Access privilege can assign queries
        if not can_view_all_queries():
            return jsonify({'error': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        user_ids = data.get('user_ids', [])
        
        if not isinstance(user_ids, list):
            return jsonify({'error': 'user_ids must be a list'}), 400
        
        # Check if resolved_at column exists
        try:
            inspector = inspect(db.engine)
            columns = [col['name'] for col in inspector.get_columns('helpdesk_queries')]
            has_resolved_at = 'resolved_at' in columns
        except Exception:
            has_resolved_at = False
        
        # Get query
        if has_resolved_at:
            q = HelpDeskQuery.query.get_or_404(query_id)
        else:
            # Use raw SQL to check if query exists
            result = db.session.execute(
                text("SELECT id, query_title FROM helpdesk_queries WHERE id = :query_id"),
                {'query_id': query_id}
            ).fetchone()
            if not result:
                return jsonify({'error': 'Query not found'}), 404
            # Create minimal query object
            class QueryObj:
                def __init__(self):
                    self.id = result.id
                    self.query_title = result.query_title
            q = QueryObj()
        
        # Remove all existing assignments
        try:
            db.session.execute(
                text("DELETE FROM query_assignees WHERE query_id = :query_id"),
                {'query_id': query_id}
            )
        except Exception as e:
            app.logger.warning(f"Could not delete existing assignments: {str(e)}")
        
        # Add new assignments
        assigned_users = []
        for user_id in user_ids:
            try:
                user = User.query.get(user_id)
                if user:
                    # Insert assignment
                    try:
                        db.session.execute(
                            text("INSERT INTO query_assignees (query_id, user_id) VALUES (:query_id, :user_id)"),
                            {'query_id': query_id, 'user_id': user_id}
                        )
                    except Exception:
                        # Try using ORM if raw SQL fails
                        if has_resolved_at:
                            q_obj = HelpDeskQuery.query.get(query_id)
                            if q_obj and user not in q_obj.assignees:
                                q_obj.assignees.append(user)
                    
                    assigned_users.append({
                        'id': user.id,
                        'username': user.username,
                        'email': user.email,
                        'firstname': user.firstname,
                        'lastname': user.lastname,
                        'userRole': user.userRole
                    })
                    
                    # Create notification for assigned user
                    try:
                        notification = Notification(
                            user_id=user.id,
                            query_id=query_id,
                            message=f"You have been assigned to query: '{q.query_title if hasattr(q, 'query_title') else 'Query #' + str(query_id)}'",
                            notification_type='assignment'
                        )
                        db.session.add(notification)
                    except Exception as e:
                        app.logger.warning(f"Could not create notification: {str(e)}")
            except Exception as e:
                app.logger.warning(f"Could not assign user {user_id}: {str(e)}")
                continue
        
        db.session.commit()
        
        return jsonify({
            'message': 'Assignments updated successfully',
            'assigned_users': assigned_users
        }), 200
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating query assignments: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Failed to update assignments'}), 500


# Update champion's school details
# Use a flexible key (library_school_id or asl_school_id) to locate the school dict
@app.route("/api/champions/<int:champ_id>/schools/<school_key>", methods=["PUT"], endpoint="update_school_by_key")
def update_school_by_key(champ_id, school_key):
    champ = ChampionSchool.query.get_or_404(champ_id)
    schools = champ.get_schools() or []

    # Normalize types to string for comparison
    key = str(school_key)
    target = None
    for s in schools:
        if str(s.get("library_school_id")) == key or str(s.get("asl_school_id")) == key:
            target = s
            break

    if not target:
        return jsonify({"error": "School not found for this champion"}), 404

    data = request.get_json() or {}
    if "school_name" in data:
        target["school_name"] = data["school_name"]
    if "asl_school_id" in data:
        target["asl_school_id"] = data["asl_school_id"]
    if "library_school_id" in data:
        target["library_school_id"] = data["library_school_id"]

    champ.set_schools(schools)
    db.session.commit()
    return jsonify(champ.to_dict())
@app.route('/api/lesotho-institutions', methods=['GET'])
def lesotho_institutions():
    conn = None
    try:
        conn = get_direct_library_conn()
        cursor = conn.cursor(pymysql.cursors.DictCursor)

        # Query: Get institutions in Lesotho with their user counts
        query = """
            SELECT 
                i.id AS institution_id,
                i.name AS institution_name,
                COUNT(u.id) AS total_users
            FROM institutions i
            LEFT JOIN users u ON u.id = i.id
            WHERE i.country = %s
            GROUP BY i.id, i.name
        """
        cursor.execute(query, ("Lesotho",))
        institutions = cursor.fetchall()

        # Calculate totals
        total_institutions = len(institutions)
        total_users = sum(inst["total_users"] for inst in institutions)

        return jsonify({
            "country": "Lesotho",
            "total_institutions": total_institutions,
            "total_users": total_users,
            "institutions": institutions
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if conn:
            conn.close()




# mobile app routes

@app.route('/api/mobile/login', methods=['POST'])
def mobile_login():
    """
    Mobile login endpoint that returns JSON responses
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'message': 'No data provided'}), 400
            
        username = data.get('username', '').strip()
        password = data.get('password', '').strip()
        
        if not username or not password:
            return jsonify({'success': False, 'message': 'Username and password are required'}), 400
        
        # Find user by username
        user = db.session.scalar(sa.select(User).where(User.username == username))
        
        if user is None or not user.check_password(password):
            return jsonify({'success': False, 'message': 'Invalid username or password'}), 401
        
        # Log the user in
        login_user(user, remember=False)
        
        # Return user data
        user_data = {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'firstname': user.firstname,
            'lastname': user.lastname,
            'userRole': user.userRole,
            'department': user.department,
            'province': user.province,
            'privileges': user.get_privileges(),
        }
        
        return jsonify({
            'success': True, 
            'message': 'Login successful',
            'user': user_data,
            'privileges': user.get_privileges(),
        }), 200
        
    except Exception as e:
        print(f"Mobile login error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500


@app.route('/api/mobile/logout', methods=['POST'])
def mobile_logout():
    """
    Mobile logout endpoint
    """
    try:
        logout_user()
        return jsonify({'success': True, 'message': 'Logout successful'}), 200
    except Exception as e:
        print(f"Mobile logout error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500


@app.route('/api/mobile/user', methods=['GET'])
@login_required
def mobile_get_user():
    """
    Get current user data for mobile
    """
    try:
        user_data = {
            'id': current_user.id,
            'username': current_user.username,
            'email': current_user.email,
            'firstname': current_user.firstname,
            'lastname': current_user.lastname,
            'userRole': current_user.userRole,
            'department': current_user.department,
            'province': current_user.province,
            'privileges': current_user.get_privileges(),
        }
        
        return jsonify({
            'success': True,
            'user': user_data,
            'privileges': current_user.get_privileges(),
        }), 200
        
    except Exception as e:
        print(f"Get user error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500


@app.route('/api/mobile/profile', methods=['GET'])
@login_required
def mobile_user_profile():
    """
    Get detailed user profile data for mobile including Brand Ambassador stats
    """
    try:
        user = current_user
        user_data = {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'firstname': user.firstname,
            'lastname': user.lastname,
            'userRole': user.userRole,
            'department': user.department,
            'province': user.province,
            'privileges': user.get_privileges(),
        }
        
        # Brand Ambassador specific data
        if user.userRole == 'Brand Ambassador':
            try:
                from app.models import ChampionSchool
                import json
                from datetime import datetime, date
                
                # Find champion record
                champ = ChampionSchool.query.filter_by(
                    firstname=user.firstname,
                    lastname=user.lastname,
                    province=user.province
                ).first()
                
                champion_data = {
                    'schools': [],
                    'stats': {
                        'school_count': 0,
                        'asl_mtd': 0,
                        'library_mtd': 0
                    }
                }
                
                if champ and champ.schools:
                    try:
                        # Parse schools JSON
                        schools_data = json.loads(champ.schools)
                        
                        # Extract school info for display
                        champion_data['schools'] = [
                            {
                                'id': s.get('asl_school_id', ''),
                                'name': s.get('school_name', 'Unknown School'),
                                'asl_id': s.get('asl_school_id'),
                                'library_id': s.get('library_school_id')
                            }
                            for s in schools_data if s.get('school_name')
                        ]
                        
                        # Calculate stats
                        asl_ids = [s.get('asl_school_id') for s in schools_data if s.get('asl_school_id')]
                        champion_data['stats']['school_count'] = len(asl_ids)
                        
                        # Calculate ASL MTD (month-to-date) student count
                        if asl_ids:
                            today = date.today()
                            start_date = today.replace(day=1)
                            
                            asl_mtd = 0
                            try:
                                conn = get_ruzivo_conn()
                                cursor = conn.cursor()
                                
                                for asl_id in asl_ids:
                                    if asl_id:
                                        query = """
                                            SELECT COUNT(*) AS student_count
                                            FROM vwstudent
                                            WHERE school_id = %s
                                            AND last_login BETWEEN %s AND %s
                                        """
                                        cursor.execute(query, (asl_id, start_date, today))
                                        row = cursor.fetchone()
                                        
                                        if row:
                                            if isinstance(row, dict):
                                                asl_mtd += row.get('student_count', 0)
                                            else:
                                                asl_mtd += row[0] or 0
                                
                                cursor.close()
                                conn.close()
                                champion_data['stats']['asl_mtd'] = asl_mtd
                            except Exception as e:
                                print(f"ASL MTD calculation error: {e}")
                        
                        # Calculate Library MTD using existing API
                        try:
                            library_ids = [s.get('library_school_id') for s in schools_data if s.get('library_school_id')]
                            if library_ids:
                                today = date.today()
                                start_date = today.replace(day=1)
                                
                                library_mtd = 0
                                conn = get_direct_library_conn()
                                if conn:
                                    import pymysql.cursors
                                    cursor = conn.cursor(pymysql.cursors.DictCursor)
                                    
                                    library_ids_int = [int(lid) for lid in library_ids if lid]
                                    if library_ids_int:
                                        in_placeholders = ', '.join(['%s'] * len(library_ids_int))
                                        query = f"""
                                            SELECT COUNT(DISTINCT la.user_id) AS active_users
                                            FROM last_activities la
                                            JOIN institution_user iu ON la.user_id = iu.user_id
                                            WHERE iu.institution_id IN ({in_placeholders})
                                              AND la.created_at BETWEEN %s AND %s
                                        """
                                        
                                        start_dt = datetime.combine(start_date, datetime.min.time())
                                        end_dt = datetime.combine(today, datetime.max.time())
                                        params = library_ids_int + [start_dt, end_dt]
                                        
                                        cursor.execute(query, params)
                                        row = cursor.fetchone()
                                        library_mtd = row["active_users"] if row else 0
                                    
                                    cursor.close()
                                    conn.close()
                                    champion_data['stats']['library_mtd'] = library_mtd
                        except Exception as e:
                            print(f"Library MTD calculation error: {e}")
                            
                    except json.JSONDecodeError as e:
                        print(f"Error parsing schools JSON: {e}")
                
                user_data['champion_data'] = champion_data
                
            except Exception as e:
                print(f"Brand Ambassador data error: {e}")
                import traceback
                traceback.print_exc()
        
        return jsonify({
            'success': True,
            'user': user_data,
            'privileges': user.get_privileges(),
        }), 200
        
    except Exception as e:
        print(f"Profile data error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': 'Server error'}), 500


@app.route('/api/mobile/help-desk', methods=['GET', 'POST'])
@login_required
def mobile_help_desk():
    """
    Mobile help desk endpoint for fetching and creating queries
    """
    from app.models import HelpDeskQuery
    from app.forms import HelpDeskForm
    
    if request.method == 'GET':
        # Fetch help desk queries for mobile
        try:
            # Show newest first; admins and privilege holders see all; others see only their own and anonymous
            if can_view_all_queries():
                queries = HelpDeskQuery.query.order_by(HelpDeskQuery.timestamp.desc()).all()
            else:
                queries = HelpDeskQuery.query.filter(
                    (HelpDeskQuery.created_by == current_user.username) | 
                    (HelpDeskQuery.created_by == 'anonymous')
                ).order_by(HelpDeskQuery.timestamp.desc()).all()
            
            # Convert to mobile-friendly format
            queries_data = []
            for q in queries:
                queries_data.append({
                    'id': q.id,
                    'query_title': q.query_title,
                    'query_description': q.query_description,
                    'query_type': q.query_type,
                    'created_by': q.created_by,
                    'timestamp': q.timestamp.isoformat() if q.timestamp else None,
                    'status': q.status or 'Not started',
                    'image_path': q.image_path
                })
            
            return jsonify({
                'success': True,
                'queries': queries_data,
                'current_user': current_user.username,
                'is_admin': can_view_all_queries()
            }), 200
            
        except Exception as e:
            print(f"Error fetching help desk queries: {e}")
            return jsonify({
                'success': False,
                'message': 'Failed to fetch queries',
                'queries': []
            }), 500
    
    elif request.method == 'POST':
        # Create new help desk query from mobile
        try:
            # Handle both form data and JSON data
            if request.is_json:
                data = request.get_json()
                query_type = data.get('query_type', '')
                query_title = data.get('query_title', '')
                query_description = data.get('query_description', '')
                image_file = None
            else:
                # Handle multipart form data (with potential file upload)
                query_type = request.form.get('query_type', '')
                query_title = request.form.get('query_title', '')
                query_description = request.form.get('query_description', '')
                image_file = request.files.get('image')
            
            # Validate required fields
            if not query_title or not query_description or not query_type:
                return jsonify({
                    'success': False,
                    'message': 'Query title, description, and type are required'
                }), 400
            
            # Handle image upload
            image_path = None
            if image_file and image_file.filename:
                try:
                    filename = secure_filename(f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{image_file.filename}")
                    save_path = os.path.join(app.config['HELP_DESK_UPLOAD_FOLDER'], filename)
                    image_file.save(save_path)
                    image_path = '/' + save_path  # make it web path
                except Exception as e:
                    print(f"Error saving image: {e}")
                    # Continue without image if upload fails
            
            # Determine created_by
            created_by = 'anonymous' if query_type == 'anonymous' else current_user.username
            
            # Create new help desk query
            q = HelpDeskQuery(
                query_title=query_title,
                query_description=query_description,
                query_type=query_type,
                created_by=created_by,
                image_path=image_path
            )
            
            db.session.add(q)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Query submitted successfully',
                'query_id': q.id
            }), 201
            
        except Exception as e:
            db.session.rollback()
            print(f"Error creating help desk query: {e}")
            return jsonify({
                'success': False,
                'message': 'Failed to submit query'
            }), 500


@app.route('/api/mobile/reports', methods=['GET', 'POST'])
@login_required
def mobile_reports():
    """
    Mobile reports endpoint for fetching and creating weekly reports
    """
    from app.models import WeeklyReport
    
    if request.method == 'GET':
        # Fetch reports data for mobile
        try:
            # Get users from same department
            users = User.query.filter_by(department=current_user.department).all()
            
            # Convert to mobile-friendly format
            users_data = []
            for user in users:
                users_data.append({
                    'id': user.id,
                    'username': user.username,
                    'firstname': user.firstname,
                    'lastname': user.lastname,
                    'department': user.department
                })
            
            return jsonify({
                'success': True,
                'users': users_data,
                'current_user': {
                    'id': current_user.id,
                    'username': current_user.username,
                    'department': current_user.department
                }
            }), 200
            
        except Exception as e:
            print(f"Error fetching reports data: {e}")
            return jsonify({
                'success': False,
                'message': 'Failed to fetch reports data',
                'users': []
            }), 500
    
    elif request.method == 'POST':
        # Create new weekly report from mobile
        try:
            data = request.get_json()
            week_start = data.get('week_start', '')
            work_done = data.get('work_done', '')
            work_next = data.get('work_next', '')
            challenges = data.get('challenges', '')
            
            # Validate required fields
            if not week_start or not work_done or not work_next:
                return jsonify({
                    'success': False,
                    'message': 'Week start, work done, and work next are required'
                }), 400
            
            # Validate week start is Monday
            from datetime import datetime
            week_start_date = datetime.strptime(week_start, "%Y-%m-%d").date()
            if week_start_date.weekday() != 0:  # Must be Monday
                return jsonify({
                    'success': False,
                    'message': 'Week start date must be a Monday'
                }), 400
            
            # Create new weekly report
            report = WeeklyReport(
                user_id=current_user.id,
                department=current_user.department,
                week_start=week_start,
                work_done=work_done,
                work_next=work_next,
                challenges=challenges
            )
            
            db.session.add(report)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Report submitted successfully',
                'report_id': report.id
            }), 201
            
        except Exception as e:
            db.session.rollback()
            print(f"Error creating weekly report: {e}")
            return jsonify({
                'success': False,
                'message': 'Failed to submit report'
            }), 500


@app.route('/api/mobile/reports/user/<int:user_id>', methods=['GET'])
@login_required
def mobile_user_reports(user_id):
    """
    Get reports for a specific user (mobile)
    """
    from app.models import WeeklyReport
    
    try:
        # Check if user exists and is in same department
        user = User.query.get_or_404(user_id)
        if user.department != current_user.department:
            return jsonify({
                'success': False,
                'message': 'You can only view reports from your department'
            }), 403
        
        # Get reports for this user
        reports = WeeklyReport.query.filter_by(user_id=user_id) \
                                   .order_by(WeeklyReport.created_at.desc()) \
                                   .all()
        
        # Convert to mobile-friendly format
        reports_data = []
        for report in reports:
            reports_data.append({
                'id': report.id,
                'user_id': report.user_id,
                'department': report.department,
                'week_start': report.week_start,
                'work_done': report.work_done,
                'work_next': report.work_next,
                'challenges': report.challenges,
                'created_at': report.created_at.isoformat() if report.created_at else None
            })
        
        return jsonify({
            'success': True,
            'reports': reports_data,
            'user': {
                'id': user.id,
                'username': user.username
            }
        }), 200
        
    except Exception as e:
        print(f"Error fetching user reports: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to fetch user reports',
            'reports': []
        }), 500


@app.route('/api/mobile/reports/<int:report_id>', methods=['PUT'])
@login_required
def mobile_update_report(report_id):
    """
    Update a weekly report (mobile)
    """
    from app.models import WeeklyReport
    from datetime import datetime, timedelta
    
    try:
        report = WeeklyReport.query.get_or_404(report_id)
        
        # Ensure user owns this report
        if report.user_id != current_user.id:
            return jsonify({
                'success': False,
                'message': 'You are not authorized to edit this report'
            }), 403
        
        # Check if report is older than 7 days
        if datetime.utcnow() > report.created_at + timedelta(days=7):
            return jsonify({
                'success': False,
                'message': 'You can only edit reports within 7 days of submission'
            }), 403
        
        data = request.get_json()
        week_start = data.get('week_start', '')
        work_done = data.get('work_done', '')
        work_next = data.get('work_next', '')
        challenges = data.get('challenges', '')
        
        # Validate required fields
        if not week_start or not work_done or not work_next:
            return jsonify({
                'success': False,
                'message': 'Week start, work done, and work next are required'
            }), 400
        
        # Validate week start is Monday
        week_start_date = datetime.strptime(week_start, "%Y-%m-%d").date()
        if week_start_date.weekday() != 0:  # Must be Monday
            return jsonify({
                'success': False,
                'message': 'Week start date must be a Monday'
            }), 400
        
        # Update report
        report.week_start = week_start
        report.work_done = work_done
        report.work_next = work_next
        report.challenges = challenges
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Report updated successfully'
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error updating report: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to update report'
        }), 500


# ===============================
# SEARCH FUNCTIONALITY ROUTES
# ===============================

@app.route('/api/search', methods=['GET'])
@login_required
def global_search():
    """
    Global search across users, projects, workspaces, reports, and other platform data
    """
    try:
        query = request.args.get('q', '').strip()
        if not query or len(query) < 2:
            return jsonify({
                'success': True,
                'results': [],
                'message': 'Please enter at least 2 characters to search'
            }), 200
        
        search_results = []
        search_term = f"%{query}%"
        
        # 1. Search Users (all authenticated users can search for other users)
        try:
            users = User.query.filter(
                db.or_(
                    User.username.ilike(search_term),
                    User.firstname.ilike(search_term),
                    User.lastname.ilike(search_term),
                    User.email.ilike(search_term),
                    User.department.ilike(search_term)
                )
            ).limit(10).all()
            
            for user in users:
                # Don't include current user in search results
                if user.id != current_user.id:
                    search_results.append({
                        'type': 'user',
                        'id': user.id,
                        'title': f"{user.firstname} {user.lastname}",
                        'subtitle': f"@{user.username} - {user.userRole}",
                        'description': f"{user.department or 'No department'} | {user.email}",
                        'url': url_for('profile', username=user.username),
                        'icon': 'user'
                    })
        except Exception as e:
            print(f"User search error: {e}")
        
        # 2. Search Workspaces
        try:
            workspaces = Workspace.query.filter(
                db.or_(
                    Workspace.name.ilike(search_term),
                    Workspace.description.ilike(search_term)
                )
            ).limit(10).all()
            
            for workspace in workspaces:
                # Show all workspaces but redirect to project planning page
                search_results.append({
                    'type': 'workspace',
                    'id': workspace.id,
                    'title': workspace.name,
                    'subtitle': 'Workspace',
                    'description': workspace.description or 'No description',
                    'url': url_for('projectmanagement'),
                    'icon': 'folder'
                })
        except Exception as e:
            print(f"Workspace search error: {e}")
        
        # 3. Search Projects (ProjectA)
        try:
            projects = ProjectA.query.filter(
                ProjectA.name.ilike(search_term)
            ).limit(10).all()
            
            for project in projects:
                if not user_can_access_project_a(project):
                    continue
                search_results.append({
                    'type': 'project',
                    'id': project.id,
                    'title': project.name,
                    'subtitle': 'Project',
                    'description': getattr(project, 'description', None) or 'Project',
                    'url': url_for('projectmanagement') + f'?project={project.id}',
                    'icon': 'clipboard'
                })
        except Exception as e:
            print(f"Project search error: {e}")
        
        # 4. Search Reports
        try:
            reports = WeeklyReport.query.filter(
                db.or_(
                    WeeklyReport.work_done.ilike(search_term),
                    WeeklyReport.work_next.ilike(search_term),
                    WeeklyReport.department.ilike(search_term)
                )
            ).limit(10).all()
            
            for report in reports:
                # Only show user's own reports or if admin
                if report.user_id == current_user.id or current_user.userRole == 'Admin':
                    search_results.append({
                        'type': 'report',
                        'id': report.id,
                        'title': f"Weekly Report - {report.week_start}",
                        'subtitle': f"Report by {report.department}",
                        'description': report.work_done[:100] + '...' if len(report.work_done) > 100 else report.work_done,
                        'url': url_for('akello_weekly_reports'),
                        'icon': 'file-text'
                    })
        except Exception as e:
            print(f"Report search error: {e}")
        
        # 5. Search Champion Schools (with proper access control)
        try:
            champions = ChampionSchool.query.filter(
                db.or_(
                    ChampionSchool.firstname.ilike(search_term),
                    ChampionSchool.lastname.ilike(search_term),
                    ChampionSchool.schools.ilike(search_term),
                    ChampionSchool.province.ilike(search_term)
                )
            ).limit(10).all()
            
            for champion in champions:
                # All users can search champion schools, but Brand Ambassadors only see their province
                if current_user.userRole == 'Brand Ambassador':
                    # Brand Ambassadors only see schools in their province
                    if current_user.province and champion.province == current_user.province:
                        search_results.append({
                            'type': 'champion',
                            'id': champion.id,
                            'title': f"{champion.firstname} {champion.lastname}",
                            'subtitle': f"Champion - {champion.province}",
                            'description': f"Province: {champion.province}",
                            'url': url_for('all_champion_details'),
                            'icon': 'award'
                        })
                else:
                    # Other roles can see all champion schools
                    search_results.append({
                        'type': 'champion',
                        'id': champion.id,
                        'title': f"{champion.firstname} {champion.lastname}",
                        'subtitle': f"Champion - {champion.province}",
                        'description': f"Province: {champion.province}",
                        'url': url_for('all_champion_details'),
                        'icon': 'award'
                    })
        except Exception as e:
            print(f"Champion school search error: {e}")
        
        # 6. Search Book Allocations (exclude Brand Ambassadors from admin features)
        if current_user.userRole != 'Brand Ambassador':
            try:
                allocations = BookAllocations.query.filter(
                    db.or_(
                        BookAllocations.school_name.ilike(search_term),
                        BookAllocations.school_province.ilike(search_term),
                        BookAllocations.books_allocated.ilike(search_term)
                    )
                ).limit(10).all()
                
                for allocation in allocations:
                    search_results.append({
                        'type': 'allocation',
                        'id': allocation.id,
                        'title': f"Books for {allocation.school_name}",
                        'subtitle': f"Book Allocation - {allocation.school_province}",
                        'description': allocation.books_allocated or 'No books specified',
                        'url': url_for('bookallocations'),
                        'icon': 'book'
                    })
            except Exception as e:
                print(f"Book allocation search error: {e}")
        
        # Sort results by relevance (exact matches first, then partial matches)
        def sort_by_relevance(item):
            title_lower = item['title'].lower()
            query_lower = query.lower()
            if query_lower == title_lower:
                return 0  # Exact match
            elif title_lower.startswith(query_lower):
                return 1  # Starts with query
            elif query_lower in title_lower:
                return 2  # Contains query
            else:
                return 3  # Other matches
        
        search_results.sort(key=sort_by_relevance)
        
        # Limit total results to prevent overwhelming UI
        search_results = search_results[:20]
        
        return jsonify({
            'success': True,
            'results': search_results,
            'total_found': len(search_results),
            'query': query
        }), 200
        
    except Exception as e:
        print(f"Global search error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'Search failed. Please try again.',
            'results': []
        }), 500

@app.route('/search')
@login_required
def search_page():
    """
    Dedicated search results page
    """
    query = request.args.get('q', '').strip()
    return render_template('search_results.html', title='Search Results', query=query)

@app.route('/api/mobile/dashboard', methods=['GET'])
@login_required
def mobile_dashboard_overview():
    """
    Get dashboard overview data for mobile
    """
    try:
        from datetime import datetime, timedelta
        
        # Get current time and ranges
        now = datetime.utcnow()
        today = now.date()
        week_ago = now - timedelta(days=7)
        month_ago = now - timedelta(days=30)
        
        # Get basic stats
        total_users = User.query.filter(User.userRole != 'Admin').count()
        
        # Get recent activity counts (if monitoring is available)
        recent_activities = 0
        active_users_today = 0
        
        try:
            from app.models import UserActivity, ActiveSession
            recent_activities = UserActivity.query.filter(
                UserActivity.timestamp >= week_ago
            ).count()
            
            active_users_today = ActiveSession.query.filter(
                ActiveSession.is_active == True,
                ActiveSession.last_seen >= today
            ).count()
        except:
            # If monitoring tables don't exist, use default values
            pass
        
        # Get performance targets
        latest_target = None
        try:
            from app.models import PerfomanceTargets
            latest_target = PerfomanceTargets.query.order_by(PerfomanceTargets.timestamp.desc()).first()
        except:
            pass
        
        # Get champion data
        champion_count = 0
        try:
            from app.models import ChampionSchool
            champion_count = ChampionSchool.query.count()
        except:
            pass
        
        # Get project data
        active_projects = 0
        total_tasks = 0
        completed_tasks = 0
        
        try:
            from app.models import ProjectA, TaskA
            active_projects = ProjectA.query.count()
            total_tasks = TaskA.query.count()
            completed_tasks = TaskA.query.filter(TaskA.progress == 100).count()
        except:
            pass
        
        dashboard_data = {
            'overview': {
                'total_users': total_users,
                'active_users_today': active_users_today,
                'recent_activities': recent_activities,
                'champion_count': champion_count
            },
            'projects': {
                'active_projects': active_projects,
                'total_tasks': total_tasks,
                'completed_tasks': completed_tasks,
                'completion_rate': round((completed_tasks / total_tasks * 100) if total_tasks > 0 else 0, 1)
            },
            'targets': {
                'has_targets': latest_target is not None,
                'latest_updated': latest_target.timestamp.isoformat() if latest_target else None,
                'updated_by': latest_target.updated_by if latest_target else None
            },
            'system_info': {
                'current_time': now.isoformat(),
                'server_status': 'online',
                'user_role': current_user.userRole,
                'user_department': current_user.department
            }
        }
        
        return jsonify({
            'success': True,
            'data': dashboard_data
        }), 200
        
    except Exception as e:
        print(f"Dashboard overview error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': 'Server error'}), 500


# ==================== MOBILE SYSTEM SETTINGS ====================
# Admin-only mobile-friendly access to a small whitelist of `AppSetting`
# rows. Anything outside this list stays editable only on the web admin to
# preserve audit trails.

_MOBILE_SETTINGS_WHITELIST = [
    {
        'key': 'checkin_reports_timezone',
        'label': 'Weekly report timezone',
        'description': 'Timezone used to compute the Sat-Fri reporting window.',
        'type': 'timezone',
        'default': 'Africa/Harare',
    },
    {
        'key': 'checkin_reports_schedule_time',
        'label': 'Friday report send time',
        'description': 'Local time (HH:MM, 24h) when weekly reports are emailed.',
        'type': 'time',
        'default': '17:00',
    },
]


def _mobile_settings_is_admin():
    role = (getattr(current_user, 'userRole', '') or '').strip().lower()
    if role in ('admin', 'super admin', 'superadmin', 'super-admin', 'manager'):
        return True
    # Fall back to any privilege-based admin check if defined.
    try:
        if hasattr(current_user, 'has_privilege') and (
            current_user.has_privilege('Admin') or
            current_user.has_privilege('Check-in Reports')
        ):
            return True
    except Exception:
        pass
    return False


def _validate_mobile_setting(key, value):
    """Returns (ok, normalized_value, error_msg)."""
    if value is None:
        return False, None, 'Value is required'
    text = str(value).strip()
    if key == 'checkin_reports_timezone':
        try:
            from zoneinfo import ZoneInfo
            ZoneInfo(text)
        except Exception:
            return False, None, f'Unknown timezone: {text}'
        return True, text, None
    if key == 'checkin_reports_schedule_time':
        import re
        if not re.fullmatch(r'^([01]?\d|2[0-3]):([0-5]\d)$', text):
            return False, None, 'Use HH:MM in 24-hour format (e.g. 17:00)'
        return True, text, None
    return True, text, None


@app.route('/api/mobile/settings', methods=['GET'])
@login_required
def api_mobile_settings_get():
    if not _mobile_settings_is_admin():
        return jsonify({'error': 'Unauthorized'}), 403
    settings = []
    for entry in _MOBILE_SETTINGS_WHITELIST:
        value = AppSetting.get_value(entry['key'], entry['default'])
        settings.append({
            'key': entry['key'],
            'label': entry['label'],
            'description': entry['description'],
            'type': entry['type'],
            'value': value,
            'default': entry['default'],
        })
    return jsonify({'data': {'settings': settings}})


@app.route('/api/mobile/settings', methods=['PATCH'])
@login_required
def api_mobile_settings_patch():
    if not _mobile_settings_is_admin():
        return jsonify({'error': 'Unauthorized'}), 403
    payload = request.get_json(silent=True) or {}
    raw_updates = payload.get('settings') or payload
    updates = raw_updates if isinstance(raw_updates, dict) else {}

    allowed_keys = {entry['key'] for entry in _MOBILE_SETTINGS_WHITELIST}
    errors = {}
    accepted = []
    for key, value in updates.items():
        if key not in allowed_keys:
            errors[key] = 'Setting is not editable from mobile'
            continue
        ok, normalized, err = _validate_mobile_setting(key, value)
        if not ok:
            errors[key] = err
            continue
        try:
            AppSetting.set_value(
                key,
                normalized,
                user_id=getattr(current_user, 'id', None),
            )
            accepted.append(key)
        except Exception as e:
            errors[key] = str(e)

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

    if errors and not accepted:
        return jsonify({'error': 'No settings updated', 'errors': errors}), 400

    # Return refreshed list so clients can replace local state in one round-trip.
    settings = []
    for entry in _MOBILE_SETTINGS_WHITELIST:
        value = AppSetting.get_value(entry['key'], entry['default'])
        settings.append({
            'key': entry['key'],
            'label': entry['label'],
            'description': entry['description'],
            'type': entry['type'],
            'value': value,
            'default': entry['default'],
        })
    return jsonify({
        'data': {'settings': settings, 'accepted': accepted},
        'errors': errors or None,
    })


# ==================== GAME SYSTEM ROUTES ====================

@app.route('/akello-game-events')
@login_required
def akello_game_events():
    """Admin page for managing game users and games"""
    # Check if user is admin or has Akello Events or Content Development privilege
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Akello Events') or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return "Unauthorized", 403
    return render_template('akello_game_events.html', title='Game Events Management')


@app.route('/game-login')
def game_login():
    """Login page for game participants"""
    return render_template('game_login.html', title='Game Login')


@app.route('/game-dashboard')
def game_dashboard():
    """Game dashboard for participants"""
    return render_template('game_dashboard.html', title='Game Dashboard')


@app.route('/play-game/<int:game_id>')
def play_game(game_id):
    """Page to play a specific game"""
    return render_template('play_game.html', title='Play Game', game_id=game_id)


# Game User API Routes
@app.route('/api/game-users/register', methods=['POST'])
@login_required
def register_game_user():
    """Register a new game user (admin, Content Development, or Akello Events privilege)"""
    # Check if user has permission to register game users
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development') or 
        current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.get_json()
        firstname = data.get('firstname', '').strip()
        surname = data.get('surname', '').strip()
        username = data.get('username', '').strip()
        password = data.get('password', '')
        age = data.get('age')
        phone_number = data.get('phone_number', '').strip() or None
        
        if not all([firstname, surname, username, password, age]):
            return jsonify({'error': 'All fields are required'}), 400
        
        # Validate age (accept all ages)
        try:
            age = int(age)
            if age < 0 or age > 150:  # Reasonable age limits
                return jsonify({'error': 'Please enter a valid age'}), 400
        except (ValueError, TypeError):
            return jsonify({'error': 'Age must be a valid number'}), 400
        
        # Check if username exists
        existing = GameUser.query.filter_by(username=username).first()
        if existing:
            return jsonify({'error': 'Username already exists'}), 400
        
        game_user = GameUser(
            firstname=firstname,
            surname=surname,
            username=username,
            age=age,
            phone_number=phone_number,
            auth_source='local',
        )
        # Automatically assign age range based on age
        game_user.age_range = game_user.determine_age_range()
        game_user.set_password(password)
        db.session.add(game_user)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'user': {
                'id': game_user.id,
                'firstname': game_user.firstname,
                'surname': game_user.surname,
                'username': game_user.username,
                'age': game_user.age,
                'age_range': game_user.age_range,
                'phone_number': game_user.phone_number,
                'auth_source': game_user.auth_source,
            }
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/register-public', methods=['POST'])
def register_game_user_public():
    """Public self-registration is disabled — use Smart Learning credentials."""
    return jsonify({
        'error': (
            'Register on Smart Learning / Ruzivo first, then sign in here '
            'with the same learner username and password.'
        )
    }), 403


@app.route('/api/game-users/login', methods=['POST'])
def game_user_login():
    """Login for game users via Smart Learning / Ruzivo student credentials."""
    try:
        from app.services.ruzivo_student_auth import (
            RuzivoAuthError,
            authenticate_student,
            sync_game_user_from_ruzivo,
            user_needs_dob,
        )

        data = request.get_json() or {}
        username = data.get('username', '').strip()
        password = data.get('password', '')

        if not username or not password:
            return jsonify({'error': 'Username and password are required'}), 400

        try:
            student = authenticate_student(username, password)
        except RuzivoAuthError as auth_err:
            return jsonify({'error': str(auth_err)}), auth_err.status_code

        game_user = sync_game_user_from_ruzivo(student)
        db.session.commit()

        needs_dob = user_needs_dob(game_user)
        return jsonify({
            'success': True,
            'needs_dob': needs_dob,
            'user': {
                'id': game_user.id,
                'firstname': game_user.firstname,
                'surname': game_user.surname,
                'username': game_user.username,
                'age': game_user.age,
                'age_range': game_user.age_range,
                'dob': game_user.dob.isoformat() if game_user.dob else None,
                'auth_source': game_user.auth_source,
                'ruzivo_student_id': game_user.ruzivo_student_id,
            }
        }), 200
    except Exception as e:
        db.session.rollback()
        app.logger.exception('Game user login failed')
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>/dob', methods=['POST'])
def set_game_user_dob(user_id):
    """Save date of birth for a game user when Ruzivo has none."""
    try:
        from app.services.ruzivo_student_auth import apply_local_dob

        data = request.get_json() or {}
        dob_value = data.get('dob')
        game_user = GameUser.query.get_or_404(user_id)
        try:
            apply_local_dob(game_user, dob_value)
        except ValueError as ve:
            return jsonify({'error': str(ve)}), 400
        db.session.commit()
        return jsonify({
            'success': True,
            'needs_dob': False,
            'user': {
                'id': game_user.id,
                'firstname': game_user.firstname,
                'surname': game_user.surname,
                'username': game_user.username,
                'age': game_user.age,
                'age_range': game_user.age_range,
                'dob': game_user.dob.isoformat() if game_user.dob else None,
                'auth_source': game_user.auth_source,
                'ruzivo_student_id': game_user.ruzivo_student_id,
            }
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users', methods=['GET'])
@login_required
def list_game_users():
    """List all game users (Admin, Content Development, or Akello Events privilege)"""
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development') or 
        current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        users = GameUser.query.order_by(GameUser.created_at.desc()).all()
        return jsonify({
            'users': [{
                'id': u.id,
                'firstname': u.firstname,
                'surname': u.surname,
                'username': u.username,
                'age': u.age,
                'phone_number': u.phone_number,
                'dob': u.dob.isoformat() if getattr(u, 'dob', None) else None,
                'auth_source': getattr(u, 'auth_source', None) or 'local',
                'ruzivo_student_id': getattr(u, 'ruzivo_student_id', None),
                'grade': getattr(u, 'grade', None),
                'created_at': u.created_at.isoformat() if u.created_at else None,
                'last_login': u.last_login.isoformat() if u.last_login else None
            } for u in users]
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>', methods=['DELETE'])
@login_required
def delete_game_user(user_id):
    """Delete a game user (Admin, Content Development, or Akello Events privilege)"""
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development') or 
        current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        game_user = GameUser.query.get_or_404(user_id)
        db.session.delete(game_user)
        db.session.commit()
        return jsonify({'success': True}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>/stats', methods=['GET'])
def get_game_user_stats(user_id):
    """Get statistics for a game user - average is total score / unique games played"""
    try:
        scores = GameScore.query.filter_by(game_user_id=user_id).all()
        if not scores:
            return jsonify({
                'stats': {
                    'total_games': 0,
                    'average_score': 0,
                    'best_score': 0,
                    'total_score': 0,
                    'games_played': []
                }
            }), 200
        
        # Get unique games played (one attempt per game)
        unique_games = {}
        for score in scores:
            if score.game_id not in unique_games:
                unique_games[score.game_id] = score
        
        total_unique_games = len(unique_games)
        total_score = sum(s.score for s in scores)
        # Average = total score / number of unique games played (not attempts)
        average_score = total_score / total_unique_games if total_unique_games > 0 else 0
        best_score = max(s.score for s in scores)
        
        # Get game titles for played games
        game_ids = list(unique_games.keys())
        games = {g.id: g.title for g in Game.query.filter(Game.id.in_(game_ids)).all()} if game_ids else {}
        games_played = [{'id': gid, 'title': games.get(gid, 'Unknown')} for gid in game_ids]
        
        return jsonify({
            'stats': {
                'total_games': total_unique_games,  # Unique games played
                'total_attempts': len(scores),  # Total attempts
                'average_score': round(average_score, 1),
                'best_score': best_score,
                'total_score': total_score,
                'games_played': games_played
            }
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/leaderboard', methods=['GET'])
@login_required
def get_leaderboard():
    """Get leaderboard with user stats (Admin or Content Development privilege)"""
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        all_users = GameUser.query.all()
        leaderboard = []
        
        for user in all_users:
            scores = GameScore.query.filter_by(game_user_id=user.id).all()
            if scores:
                # Get unique games played
                unique_games = set(s.game_id for s in scores)
                total_unique_games = len(unique_games)
                total_score = sum(s.score for s in scores)
                average_score = total_score / total_unique_games if total_unique_games > 0 else 0
                
                leaderboard.append({
                    'user_id': user.id,
                    'firstname': user.firstname,
                    'surname': user.surname,
                    'username': user.username,
                    'age': user.age,
                    'total_score': total_score,
                    'games_played': total_unique_games,
                    'average_score': round(average_score, 1),
                    'total_attempts': len(scores)
                })
            else:
                leaderboard.append({
                    'user_id': user.id,
                    'firstname': user.firstname,
                    'surname': user.surname,
                    'username': user.username,
                    'age': user.age,
                    'total_score': 0,
                    'games_played': 0,
                    'average_score': 0,
                    'total_attempts': 0
                })
        
        # Sort by average score descending
        leaderboard.sort(key=lambda x: x['average_score'], reverse=True)
        
        return jsonify({'leaderboard': leaderboard}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ruzivo-exercises', methods=['GET'])
@login_required
def browse_ruzivo_exercises():
    """Browse Ruzivo exercises available for import (staff only)."""
    has_permission = (
        current_user.userRole == 'Admin'
        or current_user.has_privilege('Content Development')
        or current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    try:
        from app.services.ruzivo_exercise_import import fetch_exercises, grade_to_age_range

        grade = request.args.get('grade', type=int)
        hs_param = (request.args.get('hs') or '').strip().lower()
        hs = True if hs_param in ('1', 'true', 'yes') else False if hs_param in ('0', 'false', 'no') else None
        limit = min(int(request.args.get('limit', 50)), 500)
        specs = []
        for is_hs in ([hs] if hs is not None else [False, True]):
            specs.extend(fetch_exercises(hs=is_hs, grade=grade, limit=limit))
        return jsonify({
            'exercises': [{
                'ex_id': s['ex_id'],
                'exercise': s['exercise'],
                'grade': s['grade'],
                'age_range': grade_to_age_range(s['grade']),
                'subject': s['subject_name'],
                'question_count': len(s['questions']),
                'ruzivo_source': s['ruzivo_source'],
            } for s in specs[:limit]],
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ruzivo-exercises/sync', methods=['POST'])
@login_required
def sync_ruzivo_exercises():
    """Import/update Ruzivo exercises into the games catalog."""
    has_permission = (
        current_user.userRole == 'Admin'
        or current_user.has_privilege('Content Development')
        or current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    try:
        from app.services.ruzivo_exercise_import import sync_ruzivo_exercises

        data = request.get_json() or {}
        grade = data.get('grade')
        if grade is not None:
            grade = int(grade)
        hs = data.get('hs')
        if hs is not None:
            hs = bool(hs)
        limit = min(int(data.get('limit_per_band', 100)), 500)
        stats = sync_ruzivo_exercises(grade=grade, hs=hs, limit_per_band=limit)
        return jsonify({'success': True, 'stats': stats}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# Game API Routes
@app.route('/api/games', methods=['GET'])
def list_games():
    """List all games (filter by active, age, subject, or user age)."""
    try:
        active_only = request.args.get('active', '').lower() == 'true'
        user_id = request.args.get('user_id')  # For filtering by user's age
        age_range_filter = (request.args.get('age_range') or '').strip()
        subject_filter = (request.args.get('subject') or '').strip()
        content_source_filter = (request.args.get('content_source') or '').strip()
        grade_filter = request.args.get('grade', type=int)
        prefer_user_grade = request.args.get('prefer_user_grade', '').lower() in ('1', 'true', 'yes')
        include_cross = request.args.get('include_cross_age', '').lower() == 'true'
        
        query = Game.query
        if active_only:
            query = query.filter_by(is_active=True)

        if content_source_filter in ('ruzivo', 'general_knowledge'):
            query = query.filter(Game.content_source == content_source_filter)

        if subject_filter:
            query = query.filter(Game.subject == subject_filter)

        if grade_filter is not None:
            query = query.filter(Game.grade == grade_filter)

        if age_range_filter:
            ages = [age_range_filter]
            if include_cross and age_range_filter in ('9-10', '11-12', '13-14', '15-16', '17-19'):
                ages.append('9-19')
            query = query.filter(Game.age_range.in_(ages))
        
        # Filter by user's age if user_id is provided
        game_user = None
        if user_id:
            try:
                game_user = GameUser.query.get(int(user_id))
                if game_user and game_user.age:
                    user_age = game_user.age
                    matching_ranges = []
                    
                    if user_age < 9:
                        matching_ranges = ['Infants']
                    elif 9 <= user_age <= 10:
                        matching_ranges = ['9-10', '9-19']
                    elif 11 <= user_age <= 12:
                        matching_ranges = ['11-12', '9-19']
                    elif 13 <= user_age <= 14:
                        matching_ranges = ['13-14', '9-19']
                    elif 15 <= user_age <= 16:
                        matching_ranges = ['15-16', '9-19']
                    elif 17 <= user_age <= 19:
                        matching_ranges = ['17-19', '9-19']
                    else:
                        matching_ranges = ['Youths & older']
                    
                    if matching_ranges:
                        query = query.filter(Game.age_range.in_(matching_ranges))
                    else:
                        query = query.filter(Game.id == -1)
                # Smart Learning: prefer learner grade when available
                if (
                    prefer_user_grade
                    and content_source_filter == 'ruzivo'
                    and game_user
                    and game_user.grade
                    and grade_filter is None
                ):
                    query = query.filter(Game.grade == int(game_user.grade))
            except (ValueError, AttributeError):
                pass
        
        games = query.order_by(Game.subject.asc(), Game.grade.asc(), Game.title.asc()).all()

        from app.games.quiz_titles import display_game_title

        return jsonify({
            'games': [{
                'id': g.id,
                'title': display_game_title(
                    g.title,
                    content_source=getattr(g, 'content_source', None),
                    grade=getattr(g, 'grade', None),
                ),
                'description': g.description,
                'max_score': g.max_score,
                'age_range': g.age_range,
                'subject': g.subject,
                'content_source': getattr(g, 'content_source', None) or 'general_knowledge',
                'grade': getattr(g, 'grade', None),
                'ruzivo_ex_id': getattr(g, 'ruzivo_ex_id', None),
                'difficulty_level': g.difficulty_level,
                'is_active': g.is_active,
                'bank_item_count': len(g.bank_items) if g.bank_items is not None else 0,
                'created_at': g.created_at.isoformat() if g.created_at else None
            } for g in games],
            'learner': {
                'grade': game_user.grade if game_user else None,
                'age': game_user.age if game_user else None,
                'needs_dob': (game_user.dob is None) if game_user else None,
            } if game_user else None,
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _can_manage_games():
    return (
        current_user.userRole == 'Admin'
        or current_user.has_privilege('Content Development')
    )


@app.route('/api/games/export', methods=['GET'])
@login_required
def export_games_pack():
    """Download General Knowledge games as a JSON pack (excludes Ruzivo imports)."""
    if not _can_manage_games():
        return jsonify({'error': 'Unauthorized'}), 403
    try:
        from app.games.game_pack import build_game_pack

        pack = build_game_pack()
        stamp = datetime.utcnow().strftime('%Y%m%d')
        filename = f'akello-gk-games-{stamp}.json'
        payload = json.dumps(pack, ensure_ascii=False, indent=2)
        response = Response(payload, mimetype='application/json; charset=utf-8')
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/games/import', methods=['POST'])
@login_required
def import_games_pack():
    """Upload a General Knowledge JSON pack and upsert games (never touches Ruzivo)."""
    if not _can_manage_games():
        return jsonify({'error': 'Unauthorized'}), 403
    try:
        from app.games.game_pack import import_game_pack

        pack = None
        upload = request.files.get('file') or request.files.get('pack')
        if upload and upload.filename:
            raw = upload.read()
            if not raw:
                return jsonify({'error': 'Uploaded file is empty.'}), 400
            pack = json.loads(raw.decode('utf-8-sig'))
        else:
            pack = request.get_json(silent=True)

        if not pack:
            return jsonify({'error': 'Provide a JSON file or JSON body.'}), 400

        stats = import_game_pack(pack, created_by=current_user.id)
        return jsonify({'success': True, **stats}), 200
    except ValueError as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400
    except json.JSONDecodeError:
        db.session.rollback()
        return jsonify({'error': 'Invalid JSON file.'}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/games', methods=['POST'])
@login_required
def create_game():
    """Create a new game (admin or Content Development privilege only)"""
    # Only Admin and Content Development can create games
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized. Only Admin and users with Content Development privilege can create games.'}), 403
    
    try:
        data = request.get_json()
        title = data.get('title', '').strip()
        description = data.get('description', '').strip()
        html_content = data.get('html_content', '').strip()
        max_score = data.get('max_score')
        age_range = data.get('age_range', '').strip()
        subject = (data.get('subject') or '').strip()
        difficulty_level = data.get('difficulty_level', '').strip()
        is_active = data.get('is_active', True)
        
        if not title or not html_content:
            return jsonify({'error': 'Title and HTML content are required'}), 400
        
        # Validate age_range if provided
        valid_age_ranges = ['Infants', '9-10', '11-12', '13-14', '15-16', '17-19', '9-19', 'Youths & older']
        if age_range and age_range not in valid_age_ranges:
            return jsonify({'error': f'Age range must be one of: {", ".join(valid_age_ranges)}'}), 400
        
        # Validate difficulty_level if provided
        valid_difficulty_levels = ['easy', 'medium', 'hard']
        if difficulty_level and difficulty_level not in valid_difficulty_levels:
            return jsonify({'error': f'Difficulty level must be one of: {", ".join(valid_difficulty_levels)}'}), 400
        
        game = Game(
            title=title,
            description=description,
            html_content=html_content,
            max_score=int(max_score) if max_score else None,
            age_range=age_range if age_range else None,
            subject=subject or None,
            content_source=(data.get('content_source') or 'general_knowledge').strip() or 'general_knowledge',
            difficulty_level=difficulty_level if difficulty_level else None,
            is_active=bool(is_active),
            created_by=current_user.id
        )
        db.session.add(game)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'game': {
                'id': game.id,
                'title': game.title,
                'description': game.description,
                'max_score': game.max_score,
                'age_range': game.age_range,
                'subject': game.subject,
                'difficulty_level': game.difficulty_level,
                'is_active': game.is_active
            }
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/games/<int:game_id>', methods=['GET'])
def get_game(game_id):
    """Get a specific game"""
    try:
        from app.games.quiz_titles import display_game_title
        game = Game.query.get_or_404(game_id)
        return jsonify({
            'game': {
                'id': game.id,
                'title': display_game_title(
                    game.title,
                    content_source=getattr(game, 'content_source', None),
                    grade=getattr(game, 'grade', None),
                ),
                'description': game.description,
                'html_content': game.html_content,
                'max_score': game.max_score,
                'age_range': game.age_range,
                'subject': game.subject,
                'content_source': getattr(game, 'content_source', None) or 'general_knowledge',
                'grade': getattr(game, 'grade', None),
                'difficulty_level': game.difficulty_level,
                'is_active': game.is_active,
                'created_at': game.created_at.isoformat() if game.created_at else None
            }
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/games/<int:game_id>', methods=['PUT'])
@login_required
def update_game(game_id):
    """Update a game (admin or Content Development privilege only)"""
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        game = Game.query.get_or_404(game_id)
        data = request.get_json()
        
        if 'title' in data:
            game.title = data['title'].strip()
        if 'description' in data:
            game.description = data.get('description', '').strip()
        if 'html_content' in data:
            game.html_content = data['html_content'].strip()
        if 'max_score' in data:
            game.max_score = int(data['max_score']) if data['max_score'] else None
        if 'age_range' in data:
            age_range = data.get('age_range', '').strip()
            valid_age_ranges = ['Infants', '9-10', '11-12', '13-14', '15-16', '17-19', '9-19', 'Youths & older']
            if age_range and age_range not in valid_age_ranges:
                return jsonify({'error': f'Age range must be one of: {", ".join(valid_age_ranges)}'}), 400
            game.age_range = age_range if age_range else None
        if 'subject' in data:
            game.subject = (data.get('subject') or '').strip() or None
        if 'difficulty_level' in data:
            difficulty_level = data.get('difficulty_level', '').strip()
            valid_difficulty_levels = ['easy', 'medium', 'hard']
            if difficulty_level and difficulty_level not in valid_difficulty_levels:
                return jsonify({'error': f'Difficulty level must be one of: {", ".join(valid_difficulty_levels)}'}), 400
            game.difficulty_level = difficulty_level if difficulty_level else None
        if 'is_active' in data:
            game.is_active = bool(data['is_active'])
        
        game.updated_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'success': True,
            'game': {
                'id': game.id,
                'title': game.title,
                'description': game.description,
                'max_score': game.max_score,
                'age_range': game.age_range,
                'subject': game.subject,
                'difficulty_level': game.difficulty_level,
                'is_active': game.is_active
            }
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-bank-items', methods=['GET'])
@login_required
def list_game_bank_items():
    """List selectable bank items for race builders (staff)."""
    has_permission = (
        current_user.userRole == 'Admin'
        or current_user.has_privilege('Content Development')
        or current_user.has_privilege('Akello Events')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    try:
        from app.models import GameBankItem
        age_range = (request.args.get('age_range') or '').strip()
        subject = (request.args.get('subject') or '').strip()
        active_only = request.args.get('active', 'true').lower() != 'false'
        include_cross = request.args.get('include_cross_age', 'true').lower() != 'false'

        query = GameBankItem.query
        if active_only:
            query = query.filter_by(is_active=True)
        if subject:
            query = query.filter(GameBankItem.subject == subject)
        if age_range:
            ages = [age_range]
            if include_cross and age_range in ('9-10', '11-12', '13-14', '15-16', '17-19'):
                ages.append('9-19')
            query = query.filter(GameBankItem.age_range.in_(ages))

        items = query.order_by(
            GameBankItem.subject.asc(),
            GameBankItem.sort_order.asc(),
            GameBankItem.title.asc(),
        ).all()
        return jsonify({
            'items': [{
                'id': item.id,
                'game_id': item.game_id,
                'subject': item.subject,
                'age_range': item.age_range,
                'item_kind': item.item_kind,
                'slug': item.slug,
                'title': item.title,
                'prompt': item.prompt,
                'payload': item.payload_json or {},
                'points_default': item.points_default,
                'sort_order': item.sort_order,
                'game_title': item.game.title if item.game else None,
            } for item in items]
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/games/<int:game_id>', methods=['DELETE'])
@login_required
def delete_game_route(game_id):
    """Delete a game (admin or Content Development privilege only)"""
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        game = Game.query.get_or_404(game_id)
        db.session.delete(game)
        db.session.commit()
        return jsonify({'success': True}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# Game Score API Routes
@app.route('/api/game-scores/submit', methods=['POST'])
def submit_game_score():
    """Submit a score for a game"""
    try:
        data = request.get_json()
        game_id = data.get('game_id')
        score = data.get('score')
        max_score = data.get('max_score')
        user_id = data.get('user_id')
        
        # Get user from request body (sent from frontend sessionStorage)
        if not user_id:
            return jsonify({'error': 'User not authenticated. Please login again.'}), 401

        game_user = GameUser.query.get(int(user_id))
        if not game_user:
            return jsonify({'error': 'User not found. Please login again.'}), 401
        from app.services.ruzivo_student_auth import user_needs_dob
        if user_needs_dob(game_user):
            return jsonify({'error': 'Please enter your date of birth before playing.', 'needs_dob': True}), 403
        
        if not game_id or score is None:
            return jsonify({'error': 'Game ID and score are required'}), 400
        
        # Get game to check max_score
        game = Game.query.get(game_id)
        if not game:
            return jsonify({'error': 'Game not found'}), 404
        
        # Use game's max_score if not provided
        if max_score is None and game.max_score:
            max_score = game.max_score
        
        # Calculate percentage
        percentage = None
        if max_score and max_score > 0:
            percentage = (score / max_score) * 100
        
        # Get attempt number
        last_attempt = GameScore.query.filter_by(
            game_user_id=user_id,
            game_id=game_id
        ).order_by(GameScore.attempt_number.desc()).first()
        
        attempt_number = (last_attempt.attempt_number + 1) if last_attempt else 1
        
        # Create score record
        game_score = GameScore(
            game_user_id=int(user_id),
            game_id=int(game_id),
            score=int(score),
            max_score=int(max_score) if max_score else None,
            percentage=percentage,
            attempt_number=attempt_number
        )
        db.session.add(game_score)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'score': {
                'id': game_score.id,
                'score': game_score.score,
                'max_score': game_score.max_score,
                'percentage': game_score.percentage,
                'attempt_number': game_score.attempt_number
            }
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>/session', methods=['GET'])
def get_game_user_session(user_id):
    """Session check for portal (DOB gate, grade, display name)."""
    try:
        from app.services.ruzivo_student_auth import user_needs_dob
        game_user = GameUser.query.get_or_404(user_id)
        needs = user_needs_dob(game_user)
        return jsonify({
            'user': {
                'id': game_user.id,
                'firstname': game_user.firstname,
                'surname': game_user.surname,
                'username': game_user.username,
                'age': game_user.age,
                'grade': game_user.grade,
                'dob': game_user.dob.isoformat() if game_user.dob else None,
                'needs_dob': needs,
                'auth_source': game_user.auth_source,
            }
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>/scores', methods=['GET'])
def get_game_user_scores(user_id):
    """Get scores for a specific game user (public endpoint for game users)"""
    try:
        scores = GameScore.query.filter_by(game_user_id=user_id).order_by(GameScore.played_at.desc()).all()
        
        # Get related data
        game_ids = list(set(s.game_id for s in scores))
        games = {g.id: g for g in Game.query.filter(Game.id.in_(game_ids)).all()} if game_ids else {}

        best_by_game = {}
        for s in scores:
            prev = best_by_game.get(s.game_id)
            if prev is None or (s.percentage or 0) > (prev.get('percentage') or 0) or (
                (s.percentage or 0) == (prev.get('percentage') or 0) and s.score >= prev.get('score', 0)
            ):
                best_by_game[s.game_id] = {
                    'game_id': s.game_id,
                    'score': s.score,
                    'max_score': s.max_score,
                    'percentage': s.percentage,
                    'attempt_number': s.attempt_number,
                    'attempts': 0,
                }
        attempt_counts = {}
        for s in scores:
            attempt_counts[s.game_id] = attempt_counts.get(s.game_id, 0) + 1
        for gid, row in best_by_game.items():
            row['attempts'] = attempt_counts.get(gid, 1)
        
        return jsonify({
            'scores': [{
                'id': s.id,
                'game_id': s.game_id,
                'game_title': games.get(s.game_id).title if s.game_id in games else 'Unknown',
                'score': s.score,
                'max_score': s.max_score,
                'percentage': s.percentage,
                'attempt_number': s.attempt_number,
                'played_at': s.played_at.isoformat() if s.played_at else None
            } for s in scores],
            'best_by_game': best_by_game,
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-users/<int:user_id>/leaderboard', methods=['GET'])
def get_age_group_leaderboard(user_id):
    """Get leaderboard filtered by user's age group (public endpoint for game users)"""
    try:
        # Get the current user to determine their age group
        game_user = GameUser.query.get(user_id)
        if not game_user or not game_user.age:
            return jsonify({'error': 'User not found or age not set'}), 404
        
        user_age = game_user.age
        
        # Determine age range based on user's age
        if user_age < 9:
            age_range = 'Infants'
        elif 9 <= user_age <= 10:
            age_range = '9-10'
        elif 11 <= user_age <= 12:
            age_range = '11-12'
        elif 13 <= user_age <= 14:
            age_range = '13-14'
        elif 15 <= user_age <= 16:
            age_range = '15-16'
        elif 17 <= user_age <= 19:
            age_range = '17-19'
        else:  # user_age > 19
            age_range = 'Youths & older'
        
        # Get all users in the same age group
        if age_range == 'Infants':
            age_group_users = GameUser.query.filter(GameUser.age < 9).all()
        elif age_range == '9-10':
            age_group_users = GameUser.query.filter(GameUser.age >= 9, GameUser.age <= 10).all()
        elif age_range == '11-12':
            age_group_users = GameUser.query.filter(GameUser.age >= 11, GameUser.age <= 12).all()
        elif age_range == '13-14':
            age_group_users = GameUser.query.filter(GameUser.age >= 13, GameUser.age <= 14).all()
        elif age_range == '15-16':
            age_group_users = GameUser.query.filter(GameUser.age >= 15, GameUser.age <= 16).all()
        elif age_range == '17-19':
            age_group_users = GameUser.query.filter(GameUser.age >= 17, GameUser.age <= 19).all()
        elif age_range == 'Youths & older':
            age_group_users = GameUser.query.filter(GameUser.age > 19).all()
        else:
            age_group_users = []
        
        user_ids = [u.id for u in age_group_users]
        
        if not user_ids:
            return jsonify({'leaderboard': [], 'age_range': age_range}), 200
        
        # Get all scores for users in this age group
        scores = GameScore.query.filter(GameScore.game_user_id.in_(user_ids)).all()
        
        # Calculate best score per user (highest score across all games)
        user_best_scores = {}
        for score in scores:
            user_id_key = score.game_user_id
            if user_id_key not in user_best_scores:
                user_best_scores[user_id_key] = {
                    'best_score': score.score,
                    'total_games': 1,
                    'total_score': score.score,
                    'username': None,
                    'firstname': None,
                    'surname': None
                }
            else:
                user_best_scores[user_id_key]['best_score'] = max(
                    user_best_scores[user_id_key]['best_score'],
                    score.score
                )
                user_best_scores[user_id_key]['total_games'] += 1
                user_best_scores[user_id_key]['total_score'] += score.score
        
        # Get user details
        users_dict = {u.id: u for u in age_group_users}
        for user_id_key in user_best_scores:
            user = users_dict.get(user_id_key)
            if user:
                user_best_scores[user_id_key]['username'] = user.username
                user_best_scores[user_id_key]['firstname'] = user.firstname
                user_best_scores[user_id_key]['surname'] = user.surname
        
        # Convert to list and sort by best score (descending)
        leaderboard = [
            {
                'user_id': uid,
                'username': data['username'],
                'firstname': data['firstname'],
                'surname': data['surname'],
                'best_score': data['best_score'],
                'total_games': data['total_games'],
                'average_score': round(data['total_score'] / data['total_games'], 1) if data['total_games'] > 0 else 0
            }
            for uid, data in user_best_scores.items()
        ]
        
        leaderboard.sort(key=lambda x: x['best_score'], reverse=True)
        
        # Add rank
        for i, entry in enumerate(leaderboard, 1):
            entry['rank'] = i
        
        return jsonify({
            'leaderboard': leaderboard,
            'age_range': age_range,
            'current_user_id': user_id
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/game-scores', methods=['GET'])
@login_required
def list_game_scores():
    """List game scores with optional filters (Admin or Content Development privilege)"""
    # Check if user has permission to view scores
    has_permission = (
        current_user.userRole == 'Admin' or 
        current_user.has_privilege('Content Development')
    )
    if not has_permission:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        game_id = request.args.get('game_id')
        user_id = request.args.get('user_id')
        
        query = GameScore.query
        
        if game_id:
            query = query.filter_by(game_id=int(game_id))
        if user_id:
            query = query.filter_by(game_user_id=int(user_id))
        
        scores = query.order_by(GameScore.played_at.desc()).all()
        
        # Get related data
        game_ids = list(set(s.game_id for s in scores))
        user_ids = list(set(s.game_user_id for s in scores))
        
        games = {g.id: g for g in Game.query.filter(Game.id.in_(game_ids)).all()} if game_ids else {}
        users = {u.id: u for u in GameUser.query.filter(GameUser.id.in_(user_ids)).all()} if user_ids else {}
        
        return jsonify({
            'scores': [{
                'id': s.id,
                'user_id': s.game_user_id,
                'user_name': f"{users.get(s.game_user_id, GameUser()).firstname} {users.get(s.game_user_id, GameUser()).surname}" if s.game_user_id in users else 'Unknown',
                'game_id': s.game_id,
                'game_title': games.get(s.game_id, Game()).title if s.game_id in games else 'Unknown',
                'score': s.score,
                'max_score': s.max_score,
                'percentage': s.percentage,
                'attempt_number': s.attempt_number,
                'played_at': s.played_at.isoformat() if s.played_at else None
            } for s in scores],
            'games': [{'id': g.id, 'title': g.title} for g in games.values()],
            'users': [{'id': u.id, 'firstname': u.firstname, 'surname': u.surname, 'username': u.username} for u in users.values()]
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _get_all_champions_school_tracker_data():
    """Fetch all schools for all champions with status. Returns list of row dicts. Raises on error."""
    def normalize_active_flag(value):
        if value is None:
            return False
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.lower() in ('yes', 'true', '1', 'active')
        if isinstance(value, int):
            return value == 1
        return False

    # 1. Fetch all champion schools
    champion_schools = ChampionSchool.query.all()
    app.logger.info(f"DEBUG: Found {len(champion_schools)} ChampionSchool records")

    def normalize_id(value):
        """Normalize IDs from either numeric or string input while preserving leading zeros."""
        if value is None:
            return ''
        s = str(value).strip()
        return s

    def normalize_school_name(value):
        """Normalize school name for matching (trim, collapse whitespace, lowercase)."""
        if value is None:
            return ''
        return ' '.join(str(value).split()).strip().lower()

    # 2. Extract unique school_ids and map to champion info
    school_id_map = {} # school_id -> list of {champion, province, fallback_name, champion_school_id}
    unique_school_ids = set()
    champion_school_ids = set()
    unique_library_school_ids = set()
    unique_asl_ids_str_raw = set()

    for cs in champion_schools:
        champion_school_ids.add(cs.id)
        current_schools_list = cs.get_schools() or []
        current_champion_name = f"{cs.firstname} {cs.lastname}"

        for school_entry in current_schools_list:
            raw_sid = school_entry.get('asl_school_id')
            raw_sid_str = normalize_id(raw_sid)
            raw_sname = school_entry.get('school_name', 'Unknown School')
            raw_lid = school_entry.get('library_school_id')

            if raw_sid_str:
                try:
                    sid_val = int(raw_sid_str)
                    if sid_val == 0:
                        continue

                    unique_school_ids.add(sid_val)
                    unique_asl_ids_str_raw.add(raw_sid_str)
                    if sid_val not in school_id_map:
                        school_id_map[sid_val] = []

                    # Library IDs are sometimes missing; keep them for timestamp matching.
                    lib_id = str(raw_lid or '').strip()
                    if lib_id:
                        unique_library_school_ids.add(lib_id)

                    school_id_map[sid_val].append({
                        'champion': current_champion_name,
                        'province': cs.province,
                        'fallback_name': raw_sname,
                        'champion_school_id': cs.id,
                        # Preserve raw string (leading zeros) so it can match ChampionSchoolRequest entries.
                        'asl_school_id': raw_sid_str,
                        'library_school_id': normalize_id(raw_lid)
                    })
                except (ValueError, TypeError):
                    app.logger.warning(f"DEBUG: Invalid school ID found: {raw_sid}")

    app.logger.info(f"DEBUG: Extracted {len(unique_school_ids)} unique school IDs")
    if not unique_school_ids:
        return []

    # 2b. Pre-compute "added" timestamps per champion+school based on approved requests
    request_added_by_key = {}  # (champion_school_id, asl_school_id, library_school_id) -> datetimes
    request_added_by_name_key = {}  # (champion_school_id, school_name_lower) -> datetimes

    unique_asl_ids_str = sorted(set([str(sid) for sid in unique_school_ids]) | unique_asl_ids_str_raw)
    if champion_school_ids:
        # Fetch all approved requests for these champion profiles.
        # We intentionally avoid over-filtering by ASL/library IDs because those may differ in formatting (e.g., leading zeros).
        req_query = ChampionSchoolRequest.query.filter(
            sa.func.lower(ChampionSchoolRequest.status) == 'approved',
            ChampionSchoolRequest.champion_school_id.in_(champion_school_ids)
        )
        approved_requests = req_query.all()
    else:
        approved_requests = []

    for req in approved_requests:
        cid = req.champion_school_id
        aid = normalize_id(req.asl_school_id)
        lid = normalize_id(req.library_school_id)
        name_lower = normalize_school_name(req.school_name)

        def id_variants(v):
            v_norm = normalize_id(v)
            if not v_norm:
                return ['']
            # Treat "0"/all-zeros as missing.
            if set(v_norm) <= {'0'}:
                return ['']
            stripped = v_norm.lstrip('0')
            if stripped and stripped != v_norm:
                return list({v_norm, stripped})
            return [v_norm]

        aid_vars = id_variants(aid)
        lid_vars = id_variants(lid)

        # Store under multiple key variants so resolve() can match either formatting.
        for aid_v in aid_vars:
            for lid_v in lid_vars:
                key = (cid, aid_v, lid_v)
                existing = request_added_by_key.get(key)
                if not existing:
                    existing = {'profile_added_at': None, 'system_added_at': None}

                if req.reviewed_at is not None:
                    if existing['profile_added_at'] is None or req.reviewed_at < existing['profile_added_at']:
                        existing['profile_added_at'] = req.reviewed_at
                if req.created_at is not None:
                    if existing['system_added_at'] is None or req.created_at < existing['system_added_at']:
                        existing['system_added_at'] = req.created_at
                request_added_by_key[key] = existing

        # Also build a name-based lookup for cases where IDs mismatch / are missing.
        if name_lower:
            name_key = (cid, name_lower)
            existing2 = request_added_by_name_key.get(name_key)
            if not existing2:
                existing2 = {'profile_added_at': None, 'system_added_at': None}
            if req.reviewed_at is not None:
                if existing2['profile_added_at'] is None or req.reviewed_at < existing2['profile_added_at']:
                    existing2['profile_added_at'] = req.reviewed_at
            if req.created_at is not None:
                if existing2['system_added_at'] is None or req.created_at < existing2['system_added_at']:
                    existing2['system_added_at'] = req.created_at
            request_added_by_name_key[name_key] = existing2

    def resolve_added_timestamps(info):
        """
        Resolve added timestamps using approved champion school requests.
        Returns (profile_added_at_dt, system_added_at_dt) as datetime objects or (None, None).
        """
        cid = info.get('champion_school_id')
        aid = normalize_id(info.get('asl_school_id'))
        lid = normalize_id(info.get('library_school_id'))
        fallback_name = normalize_school_name(info.get('fallback_name'))

        def id_variants(v):
            v_norm = normalize_id(v)
            if not v_norm:
                return ['']
            if set(v_norm) <= {'0'}:
                return ['']
            stripped = v_norm.lstrip('0')
            if stripped and stripped != v_norm:
                return list({v_norm, stripped})
            return [v_norm]

        aid_vars = id_variants(aid)
        lid_vars = id_variants(lid)

        # Most specific first: (aid,lid), then (aid,'') and ('',lid)
        for aid_v in aid_vars:
            for lid_v in lid_vars:
                key = (cid, aid_v, lid_v)
                if key in request_added_by_key:
                    entry = request_added_by_key[key]
                    pa = entry.get('profile_added_at')
                    sa = entry.get('system_added_at')
                    if pa or sa:
                        return pa, sa
                key2 = (cid, aid_v, '')
                if key2 in request_added_by_key:
                    entry = request_added_by_key[key2]
                    pa = entry.get('profile_added_at')
                    sa = entry.get('system_added_at')
                    if pa or sa:
                        return pa, sa
                key3 = (cid, '', lid_v)
                if key3 in request_added_by_key:
                    entry = request_added_by_key[key3]
                    pa = entry.get('profile_added_at')
                    sa = entry.get('system_added_at')
                    if pa or sa:
                        return pa, sa

        if fallback_name:
            name_key = (cid, fallback_name)
            if name_key in request_added_by_name_key:
                entry = request_added_by_name_key[name_key]
                pa = entry.get('profile_added_at')
                sa = entry.get('system_added_at')
                if pa or sa:
                    return pa, sa

        return None, None

    # Convert to list for SQL IN clause
    school_ids_list = list(unique_school_ids)

    # 3. Query Ruzivo/Scholarship DB
    results = []

    # Process in chunks if too many schools
    chunk_size = 1000
    for i in range(0, len(school_ids_list), chunk_size):
        chunk = school_ids_list[i:i + chunk_size]
        placeholders = ','.join(['%s'] * len(chunk))

        conn = get_ruzivo_conn()
        cursor = conn.cursor()
            
        # Fetch scholarship details
        query = f"""
            SELECT ts.school_id, sc.school_name, ts.scholarship_type,
                   ts.expiry_date, ts.active, ts.duration
            FROM tblscholarships_schools ts
            LEFT JOIN tblschools sc ON sc.school_id = ts.school_id
            WHERE ts.school_id IN ({placeholders})
        """
        cursor.execute(query, chunk)
        scholarship_rows = cursor.fetchall()

        # Fetch durations and first awards since 2025
        since_2025_query = f"""
            SELECT x.school_id, SUM(x.duration) AS total_months, MIN(x.awarded_at) AS first_awarded_at
            FROM tblscholarships_schools x
            WHERE x.school_id IN ({placeholders})
            AND x.awarded_at >= %s
            GROUP BY x.school_id
        """
        cursor.execute(since_2025_query, (*chunk, '2025-01-01'))
        since_2025_rows = cursor.fetchall()
        since_2025_map = {}
        for row in since_2025_rows:
            today_date = date.today()
            if isinstance(row, dict):
                q_sid = row['school_id']
                fa = row['first_awarded_at']
                total_m = row['total_months'] or 0
            else:
                q_sid = row[0]
                total_m = row[1] or 0
                fa = row[2]

            is_over_12 = False
            if fa:
                fa_dt = fa.date() if hasattr(fa, 'date') else fa
                if isinstance(fa_dt, date) and fa_dt <= today_date - timedelta(days=365):
                    is_over_12 = True

            since_2025_map[q_sid] = {
                "total_months": total_m,
                "first_awarded_at": fa,
                "is_over_12_months": is_over_12
            }

        # Fetch active subscriptions for expired checks
        today_now = datetime.now().date()
        first_day_of_month = today_now.replace(day=1)
        if today_now.month == 12:
            first_day_next_month = today_now.replace(year=today_now.year + 1, month=1, day=1)
        else:
            first_day_next_month = today_now.replace(month=today_now.month + 1, day=1)

        sub_query = f"""
            SELECT s.school_id, COUNT(DISTINCT p.student_id) as active_count
            FROM tblpoints_purchase p
            JOIN tblstudents s ON p.student_id = s.student_id
            WHERE s.school_id IN ({placeholders})
              AND p.expiry_date >= %s
              AND p.expiry_date < %s
            GROUP BY s.school_id
        """
        cursor.execute(sub_query, (*chunk, first_day_of_month, first_day_next_month))
        sub_rows = cursor.fetchall()
        sub_map = {}
        for row in sub_rows:
            if isinstance(row, dict):
                sub_map[row['school_id']] = row['active_count']
            else:
                sub_map[row[0]] = row[1]

        cursor.close()
        conn.close()

        # Create map for scholarship rows
        scholarship_map = {}
        for row in scholarship_rows:
            if isinstance(row, dict):
                q_sid = row['school_id']
                scholarship_map[q_sid] = row
            else:
                q_sid = row[0]
                scholarship_map[q_sid] = {
                    'school_name': row[1],
                    'scholarship_type': row[2],
                    'expiry_date': row[3],
                    'active': row[4],
                    'duration': row[5]
                }

        # Process all unique school IDs from ChampionSchool
        for current_sid in chunk:
            champ_infos = school_id_map.get(current_sid, [])
            s_data = scholarship_map.get(current_sid)

            if s_data:
                current_sname = s_data['school_name']
                current_stype = s_data['scholarship_type']
                current_sexpiry = s_data['expiry_date']
                current_sactive = s_data['active']

                # Normalize Expiry
                if current_sexpiry:
                    if isinstance(current_sexpiry, str):
                        try:
                            current_sexpiry = datetime.strptime(current_sexpiry, "%Y-%m-%d").date()
                        except Exception:
                            try:
                                current_sexpiry = datetime.strptime(current_sexpiry, "%Y-%m-%d %H:%M:%S").date()
                            except Exception:
                                pass
                    elif hasattr(current_sexpiry, 'date'):
                        current_sexpiry = current_sexpiry.date()

                current_is_expired = False
                if current_sexpiry and isinstance(current_sexpiry, date):
                    if current_sexpiry < today_now:
                        current_is_expired = True

                normalized_active = normalize_active_flag(current_sactive)

                # Determine Status
                current_status = "Inactive"
                if normalized_active:
                    if current_is_expired:
                        current_status = "Expired"
                    else:
                        current_status = "Active"

                current_active_subs = sub_map.get(current_sid, 0)
                current_over_12 = since_2025_map.get(current_sid, {}).get("is_over_12_months", False)
                current_first_awarded = since_2025_map.get(current_sid, {}).get("first_awarded_at", "N/A")
                current_total_m = since_2025_map.get(current_sid, {}).get("total_months", 0)
            else:
                # Not in scholarship table
                fallback_name_found = "Unknown School"
                if champ_infos:
                    fallback_name_found = champ_infos[0].get('fallback_name', "Unknown School")

                current_sname = fallback_name_found
                current_stype = "N/A"
                current_sexpiry = None
                current_status = "Never Assigned"
                current_active_subs = 0
                current_over_12 = False
                current_first_awarded = "N/A"
                current_total_m = 0

            # Days left to 12 months: from First Awarded date to today in days (365 days = 12 months)
            days_left_until_12 = None
            fa_raw = since_2025_map.get(current_sid, {}).get("first_awarded_at")
            if fa_raw:
                fa_d = fa_raw.date() if hasattr(fa_raw, 'date') else fa_raw
                if isinstance(fa_d, date):
                    if fa_d > today_now:
                        days_elapsed = 0
                    else:
                        days_elapsed = (today_now - fa_d).days
                    if days_elapsed < 365:
                        days_left_until_12 = 365 - days_elapsed
            for info in champ_infos:
                profile_added_at_dt, system_added_at_dt = resolve_added_timestamps(info)
                results.append({
                    "champion": info['champion'],
                    "province": info['province'],
                    "school_name": current_sname,
                    "status": current_status,
                    "expiry_date": current_sexpiry.isoformat() if current_sexpiry else None,
                    "first_awarded_at": current_first_awarded,
                    "profile_added_at": profile_added_at_dt.isoformat() if profile_added_at_dt else None,
                    "system_added_at": system_added_at_dt.isoformat() if system_added_at_dt else None,
                    "is_over_12_months": current_over_12,
                    "scholarship_type": current_stype,
                    "active_subscriptions": current_active_subs,
                    "duration": current_total_m,
                    "days_left_until_12": days_left_until_12,
                    "asl_school_id": info.get('asl_school_id') or str(current_sid),
                    "library_school_id": info.get('library_school_id') or ''
                })

    return results


@app.route('/api/all-champions-school-tracker', methods=['GET'])
@login_required
def all_champions_school_tracker():
    """Fetch all schools for all champions with status"""
    try:
        results = _get_all_champions_school_tracker_data()
        return jsonify(results), 200
    except Exception as e:
        app.logger.error(f"Error in school tracker: {e}")
        return jsonify({'error': str(e)}), 500


# Column name normalization for champion compare upload (flexible header matching)
CHAMPION_COMPARE_FILE_COLUMNS = {
    'asl school id': 'asl_school_id',
    'al id': 'al_id',
    'akello champion': 'champion',
    'school': 'school_name',
    'province': 'province',
    'asl scholarship check': 'asl_scholarship_check',
    'duration on scholarship': 'duration_on_scholarship',
    'validity of scholarship': 'validity_of_scholarship',
}


def _normalize_compare_file_columns(df):
    """Normalize DataFrame columns to expected names. Returns new df with normalized cols, or raises ValueError if required missing."""
    required = {'asl_school_id', 'asl_scholarship_check', 'duration_on_scholarship', 'validity_of_scholarship'}
    col_map = {}
    for c in df.columns:
        key = (c or '').strip().lower()
        if key in CHAMPION_COMPARE_FILE_COLUMNS:
            col_map[c] = CHAMPION_COMPARE_FILE_COLUMNS[key]
    df = df.rename(columns=col_map)
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns: {sorted(missing)}. Expected file columns: Asl School ID, AL ID, Akello Champion, School, Province, ASL Scholarship check, Duration on scholarship, Validity of scholarship.")
    return df


def _normalize_scholarship_status_for_compare(val):
    """Normalize file or system status to comparable form (active vs not)."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return 'no'
    s = str(val).strip().lower()
    if s in ('yes', 'true', '1', 'active'):
        return 'yes'
    if s in ('no', 'false', '0', 'inactive', 'expired', 'never assigned', 'n/a'):
        return 'no'
    return s


def _normalize_date_for_compare(val):
    """Return date string YYYY-MM-DD or empty string."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return ''
    if hasattr(val, 'strftime'):
        return val.strftime('%Y-%m-%d')
    s = str(val).strip()
    if not s:
        return ''
    try:
        from datetime import datetime as dt
        parsed = pd.to_datetime(val)
        return parsed.strftime('%Y-%m-%d')
    except Exception:
        return s


@app.route('/api/all-champions-compare-upload', methods=['POST'])
@login_required
def all_champions_compare_upload():
    """Accept CSV/Excel file; compare with school tracker data; return report and summary."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        file = request.files['file']
        if not file or not file.filename:
            return jsonify({'error': 'No file selected'}), 400
        ext = (file.filename or '').rsplit('.', 1)[-1].lower()
        if ext not in ('csv', 'xlsx', 'xls'):
            return jsonify({'error': 'Invalid file type (only .csv, .xlsx, .xls allowed)'}), 400

        try:
            if ext == 'csv':
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
        except Exception as e:
            return jsonify({'error': f'Could not parse file: {str(e)}'}), 400

        if df.empty:
            return jsonify({'error': 'File has no data rows'}), 400

        try:
            df = _normalize_compare_file_columns(df)
        except ValueError as e:
            return jsonify({'error': str(e)}), 400

        try:
            system_rows = _get_all_champions_school_tracker_data()
        except Exception as e:
            app.logger.error(f"Compare upload: failed to get school tracker data: {e}")
            return jsonify({'error': 'Failed to load system data for comparison'}), 500

        # Build lookup: (asl_school_id, library_school_id) -> list of system row dicts (may be multiple champions per school)
        system_by_key = {}
        for row in system_rows:
            aid = str(row.get('asl_school_id') or '').strip()
            lid = str(row.get('library_school_id') or '').strip()
            key = (aid, lid)
            if key not in system_by_key:
                system_by_key[key] = []
            system_by_key[key].append(row)

        report = []
        summary = {'total_file_rows': 0, 'matched': 0, 'mismatched': 0, 'file_only': 0, 'system_only': 0}
        matched_system_keys = set()

        for _, row in df.iterrows():
            summary['total_file_rows'] += 1
            f_asl = str(row.get('asl_school_id', '') or '').strip()
            f_al = str(row.get('al_id', '') or '').strip() if 'al_id' in row else ''
            f_champ = str(row.get('champion', '') or '') if 'champion' in row else ''
            f_school = str(row.get('school_name', '') or '') if 'school_name' in row else ''
            f_prov = str(row.get('province', '') or '') if 'province' in row else ''
            f_check = _normalize_scholarship_status_for_compare(row.get('asl_scholarship_check'))
            f_dur = row.get('duration_on_scholarship')
            if f_dur is not None and not (isinstance(f_dur, float) and pd.isna(f_dur)):
                try:
                    f_dur = int(float(f_dur))
                except (ValueError, TypeError):
                    f_dur = str(f_dur)
            else:
                f_dur = ''
            f_valid = _normalize_date_for_compare(row.get('validity_of_scholarship'))

            file_row = {
                'asl_school_id': f_asl,
                'library_school_id': f_al,
                'champion': f_champ,
                'school_name': f_school,
                'province': f_prov,
                'asl_scholarship_check': f_check,
                'duration_on_scholarship': f_dur,
                'validity_of_scholarship': f_valid,
            }
            # Match: prefer (asl, al_id); if al_id missing in file, try (asl, '') and (asl, any)
            candidates = []
            if f_asl:
                key1 = (f_asl, f_al)
                key2 = (f_asl, '')
                candidates = system_by_key.get(key1, []) or system_by_key.get(key2, [])
                if not candidates and f_al == '':
                    for (aid, lid), rows in system_by_key.items():
                        if aid == f_asl:
                            candidates.extend(rows)
                            break

            if not candidates:
                report.append({
                    'match': False,
                    'in_system': False,
                    'in_file_only': True,
                    'file_row': file_row,
                    'system_row': None,
                    'differences': [],
                    'status': 'File only',
                })
                summary['file_only'] += 1
                continue

            # Use first matching system row for comparison (same school may have multiple champions)
            sys_row = candidates[0]
            matched_system_keys.add((str(sys_row.get('asl_school_id') or ''), str(sys_row.get('library_school_id') or '')))

            s_status = (sys_row.get('status') or '').strip().lower()
            s_check = 'yes' if s_status in ('active',) else 'no'
            s_dur = sys_row.get('duration')
            if s_dur is not None:
                try:
                    s_dur = int(float(s_dur))
                except (ValueError, TypeError):
                    s_dur = str(s_dur)
            else:
                s_dur = ''
            s_valid = _normalize_date_for_compare(sys_row.get('expiry_date'))

            system_row = {
                'asl_school_id': str(sys_row.get('asl_school_id') or ''),
                'library_school_id': str(sys_row.get('library_school_id') or ''),
                'champion': sys_row.get('champion', ''),
                'school_name': sys_row.get('school_name', ''),
                'province': sys_row.get('province', ''),
                'asl_scholarship_check': s_check,
                'duration_on_scholarship': s_dur,
                'validity_of_scholarship': s_valid,
            }

            diffs = []
            if f_check != s_check:
                diffs.append('ASL Scholarship check')
            if str(f_dur) != str(s_dur):
                diffs.append('Duration on scholarship')
            if f_valid != s_valid:
                diffs.append('Validity of scholarship')

            is_match = len(diffs) == 0
            if is_match:
                summary['matched'] += 1
            else:
                summary['mismatched'] += 1

            report.append({
                'match': is_match,
                'in_system': True,
                'in_file_only': False,
                'file_row': file_row,
                'system_row': system_row,
                'differences': diffs,
                'status': 'Match' if is_match else 'Mismatch',
            })

        # Optional: system-only rows (in system but not in file)
        for (aid, lid), rows in system_by_key.items():
            if (aid, lid) in matched_system_keys:
                continue
            for sys_row in rows:
                summary['system_only'] += 1
                system_row = {
                    'asl_school_id': str(sys_row.get('asl_school_id') or ''),
                    'library_school_id': str(sys_row.get('library_school_id') or ''),
                    'champion': sys_row.get('champion', ''),
                    'school_name': sys_row.get('school_name', ''),
                    'province': sys_row.get('province', ''),
                    'asl_scholarship_check': 'yes' if (sys_row.get('status') or '').strip().lower() == 'active' else 'no',
                    'duration_on_scholarship': sys_row.get('duration') if sys_row.get('duration') is not None else '',
                    'validity_of_scholarship': _normalize_date_for_compare(sys_row.get('expiry_date')),
                }
                report.append({
                    'match': False,
                    'in_system': True,
                    'in_file_only': False,
                    'file_row': None,
                    'system_row': system_row,
                    'differences': [],
                    'status': 'System only',
                })

        return jsonify({'report': report, 'summary': summary}), 200
    except Exception as e:
        app.logger.exception(e)
        return jsonify({'error': str(e)}), 500


# --- Sales & Marketing: public routes on main app (always registered with routes.py) ---


@app.route('/connect')
def sm_connect_page():
    from app.blueprints.sales_marketing.routes import connect_page
    return connect_page()


@app.route('/connect/e/<slug>')
def sm_connect_page_by_slug(slug):
    from app.blueprints.sales_marketing.sm_roadmap_routes import connect_page_by_slug
    return connect_page_by_slug(slug)


@app.route('/interest')
def sm_interest_redirect():
    return redirect('/connect')


@app.route('/api/public/marketing-events')
def sm_api_public_marketing_events():
    from app.blueprints.sales_marketing.routes import api_public_marketing_events
    return api_public_marketing_events()


@app.route('/api/public/interest-options')
def sm_api_public_interest_options():
    from app.blueprints.sales_marketing.routes import api_public_interest_options
    return api_public_interest_options()


@app.route('/api/public/stakeholder-leads', methods=['POST'])
@csrf.exempt
def sm_api_public_submit_lead():
    from app.blueprints.sales_marketing.routes import api_public_submit_lead
    return api_public_submit_lead()


